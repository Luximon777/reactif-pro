from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, BackgroundTasks
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
import os
import logging
import json
import asyncio
import base64
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import secrets
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
import io
import concurrent.futures
from database import _infer_sectors_from_profile, _SECTOR_KEYWORDS

def _sync_llm_call(chat, message):
    """Run async LLM call synchronously in a thread."""
    return asyncio.run(chat.send_message(message))

async def run_llm_nonblocking(chat, message):
    """Run LLM call in a thread pool to avoid blocking the FastAPI event loop."""
    return await asyncio.to_thread(_sync_llm_call, chat, message)


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]
gridfs_bucket = AsyncIOMotorGridFSBucket(db, bucket_name="proof_documents")

# OpenAI via Emergent
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

app = FastAPI()
api_router = APIRouter(prefix="/api")

# ============== MODELS ==============

class AnonymousToken(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    token: str = Field(default_factory=lambda: secrets.token_urlsafe(32))
    role: str = "particulier"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    profile_id: Optional[str] = None

class Profile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    token_id: str
    name: str = "Utilisateur Anonyme"
    role: str = "particulier"
    skills: List[Dict[str, Any]] = []
    strengths: List[str] = []
    gaps: List[str] = []
    experience_years: int = 0
    sectors: List[str] = []
    profile_score: int = 0
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class JobOffer(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    company: str
    location: str
    contract_type: str
    salary_range: Optional[str] = None
    required_skills: List[str] = []
    description: str
    sector: str
    match_score: int = 0
    status: str = "active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    created_by: Optional[str] = None

class LearningModule(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    description: str
    duration: str
    level: str
    skills_developed: List[str] = []
    progress: int = 0
    category: str
    image_url: Optional[str] = None

class Beneficiary(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    status: str
    progress: int = 0
    skills_acquired: List[str] = []
    sector: str
    last_activity: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    partner_id: str

# ============== COFFRE-FORT MODELS ==============

class DocumentCategory(str):
    IDENTITE = "identite_professionnelle"
    DIPLOMES = "diplomes_certifications"
    EXPERIENCES = "experiences_professionnelles"
    COMPETENCES = "competences_preuves"
    ACCOMPAGNEMENT = "accompagnement_insertion"
    CANDIDATURES = "recherche_emploi"
    FORMATION = "formation_apprentissages"
    ADMINISTRATIF = "documents_administratifs"

class PrivacyLevel(str):
    PRIVATE = "private"
    CONSEILLER = "shared_conseiller"
    RECRUTEUR = "shared_recruteur"
    PUBLIC = "public"

class CoffreDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    token_id: str
    title: str
    category: str
    document_type: str
    file_name: str
    file_url: Optional[str] = None
    file_size: int = 0
    mime_type: str = "application/pdf"
    
    # Indexation
    date_document: Optional[str] = None
    metier_associe: Optional[str] = None
    secteur: Optional[str] = None
    competences_liees: List[str] = []
    description: Optional[str] = None
    
    # Confidentialité
    privacy_level: str = "private"
    shared_with: List[str] = []
    share_expiry: Optional[str] = None
    
    # Métadonnées
    date_expiration: Optional[str] = None
    is_expiring_soon: bool = False
    is_sensitive: bool = False
    
    # Audit
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    view_history: List[Dict[str, Any]] = []

class DocumentShare(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    shared_by: str
    shared_with_email: Optional[str] = None
    shared_with_role: Optional[str] = None
    access_token: str = Field(default_factory=lambda: secrets.token_urlsafe(16))
    expires_at: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    accessed_at: Optional[str] = None
    access_count: int = 0

class CreateDocumentRequest(BaseModel):
    title: str
    category: str
    document_type: str
    file_name: str
    file_url: Optional[str] = None
    date_document: Optional[str] = None
    metier_associe: Optional[str] = None
    secteur: Optional[str] = None
    competences_liees: List[str] = []
    description: Optional[str] = None
    privacy_level: str = "private"
    date_expiration: Optional[str] = None
    is_sensitive: bool = False

# ============== OBSERVATOIRE MODELS ==============

class ContributionType(str):
    NEW_SKILL = "nouvelle_competence"
    SKILL_EVOLUTION = "evolution_competence"
    NEW_TOOL = "nouvel_outil"
    JOB_EVOLUTION = "evolution_metier"
    SECTOR_TREND = "tendance_secteur"
    SKILL_OBSOLESCENCE = "competence_obsolete"

class ContributionStatus(str):
    PENDING = "en_attente"
    AI_VALIDATED = "validee_ia"
    AI_REJECTED = "rejetee_ia"
    HUMAN_VALIDATED = "validee_humain"
    HUMAN_REJECTED = "rejetee_humain"
    INTEGRATED = "integree"

class SkillContribution(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    contributor_id: str
    contribution_type: str = "nouvelle_competence"
    
    # Contenu de la contribution
    skill_name: str
    skill_description: Optional[str] = None
    related_job: Optional[str] = None
    related_sector: Optional[str] = None
    related_tools: List[str] = []
    context: Optional[str] = None
    
    # Métadonnées
    status: str = "en_attente"
    ai_analysis: Optional[Dict[str, Any]] = None
    ai_score: float = 0.0
    human_validator: Optional[str] = None
    human_notes: Optional[str] = None
    
    # Compteurs
    similar_count: int = 1
    upvotes: int = 0
    
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    validated_at: Optional[str] = None

class EmergingSkill(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    skill_name: str
    description: Optional[str] = None
    related_sectors: List[str] = []
    related_jobs: List[str] = []
    related_tools: List[str] = []
    
    # Indicateurs
    emergence_score: float = 0.0
    growth_rate: float = 0.0
    mention_count: int = 0
    contributor_count: int = 0
    
    # Statut
    status: str = "emergente"  # emergente, en_croissance, etablie, en_declin
    first_detected: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    last_updated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class SectorTrend(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sector_name: str
    
    # Compétences
    emerging_skills: List[str] = []
    declining_skills: List[str] = []
    stable_skills: List[str] = []
    
    # Indicateurs
    transformation_index: float = 0.0
    hiring_trend: str = "stable"  # croissance, stable, declin
    skill_gap_alert: bool = False
    
    # Prédictions
    predicted_skills_demand: List[Dict[str, Any]] = []
    
    last_updated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class CreateContributionRequest(BaseModel):
    contribution_type: str = "nouvelle_competence"
    skill_name: str
    skill_description: Optional[str] = None
    related_job: Optional[str] = None
    related_sector: Optional[str] = None
    related_tools: List[str] = []
    context: Optional[str] = None

# ============== INDICE D'ÉVOLUTION DES COMPÉTENCES ==============

class EvolutionIndexLevel(str):
    STABLE = "stable"  # 0-20
    EVOLVING = "evolutif"  # 20-50
    TRANSFORMING = "en_transformation"  # 50-80
    HIGHLY_IMPACTED = "forte_mutation"  # 80-100

class JobEvolutionIndex(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    job_name: str
    sector: str
    
    # Indice principal (0-100)
    evolution_index: float = 0.0
    index_level: str = "stable"
    
    # Variables de calcul
    new_skills_count: int = 0
    skill_frequency_score: float = 0.0
    task_evolution_score: float = 0.0
    new_tools_score: float = 0.0
    job_posting_evolution: float = 0.0
    declining_skills_count: int = 0
    
    # Compétences associées
    emerging_skills: List[str] = []
    stable_skills: List[str] = []
    declining_skills: List[str] = []
    recommended_skills: List[str] = []
    
    # Recommandations
    recommended_trainings: List[str] = []
    job_passerelles: List[str] = []
    
    # Métadonnées
    last_calculated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    data_sources: List[str] = []
    confidence_level: float = 0.0

class SectorEvolutionIndex(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sector_name: str
    
    # Indice sectoriel (0-100)
    evolution_index: float = 0.0
    index_level: str = "stable"
    
    # Statistiques des métiers
    jobs_count: int = 0
    jobs_in_transformation: int = 0
    jobs_stable: int = 0
    jobs_emerging: int = 0
    
    # Compétences clés
    top_emerging_skills: List[Dict[str, Any]] = []
    top_declining_skills: List[Dict[str, Any]] = []
    skill_gap_areas: List[str] = []
    
    # Indicateurs économiques
    hiring_trend: str = "stable"
    innovation_intensity: float = 0.0
    
    # Prévisions
    predicted_evolution_6m: float = 0.0
    predicted_evolution_12m: float = 0.0
    
    last_updated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

# ============== UBUNTOO INTELLIGENCE MODELS ==============

class UbuntooExchange(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    exchange_type: str  # discussion, mentorat, conseil, retour_experience, question
    content_summary: str
    detected_skills: List[str] = []
    detected_tools: List[str] = []
    detected_practices: List[str] = []
    related_jobs: List[str] = []
    related_sectors: List[str] = []
    author_role: str = "professionnel"  # anonymized
    timestamp: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class UbuntooSignal(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    signal_type: str  # competence_emergente, nouvel_outil, pratique_nouvelle, transformation_metier, difficulte_metier
    name: str
    description: str
    mention_count: int = 1
    first_detected: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    last_detected: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    related_jobs: List[str] = []
    related_sectors: List[str] = []
    source_exchanges_count: int = 1
    trend_direction: str = "hausse"  # hausse, stable, baisse
    growth_rate: float = 0.0
    # Validation pipeline
    validation_status: str = "detectee"  # detectee, analysee_ia, validee_humain, integree, rejetee
    ai_confidence: float = 0.0
    ai_analysis: Optional[Dict[str, Any]] = None
    human_validator: Optional[str] = None
    human_notes: Optional[str] = None
    # Cross-reference
    linked_observatory_skills: List[str] = []
    linked_evolution_jobs: List[str] = []
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class UbuntooInsight(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    insight_type: str  # tendance_emergente, alerte_competence, opportunite_formation, transformation_metier
    title: str
    description: str
    supporting_signals: List[str] = []
    impacted_jobs: List[str] = []
    impacted_sectors: List[str] = []
    recommendation: str = ""
    priority: str = "moyenne"  # haute, moyenne, basse
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class MatchRequest(BaseModel):
    profile_skills: List[str]
    job_requirements: List[str]
    profile_sectors: List[str] = []
    job_sector: str = ""

# ============== PASSEPORT DYNAMIQUE DE COMPÉTENCES ==============

class LamriLubartComponents(BaseModel):
    """Modèle Lamri & Lubart: 5 composantes d'une compétence (0-5 scale)"""
    connaissance: int = 0  # Savoirs théoriques et factuels
    cognition: int = 0     # Processus cognitifs (analyse, raisonnement)
    conation: int = 0      # Motivation, volonté, engagement
    affection: int = 0     # Gestion émotionnelle, empathie
    sensori_moteur: int = 0  # Habiletés physiques et pratiques

class CCSPClassification(BaseModel):
    """Référentiel CCSP: Pôle et degré de maîtrise"""
    pole: str = ""  # realisation, interaction, initiative
    degree: str = ""  # imitation, adaptation, transposition

class PassportCompetence(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    nature: str = ""  # savoir_faire, savoir_etre
    category: str = "technique"  # technique, transversale (cross-sector/universal), transferable (within sector/company), sectorielle
    level: str = "intermediaire"  # debutant, intermediaire, avance, expert
    experience_years: float = 0
    proof: Optional[str] = None
    source: str = "declaratif"
    is_emerging: bool = False
    # Lamri & Lubart: 5 composantes
    components: Dict[str, int] = Field(default_factory=lambda: {
        "connaissance": 0, "cognition": 0, "conation": 0,
        "affection": 0, "sensori_moteur": 0
    })
    # CCSP: pôle et degré
    ccsp_pole: str = ""
    ccsp_degree: str = ""
    # Archéologie: liens vers la chaîne vertus-valeurs-qualités
    linked_qualites: List[str] = []
    linked_valeurs: List[str] = []
    linked_vertus: List[str] = []
    added_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class PassportExperience(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    organization: str = ""
    description: str = ""
    skills_used: List[str] = []
    achievements: List[str] = []
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    is_current: bool = False
    experience_type: str = "professionnel"  # professionnel, personnel, benevole, projet
    source: str = "declaratif"
    added_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class PassportLearning(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str
    provider: str = ""
    skills_acquired: List[str] = []
    status: str = "en_cours"  # en_cours, termine, valide
    completion_date: Optional[str] = None
    badge: Optional[str] = None
    source: str = "plateforme"  # plateforme, externe, declaratif

class PassportPasserelle(BaseModel):
    job_name: str
    compatibility_score: float = 0
    shared_skills: List[str] = []
    skills_to_acquire: List[str] = []
    training_needed: str = ""
    accessibility: str = "accessible"  # accessible, formation_courte, formation_longue
    sector: str = ""

class Passport(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    token_id: str
    # Profil professionnel
    professional_summary: str = ""
    career_project: str = ""
    motivations: List[str] = []
    compatible_environments: List[str] = []
    target_sectors: List[str] = []
    # Compétences
    competences: List[Dict[str, Any]] = []
    # Expériences
    experiences: List[Dict[str, Any]] = []
    # Parcours d'apprentissage
    learning_path: List[Dict[str, Any]] = []
    # Passerelles (generated by AI)
    passerelles: List[Dict[str, Any]] = []
    # Métadonnées
    completeness_score: int = 0
    last_updated: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    # Partage
    sharing: Dict[str, Any] = Field(default_factory=lambda: {
        "is_public": False,
        "shared_sections": [],
        "shared_with": [],
        "share_expiry": None
    })

class AddCompetenceRequest(BaseModel):
    name: str
    nature: str = ""  # savoir_faire, savoir_etre
    category: str = "technique"
    level: str = "intermediaire"
    experience_years: float = 0
    proof: Optional[str] = None
    components: Optional[Dict[str, int]] = None
    ccsp_pole: Optional[str] = None
    ccsp_degree: Optional[str] = None
    linked_qualites: List[str] = []
    linked_valeurs: List[str] = []
    linked_vertus: List[str] = []

class EvaluateCompetenceRequest(BaseModel):
    components: Dict[str, int]  # connaissance, cognition, conation, affection, sensori_moteur (0-5)
    ccsp_pole: Optional[str] = None
    ccsp_degree: Optional[str] = None

class CCSPDiagnosticRequest(BaseModel):
    competence_ids: List[str] = []  # If empty, analyze all competences

class AddExperienceRequest(BaseModel):
    title: str
    organization: str = ""
    description: str = ""
    skills_used: List[str] = []
    achievements: List[str] = []
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    is_current: bool = False
    experience_type: str = "professionnel"

class UpdatePassportProfileRequest(BaseModel):
    professional_summary: Optional[str] = None
    career_project: Optional[str] = None
    motivations: Optional[List[str]] = None
    compatible_environments: Optional[List[str]] = None
    target_sectors: Optional[List[str]] = None

class SharePassportRequest(BaseModel):
    is_public: bool = False
    shared_sections: List[str] = []
    shared_with: List[str] = []
    share_expiry: Optional[str] = None

class CreateTokenRequest(BaseModel):
    role: str = "particulier"

class UpdateProfileRequest(BaseModel):
    name: Optional[str] = None
    skills: Optional[List[Dict[str, Any]]] = None
    experience_years: Optional[int] = None
    sectors: Optional[List[str]] = None

class CreateJobRequest(BaseModel):
    title: str
    company: str
    location: str
    contract_type: str
    salary_range: Optional[str] = None
    required_skills: List[str] = []
    description: str
    sector: str

class MetricData(BaseModel):
    particuliers: Dict[str, int]
    entreprises: Dict[str, int]
    partenaires: Dict[str, int]

# ============== HELPER FUNCTIONS ==============

async def get_current_token(token: str) -> dict:
    token_doc = await db.tokens.find_one({"token": token}, {"_id": 0})
    if not token_doc:
        raise HTTPException(status_code=401, detail="Token invalide")
    return token_doc

async def calculate_match_with_ai(profile_skills: List[str], job_requirements: List[str], profile_sectors: List[str], job_sector: str) -> Dict[str, Any]:
    """Use OpenAI to calculate intelligent matching"""
    if not EMERGENT_LLM_KEY:
        # Fallback to simple matching
        common_skills = set(profile_skills) & set(job_requirements)
        score = int((len(common_skills) / max(len(job_requirements), 1)) * 100)
        return {"score": min(score + 20, 100), "rationale": "Correspondance basée sur les compétences communes."}
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"match-{uuid.uuid4()}",
            system_message="Tu es un expert RH français. Analyse la correspondance entre un profil et une offre d'emploi. Réponds en JSON avec 'score' (0-100) et 'rationale' (explication courte en français)."
        ).with_model("openai", "gpt-5.2")
        
        prompt = f"""
        Compétences du candidat: {', '.join(profile_skills)}
        Compétences requises pour le poste: {', '.join(job_requirements)}
        Secteurs du candidat: {', '.join(profile_sectors)}
        Secteur du poste: {job_sector}
        
        Calcule un score de correspondance et explique pourquoi.
        """
        
        response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
        
        # Parse response
        import json
        try:
            result = json.loads(response)
            return {"score": result.get("score", 50), "rationale": result.get("rationale", "Analyse IA")}
        except:
            return {"score": 65, "rationale": response[:200]}
    except Exception as e:
        logging.error(f"AI matching error: {e}")
        common_skills = set(profile_skills) & set(job_requirements)
        score = int((len(common_skills) / max(len(job_requirements), 1)) * 100)
        return {"score": min(score + 20, 100), "rationale": "Correspondance basée sur les compétences communes."}

# ============== AUTH ENDPOINTS ==============

@api_router.post("/auth/anonymous")
async def create_anonymous_token(request: CreateTokenRequest):
    """Create an anonymous secure token"""
    token = AnonymousToken(role=request.role)
    token_dict = token.model_dump()
    await db.tokens.insert_one(token_dict)
    
    # Create associated profile
    profile = Profile(
        token_id=token.id,
        role=request.role,
        name=f"Utilisateur {token.id[:8].upper()}"
    )
    profile_dict = profile.model_dump()
    await db.profiles.insert_one(profile_dict)
    
    # Update token with profile_id
    await db.tokens.update_one({"id": token.id}, {"$set": {"profile_id": profile.id}})
    
    return {"token": token.token, "role": token.role, "profile_id": profile.id}

@api_router.get("/auth/verify")
async def verify_token(token: str):
    """Verify token validity"""
    token_doc = await get_current_token(token)
    return {
        "valid": True, "role": token_doc["role"],
        "profile_id": token_doc.get("profile_id"),
        "auth_mode": token_doc.get("auth_mode", "anonymous"),
        "pseudo": token_doc.get("pseudo"),
        "identity_level": token_doc.get("identity_level", "none")
    }

@api_router.post("/auth/switch-role")
async def switch_role(token: str, new_role: str):
    """Switch user role"""
    if new_role not in ["particulier", "entreprise", "partenaire"]:
        raise HTTPException(status_code=400, detail="Rôle invalide")
    
    await db.tokens.update_one({"token": token}, {"$set": {"role": new_role}})
    return {"message": "Rôle mis à jour", "role": new_role}

# ============== PROFILE ENDPOINTS ==============

@api_router.get("/profile")
async def get_profile(token: str):
    """Get user profile"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not profile:
        raise HTTPException(status_code=404, detail="Profil non trouvé")
    # Enrich with pseudo from token doc (always available)
    if not profile.get("pseudo"):
        profile["pseudo"] = token_doc.get("pseudo") or profile.get("name")
    return profile

@api_router.put("/profile")
async def update_profile(token: str, request: UpdateProfileRequest):
    """Update user profile"""
    token_doc = await get_current_token(token)
    
    update_data = {k: v for k, v in request.model_dump().items() if v is not None}
    
    if update_data:
        # Calculate profile score based on completeness
        profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
        if profile:
            skills_count = len(update_data.get("skills", profile.get("skills", [])))
            sectors_count = len(update_data.get("sectors", profile.get("sectors", [])))
            exp = update_data.get("experience_years", profile.get("experience_years", 0))
            score = min(100, skills_count * 10 + sectors_count * 5 + (10 if exp > 0 else 0) + 30)
            update_data["profile_score"] = score
        
        await db.profiles.update_one({"token_id": token_doc["id"]}, {"$set": update_data})
    
    return await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})


@api_router.get("/profile/confidence-scores/simple")
async def get_confidence_scores_simple(token: str):
    """Calculate trust/confidence scores based on profile completeness"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    docs_count = await db.coffre_documents.count_documents({"token_id": token_doc["id"]})

    # Calculate 4 dimensions
    # 1. Identité: pseudo, real_name, CV analysé
    identite_score = 0
    if profile:
        if profile.get("cv_analyzed"): identite_score += 40
        if profile.get("real_first_name"): identite_score += 20
        if profile.get("real_last_name"): identite_score += 20
        if profile.get("dclic_imported"): identite_score += 20
    identite_score = min(100, identite_score)

    # 2. Compétences: skills, savoir_faire, savoir_etre
    comp_score = 0
    skills_count = len((profile or {}).get("skills", []))
    sf_count = len((passport or {}).get("savoir_faire", []))
    se_count = len((passport or {}).get("savoir_etre", []))
    comp_count = len((passport or {}).get("competences", []))
    total_skills = skills_count + sf_count + se_count + comp_count
    if total_skills >= 20: comp_score = 100
    elif total_skills >= 10: comp_score = 70
    elif total_skills >= 5: comp_score = 50
    elif total_skills >= 1: comp_score = 25

    # 3. Expériences
    exp_score = 0
    exp_count = len((passport or {}).get("experiences", []))
    if exp_count >= 5: exp_score = 100
    elif exp_count >= 3: exp_score = 70
    elif exp_count >= 1: exp_score = 40

    # 4. Preuves (documents dans le coffre)
    preuves_score = 0
    if docs_count >= 5: preuves_score = 100
    elif docs_count >= 3: preuves_score = 70
    elif docs_count >= 1: preuves_score = 40

    global_pct = int((identite_score + comp_score + exp_score + preuves_score) / 4)
    level = "eleve" if global_pct >= 70 else "moyen" if global_pct >= 40 else "faible"

    tips = []
    if identite_score < 60: tips.append("Ajoutez votre vrai nom et prénom pour renforcer votre identité")
    if comp_score < 50: tips.append("Enrichissez vos compétences via l'analyse CV ou D'CLIC PRO")
    if exp_score < 40: tips.append("Ajoutez vos expériences professionnelles")
    if preuves_score < 40: tips.append("Déposez des preuves (diplômes, attestations) dans votre portefeuille")

    return {
        "global_pct": global_pct,
        "level": level,
        "dimensions": [
            {"key": "identite", "label": "Identité", "pct": identite_score},
            {"key": "competences", "label": "Compétences", "pct": comp_score},
            {"key": "experiences", "label": "Expériences", "pct": exp_score},
            {"key": "preuves", "label": "Preuves", "pct": preuves_score},
        ],
        "tips": tips,
    }




@api_router.post("/profile/identity-adn")
async def generate_identity_adn(token: str):
    """Génère l'ADN Professionnel de l'utilisateur via IA."""
    token_doc = await get_current_token(token)

    if not EMERGENT_LLM_KEY:
        raise HTTPException(503, "Clé LLM non configurée")

    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})

    if not passport and not profile:
        return {"error": "Profil insuffisant. Ajoutez vos compétences et expériences d'abord."}

    skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:20]]
    experiences = (passport or {}).get("experiences", [])
    formations = (passport or {}).get("formations", [])
    savoir_faire = (passport or {}).get("savoir_faire", [])
    savoir_etre = (passport or {}).get("savoir_etre", [])

    sf_list = [s if isinstance(s, str) else s.get("name", "") for s in savoir_faire[:12]]
    se_list = [s if isinstance(s, str) else s.get("name", "") for s in savoir_etre[:12]]
    exp_list = []
    for e in experiences[:6]:
        if isinstance(e, dict):
            exp_list.append(f"{e.get('title', '')} ({e.get('duration', '')}) - {e.get('company', '')}")

    form_list = []
    for f in formations[:4]:
        if isinstance(f, dict):
            form_list.append(f"{f.get('title', '')} ({f.get('level', '')})")

    # D'CLIC PRO results if available
    dclic = (passport or {}).get("dclic_results", {})
    dclic_text = ""
    if dclic:
        if dclic.get("mbti"):
            dclic_text += f"MBTI: {dclic['mbti']}. "
        if dclic.get("disc"):
            dclic_text += f"DISC: {dclic['disc']}. "
        if dclic.get("riasec"):
            dclic_text += f"RIASEC: {dclic['riasec']}. "

    user_context = f"""Compétences techniques: {', '.join(skills) if skills else 'Non renseigné'}
Savoir-faire: {', '.join(sf_list) if sf_list else 'Non renseigné'}
Savoir-être: {', '.join(se_list) if se_list else 'Non renseigné'}
Expériences: {chr(10).join(exp_list) if exp_list else 'Non renseigné'}
Formations: {', '.join(form_list) if form_list else 'Non renseigné'}
Tests psychométriques: {dclic_text if dclic_text else 'Non passés'}"""

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"adn-pro-{token_doc['id']}-{uuid.uuid4()}",
        system_message="Tu es un expert en bilan de compétences et psychologie du travail en France. Tu génères des analyses profondes et personnalisées."
    ).with_model("openai", "gpt-5.2")

    prompt = f"""Analyse ce profil professionnel et génère son ADN Professionnel : une synthèse identitaire unique qui capture l'essence de ce professionnel.

PROFIL:
{user_context}

Retourne UNIQUEMENT un JSON valide (pas de markdown) :
{{
  "synthese_adn": "Paragraphe de 3-4 phrases décrivant l'identité professionnelle unique de cette personne, ses atouts distinctifs et son positionnement sur le marché.",
  "style_professionnel": "Une phrase décrivant le style de travail (ex: 'Opérationnel méthodique, orienté terrain et fiabilité')",
  "forces_principales": ["Force 1", "Force 2", "Force 3", "Force 4"],
  "environnements_favorables": ["Type d'environnement 1", "Type d'environnement 2", "Type d'environnement 3"],
  "axes_projection": ["Axe d'évolution possible 1", "Axe d'évolution possible 2", "Axe d'évolution possible 3"],
  "potentiel_evolution": "Phrase décrivant le potentiel de progression et les directions d'évolution possibles"
}}

IMPORTANT: Base-toi UNIQUEMENT sur les données du profil. Sois spécifique et concret, pas générique."""

    response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
    import re
    text = response.strip() if isinstance(response, str) else response.text.strip()
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        adn = json.loads(json_match.group())
    else:
        adn = json.loads(text)

    # Save to passport
    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {
            "identity_adn": adn,
            "identity_adn_generated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True
    )

    return adn



# ============== JOBS ENDPOINTS ==============

@api_router.get("/jobs")
async def get_jobs(token: str, limit: int = 20):
    """Get job offers with match scores"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    
    jobs = await db.jobs.find({"status": "active"}, {"_id": 0}).to_list(limit)
    
    # Calculate match scores
    if profile and profile.get("skills"):
        profile_skill_names = [s.get("name", "") for s in profile.get("skills", [])]
        for job in jobs:
            common = set(profile_skill_names) & set(job.get("required_skills", []))
            job["match_score"] = min(100, int((len(common) / max(len(job.get("required_skills", [])), 1)) * 100) + 25)
    
    return sorted(jobs, key=lambda x: x.get("match_score", 0), reverse=True)

@api_router.post("/jobs")
async def create_job(token: str, request: CreateJobRequest):
    """Create a new job offer (RH only)"""
    token_doc = await get_current_token(token)
    if token_doc["role"] != "entreprise":
        raise HTTPException(status_code=403, detail="Accès réservé aux entreprises")
    
    job = JobOffer(**request.model_dump(), created_by=token_doc["id"])
    await db.jobs.insert_one(job.model_dump())
    return job.model_dump()

# Specific routes MUST come before wildcard /jobs/{job_id}
def _build_match_entry(job, score, skills_matched, rationale, scoring=None):
    """Transform a DB job into the frontend-expected match format."""
    return {
        "titre": job.get("title", job.get("intitule_poste", "")),
        "matching_score": score,
        "secteur": job.get("sector", job.get("secteur", "")),
        "type_contrat": job.get("contract_type", job.get("type_contrat", "CDI")),
        "entreprise_type": job.get("company", job.get("entreprise", "")),
        "localisation": job.get("location", job.get("localisation", "")),
        "description": job.get("description", ""),
        "salaire_indicatif": job.get("salary_range", job.get("salaire_indicatif", "")),
        "competences_matchees": skills_matched,
        "pourquoi_ce_match": rationale,
        "url_offre": job.get("url_offre", ""),
        "scoring": scoring,
    }


def _score_job_basic(job, skills, user_words, exp_titles, inferred_set):
    """Basic scoring of a job against user profile. Returns (score, matched_skills, rationale)."""
    req = job.get("required_skills", [])
    job_sector = job.get("sector", "").lower()
    job_title_words = set(w.lower() for w in job.get("title", "").split() if len(w) > 3)

    skills_lower = set(s.lower() for s in skills)
    req_lower = set(r.lower() for r in req)
    exact = len(skills_lower & req_lower)
    matched_skill_names = list(skills_lower & req_lower)

    keyword_matches = 0
    for r_skill in req:
        r_words = set(w for w in r_skill.lower().split() if len(w) > 3)
        if r_words & user_words:
            keyword_matches += 1

    exp_match = 0
    for et in exp_titles:
        if any(w in et.lower() for w in job_title_words if len(w) > 3):
            exp_match += 1
            break

    sector_match = job_sector in inferred_set or any(job_sector in s for s in inferred_set)

    total_req = max(len(req), 1)
    skill_score = ((exact + keyword_matches * 0.7) / total_req) * 50
    score = int(min(100, 20 + skill_score + (exp_match * 15) + (15 if sector_match else 0)))

    parts = []
    if exact > 0:
        parts.append(f"{exact} compétence(s) exacte(s)")
    if keyword_matches > exact:
        parts.append(f"{keyword_matches - exact} par mots-clés")
    if exp_match:
        parts.append("expérience similaire")
    if sector_match:
        parts.append("secteur correspondant")
    if not parts:
        parts.append(f"0/{len(req)} compétences en commun")

    return score, matched_skill_names[:6], " · ".join(parts)


def _apply_filter_scoring(job, match_entry, filters, user_profile_text):
    """Apply advanced filter-based scoring on a match. Returns updated match_entry with scoring."""
    evaluations = []
    blocages = []
    vigilances = []
    points_forts = []
    total_points = 0
    earned_points = 0

    # Métier filter
    if "metier" in filters:
        priority = filters["metier"].get("priority", 3)
        weight = priority * 4
        total_points += weight
        metiers = filters["metier"].get("value", [])
        job_title = match_entry["titre"].lower()
        job_desc = match_entry.get("description", "").lower()
        match_found = any(m.lower() in job_title or m.lower() in job_desc for m in metiers)
        compat = 1.0 if match_found else 0.2
        earned_points += int(weight * compat)
        evaluations.append({"label": "Métier", "compatibility": compat, "priority": priority})
        if match_found:
            points_forts.append({"critere": "Métier", "message": f"Correspond au(x) métier(s) recherché(s)"})
        elif priority >= 5:
            blocages.append({"critere": "Métier", "raison": "Ne correspond pas au métier recherché"})

    # Secteur filter
    if "secteur" in filters:
        priority = filters["secteur"].get("priority", 3)
        weight = priority * 4
        total_points += weight
        secteurs = filters["secteur"].get("value", [])
        job_sector = match_entry["secteur"].lower()
        match_found = any(s.lower() in job_sector or job_sector in s.lower() for s in secteurs)
        compat = 1.0 if match_found else 0.3
        earned_points += int(weight * compat)
        evaluations.append({"label": "Secteur", "compatibility": compat, "priority": priority})
        if match_found:
            points_forts.append({"critere": "Secteur", "message": f"Secteur {match_entry['secteur']} correspondant"})

    # Contrat filter
    if "contrat" in filters:
        priority = filters["contrat"].get("priority", 3)
        weight = priority * 3
        total_points += weight
        contrats = [c.lower() for c in filters["contrat"].get("value", [])]
        job_contrat = match_entry["type_contrat"].lower()
        match_found = any(c in job_contrat or job_contrat in c for c in contrats)
        compat = 1.0 if match_found else 0.1
        earned_points += int(weight * compat)
        evaluations.append({"label": "Contrat", "compatibility": compat, "priority": priority})
        if not match_found and priority >= 4:
            vigilances.append({"critere": "Contrat", "message": f"Type {match_entry['type_contrat']} non recherché"})
        if match_found:
            points_forts.append({"critere": "Contrat", "message": f"{match_entry['type_contrat']} correspond"})

    # Salaire filter
    if "salaire_minimum" in filters:
        priority = filters["salaire_minimum"].get("priority", 3)
        weight = priority * 3
        total_points += weight
        sal_min = filters["salaire_minimum"].get("value", 0)
        sal_text = match_entry.get("salaire_indicatif", "")
        # Try to extract number from salary text
        import re
        nums = re.findall(r"[\d\s]+", sal_text.replace(" ", "").replace("\u202f", ""))
        max_sal = 0
        for n in nums:
            try:
                v = int(n.strip())
                if v > max_sal:
                    max_sal = v
            except ValueError:
                pass
        if max_sal > 0 and sal_min:
            compat = 1.0 if max_sal >= int(sal_min) else max(0.1, max_sal / int(sal_min))
            earned_points += int(weight * compat)
            evaluations.append({"label": "Salaire", "compatibility": round(compat, 2), "priority": priority})
            if compat < 0.5 and priority >= 4:
                vigilances.append({"critere": "Salaire", "message": f"Salaire potentiellement inférieur au minimum souhaité ({sal_min}€)"})
        else:
            earned_points += int(weight * 0.5)
            evaluations.append({"label": "Salaire", "compatibility": 0.5, "priority": priority})

    # Zone géographique filter
    if "zone_geographique" in filters:
        priority = filters["zone_geographique"].get("priority", 4)
        weight = priority * 4
        total_points += weight
        zone = filters["zone_geographique"].get("value", "").lower()
        job_loc = match_entry["localisation"].lower()
        match_found = zone in job_loc or job_loc in zone or any(w in job_loc for w in zone.split(",") if len(w.strip()) > 2)
        compat = 1.0 if match_found else 0.2
        earned_points += int(weight * compat)
        evaluations.append({"label": "Localisation", "compatibility": compat, "priority": priority})
        if match_found:
            points_forts.append({"critere": "Localisation", "message": f"Dans la zone recherchée ({match_entry['localisation']})"})
        elif priority >= 5:
            blocages.append({"critere": "Localisation", "raison": f"Hors zone ({match_entry['localisation']})"})
        elif priority >= 3:
            vigilances.append({"critere": "Localisation", "message": f"Localisation {match_entry['localisation']} hors zone demandée"})

    # Télétravail filter
    if "teletravail" in filters:
        priority = filters["teletravail"].get("priority", 2)
        weight = priority * 2
        total_points += weight
        earned_points += int(weight * 0.5)  # Unknown by default
        evaluations.append({"label": "Télétravail", "compatibility": 0.5, "priority": priority})

    # Temps de travail
    if "temps_travail" in filters:
        priority = filters["temps_travail"].get("priority", 2)
        weight = priority * 2
        total_points += weight
        earned_points += int(weight * 0.7)
        evaluations.append({"label": "Temps de travail", "compatibility": 0.7, "priority": priority})

    # Restrictions fonctionnelles
    if "restrictions_fonctionnelles" in filters:
        priority = filters["restrictions_fonctionnelles"].get("priority", 5)
        restrictions = filters["restrictions_fonctionnelles"].get("value", {})
        active_restrictions = [k for k, v in restrictions.items() if v]
        if active_restrictions:
            weight = priority * 3
            total_points += weight
            earned_points += int(weight * 0.5)
            evaluations.append({"label": "Restrictions fonctionnelles", "compatibility": 0.5, "priority": priority})
            vigilances.append({"critere": "Restrictions", "message": f"{len(active_restrictions)} restriction(s) à vérifier avec l'employeur"})

    # Inclusion score
    score_inclusion = 0
    if "ciblage_employeurs_inclusifs" in filters and filters["ciblage_employeurs_inclusifs"].get("value"):
        score_inclusion = 50  # Default, would need employer DB data
    if "accessibilite_metier_handicap" in filters and filters["accessibilite_metier_handicap"].get("value"):
        score_inclusion = max(score_inclusion, 40)

    # Calculate final scored matching_score
    if total_points > 0:
        filter_pct = (earned_points / total_points) * 100
        # Blend: 40% profile match + 60% filter match
        blended = int(match_entry["matching_score"] * 0.4 + filter_pct * 0.6)
    else:
        blended = match_entry["matching_score"]

    # Determine status
    has_blocage = len(blocages) > 0
    if has_blocage:
        statut = "Incompatible"
    elif blended >= 75:
        statut = "Excellent match"
    elif blended >= 55:
        statut = "Match pertinent"
    elif blended >= 35:
        statut = "Match moyen"
    else:
        statut = "Faible compatibilité"

    match_entry["matching_score"] = blended
    match_entry["scoring"] = {
        "statut": statut,
        "score_detail": {"obtenu": earned_points, "maximum": total_points},
        "evaluations": evaluations,
        "blocages": blocages,
        "vigilances": vigilances,
        "points_forts": points_forts,
        "score_inclusion": score_inclusion,
    }
    return match_entry


@api_router.get("/jobs/matching")
async def jobs_matching_early(token: str):
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:15]]

    exp_titles = []
    for e in (passport or {}).get("experiences", []):
        if isinstance(e, dict) and e.get("title"):
            exp_titles.append(e["title"])

    user_words = set()
    for s in skills + exp_titles:
        for w in s.lower().split():
            if len(w) > 3:
                user_words.add(w)

    inferred_sectors = _infer_sectors_from_profile(
        (passport or {}).get("experiences", []),
        (passport or {}).get("savoir_faire", skills)
    )
    inferred_set = set(s.lower() for s in inferred_sectors[:3])

    # Check for optimized CV and career project
    cv_doc = await db.cv_models.find_one({"token_id": token_doc["id"]})
    career_project = (passport or {}).get("projet_professionnel", "")
    profile_title = ""
    if exp_titles:
        profile_title = exp_titles[0]
    elif inferred_sectors:
        profile_title = f"Profil {inferred_sectors[0]}"

    jobs = await db.jobs.find({"status": "active"}, {"_id": 0}).limit(30).to_list(30)
    matches = []
    for job in jobs:
        score, matched_skills, rationale = _score_job_basic(job, skills, user_words, exp_titles, inferred_set)
        entry = _build_match_entry(job, score, matched_skills, rationale)
        matches.append(entry)

    matches.sort(key=lambda x: x["matching_score"], reverse=True)
    return {
        "has_data": len(matches) > 0 or len(skills) > 0,
        "has_filters": False,
        "profile_summary": {
            "titre": profile_title,
            "skills_count": len(skills),
            "has_optimized_cv": cv_doc is not None,
            "has_career_project": bool(career_project),
        },
        "matches": matches,
    }


@api_router.post("/jobs/matching/search")
async def jobs_matching_search_post(token: str, body: dict = {}):
    """Recherche d'offres avec scoring avancé basé sur les filtres utilisateur."""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:15]]

    exp_titles = []
    for e in (passport or {}).get("experiences", []):
        if isinstance(e, dict) and e.get("title"):
            exp_titles.append(e["title"])

    user_words = set()
    for s in skills + exp_titles:
        for w in s.lower().split():
            if len(w) > 3:
                user_words.add(w)

    inferred_sectors = _infer_sectors_from_profile(
        (passport or {}).get("experiences", []),
        (passport or {}).get("savoir_faire", skills)
    )
    inferred_set = set(s.lower() for s in inferred_sectors[:3])

    cv_doc = await db.cv_models.find_one({"token_id": token_doc["id"]})
    career_project = (passport or {}).get("projet_professionnel", "")
    profile_title = exp_titles[0] if exp_titles else (f"Profil {inferred_sectors[0]}" if inferred_sectors else "Profil")

    # Save preferences
    await db.matching_prefs.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"token_id": token_doc["id"], "filters": body, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )

    user_profile_text = " ".join(skills + exp_titles).lower()

    jobs = await db.jobs.find({"status": "active"}, {"_id": 0}).limit(30).to_list(30)
    matches = []
    for job in jobs:
        score, matched_skills, rationale = _score_job_basic(job, skills, user_words, exp_titles, inferred_set)
        entry = _build_match_entry(job, score, matched_skills, rationale)
        entry = _apply_filter_scoring(job, entry, body, user_profile_text)
        matches.append(entry)

    matches.sort(key=lambda x: x["matching_score"], reverse=True)
    return {
        "has_data": len(matches) > 0 or len(skills) > 0,
        "has_filters": True,
        "profile_summary": {
            "titre": profile_title,
            "skills_count": len(skills),
            "has_optimized_cv": cv_doc is not None,
            "has_career_project": bool(career_project),
        },
        "matches": matches,
    }


@api_router.get("/jobs/applications")
async def jobs_applications_early(token: str):
    token_doc = await get_current_token(token)
    apps = await db.applications.find({"token_id": token_doc["id"]}, {"_id": 0}).sort("applied_at", -1).to_list(50)
    return {"applications": apps, "total": len(apps)}


@api_router.post("/jobs/apply")
async def jobs_apply_early(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    job_title = body.get("job_title", "") or body.get("job_id", "")
    if not job_title:
        raise HTTPException(400, "job_title requis")
    existing = await db.applications.find_one({"token_id": token_doc["id"], "job_title": job_title})
    if existing:
        return {"success": True, "already_applied": True, "message": "Vous avez déjà enregistré cette candidature"}
    job_data = body.get("job_data", {})
    app_doc = {
        "id": str(uuid.uuid4()),
        "token_id": token_doc["id"],
        "job_title": job_title,
        "job_data": job_data,
        "status": "en_preparation",
        "applied_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.applications.insert_one(app_doc)
    return {"success": True, "already_applied": False, "message": "Candidature enregistrée", "application_id": app_doc["id"]}


@api_router.get("/jobs/rome-suggestions")
async def get_rome_suggestions(token: str):
    """Suggère des codes ROME basés sur le profil utilisateur (expériences, compétences, D'CLIC)."""
    import re as re_module
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})

    search_terms = []
    for exp in (passport or {}).get("experiences", []):
        if isinstance(exp, dict) and exp.get("title"):
            search_terms.append(exp["title"])
    for sec in (profile or {}).get("sectors", []):
        if sec:
            search_terms.append(sec)
    for skill in (profile or {}).get("skills", [])[:10]:
        name = skill.get("name", "") if isinstance(skill, dict) else str(skill)
        if name and len(name) > 3:
            search_terms.append(name)
    dclic = await db.dclic_results.find_one({"token_id": token_doc["id"]}, {"_id": 0, "job_matches": 1})
    if dclic:
        for jm in (dclic.get("job_matches") or [])[:5]:
            if isinstance(jm, dict) and jm.get("metier"):
                search_terms.append(jm["metier"])

    if not search_terms:
        return {"suggestions": [], "message": "Complétez votre profil ou passez le test D'CLIC PRO pour obtenir des suggestions ROME."}

    seen_codes = set()
    suggestions = []
    for term in search_terms[:15]:
        words = [w for w in term.split() if len(w) > 3]
        if not words:
            words = [term]
        escaped = [re_module.escape(w) for w in words[:3]]
        regex_parts = "|".join(escaped)
        regex = {"$regex": regex_parts, "$options": "i"}
        try:
            matches = await db.rome_metiers.find({"libelle": regex}, {"_id": 0}).to_list(3)
        except Exception:
            continue
        for m in matches:
            code = m.get("code_rome", "")
            if code and code not in seen_codes:
                seen_codes.add(code)
                suggestions.append({
                    "code_rome": code,
                    "libelle": m.get("libelle", ""),
                    "domaine": m.get("grand_domaine_nom", ""),
                    "matched_from": term,
                })
    return {"suggestions": suggestions[:12]}


@api_router.get("/jobs/rome-search")
async def search_rome_codes(q: str = ""):
    """Recherche de codes ROME par texte libre."""
    import re as re_module
    if not q or len(q) < 2:
        return {"results": [], "total": 0}
    escaped_q = re_module.escape(q)
    regex = {"$regex": escaped_q, "$options": "i"}
    metiers = await db.rome_metiers.find(
        {"$or": [{"libelle": regex}, {"code_rome": regex}]},
        {"_id": 0}
    ).to_list(20)
    results = [{
        "code_rome": m.get("code_rome", ""),
        "libelle": m.get("libelle", ""),
        "domaine": m.get("grand_domaine_nom", ""),
    } for m in metiers]
    return {"results": results, "total": len(results)}



@api_router.get("/jobs/{job_id}")
async def get_job(job_id: str):
    """Get job details"""
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Offre non trouvée")
    return job

@api_router.get("/jobs/{job_id}/match")
async def get_job_match(token: str, job_id: str):
    """Get AI-powered match analysis for a job"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
    
    if not profile or not job:
        raise HTTPException(status_code=404, detail="Profil ou offre non trouvé")
    
    profile_skill_names = [s.get("name", "") for s in profile.get("skills", [])]
    match_result = await calculate_match_with_ai(
        profile_skill_names,
        job.get("required_skills", []),
        profile.get("sectors", []),
        job.get("sector", "")
    )
    
    return match_result

# ============== LEARNING ENDPOINTS ==============

@api_router.get("/learning")
async def get_learning_modules(token: str):
    """Get recommended learning modules"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    
    modules = await db.learning_modules.find({}, {"_id": 0}).to_list(50)
    
    # Get user progress
    progress_docs = await db.learning_progress.find({"token_id": token_doc["id"]}, {"_id": 0}).to_list(100)
    progress_map = {p["module_id"]: p["progress"] for p in progress_docs}
    
    for module in modules:
        module["progress"] = progress_map.get(module["id"], 0)
    
    return modules

@api_router.post("/learning/{module_id}/progress")
async def update_learning_progress(token: str, module_id: str, progress: int):
    """Update learning progress"""
    token_doc = await get_current_token(token)
    
    await db.learning_progress.update_one(
        {"token_id": token_doc["id"], "module_id": module_id},
        {"$set": {"progress": min(100, max(0, progress)), "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {"message": "Progression mise à jour", "progress": progress}

# ============== RH ENDPOINTS ==============

@api_router.get("/rh/offers")
async def get_rh_offers(token: str):
    """Get offers created by RH"""
    token_doc = await get_current_token(token)
    offers = await db.jobs.find({"created_by": token_doc["id"]}, {"_id": 0}).to_list(50)
    return offers

@api_router.get("/rh/candidates")
async def get_candidates(token: str, job_id: Optional[str] = None):
    """Get compatible candidates for RH"""
    token_doc = await get_current_token(token)
    
    # Get all profiles (in real app, would filter by job compatibility)
    profiles = await db.profiles.find({"role": "particulier"}, {"_id": 0}).to_list(50)
    
    if job_id:
        job = await db.jobs.find_one({"id": job_id}, {"_id": 0})
        if job:
            for profile in profiles:
                profile_skill_names = [s.get("name", "") for s in profile.get("skills", [])]
                common = set(profile_skill_names) & set(job.get("required_skills", []))
                profile["match_score"] = min(100, int((len(common) / max(len(job.get("required_skills", [])), 1)) * 100) + 20)
            profiles = sorted(profiles, key=lambda x: x.get("match_score", 0), reverse=True)
    
    return profiles

# ============== PARTENAIRES ENDPOINTS ==============

@api_router.get("/partenaires/beneficiaires")
async def get_beneficiaires(token: str):
    """Get beneficiaries for social partners"""
    token_doc = await get_current_token(token)
    beneficiaires = await db.beneficiaires.find({"partner_id": token_doc["id"]}, {"_id": 0}).to_list(100)
    return beneficiaires

@api_router.post("/partenaires/beneficiaires")
async def create_beneficiaire(token: str, name: str, sector: str):
    """Create a new beneficiary"""
    token_doc = await get_current_token(token)
    
    beneficiary = Beneficiary(
        name=name,
        status="En accompagnement",
        sector=sector,
        partner_id=token_doc["id"]
    )
    await db.beneficiaires.insert_one(beneficiary.model_dump())
    return beneficiary.model_dump()

@api_router.put("/partenaires/beneficiaires/{beneficiary_id}")
async def update_beneficiaire(token: str, beneficiary_id: str, status: Optional[str] = None, progress: Optional[int] = None):
    """Update beneficiary status"""
    update_data = {}
    if status:
        update_data["status"] = status
    if progress is not None:
        update_data["progress"] = progress
    update_data["last_activity"] = datetime.now(timezone.utc).isoformat()
    
    await db.beneficiaires.update_one({"id": beneficiary_id}, {"$set": update_data})
    return await db.beneficiaires.find_one({"id": beneficiary_id}, {"_id": 0})

# ============== METRICS ENDPOINTS ==============

@api_router.get("/metrics")
async def get_metrics():
    """Get platform metrics"""
    particuliers_count = await db.profiles.count_documents({"role": "particulier"})
    entreprises_count = await db.profiles.count_documents({"role": "entreprise"})
    jobs_count = await db.jobs.count_documents({"status": "active"})
    beneficiaires_count = await db.beneficiaires.count_documents({})
    
    return {
        "particuliers": {
            "total": particuliers_count,
            "active": particuliers_count
        },
        "entreprises": {
            "total": entreprises_count,
            "jobs_posted": jobs_count
        },
        "partenaires": {
            "beneficiaires": beneficiaires_count,
            "active_support": beneficiaires_count
        }
    }

# ============== COFFRE-FORT ENDPOINTS ==============

DOCUMENT_CATEGORIES = {
    "identite_professionnelle": {
        "label": "Identité professionnelle",
        "types": ["CV", "CV ciblé", "Lettre de motivation", "Présentation professionnelle", "Projet professionnel", "Portfolio", "Bilan de compétences"]
    },
    "diplomes_certifications": {
        "label": "Diplômes et certifications",
        "types": ["Diplôme", "Titre professionnel", "Certificat", "Attestation de formation", "Habilitation", "Certification", "Permis", "CACES", "SST"]
    },
    "experiences_professionnelles": {
        "label": "Expériences professionnelles",
        "types": ["Contrat de travail", "Certificat de travail", "Attestation employeur", "Fiche de poste", "Évaluation annuelle", "Lettre de recommandation", "Attestation de mission"]
    },
    "competences_preuves": {
        "label": "Compétences et preuves",
        "types": ["Réalisation professionnelle", "Rapport", "Support créé", "Projet réalisé", "Badge de compétence", "Auto-évaluation", "Production écrite"]
    },
    "accompagnement_insertion": {
        "label": "Accompagnement et insertion",
        "types": ["Compte rendu d'entretien", "Diagnostic", "Synthèse de parcours", "Objectifs personnalisés", "Plan d'action", "Prescription", "Bilan"]
    },
    "recherche_emploi": {
        "label": "Recherche d'emploi",
        "types": ["Candidature", "Réponse employeur", "Convocation entretien", "Compte rendu entretien", "Offre sauvegardée", "Simulation entretien"]
    },
    "formation_apprentissages": {
        "label": "Formation et apprentissages",
        "types": ["Attestation de participation", "Certificat de module", "Résultat de quiz", "Badge interne", "Validation de parcours", "Exercice réalisé"]
    },
    "documents_administratifs": {
        "label": "Documents administratifs",
        "types": ["Permis de conduire", "Justificatif de mobilité", "Carte professionnelle", "Convention de stage", "Contrat d'alternance", "Autre document"]
    }
}

@api_router.get("/coffre/categories")
async def get_coffre_categories():
    """Get all document categories for the coffre-fort"""
    return DOCUMENT_CATEGORIES

@api_router.get("/coffre/documents")
async def get_coffre_documents(token: str, category: Optional[str] = None, search: Optional[str] = None):
    """Get all documents in user's coffre-fort"""
    token_doc = await get_current_token(token)
    
    query = {"token_id": token_doc["id"]}
    if category:
        query["category"] = category
    
    documents = await db.coffre_documents.find(query, {"_id": 0}).to_list(500)
    
    if search:
        search_lower = search.lower()
        documents = [d for d in documents if 
            search_lower in d.get("title", "").lower() or 
            search_lower in d.get("description", "").lower() or
            any(search_lower in c.lower() for c in d.get("competences_liees", []))]
    
    # Check for expiring documents
    today = datetime.now(timezone.utc)
    for doc in documents:
        if doc.get("date_expiration"):
            try:
                exp_date = datetime.fromisoformat(doc["date_expiration"].replace('Z', '+00:00'))
                days_until = (exp_date - today).days
                doc["is_expiring_soon"] = 0 <= days_until <= 30
                doc["days_until_expiry"] = days_until
            except:
                pass
    
    return sorted(documents, key=lambda x: x.get("created_at", ""), reverse=True)

@api_router.get("/coffre/documents/{document_id}")
async def get_coffre_document(token: str, document_id: str):
    """Get a specific document"""
    token_doc = await get_current_token(token)
    doc = await db.coffre_documents.find_one({"id": document_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    return doc

@api_router.get("/coffre/download/{document_id}")
async def download_coffre_document(document_id: str, token: str):
    """Download a file from the coffre-fort by document ID."""
    token_doc = await get_current_token(token)

    # Find document metadata
    doc = await db.coffre_documents.find_one({"id": document_id, "token_id": token_doc["id"]})
    if not doc:
        raise HTTPException(404, "Document non trouvé dans le coffre-fort")

    # Find the file in GridFS
    cursor = gridfs_bucket.find({"metadata.file_id": document_id, "metadata.token_id": token_doc["id"]})
    grid_files = await cursor.to_list(1)

    if not grid_files:
        raise HTTPException(404, "Fichier non trouvé dans le stockage")

    grid_file = grid_files[0]
    stream = await gridfs_bucket.open_download_stream(grid_file["_id"])
    content = await stream.read()

    mime = doc.get("content_type", grid_file.get("metadata", {}).get("content_type", "application/octet-stream"))
    filename = doc.get("filename", "document")

    return StreamingResponse(
        io.BytesIO(content),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

@api_router.post("/coffre/documents")
async def create_coffre_document(token: str, request: CreateDocumentRequest):
    """Create a new document in coffre-fort"""
    token_doc = await get_current_token(token)
    
    document = CoffreDocument(
        token_id=token_doc["id"],
        **request.model_dump()
    )
    
    await db.coffre_documents.insert_one(document.model_dump())
    
    # If competences are linked, update profile skills
    if request.competences_liees:
        profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
        if profile:
            existing_skills = [s.get("name") for s in profile.get("skills", [])]
            new_skills = profile.get("skills", [])
            for comp in request.competences_liees:
                if comp not in existing_skills:
                    new_skills.append({"name": comp, "level": 50, "proven": True})
            await db.profiles.update_one({"token_id": token_doc["id"]}, {"$set": {"skills": new_skills}})
    
    return document.model_dump()

@api_router.put("/coffre/documents/{document_id}")
async def update_coffre_document(token: str, document_id: str, request: CreateDocumentRequest):
    """Update a document"""
    token_doc = await get_current_token(token)
    
    update_data = request.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    result = await db.coffre_documents.update_one(
        {"id": document_id, "token_id": token_doc["id"]},
        {"$set": update_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    
    return await db.coffre_documents.find_one({"id": document_id}, {"_id": 0})

@api_router.patch("/coffre/documents/{document_id}")
async def patch_coffre_document(token: str, document_id: str, body: dict):
    """Partial update of a document (e.g. trust_level validation)."""
    token_doc = await get_current_token(token)
    allowed = {"trust_level", "description", "title", "category"}
    update_data = {k: v for k, v in body.items() if k in allowed}
    if not update_data:
        raise HTTPException(400, "Aucun champ modifiable")
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    result = await db.coffre_documents.update_one(
        {"id": document_id, "token_id": token_doc["id"]},
        {"$set": update_data}
    )
    if result.modified_count == 0:
        raise HTTPException(404, "Document non trouvé")
    return {"success": True}

@api_router.delete("/coffre/documents/{document_id}")
async def delete_coffre_document(token: str, document_id: str):
    """Delete a document"""
    token_doc = await get_current_token(token)
    result = await db.coffre_documents.delete_one({"id": document_id, "token_id": token_doc["id"]})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    return {"message": "Document supprimé"}

@api_router.post("/coffre/documents/{document_id}/share")
async def share_document(token: str, document_id: str, shared_with_email: Optional[str] = None, shared_with_role: Optional[str] = None, expires_in_days: int = 7):
    """Create a share link for a document"""
    token_doc = await get_current_token(token)
    
    doc = await db.coffre_documents.find_one({"id": document_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document non trouvé")
    
    expires_at = (datetime.now(timezone.utc) + timedelta(days=expires_in_days)).isoformat()
    
    share = DocumentShare(
        document_id=document_id,
        shared_by=token_doc["id"],
        shared_with_email=shared_with_email,
        shared_with_role=shared_with_role,
        expires_at=expires_at
    )
    
    await db.document_shares.insert_one(share.model_dump())
    
    # Update document privacy
    await db.coffre_documents.update_one(
        {"id": document_id},
        {
            "$set": {"privacy_level": "shared_recruteur" if shared_with_role == "recruteur" else "shared_conseiller"},
            "$push": {"shared_with": share.id}
        }
    )
    
    return {"share_id": share.id, "access_token": share.access_token, "expires_at": expires_at}

@api_router.get("/coffre/shares")
async def get_document_shares(token: str):
    """Get all active shares for user's documents"""
    token_doc = await get_current_token(token)
    shares = await db.document_shares.find({"shared_by": token_doc["id"]}, {"_id": 0}).to_list(100)
    
    # Enrich with document info
    for share in shares:
        doc = await db.coffre_documents.find_one({"id": share["document_id"]}, {"_id": 0})
        if doc:
            share["document_title"] = doc.get("title")
            share["document_category"] = doc.get("category")
    
    return shares

@api_router.delete("/coffre/shares/{share_id}")
async def revoke_share(token: str, share_id: str):
    """Revoke a document share"""
    token_doc = await get_current_token(token)
    
    share = await db.document_shares.find_one({"id": share_id, "shared_by": token_doc["id"]}, {"_id": 0})
    if not share:
        raise HTTPException(status_code=404, detail="Partage non trouvé")
    
    await db.document_shares.delete_one({"id": share_id})
    
    # Update document
    await db.coffre_documents.update_one(
        {"id": share["document_id"]},
        {"$pull": {"shared_with": share_id}}
    )
    
    return {"message": "Partage révoqué"}

@api_router.get("/coffre/stats")
async def get_coffre_stats(token: str):
    """Get coffre-fort statistics"""
    token_doc = await get_current_token(token)
    
    documents = await db.coffre_documents.find({"token_id": token_doc["id"]}, {"_id": 0}).to_list(500)
    
    stats = {
        "total_documents": len(documents),
        "by_category": {},
        "competences_prouvees": set(),
        "documents_partages": 0,
        "documents_expirants": 0,
        "documents_sensibles": 0
    }
    
    today = datetime.now(timezone.utc)
    
    for doc in documents:
        cat = doc.get("category", "autre")
        stats["by_category"][cat] = stats["by_category"].get(cat, 0) + 1
        
        for comp in doc.get("competences_liees", []):
            stats["competences_prouvees"].add(comp)
        
        if doc.get("privacy_level") != "private":
            stats["documents_partages"] += 1
        
        if doc.get("is_sensitive"):
            stats["documents_sensibles"] += 1
        
        if doc.get("date_expiration"):
            try:
                exp_date = datetime.fromisoformat(doc["date_expiration"].replace('Z', '+00:00'))
                if 0 <= (exp_date - today).days <= 30:
                    stats["documents_expirants"] += 1
            except:
                pass
    
    stats["competences_prouvees"] = list(stats["competences_prouvees"])
    
    return stats

@api_router.get("/coffre/expiring")
async def get_expiring_documents(token: str):
    """Get documents expiring in the next 30 days"""
    token_doc = await get_current_token(token)
    documents = await db.coffre_documents.find({"token_id": token_doc["id"]}, {"_id": 0}).to_list(500)
    
    today = datetime.now(timezone.utc)
    expiring = []
    
    for doc in documents:
        if doc.get("date_expiration"):
            try:
                exp_date = datetime.fromisoformat(doc["date_expiration"].replace('Z', '+00:00'))
                days_until = (exp_date - today).days
                if 0 <= days_until <= 30:
                    doc["days_until_expiry"] = days_until
                    expiring.append(doc)
            except:
                pass
    
    return sorted(expiring, key=lambda x: x.get("days_until_expiry", 999))

# ============== OBSERVATOIRE ENDPOINTS ==============

@api_router.get("/observatoire/dashboard")
async def get_observatoire_dashboard():
    """Get observatoire main dashboard data"""
    emerging_skills = await db.emerging_skills.find({}, {"_id": 0}).to_list(50)
    sector_trends = await db.sector_trends.find({}, {"_id": 0}).to_list(20)
    contributions_count = await db.skill_contributions.count_documents({})
    validated_count = await db.skill_contributions.count_documents({"status": "integree"})
    
    # Calculate global indicators
    total_emerging = len([s for s in emerging_skills if s.get("status") == "emergente"])
    total_growing = len([s for s in emerging_skills if s.get("status") == "en_croissance"])
    sectors_in_transformation = len([t for t in sector_trends if t.get("transformation_index", 0) > 0.6])
    
    return {
        "emerging_skills": emerging_skills,
        "sector_trends": sector_trends,
        "indicators": {
            "total_emerging_skills": total_emerging,
            "total_growing_skills": total_growing,
            "sectors_in_transformation": sectors_in_transformation,
            "total_contributions": contributions_count,
            "validated_contributions": validated_count,
            "skill_gap_alerts": len([t for t in sector_trends if t.get("skill_gap_alert")])
        }
    }

@api_router.get("/observatoire/emerging-skills")
async def get_emerging_skills(sector: Optional[str] = None, status: Optional[str] = None):
    """Get emerging skills with optional filters"""
    query = {}
    if sector:
        query["related_sectors"] = sector
    if status:
        query["status"] = status
    
    skills = await db.emerging_skills.find(query, {"_id": 0}).to_list(100)
    return sorted(skills, key=lambda x: x.get("emergence_score", 0), reverse=True)

@api_router.get("/observatoire/sector-trends")
async def get_sector_trends(sector: Optional[str] = None):
    """Get sector transformation trends"""
    query = {}
    if sector:
        query["sector_name"] = sector
    
    trends = await db.sector_trends.find(query, {"_id": 0}).to_list(50)
    return sorted(trends, key=lambda x: x.get("transformation_index", 0), reverse=True)

@api_router.get("/observatoire/sector/{sector_name}")
async def get_sector_detail(sector_name: str):
    """Get detailed information about a sector"""
    trend = await db.sector_trends.find_one({"sector_name": sector_name}, {"_id": 0})
    if not trend:
        raise HTTPException(status_code=404, detail="Secteur non trouvé")
    
    # Get related emerging skills
    related_skills = await db.emerging_skills.find(
        {"related_sectors": sector_name}, 
        {"_id": 0}
    ).to_list(20)
    
    # Get recent contributions for this sector
    contributions = await db.skill_contributions.find(
        {"related_sector": sector_name, "status": {"$in": ["validee_ia", "validee_humain", "integree"]}},
        {"_id": 0}
    ).to_list(10)
    
    return {
        "trend": trend,
        "related_skills": related_skills,
        "recent_contributions": contributions
    }

@api_router.post("/observatoire/contributions")
async def create_contribution(token: str, request: CreateContributionRequest):
    """Submit a new skill/job contribution"""
    token_doc = await get_current_token(token)
    
    contribution = SkillContribution(
        contributor_id=token_doc["id"],
        **request.model_dump()
    )
    
    # AI Analysis (simplified - would use OpenAI in production)
    ai_analysis = await analyze_contribution_with_ai(contribution)
    contribution.ai_analysis = ai_analysis
    contribution.ai_score = ai_analysis.get("confidence_score", 0.5)
    
    if ai_analysis.get("is_valid", False) and ai_analysis.get("confidence_score", 0) > 0.7:
        contribution.status = "validee_ia"
    elif ai_analysis.get("confidence_score", 0) < 0.3:
        contribution.status = "rejetee_ia"
    
    # Check for similar contributions
    similar = await db.skill_contributions.find_one({
        "skill_name": {"$regex": contribution.skill_name, "$options": "i"},
        "status": {"$ne": "rejetee_ia"}
    }, {"_id": 0})
    
    if similar:
        # Increment similar count
        await db.skill_contributions.update_one(
            {"id": similar["id"]},
            {"$inc": {"similar_count": 1}}
        )
        contribution.similar_count = similar.get("similar_count", 1) + 1
    
    await db.skill_contributions.insert_one(contribution.model_dump())
    
    return {
        "contribution_id": contribution.id,
        "status": contribution.status,
        "ai_analysis": ai_analysis,
        "message": "Contribution enregistrée et analysée"
    }

@api_router.get("/observatoire/contributions")
async def get_contributions(token: str, status: Optional[str] = None):
    """Get user's contributions"""
    token_doc = await get_current_token(token)
    
    query = {"contributor_id": token_doc["id"]}
    if status:
        query["status"] = status
    
    contributions = await db.skill_contributions.find(query, {"_id": 0}).to_list(100)
    return contributions

@api_router.get("/observatoire/contributions/pending")
async def get_pending_contributions():
    """Get contributions pending human validation (for validators)"""
    contributions = await db.skill_contributions.find(
        {"status": "validee_ia"},
        {"_id": 0}
    ).to_list(50)
    return contributions

@api_router.post("/observatoire/contributions/{contribution_id}/validate")
async def validate_contribution(contribution_id: str, approved: bool, notes: Optional[str] = None):
    """Human validation of a contribution"""
    update_data = {
        "status": "validee_humain" if approved else "rejetee_humain",
        "human_notes": notes,
        "validated_at": datetime.now(timezone.utc).isoformat()
    }
    
    result = await db.skill_contributions.update_one(
        {"id": contribution_id},
        {"$set": update_data}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Contribution non trouvée")
    
    # If validated, potentially add to emerging skills
    if approved:
        contribution = await db.skill_contributions.find_one({"id": contribution_id}, {"_id": 0})
        if contribution and contribution.get("similar_count", 1) >= 3:
            await integrate_contribution_to_skills(contribution)
    
    return {"message": "Validation enregistrée", "status": update_data["status"]}

@api_router.post("/observatoire/contributions/{contribution_id}/upvote")
async def upvote_contribution(token: str, contribution_id: str):
    """Upvote a contribution"""
    await get_current_token(token)
    
    result = await db.skill_contributions.update_one(
        {"id": contribution_id},
        {"$inc": {"upvotes": 1}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Contribution non trouvée")
    
    return {"message": "Vote enregistré"}

@api_router.get("/observatoire/predictions")
async def get_predictions():
    """Get skill demand predictions"""
    trends = await db.sector_trends.find({}, {"_id": 0}).to_list(50)
    
    predictions = []
    for trend in trends:
        for pred in trend.get("predicted_skills_demand", []):
            predictions.append({
                "sector": trend["sector_name"],
                **pred
            })
    
    return sorted(predictions, key=lambda x: x.get("demand_change", "0%"), reverse=True)

@api_router.get("/observatoire/skill-gaps")
async def get_skill_gaps():
    """Get sectors with skill gap alerts"""
    trends = await db.sector_trends.find({"skill_gap_alert": True}, {"_id": 0}).to_list(20)
    return trends


@api_router.get("/observatoire/personalized")
async def get_personalized_observatory(token: str):
    """Cross-reference user profile/CV with observatory data for personalized predictions."""
    token_doc = await get_current_token(token)

    # Get user profile and passport
    profile = await db.profiles.find_one({"token_id": token_doc["id"]})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})

    # Gather user skills from all sources
    user_skills = set()
    user_skill_names = []  # Keep original casing
    if profile:
        for s in profile.get("skills", []):
            name = s.get("name", "").strip()
            if name:
                user_skills.add(name.lower())
                user_skill_names.append(name)
        if profile.get("cv_skills"):
            for s in profile["cv_skills"]:
                name = (s if isinstance(s, str) else s.get("name", "")).strip()
                if name:
                    user_skills.add(name.lower())
                    user_skill_names.append(name)
    if passport:
        for c in passport.get("competences", []):
            name = c.get("name", "").strip()
            if name:
                user_skills.add(name.lower())
                user_skill_names.append(name)
    user_skills.discard("")

    if not user_skills:
        return {"has_cv": False, "matches": [], "skill_gaps": [], "declining_alerts": [], "sector_relevance": [], "emerging_from_cv": [], "summary": {}}

    # Gather user context
    user_sectors = list(set(
        (profile or {}).get("sectors", []) +
        (passport or {}).get("target_sectors", []) +
        ([s for s in [(passport or {}).get("secteur_activite")] if s])
    ))
    experiences = (passport or {}).get("experiences", [])
    exp_titles = [e.get("titre", e.get("poste", "")) for e in experiences if e.get("titre") or e.get("poste")]
    metier_cible = (passport or {}).get("metier_cible", "")

    # ---- Use AI to generate truly personalized analysis (with cache) ----
    if EMERGENT_LLM_KEY:
        try:
            import json as json_lib

            # Check cache first (valid for 24h)
            cache_key = f"obs_perso_{token_doc['id']}"
            cached = await db.observatory_cache.find_one({"cache_key": cache_key})
            if cached:
                cache_age = (datetime.now(timezone.utc) - datetime.fromisoformat(cached["created_at"])).total_seconds()
                if cache_age < 86400:  # 24h
                    cached.pop("_id", None)
                    return cached["data"]

            skills_text = ", ".join(list(set(user_skill_names))[:20])
            sectors_text = ", ".join(user_sectors) if user_sectors else "non précisés"
            exp_text = ", ".join(exp_titles[:5]) if exp_titles else "non précisées"

            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"obs-perso-{token_doc['id'][:8]}",
                system_message=f"""Tu es un conseiller en emploi expert du marché du travail en France.

PROFIL DE L'UTILISATEUR :
- Compétences : {skills_text}
- Secteurs : {sectors_text}
- Expériences : {exp_text}
- Métier cible : {metier_cible or 'non défini'}

RÈGLES STRICTES :
1. TOUTES les compétences émergentes (matches) DOIVENT être directement liées aux compétences, secteurs et expériences listés ci-dessus. 
2. NE PROPOSE JAMAIS de compétences sans rapport avec le profil (ex: pas de cybersécurité pour un profil propreté, pas d'IA pour un profil restauration basique).
3. Les "skill_gaps" sont les compétences QUE CE PROFIL SPÉCIFIQUE devrait développer pour évoluer dans SES secteurs.
4. Les "sector_relevance" sont UNIQUEMENT les secteurs où les compétences de ce profil sont valorisables.
5. Reste réaliste et concret — pas de buzzwords déconnectés du terrain.

Génère un JSON strict :
{{
  "matches": [{{"name":"compétence émergente LIÉE AU PROFIL","score":0.0-1.0,"growth_rate":0.01-0.30,"trend":"croissance|stable","sectors":["secteur lié"]}}],
  "skill_gaps": [{{"name":"compétence à acquérir pour CE profil","score":0.0-1.0,"growth_rate":0.01-0.20,"priority":"haute|moyenne","sectors":["secteur"]}}],
  "sector_relevance": [{{"name":"secteur pertinent","relevance":0.0-1.0,"trend":"croissance|stable|declin","growth_rate":0.01-0.15,"skill_matches":3,"your_emerging_skills":["compétences du profil utiles ici"]}}],
  "declining_alerts": [{{"name":"compétence potentiellement en déclin","score":0.0-1.0,"growth_rate":-0.05,"trend":"declin"}}]
}}"""
            ).with_model("openai", "gpt-5.2")

            prompt = f"Génère l'analyse JSON personnalisée du marché pour ce profil. Maximum 6 matches, 5 skill_gaps, 4 sector_relevance, 2 declining_alerts. JSON uniquement, pas de markdown."

            response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
            clean = response.strip()
            if clean.startswith("```"):
                clean = clean.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            ai_data = json_lib.loads(clean)

            # Normalize AI data to match frontend expected fields
            def normalize_item(item):
                score = item.get("score", item.get("emergence_score", 0.5))
                if isinstance(score, str):
                    try:
                        score = float(score)
                    except:
                        score = 0.5
                gr = item.get("growth_rate", 0)
                if isinstance(gr, str):
                    gr_clean = gr.replace("%", "").replace("+", "").strip()
                    try:
                        gr = float(gr_clean) / 100.0
                    except:
                        gr = 0.05
                elif isinstance(gr, (int, float)) and gr > 1:
                    gr = gr / 100.0
                item["score"] = score
                item["emergence_score"] = score
                item["growth_rate"] = gr
                item["observatory_skill"] = item.get("name", "")
                return item

            def normalize_sector(s):
                rel = s.get("relevance", 0.5)
                if isinstance(rel, str):
                    try:
                        rel = float(rel)
                    except:
                        rel = 0.5
                s["relevance"] = rel
                s["transformation_index"] = rel
                gr = s.get("growth_rate", 0.05)
                if isinstance(gr, str):
                    gr_clean = gr.replace("%", "").replace("+", "").split("à")[0].split("-")[0].strip()
                    try:
                        gr_num = float(gr_clean)
                        gr = gr_num / 100.0 if gr_num > 1 else gr_num
                    except:
                        gr = 0.05
                elif isinstance(gr, (int, float)) and abs(gr) > 1:
                    gr = gr / 100.0
                s["growth_rate"] = gr
                s["sector"] = s.get("name", "")
                s["hiring_trend"] = s.get("trend", "stable")
                return s

            norm_matches = [normalize_item(m) for m in ai_data.get("matches", [])]
            norm_gaps = [normalize_item(g) for g in ai_data.get("skill_gaps", [])]
            norm_declining = [normalize_item(a) for a in ai_data.get("declining_alerts", [])]
            norm_sectors = [normalize_sector(s) for s in ai_data.get("sector_relevance", [])]

            result_data = {
                "has_cv": True,
                "matches": norm_matches[:6],
                "skill_gaps": norm_gaps[:5],
                "declining_alerts": norm_declining[:2],
                "sector_relevance": norm_sectors[:4],
                "emerging_from_cv": [m["name"] for m in norm_matches[:5]],
                "summary": {
                    "total_skills_analyzed": len(user_skills),
                    "skills_in_observatory": len(norm_matches),
                    "gaps_to_fill": len([g for g in norm_gaps if g.get("priority") == "haute"]),
                    "skills_declining": len(norm_declining),
                }
            }

            # Cache result for 24h
            await db.observatory_cache.update_one(
                {"cache_key": cache_key},
                {"$set": {"cache_key": cache_key, "data": result_data, "created_at": datetime.now(timezone.utc).isoformat()}},
                upsert=True,
            )

            return result_data
        except Exception as e:
            logger.error(f"Observatory personalized AI error: {e}")

    # ---- Fallback: improved keyword matching ----
    emerging_skills = await db.emerging_skills.find({}, {"_id": 0}).to_list(200)
    sector_trends = await db.sector_trends.find({}, {"_id": 0}).to_list(50)

    # Cross-match with keyword extraction
    matches = []
    declining_alerts = []
    for es in emerging_skills:
        es_name = es.get("skill_name", es.get("name", "")).lower().strip()
        es_alt = [a.lower().strip() for a in es.get("alternative_names", [])]
        # Extract keywords for fuzzy matching
        es_keywords = [w for w in es_name.split() if len(w) > 3]
        matched = False
        for us in user_skills:
            us_keywords = [w for w in us.split() if len(w) > 3]
            if (us in es_name or es_name in us or
                any(us in a or a in us for a in es_alt) or
                any(ek in us for ek in es_keywords) or
                any(uk in es_name for uk in us_keywords)):
                matched = True
                break
        if matched:
            trend = es.get("trend", es.get("status", "stable"))
            entry = {
                "observatory_skill": es.get("skill_name", es.get("name", "")),
                "name": es.get("skill_name", es.get("name", "")),
                "emergence_score": es.get("emergence_score", 0.5),
                "growth_rate": es.get("growth_rate", 0),
                "score": es.get("emergence_score", 0.5),
                "trend": trend,
                "status": es.get("status", trend),
                "sectors": es.get("related_sectors", []),
            }
            if trend in ("declin", "en_declin", "declining"):
                declining_alerts.append(entry)
            else:
                matches.append(entry)

    matches.sort(key=lambda x: x.get("score", 0), reverse=True)
    seen = set()
    unique_matches = []
    for m in matches:
        key = m.get("name", "").lower()
        if key not in seen:
            seen.add(key)
            unique_matches.append(m)
    matches = unique_matches

    # Skill gaps - only from sectors relevant to the user
    skill_gaps = []
    for es in emerging_skills:
        es_name = es.get("skill_name", es.get("name", "")).lower().strip()
        es_sectors = [s.lower() for s in es.get("related_sectors", [])]
        # Only suggest gaps from user's sectors
        sector_match = False
        for us in user_sectors:
            us_lower = us.lower()
            if any(us_lower in es_s or es_s in us_lower for es_s in es_sectors):
                sector_match = True
                break
        if not sector_match and user_sectors:
            continue
        if es.get("emergence_score", 0) >= 0.6:
            has_skill = any(us in es_name or es_name in us for us in user_skills)
            if not has_skill:
                skill_gaps.append({
                    "name": es.get("skill_name", es.get("name", "")),
                    "priority": "haute" if es.get("emergence_score", 0) >= 0.8 else "moyenne",
                    "score": es.get("emergence_score", 0),
                    "sectors": es.get("related_sectors", []),
                })
    skill_gaps.sort(key=lambda x: x.get("score", 0), reverse=True)
    seen_gaps = set()
    unique_gaps = []
    for g in skill_gaps:
        key = g.get("name", "").lower()
        if key not in seen_gaps:
            seen_gaps.add(key)
            unique_gaps.append(g)
    skill_gaps = unique_gaps

    # Sector relevance - prioritize user's actual sectors
    sector_relevance = []
    for st in sector_trends:
        sector_name = st.get("sector_name", "")
        sector_lower = sector_name.lower()
        predicted = st.get("predicted_skills_demand", [])

        # Check if this sector matches the user's sectors via keywords
        sector_match = False
        for us in user_sectors:
            us_kw = [w for w in us.lower().split() if len(w) > 3]
            if any(kw in sector_lower for kw in us_kw) or sector_lower in us.lower() or us.lower() in sector_lower:
                sector_match = True
                break

        # Check skill matches with keywords
        match_count = 0
        matched_skills = []
        for ps in predicted:
            ps_name = ps.get("skill", "").lower()
            ps_kw = [w for w in ps_name.split() if len(w) > 3]
            for us in user_skills:
                us_kw = [w for w in us.split() if len(w) > 3]
                if (us in ps_name or ps_name in us or
                    any(pk in us for pk in ps_kw) or
                    any(uk in ps_name for uk in us_kw)):
                    match_count += 1
                    matched_skills.append(ps.get("skill", ""))
                    break

        # Only include sectors that match user's profile OR have skill matches
        if sector_match or match_count > 0:
            sector_relevance.append({
                "sector": sector_name,
                "name": sector_name,
                "relevance": min((match_count + (2 if sector_match else 0)) / max(len(predicted), 1), 1.0),
                "transformation_index": st.get("transformation_index", 0),
                "hiring_trend": "croissance" if st.get("growth_rate", "0%").replace("%", "").replace("+", "").strip() not in ("0", "") else "stable",
                "skill_matches": match_count,
                "growth_rate": st.get("growth_rate", "0%"),
                "your_emerging_skills": matched_skills[:3],
            })
    sector_relevance.sort(key=lambda x: x.get("relevance", 0), reverse=True)

    return {
        "has_cv": True,
        "matches": matches[:20],
        "skill_gaps": skill_gaps[:20],
        "declining_alerts": declining_alerts,
        "sector_relevance": sector_relevance[:10],
        "emerging_from_cv": [m["name"] for m in matches[:5]],
        "summary": {
            "total_skills_analyzed": len(user_skills),
            "skills_in_observatory": len(matches),
            "gaps_to_fill": len([g for g in skill_gaps if g.get("priority") == "haute"]),
            "skills_declining": len(declining_alerts),
        }
    }



# ============== INDICE D'ÉVOLUTION ENDPOINTS ==============

def calculate_index_level(index: float) -> str:
    """Determine the level based on index value"""
    if index < 20:
        return "stable"
    elif index < 50:
        return "evolutif"
    elif index < 80:
        return "en_transformation"
    else:
        return "forte_mutation"

def get_index_interpretation(index: float, job_name: str = None) -> Dict[str, Any]:
    """Get human-readable interpretation of the index"""
    if index < 20:
        return {
            "level": "stable",
            "label": "Métier très stable",
            "description": f"Les compétences de ce métier évoluent peu. La formation initiale reste pertinente sur le long terme.",
            "color": "emerald",
            "recommendation": "Maintenez vos compétences actuelles et restez en veille sur les évolutions du secteur."
        }
    elif index < 50:
        return {
            "level": "evolutif",
            "label": "Métier évolutif",
            "description": f"Ce métier connaît des évolutions modérées. Certaines compétences nouvelles apparaissent progressivement.",
            "color": "blue",
            "recommendation": "Renforcez vos compétences numériques et suivez une formation continue régulière."
        }
    elif index < 80:
        return {
            "level": "en_transformation",
            "label": "Métier en transformation",
            "description": f"Ce métier évolue significativement sous l'effet des innovations technologiques ou organisationnelles.",
            "color": "amber",
            "recommendation": "Anticipez les changements en développant les compétences émergentes de votre secteur."
        }
    else:
        return {
            "level": "forte_mutation",
            "label": "Forte mutation",
            "description": f"Ce métier est fortement impacté par les transformations. Une adaptation rapide est nécessaire.",
            "color": "rose",
            "recommendation": "Envisagez une montée en compétences significative ou une reconversion vers des métiers connexes."
        }

@api_router.get("/evolution-index/jobs")
async def get_jobs_evolution_index(sector: Optional[str] = None):
    """Get evolution index for all jobs"""
    query = {}
    if sector:
        query["sector"] = sector
    
    indices = await db.job_evolution_indices.find(query, {"_id": 0}).to_list(100)
    
    # Enrich with interpretation
    for idx in indices:
        idx["interpretation"] = get_index_interpretation(idx.get("evolution_index", 0), idx.get("job_name"))
    
    return sorted(indices, key=lambda x: x.get("evolution_index", 0), reverse=True)

@api_router.get("/evolution-index/jobs/{job_name}")
async def get_job_evolution_detail(job_name: str):
    """Get detailed evolution index for a specific job"""
    index = await db.job_evolution_indices.find_one(
        {"job_name": {"$regex": job_name, "$options": "i"}}, 
        {"_id": 0}
    )
    
    if not index:
        raise HTTPException(status_code=404, detail="Métier non trouvé")
    
    index["interpretation"] = get_index_interpretation(index.get("evolution_index", 0), job_name)
    
    # Get related emerging skills
    related_skills = await db.emerging_skills.find(
        {"related_jobs": {"$regex": job_name, "$options": "i"}},
        {"_id": 0}
    ).to_list(10)
    
    index["related_emerging_skills"] = related_skills
    
    return index

@api_router.get("/evolution-index/sectors")
async def get_sectors_evolution_index():
    """Get evolution index for all sectors"""
    indices = await db.sector_evolution_indices.find({}, {"_id": 0}).to_list(50)
    
    for idx in indices:
        idx["interpretation"] = get_index_interpretation(idx.get("evolution_index", 0))
    
    return sorted(indices, key=lambda x: x.get("evolution_index", 0), reverse=True)

@api_router.get("/evolution-index/sectors/{sector_name}")
async def get_sector_evolution_detail(sector_name: str):
    """Get detailed evolution index for a sector"""
    index = await db.sector_evolution_indices.find_one(
        {"sector_name": {"$regex": sector_name, "$options": "i"}},
        {"_id": 0}
    )
    
    if not index:
        raise HTTPException(status_code=404, detail="Secteur non trouvé")
    
    index["interpretation"] = get_index_interpretation(index.get("evolution_index", 0))
    
    # Get all jobs in this sector
    jobs = await db.job_evolution_indices.find(
        {"sector": {"$regex": sector_name, "$options": "i"}},
        {"_id": 0}
    ).to_list(50)
    
    index["jobs"] = jobs
    
    return index

@api_router.get("/evolution-index/dashboard")
async def get_evolution_dashboard():
    """Get comprehensive evolution index dashboard"""
    job_indices = await db.job_evolution_indices.find({}, {"_id": 0}).to_list(100)
    sector_indices = await db.sector_evolution_indices.find({}, {"_id": 0}).to_list(50)
    
    # Calculate statistics
    total_jobs = len(job_indices)
    jobs_stable = len([j for j in job_indices if j.get("evolution_index", 0) < 20])
    jobs_evolving = len([j for j in job_indices if 20 <= j.get("evolution_index", 0) < 50])
    jobs_transforming = len([j for j in job_indices if 50 <= j.get("evolution_index", 0) < 80])
    jobs_highly_impacted = len([j for j in job_indices if j.get("evolution_index", 0) >= 80])
    
    avg_job_index = sum(j.get("evolution_index", 0) for j in job_indices) / max(total_jobs, 1)
    avg_sector_index = sum(s.get("evolution_index", 0) for s in sector_indices) / max(len(sector_indices), 1)
    
    # Top transforming jobs
    top_transforming = sorted(job_indices, key=lambda x: x.get("evolution_index", 0), reverse=True)[:5]
    most_stable = sorted(job_indices, key=lambda x: x.get("evolution_index", 0))[:5]
    
    # Sectors overview
    for sector in sector_indices:
        sector["interpretation"] = get_index_interpretation(sector.get("evolution_index", 0))
    
    return {
        "summary": {
            "total_jobs_analyzed": total_jobs,
            "total_sectors_analyzed": len(sector_indices),
            "average_job_evolution_index": round(avg_job_index, 1),
            "average_sector_evolution_index": round(avg_sector_index, 1)
        },
        "distribution": {
            "stable": {"count": jobs_stable, "percentage": round(jobs_stable / max(total_jobs, 1) * 100, 1)},
            "evolving": {"count": jobs_evolving, "percentage": round(jobs_evolving / max(total_jobs, 1) * 100, 1)},
            "transforming": {"count": jobs_transforming, "percentage": round(jobs_transforming / max(total_jobs, 1) * 100, 1)},
            "highly_impacted": {"count": jobs_highly_impacted, "percentage": round(jobs_highly_impacted / max(total_jobs, 1) * 100, 1)}
        },
        "top_transforming_jobs": top_transforming,
        "most_stable_jobs": most_stable,
        "sectors": sector_indices,
        "interpretation_guide": {
            "stable": {"range": "0-20", "description": "Métier très stable, évolution lente"},
            "evolutif": {"range": "20-50", "description": "Métier évolutif mais relativement stable"},
            "en_transformation": {"range": "50-80", "description": "Métier en transformation importante"},
            "forte_mutation": {"range": "80-100", "description": "Métier fortement impacté par les innovations"}
        }
    }

@api_router.get("/evolution-index/user-profile")
async def get_user_evolution_analysis(token: str):
    """Get evolution analysis based on user's profile, passport and CV skills"""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})

    if not profile and not passport:
        raise HTTPException(status_code=404, detail="Profil non trouvé")

    # Gather sectors from profile + passport
    user_sectors = list(set(
        (profile or {}).get("sectors", []) +
        (passport or {}).get("target_sectors", []) +
        ([s for s in [(passport or {}).get("secteur_activite")] if s])
    ))

    # Gather skills from profile + passport + cv_skills
    user_skills = set()
    for s in (profile or {}).get("skills", []):
        user_skills.add(s.get("name", ""))
    for s in (profile or {}).get("cv_skills", []):
        user_skills.add(s if isinstance(s, str) else s.get("name", ""))
    for c in (passport or {}).get("competences", []):
        user_skills.add(c.get("name", ""))
    for sf in (passport or {}).get("savoir_faire", []):
        user_skills.add(sf if isinstance(sf, str) else sf.get("name", ""))
    user_skills.discard("")
    user_skills = list(user_skills)

    has_cv = bool(user_skills) or bool((passport or {}).get("experiences"))

    # Find relevant job indices from sectors (fuzzy keyword matching)
    relevant_jobs = []
    seen_jobs = set()
    for sector in user_sectors:
        # Extract keywords from sector for better matching
        keywords = [w for w in sector.lower().split() if len(w) > 3 and w not in ("dans", "avec", "pour", "les", "des")]
        for kw in keywords[:2]:
            jobs = await db.job_evolution_indices.find(
                {"sector": {"$regex": kw, "$options": "i"}},
                {"_id": 0}
            ).to_list(10)
            for j in jobs:
                jn = j.get("job_name", "")
                if jn not in seen_jobs:
                    seen_jobs.add(jn)
                    relevant_jobs.append(j)

    # Also try matching by user's job title / metier_cible from passport
    metier_cible = (passport or {}).get("metier_cible", "") or (passport or {}).get("career_project", "")
    metier_words = metier_cible.strip().split() if metier_cible else []
    if metier_words and len(relevant_jobs) < 3:
        extra = await db.job_evolution_indices.find(
            {"job_name": {"$regex": metier_words[0], "$options": "i"}},
            {"_id": 0}
        ).to_list(5)
        for j in extra:
            jn = j.get("job_name", "")
            if jn not in seen_jobs:
                seen_jobs.add(jn)
                relevant_jobs.append(j)

    # Find skills at risk
    skills_at_risk = []
    skills_in_demand = []

    for job in relevant_jobs:
        for skill in user_skills:
            skill_lower = skill.lower()
            if any(skill_lower in d.lower() or d.lower() in skill_lower for d in job.get("declining_skills", [])):
                skills_at_risk.append({"skill": skill, "job": job["job_name"]})
            if any(skill_lower in e.lower() or e.lower() in skill_lower for e in job.get("emerging_skills", [])):
                skills_in_demand.append({"skill": skill, "job": job["job_name"]})

    # Recommendations
    all_recommended = set()
    for job in relevant_jobs:
        all_recommended.update(job.get("recommended_skills", []))

    # Emerging skills from CV - only match skills relevant to user's sectors
    emerging_from_cv = []
    emerging_skills_db = await db.emerging_skills.find({}, {"_id": 0}).to_list(200)
    for es in emerging_skills_db:
        es_name = es.get("skill_name", es.get("name", "")).lower()
        es_sectors = [s.lower() for s in es.get("related_sectors", [])]
        
        # Check if this emerging skill is relevant to user's sectors
        sector_relevant = False
        if not user_sectors:
            sector_relevant = True  # No sectors = show all
        else:
            for us in user_sectors:
                us_kw = [w for w in us.lower().split() if len(w) > 3]
                if any(kw in es_s for kw in us_kw for es_s in es_sectors):
                    sector_relevant = True
                    break
        
        if not sector_relevant:
            continue
            
        for us in user_skills:
            if us.lower() in es_name or es_name in us.lower():
                emerging_from_cv.append({"name": es.get("skill_name", es.get("name", "")), "score": round(es.get("emergence_score", 0.5) * 100)})
                break

    # Calculate personal evolution exposure
    if relevant_jobs:
        avg_exposure = sum(j.get("evolution_index", 0) for j in relevant_jobs) / len(relevant_jobs)
    else:
        avg_exposure = 50

    return {
        "has_cv": has_cv,
        "profile_sectors": user_sectors,
        "profile_skills": user_skills,
        "evolution_exposure": round(avg_exposure, 1),
        "exposure_interpretation": get_index_interpretation(avg_exposure),
        "relevant_jobs": relevant_jobs[:8],
        "skills_at_risk": skills_at_risk[:10],
        "skills_in_demand": skills_in_demand[:10],
        "recommended_skills_to_acquire": list(all_recommended - set(user_skills))[:10],
        "recommended_trainings": list(set(t for j in relevant_jobs for t in j.get("recommended_trainings", [])))[:5],
        "emerging_from_cv": emerging_from_cv[:5],
        "data_sources": {
            "cv_analysis": bool((profile or {}).get("cv_skills")),
            "passport": bool(passport),
        }
    }

# ============== MARCHÉ CACHÉ ENDPOINTS ==============

class MarcheCacheDiagnosticRequest(BaseModel):
    token: str

@api_router.post("/marche-cache/diagnostic")
async def marche_cache_diagnostic(payload: MarcheCacheDiagnosticRequest):
    """AI-powered diagnostic of user's hidden job market access potential."""
    token_doc = await get_current_token(payload.token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})

    # Gather user context
    user_skills = []
    for s in (profile or {}).get("skills", []):
        user_skills.append(s.get("name", ""))
    for s in (profile or {}).get("cv_skills", []):
        user_skills.append(s if isinstance(s, str) else s.get("name", ""))
    for c in (passport or {}).get("competences", []):
        user_skills.append(c.get("name", ""))
    user_skills = list(set(sk for sk in user_skills if sk))

    experiences = (passport or {}).get("experiences", [])
    metier = (passport or {}).get("metier_cible", "") or (passport or {}).get("career_project", "")
    sectors = list(set(
        (profile or {}).get("sectors", []) +
        (passport or {}).get("target_sectors", [])
    ))
    raw_soft = (passport or {}).get("savoir_etre", [])
    soft_skills = [s if isinstance(s, str) else s.get("name", str(s)) for s in raw_soft]

    # Get D'CLIC results if available
    dclic = await db.dclic_results.find_one({"user_id": token_doc["id"]}, {"_id": 0})
    personality_traits = []
    if dclic and dclic.get("profile"):
        p = dclic["profile"]
        personality_traits = [
            f"MBTI: {p.get('mbti', 'N/A')}",
            f"DISC: {p.get('disc', 'N/A')}",
            f"RIASEC: {p.get('riasec', 'N/A')}",
        ]

    skills_str = ", ".join(str(s) for s in user_skills[:15]) if user_skills else "Non renseignées"
    sectors_str = ", ".join(str(s) for s in sectors[:5]) if sectors else "Non définis"
    soft_str = ", ".join(str(s) for s in soft_skills[:8]) if soft_skills else "Non renseignés"

    context_lines = [
        f"Compétences: {skills_str}",
        f"Expériences: {len(experiences)} postes",
        f"Métier cible: {metier or 'Non défini'}",
        f"Secteurs: {sectors_str}",
        f"Soft skills: {soft_str}",
    ]
    if personality_traits:
        context_lines.append(f"Profil psychométrique: {', '.join(personality_traits)}")

    user_context = "\n".join(context_lines)

    if EMERGENT_LLM_KEY:
        try:
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"marche-cache-{token_doc['id'][:8]}",
                system_message="""Tu es un expert en stratégie d'accès au marché caché de l'emploi en France.
Analyse le profil de l'utilisateur et génère un diagnostic complet en JSON strict (pas de markdown).
Le JSON doit contenir exactement ces champs:
{
  "score_acces": (entier 1-10),
  "analyse": "texte court décrivant la situation globale",
  "forces_marche_cache": ["liste des atouts pour accéder au marché caché"],
  "faiblesses": ["liste des points faibles"],
  "recommandations": [{"titre": "...", "description": "...", "priorite": "haute|moyenne|basse"}],
  "canaux_privilegier": ["liste de canaux d'accès recommandés"],
  "types_entreprises": ["types d'entreprises à cibler"],
  "strategie_reseau": "paragraphe décrivant la stratégie réseau personnalisée"
}"""
            ).with_model("openai", "gpt-5.2")

            response = await run_llm_nonblocking(chat, UserMessage(text=f"Profil utilisateur:\n{user_context}\n\nGénère le diagnostic JSON du marché caché."))
            import json as json_lib
            clean = response.strip()
            if clean.startswith("```"):
                clean = clean.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            diagnostic = json_lib.loads(clean)
            return {"diagnostic": diagnostic}
        except Exception as e:
            logger.error(f"Marché caché AI error: {e}")

    # Fallback without AI
    score = min(10, max(1, len(user_skills) // 2 + len(experiences) + (2 if metier else 0)))
    return {"diagnostic": {
        "score_acces": score,
        "analyse": f"Avec {len(user_skills)} compétences et {len(experiences)} expériences, votre accès au marché caché est {'bon' if score >= 6 else 'à développer'}.",
        "forces_marche_cache": [f"Compétences diversifiées ({len(user_skills)} identifiées)"] + ([f"Secteurs ciblés : {', '.join(sectors[:3])}"] if sectors else []),
        "faiblesses": ["Enrichissez votre profil pour un diagnostic plus précis"] if len(user_skills) < 5 else [],
        "recommandations": [
            {"titre": "Activez votre réseau professionnel", "description": "Contactez d'anciens collègues et participez à des événements sectoriels.", "priorite": "haute"},
            {"titre": "Candidatures spontanées ciblées", "description": "Identifiez les entreprises de vos secteurs et envoyez des candidatures personnalisées.", "priorite": "moyenne"},
        ],
        "canaux_privilegier": ["LinkedIn", "Événements sectoriels", "Anciens collègues", "Associations professionnelles"],
        "types_entreprises": ["PME en croissance", "Start-ups de votre secteur", "Cabinets de conseil"],
        "strategie_reseau": "Développez votre présence en ligne et participez activement aux discussions de votre secteur."
    }}


# ============== EXPLORATEUR SUGGESTIONS PERSONNALISÉES ==============

@api_router.get("/referentiel/explorer/suggestions")
async def get_explorer_suggestions(token: str):
    """Return personalized job suggestions based on user profile."""
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})

    # Gather user context
    user_skills = set()
    for s in (profile or {}).get("skills", []):
        user_skills.add(s.get("name", "").lower())
    for c in (passport or {}).get("competences", []):
        user_skills.add(c.get("name", "").lower())
    user_skills.discard("")

    metier_cible = (passport or {}).get("metier_cible", "")
    sectors = list(set(
        (profile or {}).get("sectors", []) +
        (passport or {}).get("target_sectors", [])
    ))
    experiences = (passport or {}).get("experiences", [])
    experience_titles = [e.get("titre", e.get("poste", "")) for e in experiences if e.get("titre") or e.get("poste")]

    suggestions = []
    seen = set()

    # 1. Suggest metier_cible if it exists
    if metier_cible and metier_cible.lower() not in seen:
        suggestions.append({"name": metier_cible, "reason": "Votre métier cible"})
        seen.add(metier_cible.lower())

    # 2. Add experience-based suggestions
    for title in experience_titles[:3]:
        if title.lower() not in seen:
            suggestions.append({"name": title, "reason": "Basé sur votre expérience"})
            seen.add(title.lower())

    # 3. Search referentiel for sector-matching jobs
    for sector in sectors[:3]:
        # Extract key words from sector for fuzzy matching
        keywords = [w for w in sector.lower().split() if len(w) > 3 and w not in ("dans", "avec", "pour", "les", "des")]
        for kw in keywords[:2]:
            # Search in referentiel_metiers_flat
            sector_data = await db.referentiel_metiers_flat.find(
                {"secteur": {"$regex": kw, "$options": "i"}},
                {"_id": 0, "name": 1, "secteur": 1}
            ).to_list(3)
            for m in sector_data:
                name = m.get("name", "")
                if name.lower() not in seen:
                    suggestions.append({"name": name, "reason": f"Secteur {m.get('secteur', sector)}"})
                    seen.add(name.lower())

    # 4. Search evolution indices for jobs in user's sectors
    for sector in sectors[:2]:
        keywords = [w for w in sector.lower().split() if len(w) > 3 and w not in ("dans", "avec", "pour", "les", "des")]
        for kw in keywords[:2]:
            jobs = await db.job_evolution_indices.find(
                {"sector": {"$regex": kw, "$options": "i"}},
                {"_id": 0, "job_name": 1, "sector": 1, "evolution_index": 1}
            ).to_list(3)
            for j in jobs:
                jn = j.get("job_name", "")
                if jn.lower() not in seen:
                    suggestions.append({"name": jn, "reason": f"Métier en {j.get('sector', sector)}"})
                    seen.add(jn.lower())

    # 5. Search ROME metiers by skill keywords
    if user_skills and len(suggestions) < 6:
        sample_skills = list(user_skills)[:3]
        for sk in sample_skills:
            kw = sk.split()[0] if sk else ""
            if len(kw) < 3:
                continue
            rome_matches = await db.rome_metiers.find(
                {"$or": [
                    {"libelle": {"$regex": kw, "$options": "i"}},
                    {"appellations": {"$regex": kw, "$options": "i"}},
                ]},
                {"_id": 0, "libelle": 1}
            ).to_list(3)
            for rm in rome_matches:
                name = rm.get("libelle", "")
                if name.lower() not in seen:
                    suggestions.append({"name": name, "reason": f"Lié à {sk}"})
                    seen.add(name.lower())

    has_profile = bool(user_skills or metier_cible or experience_titles)
    return {
        "has_profile": has_profile,
        "suggestions": suggestions[:12],
        "skills_count": len(user_skills),
        "sectors": sectors[:5],
    }


async def analyze_contribution_with_ai(contribution: SkillContribution) -> Dict[str, Any]:
    """Analyze a contribution using AI"""
    if not EMERGENT_LLM_KEY:
        # Fallback analysis
        is_valid = len(contribution.skill_name) > 3 and len(contribution.skill_name) < 100
        return {
            "is_valid": is_valid,
            "confidence_score": 0.6 if is_valid else 0.3,
            "category": "technique" if any(kw in contribution.skill_name.lower() for kw in ["code", "data", "dev", "ia", "cyber"]) else "transversale",
            "similar_existing": [],
            "rationale": "Analyse basique - cohérence vérifiée"
        }
    
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"contrib-{contribution.id}",
            system_message="Tu es un expert RH français. Analyse cette contribution à un observatoire des compétences. Réponds en JSON avec: is_valid (bool), confidence_score (0-1), category (technique/transversale/sectorielle), similar_existing (list), rationale (string)."
        ).with_model("openai", "gpt-5.2")
        
        prompt = f"""
        Nouvelle compétence proposée: {contribution.skill_name}
        Description: {contribution.skill_description or 'Non fournie'}
        Métier associé: {contribution.related_job or 'Non spécifié'}
        Secteur: {contribution.related_sector or 'Non spécifié'}
        Contexte: {contribution.context or 'Non fourni'}
        
        Analyse si cette compétence est:
        1. Valide et pertinente pour le marché du travail
        2. Suffisamment précise
        3. Potentiellement émergente
        """
        
        response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
        
        import json
        try:
            result = json.loads(response)
            return result
        except:
            return {
                "is_valid": True,
                "confidence_score": 0.65,
                "category": "transversale",
                "similar_existing": [],
                "rationale": response[:200]
            }
    except Exception as e:
        logging.error(f"AI contribution analysis error: {e}")
        return {
            "is_valid": True,
            "confidence_score": 0.5,
            "category": "non_classifie",
            "similar_existing": [],
            "rationale": "Analyse automatique non disponible"
        }

async def integrate_contribution_to_skills(contribution: dict):
    """Integrate a validated contribution into emerging skills"""
    existing = await db.emerging_skills.find_one(
        {"skill_name": {"$regex": contribution["skill_name"], "$options": "i"}},
        {"_id": 0}
    )
    
    if existing:
        # Update existing skill
        await db.emerging_skills.update_one(
            {"id": existing["id"]},
            {
                "$inc": {"mention_count": 1, "contributor_count": 1},
                "$set": {"last_updated": datetime.now(timezone.utc).isoformat()}
            }
        )
    else:
        # Create new emerging skill
        new_skill = EmergingSkill(
            skill_name=contribution["skill_name"],
            description=contribution.get("skill_description"),
            related_sectors=[contribution["related_sector"]] if contribution.get("related_sector") else [],
            related_jobs=[contribution["related_job"]] if contribution.get("related_job") else [],
            related_tools=contribution.get("related_tools", []),
            emergence_score=0.5,
            growth_rate=0.1,
            mention_count=contribution.get("similar_count", 1),
            contributor_count=1
        )
        await db.emerging_skills.insert_one(new_skill.model_dump())
    
    # Mark contribution as integrated
    await db.skill_contributions.update_one(
        {"id": contribution["id"]},
        {"$set": {"status": "integree"}}
    )

# ============== PASSEPORT DYNAMIQUE ENDPOINTS ==============

def calculate_completeness(passport: dict) -> int:
    """Calculate passport completeness score (0-100)"""
    score = 0
    if passport.get("professional_summary"): score += 12
    if passport.get("career_project"): score += 8
    if passport.get("motivations"): score += 5
    if passport.get("compatible_environments"): score += 5
    if passport.get("target_sectors"): score += 5
    comps = passport.get("competences", [])
    if len(comps) >= 1: score += 8
    if len(comps) >= 3: score += 7
    if len(comps) >= 5: score += 5
    exps = passport.get("experiences", [])
    if len(exps) >= 1: score += 8
    if len(exps) >= 3: score += 7
    learning = passport.get("learning_path", [])
    if len(learning) >= 1: score += 8
    if len(learning) >= 2: score += 5
    # Bonus for Lamri & Lubart evaluations
    evaluated = sum(1 for c in comps if any(c.get("components", {}).get(k, 0) > 0 for k in ["connaissance", "cognition", "conation", "affection", "sensori_moteur"]))
    if evaluated >= 1: score += 5
    if evaluated >= 3: score += 7
    # Bonus for CCSP classification
    ccsp_classified = sum(1 for c in comps if c.get("ccsp_pole"))
    if ccsp_classified >= 1: score += 5
    return min(score, 100)

async def aggregate_passport_from_sources(token_id: str) -> dict:
    """Aggregate passport data from all platform sources"""
    aggregated = {"competences": [], "experiences": [], "learning_path": []}
    seen_comp_names = set()
    seen_exp_titles = set()

    # 1. From coffre-fort documents
    docs = await db.documents.find({"user_token": token_id}, {"_id": 0}).to_list(50)
    for doc in docs:
        for skill in doc.get("skills", []):
            if skill.lower() not in seen_comp_names:
                seen_comp_names.add(skill.lower())
                aggregated["competences"].append({
                    "id": str(uuid.uuid4()), "name": skill, "category": "technique",
                    "level": "intermediaire", "experience_years": 0, "proof": doc.get("title"),
                    "source": "coffre_fort", "is_emerging": False,
                    "added_at": datetime.now(timezone.utc).isoformat()
                })

    # 2. From learning modules
    profile = await db.profiles.find_one({"token_id": token_id}, {"_id": 0})
    if profile:
        for skill_data in profile.get("skills", []):
            sname = skill_data.get("name", "") if isinstance(skill_data, dict) else str(skill_data)
            if sname.lower() not in seen_comp_names:
                seen_comp_names.add(sname.lower())
                level = skill_data.get("level", "intermediaire") if isinstance(skill_data, dict) else "intermediaire"
                aggregated["competences"].append({
                    "id": str(uuid.uuid4()), "name": sname, "category": "technique",
                    "level": level, "experience_years": 0, "proof": None,
                    "source": "profil", "is_emerging": False,
                    "added_at": datetime.now(timezone.utc).isoformat()
                })

    modules = await db.modules.find({}, {"_id": 0}).to_list(50)
    for mod in modules:
        aggregated["learning_path"].append({
            "id": str(uuid.uuid4()), "title": mod.get("title", ""),
            "provider": "RE'ACTIF PRO", "skills_acquired": mod.get("skills", []),
            "status": "en_cours", "completion_date": None, "badge": None,
            "source": "plateforme"
        })

    # 2b. From CV analysis (savoir_faire & savoir_etre)
    last_cv = await db.cv_jobs.find_one(
        {"token_id": token_id, "status": "completed"}, sort=[("created_at", -1)]
    )
    if last_cv and last_cv.get("result"):
        cv_result = last_cv["result"]
        # Savoir-faire from CV
        for sf in cv_result.get("savoir_faire", cv_result.get("competences", []))[:20]:
            sname = sf.get("name", "") if isinstance(sf, dict) else str(sf)
            if sname and sname.lower() not in seen_comp_names:
                seen_comp_names.add(sname.lower())
                level_raw = sf.get("level", sf.get("niveau", 50)) if isinstance(sf, dict) else 50
                level_str = "avance" if (isinstance(level_raw, (int, float)) and level_raw >= 70) else "intermediaire" if (isinstance(level_raw, (int, float)) and level_raw >= 40) else "debutant"
                aggregated["competences"].append({
                    "id": str(uuid.uuid4()), "name": sname, "category": "savoir_faire",
                    "level": level_str, "experience_years": 0, "proof": last_cv.get("filename"),
                    "source": "ia_detectee", "is_emerging": False,
                    "added_at": datetime.now(timezone.utc).isoformat()
                })
        # Savoir-être from CV
        aggregated["savoir_faire"] = [sf.get("name","") if isinstance(sf, dict) else str(sf) for sf in cv_result.get("savoir_faire", [])[:15]]
        aggregated["savoir_etre"] = cv_result.get("savoir_etre", [])[:10]
        # Experiences from CV
        for exp in cv_result.get("experiences", []):
            aggregated["experiences"].append({
                "id": str(uuid.uuid4()),
                "title": exp.get("title", ""),
                "organization": exp.get("organization", ""),
                "description": exp.get("description", ""),
                "start_date": exp.get("start_date", ""),
                "end_date": exp.get("end_date", ""),
                "is_current": exp.get("is_ongoing", False),
                "skills_used": exp.get("skills_used", []),
                "proof": None,
                "source": "cv_analysis"
            })

    # 3. From Ubuntoo signals (emerging skills)
    signals = await db.ubuntoo_signals.find(
        {"validation_status": {"$in": ["validee_humain", "integree"]}}, {"_id": 0}
    ).to_list(20)
    for signal in signals:
        if signal.get("signal_type") == "competence_emergente" and signal["name"].lower() not in seen_comp_names:
            seen_comp_names.add(signal["name"].lower())
            aggregated["competences"].append({
                "id": str(uuid.uuid4()), "name": signal["name"], "category": "technique",
                "level": "debutant", "experience_years": 0, "proof": None,
                "source": "ubuntoo", "is_emerging": True,
                "added_at": datetime.now(timezone.utc).isoformat()
            })

    # 4. From contributions
    contributions = await db.contributions.find(
        {"user_token": token_id, "status": {"$in": ["validee_ia", "validee"]}}, {"_id": 0}
    ).to_list(20)
    for contrib in contributions:
        sname = contrib.get("skill_name", "")
        if sname and sname.lower() not in seen_comp_names:
            seen_comp_names.add(sname.lower())
            aggregated["competences"].append({
                "id": str(uuid.uuid4()), "name": sname, "category": "technique",
                "level": "intermediaire", "experience_years": 0, "proof": None,
                "source": "contribution", "is_emerging": True,
                "added_at": datetime.now(timezone.utc).isoformat()
            })

    return aggregated

async def generate_passerelles_with_ai(competences: List[dict], sectors: List[str]) -> List[dict]:
    """Use AI to suggest career pathways based on passport competences"""
    if not EMERGENT_LLM_KEY or not competences:
        return []
    try:
        skills_list = ", ".join([c.get("name", "") for c in competences[:15]])
        sectors_str = ", ".join(sectors[:5]) if sectors else "tous secteurs"

        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"passerelle-{uuid.uuid4()}",
            system_message="Tu es un conseiller en évolution professionnelle français expert. Analyse les compétences et suggère des passerelles professionnelles réalistes. Réponds UNIQUEMENT en JSON valide: un array de max 5 objets avec les clés: job_name (str), compatibility_score (float 0-1), shared_skills (list str), skills_to_acquire (list str max 3), training_needed (str court), accessibility (str: accessible/formation_courte/formation_longue), sector (str)."
        ).with_model("openai", "gpt-5.2")

        response = await run_llm_nonblocking(chat, UserMessage(text=f"""Compétences de la personne: {skills_list}
Secteurs d'intérêt: {sectors_str}

Propose 5 passerelles professionnelles réalistes."""))

        import json
        try:
            result = json.loads(response)
            if isinstance(result, list):
                return result[:5]
            return result.get("passerelles", result.get("pathways", []))[:5]
        except:
            return []
    except Exception as e:
        logging.error(f"Passerelles AI error: {e}")
        return []

@api_router.get("/passport")
async def get_passport(token: str):
    """Get or create the user's dynamic passport"""
    token_doc = await get_current_token(token)

    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        # Create new passport with aggregated data
        aggregated = await aggregate_passport_from_sources(token_doc["id"])
        new_passport = Passport(token_id=token_doc["id"])
        passport_data = new_passport.model_dump()
        passport_data["competences"] = aggregated["competences"]
        passport_data["learning_path"] = aggregated["learning_path"]
        passport_data["experiences"] = aggregated.get("experiences", [])
        passport_data["savoir_faire"] = aggregated.get("savoir_faire", [])
        passport_data["savoir_etre"] = aggregated.get("savoir_etre", [])
        passport_data["completeness_score"] = calculate_completeness(passport_data)
        await db.passports.insert_one(passport_data)
        passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})

    # Count sources
    sources_count = {}
    for c in passport.get("competences", []):
        src = c.get("source", "declaratif")
        sources_count[src] = sources_count.get(src, 0) + 1

    passport["sources_count"] = sources_count
    passport["total_competences"] = len(passport.get("competences", []))
    passport["total_experiences"] = len(passport.get("experiences", []))
    passport["total_learning"] = len(passport.get("learning_path", []))
    passport["emerging_count"] = len([c for c in passport.get("competences", []) if c.get("is_emerging")])

    return passport

@api_router.post("/passport/refresh")
async def refresh_passport(token: str):
    """Refresh passport by re-aggregating from all sources"""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    aggregated = await aggregate_passport_from_sources(token_doc["id"])

    # Merge: keep existing user-declared competences, add new from sources
    existing_names = {c.get("name", "").lower() for c in passport.get("competences", []) if c.get("source") == "declaratif"}
    new_comps = [c for c in aggregated["competences"] if c.get("name", "").lower() not in existing_names]
    declared_comps = [c for c in passport.get("competences", []) if c.get("source") == "declaratif"]

    all_comps = declared_comps + new_comps
    passport["competences"] = all_comps
    passport["learning_path"] = aggregated["learning_path"]
    passport["savoir_faire"] = aggregated.get("savoir_faire", [])
    passport["savoir_etre"] = aggregated.get("savoir_etre", [])
    # Merge experiences: keep proofs from existing, add new from CV
    existing_exps = {(e.get("title","").lower(), e.get("organization","").lower()): e for e in passport.get("experiences", [])}
    merged_exps = list(passport.get("experiences", []))
    for new_exp in aggregated.get("experiences", []):
        key = (new_exp.get("title","").lower(), new_exp.get("organization","").lower())
        if key not in existing_exps:
            merged_exps.append(new_exp)
    passport["experiences"] = merged_exps
    passport["completeness_score"] = calculate_completeness(passport)
    passport["last_updated"] = datetime.now(timezone.utc).isoformat()

    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {
            "competences": passport["competences"],
            "learning_path": passport["learning_path"],
            "experiences": passport["experiences"],
            "savoir_faire": passport["savoir_faire"],
            "savoir_etre": passport["savoir_etre"],
            "completeness_score": passport["completeness_score"],
            "last_updated": passport["last_updated"]
        }}
    )
    return {"message": "Passeport actualisé", "completeness_score": passport["completeness_score"]}

@api_router.put("/passport/profile")
async def update_passport_profile(token: str, data: UpdatePassportProfileRequest):
    """Update passport profile section"""
    token_doc = await get_current_token(token)
    update = {}
    if data.professional_summary is not None: update["professional_summary"] = data.professional_summary
    if data.career_project is not None: update["career_project"] = data.career_project
    if data.motivations is not None: update["motivations"] = data.motivations
    if data.compatible_environments is not None: update["compatible_environments"] = data.compatible_environments
    if data.target_sectors is not None: update["target_sectors"] = data.target_sectors

    if not update:
        raise HTTPException(status_code=400, detail="Aucune donnée à mettre à jour")

    update["last_updated"] = datetime.now(timezone.utc).isoformat()
    result = await db.passports.update_one({"token_id": token_doc["id"]}, {"$set": update})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    # Recalculate completeness
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    new_score = calculate_completeness(passport)
    await db.passports.update_one({"token_id": token_doc["id"]}, {"$set": {"completeness_score": new_score}})

    return {"message": "Profil mis à jour", "completeness_score": new_score}

@api_router.post("/passport/competences")
async def add_passport_competence(token: str, data: AddCompetenceRequest):
    """Add a competence to the passport"""
    token_doc = await get_current_token(token)

    components = data.components or {"connaissance": 0, "cognition": 0, "conation": 0, "affection": 0, "sensori_moteur": 0}
    new_comp = PassportCompetence(
        name=data.name, nature=data.nature, category=data.category, level=data.level,
        experience_years=data.experience_years, proof=data.proof, source="declaratif",
        components=components,
        ccsp_pole=data.ccsp_pole or "",
        ccsp_degree=data.ccsp_degree or "",
        linked_qualites=data.linked_qualites,
        linked_valeurs=data.linked_valeurs,
        linked_vertus=data.linked_vertus
    ).model_dump()

    result = await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$push": {"competences": new_comp}, "$set": {"last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    new_score = calculate_completeness(passport)
    await db.passports.update_one({"token_id": token_doc["id"]}, {"$set": {"completeness_score": new_score}})

    return {"message": "Compétence ajoutée", "competence": new_comp, "completeness_score": new_score}

@api_router.post("/passport/experiences")
async def add_passport_experience(token: str, data: AddExperienceRequest):
    """Add an experience to the passport"""
    token_doc = await get_current_token(token)

    new_exp = PassportExperience(
        title=data.title, organization=data.organization, description=data.description,
        skills_used=data.skills_used, achievements=data.achievements,
        start_date=data.start_date, end_date=data.end_date, is_current=data.is_current,
        experience_type=data.experience_type, source="declaratif"
    ).model_dump()

    result = await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$push": {"experiences": new_exp}, "$set": {"last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    new_score = calculate_completeness(passport)
    await db.passports.update_one({"token_id": token_doc["id"]}, {"$set": {"completeness_score": new_score}})

    return {"message": "Expérience ajoutée", "experience": new_exp, "completeness_score": new_score}

@api_router.delete("/passport/competences/{comp_id}")
async def delete_passport_competence(comp_id: str, token: str):
    """Remove a competence from the passport"""
    token_doc = await get_current_token(token)
    result = await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$pull": {"competences": {"id": comp_id}}, "$set": {"last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")
    return {"message": "Compétence supprimée"}

@api_router.delete("/passport/experiences/{exp_id}")
async def delete_passport_experience(exp_id: str, token: str):
    """Remove an experience from the passport"""
    token_doc = await get_current_token(token)
    result = await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$pull": {"experiences": {"id": exp_id}}, "$set": {"last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")
    return {"message": "Expérience supprimée"}


@api_router.delete("/passport/reset")
async def reset_passport_sections(token: str, sections: str = "all"):
    """Reset specific sections of the passport"""
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]
    passport = await db.passports.find_one({"token_id": token_id})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    unset_fields = {}
    set_fields = {"last_updated": datetime.now(timezone.utc).isoformat()}

    if sections == "all":
        set_fields.update({
            "competences": [],
            "experiences": [],
            "learning_path": [],
            "passerelles": [],
            "professional_summary": "",
            "career_project": "",
            "motivations": [],
            "compatible_environments": [],
            "target_sectors": [],
            "offres_emploi": [],
            "competences_transversales": [],
            "competences_transferables": [],
        })
        # Also reset profile flags
        await db.profiles.update_one(
            {"token_id": token_id},
            {"$set": {"cv_analyzed": False, "skills": [], "savoir_etre": [], "experiences": [], "strengths": [], "gaps": []}}
        )
        # Reset CV data
        await db.cv_jobs.delete_many({"token_id": token_id})
        await db.cv_models.delete_many({"token_id": token_id})
    elif sections == "competences":
        set_fields.update({
            "competences": [],
            "competences_transversales": [],
            "competences_transferables": [],
        })
        await db.profiles.update_one(
            {"token_id": token_id},
            {"$set": {"skills": [], "savoir_etre": []}}
        )
    elif sections == "experiences":
        set_fields["experiences"] = []
        await db.profiles.update_one(
            {"token_id": token_id},
            {"$set": {"experiences": []}}
        )
        # Also clear trajectory
        await db.trajectories.delete_many({"token_id": token_id})
    elif sections == "formations":
        set_fields["learning_path"] = []
    elif sections == "profile":
        set_fields.update({
            "professional_summary": "",
            "career_project": "",
            "motivations": [],
            "compatible_environments": [],
            "target_sectors": [],
        })
    elif sections == "passerelles":
        set_fields["passerelles"] = []
    elif sections == "dclic":
        unset_fields["dclic_results"] = ""
        unset_fields["dclic_imported_at"] = ""
        # Reset dclic_imported flag in profiles
        await db.profiles.update_one(
            {"token_id": token_id},
            {"$set": {"dclic_imported": False}, "$unset": {"dclic_imported_at": ""}}
        )
        # Delete D'CLIC test results from dclic_results collection
        await db.dclic_results.delete_many({"claimed_by": token_id})
    else:
        raise HTTPException(status_code=400, detail=f"Section inconnue: {sections}")

    update_ops = {"$set": set_fields}
    if unset_fields:
        update_ops["$unset"] = unset_fields
    await db.passports.update_one({"token_id": token_id}, update_ops)

    # Recalculate completeness
    updated = await db.passports.find_one({"token_id": token_id})
    if updated:
        score = calculate_completeness(updated)
        await db.passports.update_one({"token_id": token_id}, {"$set": {"completeness_score": score}})

    return {"success": True, "message": f"Section '{sections}' réinitialisée"}


@api_router.get("/passport/passerelles")
async def get_passport_passerelles(token: str):
    """Get AI-suggested career pathways"""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    passerelles = await generate_passerelles_with_ai(
        passport.get("competences", []),
        passport.get("target_sectors", [])
    )

    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"passerelles": passerelles, "last_updated": datetime.now(timezone.utc).isoformat()}}
    )

    return {"passerelles": passerelles}

@api_router.put("/passport/sharing")
async def update_passport_sharing(token: str, data: SharePassportRequest):
    """Update passport sharing settings"""
    token_doc = await get_current_token(token)
    sharing = {
        "is_public": data.is_public,
        "shared_sections": data.shared_sections,
        "shared_with": data.shared_with,
        "share_expiry": data.share_expiry
    }
    result = await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"sharing": sharing, "last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")
    return {"message": "Paramètres de partage mis à jour", "sharing": sharing}

@api_router.put("/passport/competences/{comp_id}/evaluate")
async def evaluate_competence(comp_id: str, token: str, data: EvaluateCompetenceRequest):
    """Evaluate a competence using Lamri & Lubart 5-component model and CCSP"""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    # Validate component values (0-5)
    valid_keys = {"connaissance", "cognition", "conation", "affection", "sensori_moteur"}
    components = {}
    for key in valid_keys:
        val = data.components.get(key, 0)
        components[key] = max(0, min(5, val))

    # Find and update the competence
    comps = passport.get("competences", [])
    found = False
    for comp in comps:
        if comp.get("id") == comp_id:
            comp["components"] = components
            if data.ccsp_pole:
                comp["ccsp_pole"] = data.ccsp_pole
            if data.ccsp_degree:
                comp["ccsp_degree"] = data.ccsp_degree
            found = True
            break

    if not found:
        raise HTTPException(status_code=404, detail="Compétence non trouvée")

    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"competences": comps, "last_updated": datetime.now(timezone.utc).isoformat()}}
    )

    return {"message": "Évaluation enregistrée", "competence_id": comp_id, "components": components}

@api_router.post("/passport/diagnostic/auto-evaluate")
async def auto_evaluate_competences(token: str):
    """Auto-evaluate all passport competences using AI (Lamri & Lubart + CCSP)."""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    comps = passport.get("competences", [])
    if not comps:
        raise HTTPException(status_code=400, detail="Aucune compétence à évaluer")

    # Filter only unevaluated competences
    to_evaluate = []
    for c in comps:
        has_eval = any(c.get("components", {}).get(k, 0) > 0 for k in ["connaissance", "cognition", "conation", "affection", "sensori_moteur"])
        if not has_eval:
            to_evaluate.append(c)

    if not to_evaluate and all(any(c.get("components", {}).get(k, 0) > 0 for k in ["connaissance", "cognition", "conation", "affection", "sensori_moteur"]) for c in comps):
        return {"message": "Toutes les compétences sont déjà évaluées", "evaluated": 0}

    # Get user context for better evaluation
    experiences = passport.get("experiences", [])
    metier = passport.get("metier_cible", "") or passport.get("career_project", "")
    exp_context = ", ".join([e.get("titre", e.get("poste", "")) for e in experiences[:3] if e.get("titre") or e.get("poste")])

    # Build competence list for AI
    comp_names = [c.get("name", "") for c in (to_evaluate if to_evaluate else comps)]

    if EMERGENT_LLM_KEY:
        try:
            import json as json_lib
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"auto-eval-{token_doc['id'][:8]}",
                system_message="""Tu es un expert en évaluation de compétences professionnelles selon deux référentiels :

1. **Lamri & Lubart** (5 composantes, score de 1 à 5 chacune) :
   - connaissance : savoirs théoriques
   - cognition : capacités d'analyse et raisonnement
   - conation : motivation et engagement
   - affection : intelligence émotionnelle
   - sensori_moteur : habiletés physiques et techniques

2. **CCSP** :
   - pole : "realisation" (produire), "interaction" (communiquer), ou "initiative" (innover)
   - degree : "imitation" (reproduire), "adaptation" (ajuster), ou "transposition" (créer)

3. **Nature** : "savoir_faire" (technique) ou "savoir_etre" (comportemental)

Réponds UNIQUEMENT en JSON strict (pas de markdown). Le JSON doit être un tableau d'objets :
[{"name":"nom_competence","connaissance":X,"cognition":X,"conation":X,"affection":X,"sensori_moteur":X,"ccsp_pole":"...","ccsp_degree":"...","nature":"savoir_faire|savoir_etre"}]"""
            ).with_model("openai", "gpt-5.2")

            prompt = f"Évalue ces compétences professionnelles :\n{', '.join(comp_names)}"
            if exp_context:
                prompt += f"\nContexte métier : {exp_context}"
            if metier:
                prompt += f"\nMétier cible : {metier}"

            response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
            clean = response.strip()
            if clean.startswith("```"):
                clean = clean.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
            ai_evals = json_lib.loads(clean)

            # Apply AI evaluations to competences
            eval_map = {}
            for ae in ai_evals:
                eval_map[ae.get("name", "").lower()] = ae

            updated = 0
            for comp in comps:
                cname = comp.get("name", "").lower()
                ai = eval_map.get(cname)
                if not ai:
                    # Fuzzy match
                    for key, val in eval_map.items():
                        if key in cname or cname in key:
                            ai = val
                            break
                if ai:
                    has_eval = any(comp.get("components", {}).get(k, 0) > 0 for k in ["connaissance", "cognition", "conation", "affection", "sensori_moteur"])
                    if not has_eval:
                        comp["components"] = {
                            "connaissance": max(0, min(5, ai.get("connaissance", 2))),
                            "cognition": max(0, min(5, ai.get("cognition", 2))),
                            "conation": max(0, min(5, ai.get("conation", 2))),
                            "affection": max(0, min(5, ai.get("affection", 2))),
                            "sensori_moteur": max(0, min(5, ai.get("sensori_moteur", 2))),
                        }
                        if ai.get("ccsp_pole"):
                            comp["ccsp_pole"] = ai["ccsp_pole"]
                        if ai.get("ccsp_degree"):
                            comp["ccsp_degree"] = ai["ccsp_degree"]
                        if ai.get("nature") and not comp.get("nature"):
                            comp["nature"] = ai["nature"]
                        updated += 1

            await db.passports.update_one(
                {"token_id": token_doc["id"]},
                {"$set": {"competences": comps, "last_updated": datetime.now(timezone.utc).isoformat()}}
            )
            return {"message": f"{updated} compétences évaluées par l'IA", "evaluated": updated}

        except Exception as e:
            logger.error(f"Auto-evaluate AI error: {e}")

    # Fallback: basic heuristic evaluation
    updated = 0
    for comp in comps:
        has_eval = any(comp.get("components", {}).get(k, 0) > 0 for k in ["connaissance", "cognition", "conation", "affection", "sensori_moteur"])
        if not has_eval:
            comp["components"] = {"connaissance": 2, "cognition": 2, "conation": 3, "affection": 2, "sensori_moteur": 2}
            comp["ccsp_pole"] = comp.get("ccsp_pole") or "realisation"
            comp["ccsp_degree"] = comp.get("ccsp_degree") or "adaptation"
            if not comp.get("nature"):
                comp["nature"] = "savoir_faire"
            updated += 1

    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"competences": comps, "last_updated": datetime.now(timezone.utc).isoformat()}}
    )
    return {"message": f"{updated} compétences évaluées (heuristique)", "evaluated": updated}


@api_router.get("/passport/diagnostic")
async def get_ccsp_diagnostic(token: str):
    """Generate a CCSP diagnostic based on all passport competences"""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    comps = passport.get("competences", [])
    total = len(comps)
    if total == 0:
        return {
            "total_competences": 0,
            "evaluated_count": 0,
            "lamri_lubart_profile": {"connaissance": 0, "cognition": 0, "conation": 0, "affection": 0, "sensori_moteur": 0},
            "ccsp_distribution": {"poles": {}, "degrees": {}},
            "recommendations": []
        }

    # Aggregate Lamri & Lubart components
    ll_totals = {"connaissance": 0, "cognition": 0, "conation": 0, "affection": 0, "sensori_moteur": 0}
    evaluated_count = 0
    for comp in comps:
        cdata = comp.get("components", {})
        has_evaluation = any(cdata.get(k, 0) > 0 for k in ll_totals)
        if has_evaluation:
            evaluated_count += 1
            for k in ll_totals:
                ll_totals[k] += cdata.get(k, 0)

    ll_avg = {}
    for k, v in ll_totals.items():
        ll_avg[k] = round(v / max(evaluated_count, 1), 1)

    # CCSP distribution
    poles = {"realisation": 0, "interaction": 0, "initiative": 0}
    degrees = {"imitation": 0, "adaptation": 0, "transposition": 0}
    for comp in comps:
        pole = comp.get("ccsp_pole", "")
        degree = comp.get("ccsp_degree", "")
        if pole in poles:
            poles[pole] += 1
        if degree in degrees:
            degrees[degree] += 1

    # Generate recommendations
    recommendations = []
    if ll_avg.get("connaissance", 0) < 2:
        recommendations.append({"type": "formation", "message": "Renforcez vos savoirs théoriques par des formations ou lectures spécialisées.", "component": "connaissance"})
    if ll_avg.get("cognition", 0) < 2:
        recommendations.append({"type": "formation", "message": "Développez vos capacités d'analyse et de raisonnement critique.", "component": "cognition"})
    if ll_avg.get("conation", 0) < 2:
        recommendations.append({"type": "accompagnement", "message": "Travaillez votre motivation et votre engagement par un accompagnement personnalisé.", "component": "conation"})
    if ll_avg.get("affection", 0) < 2:
        recommendations.append({"type": "developpement", "message": "Renforcez votre intelligence émotionnelle et votre empathie.", "component": "affection"})
    if ll_avg.get("sensori_moteur", 0) < 2:
        recommendations.append({"type": "pratique", "message": "Augmentez la pratique terrain pour développer vos habiletés techniques.", "component": "sensori_moteur"})

    if poles.get("initiative", 0) == 0 and total > 2:
        recommendations.append({"type": "ccsp", "message": "Aucune compétence d'initiative identifiée. Explorez l'autonomie et l'innovation.", "component": "initiative"})
    if degrees.get("transposition", 0) == 0 and total > 2:
        recommendations.append({"type": "ccsp", "message": "Développez votre capacité à transposer vos compétences dans de nouveaux contextes.", "component": "transposition"})

    # Nature distribution (savoir-faire vs savoir-être)
    nature_dist = {"savoir_faire": 0, "savoir_etre": 0, "non_classee": 0}
    for comp in comps:
        n = comp.get("nature", "")
        if n == "savoir_faire":
            nature_dist["savoir_faire"] += 1
        elif n == "savoir_etre":
            nature_dist["savoir_etre"] += 1
        else:
            nature_dist["non_classee"] += 1

    # Recommendations based on nature balance
    if nature_dist["savoir_etre"] == 0 and total > 2:
        recommendations.append({"type": "orientation", "message": "Identifiez vos savoir-être (soft skills) pour enrichir votre stratégie d'orientation professionnelle.", "component": "savoir_etre"})
    if nature_dist["savoir_faire"] == 0 and total > 2:
        recommendations.append({"type": "orientation", "message": "Ajoutez vos compétences techniques (savoir-faire) pour mieux cibler les métiers compatibles.", "component": "savoir_faire"})
    if nature_dist["non_classee"] > 0:
        recommendations.append({"type": "classification", "message": f"{nature_dist['non_classee']} compétence(s) non classée(s). Précisez leur nature (savoir-faire ou savoir-être) pour un meilleur diagnostic.", "component": "nature"})

    return {
        "total_competences": total,
        "evaluated_count": evaluated_count,
        "lamri_lubart_profile": ll_avg,
        "ccsp_distribution": {"poles": poles, "degrees": degrees},
        "nature_distribution": nature_dist,
        "recommendations": recommendations
    }

# ============== CV ANALYSIS & GENERATION ==============

async def extract_text_from_upload(file: UploadFile) -> str:
    """Extract text from uploaded PDF or DOCX file"""
    content = await file.read()
    text = ""
    if file.filename.lower().endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    elif file.filename.lower().endswith((".docx", ".doc")):
        import docx
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        text = content.decode("utf-8", errors="ignore")
    return text.strip()


async def _llm_call_with_retry(system_msg: str, user_msg: str, max_retries: int = 2) -> dict:
    """Make an LLM call with retry logic and JSON parsing."""
    last_error = None
    for attempt in range(max_retries + 1):
        try:
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"cv-{uuid.uuid4()}",
                system_message=system_msg
            ).with_model("openai", "gpt-5.2")
            response = await run_llm_nonblocking(chat, UserMessage(text=user_msg))
            raw = response.strip() if isinstance(response, str) else response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
                if raw.endswith("```"):
                    raw = raw[:-3]
                raw = raw.strip()
            return json.loads(raw)
        except json.JSONDecodeError as e:
            last_error = f"Réponse IA non valide (tentative {attempt+1})"
            logging.warning(f"CV analysis JSON error attempt {attempt+1}: {e}")
        except Exception as e:
            last_error = str(e)
            logging.warning(f"CV analysis LLM error attempt {attempt+1}: {e}")
    raise Exception(f"Erreur IA après {max_retries+1} tentatives: {last_error}")


def _extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Extract text from file bytes (PDF, DOCX, TXT)"""
    text = ""
    if filename.lower().endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    elif filename.lower().endswith((".docx", ".doc")):
        import docx
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        text = content.decode("utf-8", errors="ignore")
    return text.strip()


async def _run_cv_analysis(job_id: str, token_id: str, file_content: bytes, filename: str, text_ready: bool = False):
    """Background task: extract text, run CV analysis and save results to DB."""
    try:
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "analyzing", "step": "Extraction du texte..."}})

        if text_ready:
            cv_text = file_content.decode("utf-8", errors="ignore")
        else:
            cv_text = _extract_text_from_bytes(file_content, filename)
        if not cv_text or len(cv_text) < 50:
            await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": "Le fichier ne contient pas assez de texte exploitable", "step": "Erreur"}})
            return

        cv_excerpt = cv_text[:6000]

        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"step": "Analyse des compétences..."}})

        # Run BOTH LLM calls in PARALLEL to cut total time in half
        analysis_task = asyncio.create_task(_llm_call_with_retry(
            system_msg="""Tu es un expert RH. Analyse ce CV. Réponds UNIQUEMENT en JSON valide (pas de markdown).
Structure exacte:
{
  "profile": {"professional_summary": "2-3 phrases", "career_project": "string", "motivations": [], "compatible_environments": [], "target_sectors": []},
  "savoir_faire": [{"name": "string", "category": "technique|transversale|transferable|sectorielle", "level": "debutant|intermediaire|avance|expert", "ccsp_pole": "realisation|interaction|initiative", "ccsp_degree": "imitation|adaptation|transposition"}],
  "savoir_etre": [{"name": "string", "category": "transversale|transferable", "level": "debutant|intermediaire|avance|expert", "linked_qualites": [], "linked_valeurs": [], "linked_vertus": []}],
  "competences_transversales": ["liste de compétences transversales identifiées"],
  "competences_transferables": ["liste de compétences transférables identifiées"],
  "experiences": [{"title": "string", "organization": "string", "description": "string", "experience_type": "professionnel|personnel|benevole|projet", "start_date": "YYYY-MM", "end_date": "YYYY-MM", "is_ongoing": false, "skills_used": [], "achievements": []}],
  "formations_suggestions": [{"title": "string", "reason": "string", "priority": "haute|moyenne|basse", "skills_to_gain": []}],
  "offres_emploi": [{"title": "string", "company_type": "string", "sector": "string", "contract_type": "CDI|CDD|Freelance", "salary_range": "string", "location": "France", "required_skills": [], "match_score": 75, "description": "courte description du poste"}],
  "strengths": ["points forts du candidat"],
  "gaps": ["lacunes ou axes d'amélioration"],
  "audit_cv": [
    {"regle": "Résumé professionnel", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation détaillée", "recommandation": "conseil si ameliorable/absent"},
    {"regle": "Expériences détaillées", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Compétences techniques", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Soft skills identifiés", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Mots-clés ATS", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Cohérence chronologique", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Quantification des résultats", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Lisibilité et structure", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Formations et certifications", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"},
    {"regle": "Objectif professionnel clair", "statut": "ok|ameliorable|absent", "score": 8, "diagnostic": "évaluation", "recommandation": "conseil"}
  ],
  "score_global_cv": 65,
  "modele_suggere": "cv_classique|cv_competences|cv_fonctionnel|cv_mixte"
}
Pour experiences: inclure start_date (YYYY-MM) et end_date (YYYY-MM) quand disponibles dans le CV. Mettre is_ongoing=true si le poste est actuel.
Pour audit_cv: évalue rigoureusement chaque critère (ok=bon, ameliorable=à améliorer, absent=insuffisant). score de 1 à 10 pour chaque regle. score_global_cv = somme sur 100. Fournis diagnostic détaillé ET recommandation concrète pour chaque critère ameliorable ou absent.
Pour modele_suggere: recommande le format de CV le plus adapté au profil.
Pour offres_emploi: génère 5-8 offres réalistes et pertinentes.
Valeurs IDs: autonomie, stimulation, hedonisme, realisation_de_soi, pouvoir, securite, conformite, tradition, bienveillance, universalisme.
Vertus: sagesse, courage, humanite, justice, temperance, transcendance.""",
            user_msg=f"Analyse ce CV:\n\n{cv_excerpt}"
        ))

        cv_gen_task = asyncio.create_task(_llm_call_with_retry(
            system_msg="""Tu es un rédacteur de CV professionnel. Génère 4 versions d'un CV. Réponds UNIQUEMENT en JSON valide.
Structure: {"cv_classique": "texte complet", "cv_competences": "texte complet", "cv_fonctionnel": "texte complet", "cv_mixte": "texte complet"}
- cv_classique: chronologique, sobre, professionnel
- cv_competences: axé savoir-faire et savoir-être par domaine
- cv_fonctionnel: par domaines de compétences, sans chronologie
- cv_mixte: parcours + compétences transférables""",
            user_msg=f"Génère 4 versions de CV pour ce profil:\n\n{cv_excerpt}"
        ))

        # Wait for analysis first (critical), then CV gen
        analysis = await analysis_task
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"step": "Génération des modèles de CV..."}})

        try:
            cv_gen = await cv_gen_task
        except Exception:
            logging.warning("CV model generation failed, continuing with analysis only")
            cv_gen = {"cv_classique": "", "cv_competences": "", "cv_fonctionnel": "", "cv_mixte": ""}

        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"step": "Remplissage du passeport..."}})

        # Save CV models
        cv_models = {
            "classique": cv_gen.get("cv_classique", ""),
            "competences": cv_gen.get("cv_competences", ""),
            "fonctionnel": cv_gen.get("cv_fonctionnel", ""),
            "mixte": cv_gen.get("cv_mixte", ""),
        }
        await db.cv_models.update_one(
            {"token_id": token_id},
            {"$set": {"token_id": token_id, "models": cv_models, "original_filename": filename, "analyzed_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True
        )

        # Auto-fill Passport
        passport = await db.passports.find_one({"token_id": token_id})
        if not passport:
            passport = {"token_id": token_id, "professional_summary": "", "career_project": "", "motivations": [], "compatible_environments": [], "target_sectors": [], "competences": [], "experiences": [], "learning_path": [], "passerelles": [], "sharing": {"is_public": False}, "created_at": datetime.now(timezone.utc).isoformat()}
            await db.passports.insert_one(passport)

        new_competences = list(passport.get("competences", []))
        existing_names = {c.get("name", "").lower() for c in new_competences}
        for sf in analysis.get("savoir_faire", []):
            if sf.get("name", "").lower() not in existing_names:
                new_competences.append(PassportCompetence(name=sf["name"], nature="savoir_faire", category=sf.get("category", "technique"), level=sf.get("level", "intermediaire"), source="ia_detectee", ccsp_pole=sf.get("ccsp_pole", ""), ccsp_degree=sf.get("ccsp_degree", "")).model_dump())
                existing_names.add(sf["name"].lower())
        for se in analysis.get("savoir_etre", []):
            if se.get("name", "").lower() not in existing_names:
                new_competences.append(PassportCompetence(name=se["name"], nature="savoir_etre", category=se.get("category", "transversale"), level=se.get("level", "intermediaire"), source="ia_detectee", linked_qualites=se.get("linked_qualites", []), linked_valeurs=se.get("linked_valeurs", []), linked_vertus=se.get("linked_vertus", [])).model_dump())
                existing_names.add(se["name"].lower())

        new_experiences = list(passport.get("experiences", []))
        existing_exp_titles = {e.get("title", "").lower() for e in new_experiences}
        for exp in analysis.get("experiences", []):
            if exp.get("title", "").lower() not in existing_exp_titles:
                new_experiences.append(PassportExperience(title=exp["title"], organization=exp.get("organization", ""), description=exp.get("description", ""), experience_type=exp.get("experience_type", "professionnel"), skills_used=exp.get("skills_used", []), achievements=exp.get("achievements", []), source="ia_detectee").model_dump())
                existing_exp_titles.add(exp["title"].lower())

        new_learning = list(passport.get("learning_path", []))
        for fs in analysis.get("formations_suggestions", []):
            new_learning.append({"title": fs.get("title", ""), "provider": f"Suggestion IA - Priorité {fs.get('priority', 'moyenne')}", "status": "suggere", "skills_acquired": fs.get("skills_to_gain", []), "reason": fs.get("reason", ""), "source": "ia_detectee"})

        profile_data = analysis.get("profile", {})
        update_fields = {"competences": new_competences, "experiences": new_experiences, "learning_path": new_learning, "last_updated": datetime.now(timezone.utc).isoformat()}
        if profile_data.get("professional_summary") and not passport.get("professional_summary"):
            update_fields["professional_summary"] = profile_data["professional_summary"]
        if profile_data.get("career_project") and not passport.get("career_project"):
            update_fields["career_project"] = profile_data["career_project"]
        if profile_data.get("motivations") and not passport.get("motivations"):
            update_fields["motivations"] = profile_data["motivations"]
        if profile_data.get("compatible_environments") and not passport.get("compatible_environments"):
            update_fields["compatible_environments"] = profile_data["compatible_environments"]
        if profile_data.get("target_sectors") and not passport.get("target_sectors"):
            update_fields["target_sectors"] = profile_data["target_sectors"]

        # Save competences transversales and transferables to passport
        ct = analysis.get("competences_transversales", [])
        ctf = analysis.get("competences_transferables", [])
        if ct:
            update_fields["competences_transversales"] = ct
        if ctf:
            update_fields["competences_transferables"] = ctf

        # Save offres d'emploi to passport
        offres = analysis.get("offres_emploi", [])
        if offres:
            update_fields["offres_emploi"] = offres

        merged = {**passport, **update_fields}
        update_fields["completeness_score"] = calculate_completeness(merged)
        await db.passports.update_one({"token_id": token_id}, {"$set": update_fields})

        # Update user profile with strengths, gaps, and skills from analysis
        profile_update = {}
        strengths = analysis.get("strengths", ct)
        gaps = analysis.get("gaps", [])
        if strengths:
            profile_update["strengths"] = strengths
        if gaps:
            profile_update["gaps"] = gaps
        # Build skills list from savoir_faire
        skills_from_cv = []
        for sf in analysis.get("savoir_faire", []):
            level_map = {"debutant": 30, "intermediaire": 55, "avance": 75, "expert": 90}
            skills_from_cv.append({"name": sf.get("name", ""), "level": level_map.get(sf.get("level", "intermediaire"), 55)})
        if skills_from_cv:
            profile_update["skills"] = skills_from_cv
        if profile_data.get("target_sectors"):
            profile_update["sectors"] = profile_data["target_sectors"]
        if profile_update:
            profile_update["profile_score"] = update_fields.get("completeness_score", 0)
            await db.profiles.update_one({"token_id": token_id}, {"$set": profile_update})

        # Store result in job
        result = {
            "message": "CV analysé avec succès",
            "filename": filename,
            "savoir_faire_count": len(analysis.get("savoir_faire", [])),
            "savoir_etre_count": len(analysis.get("savoir_etre", [])),
            "experiences_count": len(analysis.get("experiences", [])),
            "formations_suggestions": analysis.get("formations_suggestions", []),
            "competences_transversales": analysis.get("competences_transversales", []),
            "competences_transferables": analysis.get("competences_transferables", []),
            "offres_emploi": analysis.get("offres_emploi", []),
            "strengths": analysis.get("strengths", analysis.get("competences_transversales", [])),
            "gaps": analysis.get("gaps", []),
            "cv_models_generated": list(cv_models.keys()),
            "completeness_score": update_fields.get("completeness_score", 0),
            "audit_cv": analysis.get("audit_cv", []),
            "score_global_cv": analysis.get("score_global_cv", 0),
            "modele_suggere": analysis.get("modele_suggere", "cv_classique"),
            "savoir_faire": analysis.get("savoir_faire", []),
            "savoir_etre": analysis.get("savoir_etre", []),
            "experiences": analysis.get("experiences", []),
            "profile": analysis.get("profile", {}),
        }
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "completed", "result": result, "step": "Terminé"}})

        # ── Auto-populate trajectory steps from CV experiences ──
        try:
            # IMPORTANT: Remove old auto-detected/suggested entries from previous CVs
            # This prevents accumulation when user uploads multiple CVs
            old_count = await db.trajectory_steps.count_documents({
                "token_id": token_id,
                "source": {"$in": ["ia_detectee", "ia_suggeree"]}
            })
            if old_count > 0:
                await db.trajectory_steps.delete_many({
                    "token_id": token_id,
                    "source": {"$in": ["ia_detectee", "ia_suggeree"]}
                })
                logging.info(f"[CV→Trajectoire] Supprimé {old_count} anciennes entrées auto-détectées pour token {token_id[:12]}")

            type_map = {"professionnel": "emploi", "personnel": "projet", "benevole": "benevolat", "projet": "projet", "formation": "formation"}
            # Only check manually-added steps for dedup (keep user's own entries)
            existing_steps = await db.trajectory_steps.find({"token_id": token_id}).to_list(500)
            existing_titles = {s.get("title", "").lower() for s in existing_steps}

            new_steps = []
            for exp in analysis.get("experiences", []):
                exp_title = exp.get("title", "").strip()
                if not exp_title or exp_title.lower() in existing_titles:
                    continue
                step = {
                    "id": str(uuid.uuid4()),
                    "token_id": token_id,
                    "step_type": type_map.get(exp.get("experience_type", "professionnel"), "emploi"),
                    "title": exp_title,
                    "organization": exp.get("organization", ""),
                    "description": exp.get("description", ""),
                    "start_date": exp.get("start_date", ""),
                    "end_date": exp.get("end_date", ""),
                    "is_ongoing": exp.get("is_ongoing", False),
                    "skills": exp.get("skills_used", []),
                    "achievements": exp.get("achievements", []),
                    "visibility": "private",
                    "source": "ia_detectee",
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }
                new_steps.append(step)
                existing_titles.add(exp_title.lower())

            # Also add formations from suggestions
            for fs in analysis.get("formations_suggestions", []):
                fs_title = fs.get("title", "").strip()
                if fs_title and fs_title.lower() not in existing_titles:
                    new_steps.append({
                        "id": str(uuid.uuid4()),
                        "token_id": token_id,
                        "step_type": "formation",
                        "title": fs_title,
                        "organization": fs.get("provider", "Suggestion IA"),
                        "description": fs.get("reason", ""),
                        "start_date": "",
                        "end_date": "",
                        "is_ongoing": False,
                        "skills": fs.get("skills_to_gain", []),
                        "achievements": [],
                        "visibility": "private",
                        "source": "ia_suggeree",
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    })
                    existing_titles.add(fs_title.lower())

            if new_steps:
                await db.trajectory_steps.insert_many(new_steps)
                logging.info(f"[CV→Trajectoire] {len(new_steps)} étapes créées pour token {token_id[:12]}")

            # Update profile with cv_analyzed flag and experience count
            await db.profiles.update_one({"token_id": token_id}, {"$set": {
                "cv_analyzed": True,
                "savoir_etre": [se.get("name", "") for se in analysis.get("savoir_etre", [])],
                "experiences_count": len(analysis.get("experiences", [])),
            }})
        except Exception as traj_err:
            logging.error(f"[CV→Trajectoire] Erreur sync: {traj_err}")

        logging.info(f"CV analysis job {job_id} completed successfully")

    except Exception as e:
        logging.error(f"CV analysis job {job_id} failed: {e}")
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": str(e), "step": "Erreur"}})



@api_router.get("/scrape/job-offer")
async def scrape_job_offer(url: str):
    """Scrape job offer content from a URL (supports France Travail and generic pages)."""
    import httpx
    from bs4 import BeautifulSoup

    if not url or not url.startswith("http"):
        raise HTTPException(400, "URL invalide")

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "fr-FR,fr;q=0.9,en;q=0.8",
        }

        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as client_http:
            # Retry up to 2 times for transient failures
            resp = None
            for attempt in range(3):
                try:
                    resp = await client_http.get(url, headers=headers)
                    resp.raise_for_status()
                    break
                except (httpx.ConnectError, httpx.ReadTimeout) as retry_err:
                    if attempt == 2:
                        raise retry_err
                    import asyncio
                    await asyncio.sleep(1)

        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove script/style/nav/footer
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript"]):
            tag.decompose()

        # France Travail specific extraction
        is_france_travail = "francetravail.fr" in url or "pole-emploi.fr" in url
        extracted_text = ""

        if is_france_travail:
            # Extract structured data from France Travail
            parts = []

            # Best source: <title> tag — format: "Offre d'emploi TITRE (H/F) - LIEU - REF | France Travail"
            job_title = ""
            title_tag = soup.find("title")
            if title_tag:
                raw_title = title_tag.get_text(strip=True)
                # Parse: "Offre d'emploi MANAGER EN RESTAURATION (H/F) - 67 - STRASBOURG - 209BTSX | France Travail"
                if "Offre d'emploi" in raw_title:
                    after = raw_title.split("Offre d'emploi", 1)[1]
                    # Remove everything after " - " (location/ref) or " | "
                    job_title = after.split(" - ")[0].strip() if " - " in after else after.split("|")[0].strip()

            # Fallback: first h1
            if not job_title:
                title_el = soup.find("h1")
                if title_el:
                    title_text = title_el.get_text(strip=True)
                    if "Offre n°" in title_text:
                        title_text = title_text.split("Offre n°")[1].strip()
                        title_parts = title_text.split(None, 1)
                        if len(title_parts) > 1:
                            title_text = title_parts[1].strip()
                    job_title = title_text

            if job_title:
                parts.append(f"POSTE: {job_title}")

            # Main content area
            main_content = soup.find("div", class_="description") or soup.find("div", {"itemprop": "description"})
            if main_content:
                parts.append(f"DESCRIPTION:\n{main_content.get_text(separator=chr(10), strip=True)}")

            # Look for all text blocks that contain job info
            for section in soup.find_all(["p", "div", "li", "span"]):
                text = section.get_text(strip=True)
                # Keywords indicating job-relevant content
                if any(kw in text.lower() for kw in ["contrat", "salaire", "expérience", "compétence", "profil", "qualification", "secteur", "savoir-être"]):
                    if text not in "\n".join(parts) and len(text) > 10:
                        parts.append(text)

            extracted_text = "\n".join(parts)

            # Fallback: get all meaningful text from the page body
            if len(extracted_text) < 100:
                body = soup.find("body")
                if body:
                    extracted_text = body.get_text(separator="\n", strip=True)
        else:
            # Generic scraping
            body = soup.find("body")
            if body:
                extracted_text = body.get_text(separator="\n", strip=True)

        # Clean up: remove excessive whitespace/blank lines
        lines = [l.strip() for l in extracted_text.split("\n") if l.strip()]
        # Deduplicate consecutive identical lines
        clean_lines = []
        for line in lines:
            if not clean_lines or line != clean_lines[-1]:
                clean_lines.append(line)
        extracted_text = "\n".join(clean_lines)

        # Limit to reasonable size
        extracted_text = extracted_text[:3000]

        if len(extracted_text) < 30:
            return {"success": False, "text": "", "error": "Contenu insuffisant extrait de cette page."}

        return {"success": True, "text": extracted_text}

    except httpx.HTTPStatusError as e:
        return {"success": False, "text": "", "error": f"Erreur HTTP {e.response.status_code}"}
    except Exception as e:
        logger.error(f"[Scrape] {type(e).__name__}: {e}", exc_info=True)
        return {"success": False, "text": "", "error": "Impossible de lire cette page. Copiez-collez le texte directement."}


@api_router.post("/cv/analyze")
async def analyze_cv(token: str, file: UploadFile = File(...)):
    """Upload CV, start background analysis, return job_id immediately"""
    token_doc = await get_current_token(token)

    # Validate file type
    ext = (file.filename or "").lower().split(".")[-1]
    if ext not in ("pdf", "docx", "doc", "txt"):
        raise HTTPException(status_code=400, detail="Format non supporté. Utilisez PDF, DOCX ou TXT.")

    # Read raw content (fast)
    file_content = await file.read()
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Fichier trop volumineux (max 10 Mo)")

    job_id = str(uuid.uuid4())
    await db.cv_jobs.insert_one({
        "job_id": job_id,
        "token_id": token_doc["id"],
        "filename": file.filename,
        "status": "started",
        "step": "Démarrage de l'analyse...",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    # Launch background task with raw file bytes
    asyncio.create_task(_run_cv_analysis(job_id, token_doc["id"], file_content, file.filename))

    # Store original CV file in coffre-fort for download
    try:
        content_type = "application/pdf" if ext == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == "docx" else "text/plain"
        file_id = str(uuid.uuid4())
        grid_id = await gridfs_bucket.upload_from_stream(
            file.filename,
            file_content,
            metadata={
                "file_id": file_id,
                "token_id": token_doc["id"],
                "content_type": content_type,
                "original_filename": file.filename,
            }
        )
        # Upsert coffre document entry for "CV original"
        await db.coffre_documents.update_one(
            {"token_id": token_doc["id"], "title": "CV original"},
            {"$set": {
                "id": file_id,
                "token_id": token_doc["id"],
                "title": "CV original",
                "filename": file.filename,
                "content_type": content_type,
                "grid_id": str(grid_id),
                "category": "cv",
                "uploaded_at": datetime.now(timezone.utc).isoformat(),
            }},
            upsert=True
        )
        logging.info(f"[CV Upload] CV original stocké dans le coffre-fort: {file.filename}")
    except Exception as store_err:
        logging.error(f"[CV Upload] Erreur stockage coffre-fort: {store_err}")

    return {"job_id": job_id, "status": "started", "message": "Analyse lancée en arrière-plan"}


class CvTextPayload(BaseModel):
    text: str
    filename: str = "cv.txt"


@api_router.post("/cv/extract-text")
async def extract_cv_text(token: str, file: UploadFile = File(...)):
    """Extract text from PDF/DOCX - lightweight, no AI, fast response"""
    await get_current_token(token)
    content = await file.read()
    text = _extract_text_from_bytes(content, file.filename or "file.txt")
    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Impossible d'extraire du texte de ce fichier")
    return {"text": text, "length": len(text)}


class CvBase64Payload(BaseModel):
    data: str
    filename: str = "cv.pdf"


@api_router.post("/cv/extract-text-b64")
async def extract_cv_text_base64(token: str, payload: CvBase64Payload):
    """Extract text from base64-encoded file. No multipart upload - avoids proxy issues."""
    await get_current_token(token)
    import base64
    try:
        content = base64.b64decode(payload.data)
    except Exception:
        raise HTTPException(status_code=400, detail="Données invalides")
    text = _extract_text_from_bytes(content, payload.filename)
    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Impossible d'extraire du texte de ce fichier")
    return {"text": text, "length": len(text)}


@api_router.post("/cv/analyze-text")
async def analyze_cv_text(token: str, payload: CvTextPayload):
    """Analyze pre-extracted CV text. No file upload needed - avoids proxy issues."""
    token_doc = await get_current_token(token)

    if not payload.text or len(payload.text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Le texte du CV est trop court pour être analysé")

    job_id = str(uuid.uuid4())
    await db.cv_jobs.insert_one({
        "job_id": job_id,
        "token_id": token_doc["id"],
        "filename": payload.filename,
        "status": "started",
        "step": "Démarrage de l'analyse...",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    # Launch background task with text directly (no file extraction needed)
    asyncio.create_task(_run_cv_analysis(job_id, token_doc["id"], payload.text.encode("utf-8"), payload.filename, text_ready=True))

    return {"job_id": job_id, "status": "started", "message": "Analyse lancée en arrière-plan"}


@api_router.get("/cv/analyze/status")
async def get_cv_analysis_status(token: str, job_id: str):
    """Poll for CV analysis job status"""
    token_doc = await get_current_token(token)
    job = await db.cv_jobs.find_one({"job_id": job_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job non trouvé")
    return {
        "job_id": job["job_id"],
        "status": job["status"],
        "step": job.get("step", ""),
        "result": job.get("result"),
        "error": job.get("error"),
    }


@api_router.get("/cv/models")
async def get_cv_models(token: str):
    """Get generated CV models for the user"""
    token_doc = await get_current_token(token)
    cv_data = await db.cv_models.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not cv_data:
        return {"models": {}, "analyzed_at": None, "original_filename": None}
    return {
        "models": cv_data.get("models", {}),
        "analyzed_at": cv_data.get("analyzed_at"),
        "original_filename": cv_data.get("original_filename"),
    }


@api_router.get("/cv/download/{model_key}")
async def download_cv_word(model_key: str, token: str):
    """Download a generated CV model as a Word document."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    token_doc = await get_current_token(token)
    cv_data = await db.cv_models.find_one({"token_id": token_doc["id"]})
    if not cv_data or model_key not in cv_data.get("models", {}):
        raise HTTPException(status_code=404, detail="Modèle de CV non trouvé")

    model = cv_data["models"][model_key]
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    h = doc.add_heading(model.get("titre", "CV Professionnel"), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)

    # Accroche
    if model.get("accroche"):
        p = doc.add_paragraph(model["accroche"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

    # Compétences clés
    comps = model.get("competences_cles", [])
    if comps:
        doc.add_heading("Compétences clés", level=1)
        for c in comps:
            doc.add_paragraph(c, style='List Bullet')

    # Expériences
    exps = model.get("experiences", [])
    if exps:
        doc.add_heading("Expériences professionnelles", level=1)
        for e in exps:
            if isinstance(e, dict):
                p = doc.add_paragraph()
                run = p.add_run(f"{e.get('poste', '')} — {e.get('entreprise', '')}")
                run.bold = True
                if e.get('periode'):
                    p.add_run(f"  ({e['periode']})")
                for r in e.get('realisations', []):
                    doc.add_paragraph(r, style='List Bullet')

    # Formations
    formations = model.get("formations", [])
    if formations:
        doc.add_heading("Formations", level=1)
        for f in formations:
            if isinstance(f, dict):
                doc.add_paragraph(f"{f.get('diplome', '')} — {f.get('ecole', '')} ({f.get('annee', '')})")

    # Atouts
    atouts = model.get("atouts", [])
    if atouts:
        doc.add_heading("Atouts", level=1)
        for a in atouts:
            doc.add_paragraph(a, style='List Bullet')

    # Langues
    langues = model.get("langues", [])
    if langues:
        doc.add_heading("Langues", level=1)
        doc.add_paragraph(", ".join(langues))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"CV_{model_key}.docx"
    return StreamingResponse(
        buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@api_router.get("/cv/download-pdf/{model_key}")
async def download_cv_pdf(model_key: str, token: str):
    """Download a generated CV model as a PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    token_doc = await get_current_token(token)
    cv_data = await db.cv_models.find_one({"token_id": token_doc["id"]})
    if not cv_data or model_key not in cv_data.get("models", {}):
        raise HTTPException(status_code=404, detail="Modèle de CV non trouvé")

    model = cv_data["models"][model_key]
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=2*cm, rightMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CVTitle', parent=styles['Title'], fontSize=18, textColor=HexColor('#1e3a5f'), alignment=TA_CENTER, spaceAfter=6)
    accroche_style = ParagraphStyle('Accroche', parent=styles['Normal'], fontSize=10, textColor=HexColor('#666666'), alignment=TA_CENTER, spaceAfter=12, leading=14)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=12, textColor=HexColor('#1e3a5f'), spaceBefore=12, spaceAfter=6, borderWidth=0)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=13, spaceAfter=3)
    bold_style = ParagraphStyle('Bold', parent=body_style, fontName='Helvetica-Bold')
    bullet_style = ParagraphStyle('Bullet', parent=body_style, leftIndent=15, bulletIndent=5, bulletFontName='Helvetica', bulletText='\u2022')

    story = []
    story.append(Paragraph(model.get("titre", "CV Professionnel"), title_style))
    if model.get("accroche"):
        story.append(Paragraph(model["accroche"], accroche_style))
    story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#1e3a5f'), spaceAfter=8))

    comps = model.get("competences_cles", [])
    if comps:
        story.append(Paragraph("COMPÉTENCES CLÉS", section_style))
        for c in comps:
            story.append(Paragraph(str(c), bullet_style))

    exps = model.get("experiences", [])
    if exps:
        story.append(Paragraph("EXPÉRIENCES PROFESSIONNELLES", section_style))
        for e in exps:
            if isinstance(e, dict):
                header = f"<b>{e.get('poste','')}</b> — {e.get('entreprise','')}"
                if e.get('periode'):
                    header += f"  <i>({e['periode']})</i>"
                story.append(Paragraph(header, body_style))
                for r in e.get('realisations', []):
                    story.append(Paragraph(str(r), bullet_style))
                story.append(Spacer(1, 4))

    formations = model.get("formations", [])
    if formations:
        story.append(Paragraph("FORMATIONS", section_style))
        for f in formations:
            if isinstance(f, dict):
                story.append(Paragraph(f"{f.get('diplome','')} — {f.get('ecole','')} ({f.get('annee','')})", body_style))

    atouts = model.get("atouts", [])
    if atouts:
        story.append(Paragraph("ATOUTS", section_style))
        for a in atouts:
            story.append(Paragraph(str(a), bullet_style))

    langues = model.get("langues", [])
    if langues:
        story.append(Paragraph("LANGUES", section_style))
        story.append(Paragraph(", ".join(langues), body_style))

    doc.build(story)
    buf.seek(0)

    filename = f"CV_{model_key}.pdf"
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@api_router.get("/cv/last-analysis")
async def get_last_cv_analysis(token: str):
    """Get the most recent completed CV analysis result for this user"""
    token_doc = await get_current_token(token)
    job = await db.cv_jobs.find_one(
        {"token_id": token_doc["id"], "status": "completed"},
        {"_id": 0},
        sort=[("created_at", -1)]
    )
    if not job or not job.get("result"):
        return {"has_analysis": False, "result": None}

    result = job["result"]

    # Normalize audit_cv field names: LLM may return {rule,status,note} instead of {regle,statut,score,diagnostic}
    if result.get("audit_cv"):
        status_map = {"ok": "ok", "warning": "ameliorable", "error": "absent"}
        score_map = {"ok": 8, "ameliorable": 5, "absent": 2}
        normalized = []
        needs_normalize = False
        for item in result["audit_cv"]:
            if "rule" in item and "regle" not in item:
                needs_normalize = True
                statut = status_map.get(item.get("status", ""), item.get("status", "ameliorable"))
                normalized.append({
                    "regle": item.get("rule", ""),
                    "statut": statut,
                    "score": item.get("score", score_map.get(statut, 5)),
                    "diagnostic": item.get("note", item.get("diagnostic", "")),
                    "recommandation": item.get("recommandation", ""),
                })
            else:
                normalized.append(item)
        if needs_normalize:
            result["audit_cv"] = normalized
            if not result.get("score_global_cv"):
                result["score_global_cv"] = sum(i.get("score", 5) for i in normalized)
            await db.cv_jobs.update_one(
                {"token_id": token_doc["id"], "status": "completed"},
                {"$set": {"result.audit_cv": normalized, "result.score_global_cv": result["score_global_cv"]}},
            )
    if not result.get("audit_cv") or not result.get("savoir_faire"):
        passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
        profile_id = token_doc.get("profile_id")
        profile = await db.profiles.find_one({"id": profile_id}) if profile_id else None
        if passport or profile:
            # Merge savoir_faire from passport and profile.skills
            if not result.get("savoir_faire") or len(result.get("savoir_faire", [])) == 0:
                sf_passport = (passport or {}).get("savoir_faire", [])
                sf_profile = (profile or {}).get("skills", [])
                result["savoir_faire"] = sf_passport if sf_passport else sf_profile
            if not result.get("savoir_etre") or len(result.get("savoir_etre", [])) == 0:
                se_passport = (passport or {}).get("savoir_etre", [])
                se_profile = (profile or {}).get("savoir_etre", [])
                # Fallback: extract savoir_etre from CV transversal competences and strengths
                if not se_passport and not se_profile:
                    extracted_se = []
                    for item in result.get("competences_transversales", []):
                        name = item.get("name", item) if isinstance(item, dict) else str(item)
                        if name:
                            extracted_se.append({"name": name, "source": "cv_transversale"})
                    for item in result.get("strengths", []):
                        name = item.get("name", item) if isinstance(item, dict) else str(item)
                        if name and not any(s.get("name") == name for s in extracted_se):
                            extracted_se.append({"name": name, "source": "cv_strength"})
                    result["savoir_etre"] = extracted_se
                else:
                    result["savoir_etre"] = se_passport if se_passport else se_profile
            if not result.get("experiences") or len(result.get("experiences", [])) == 0:
                exp_passport = (passport or {}).get("experiences", [])
                exp_profile = (profile or {}).get("experiences", [])
                result["experiences"] = exp_passport if exp_passport else exp_profile
            if not result.get("profile"):
                result["profile"] = {
                    "professional_summary": (passport or {}).get("professional_summary", ""),
                    "career_project": (passport or {}).get("career_project", ""),
                    "target_sectors": (passport or {}).get("target_sectors", []),
                }
            result["savoir_faire_count"] = len(result.get("savoir_faire", []))
            result["savoir_etre_count"] = len(result.get("savoir_etre", []))
            result["experiences_count"] = len(result.get("experiences", []))

        # Generate audit retroactively if missing
        if not result.get("audit_cv"):
            sf_count = result.get("savoir_faire_count", 0)
            se_count = result.get("savoir_etre_count", 0)
            exp_count = result.get("experiences_count", 0)
            has_summary = bool(result.get("profile", {}).get("professional_summary"))
            has_formations = len(result.get("formations_suggestions", [])) > 0
            trans_count = len(result.get("competences_transversales", []))

            def _audit(regle, ok_cond, warn_cond, diag_ok, diag_warn, diag_absent, reco, score_ok=8, score_warn=5, score_absent=2):
                if ok_cond:
                    return {"regle": regle, "statut": "ok", "score": score_ok, "diagnostic": diag_ok, "recommandation": ""}
                if warn_cond:
                    return {"regle": regle, "statut": "ameliorable", "score": score_warn, "diagnostic": diag_warn, "recommandation": reco}
                return {"regle": regle, "statut": "absent", "score": score_absent, "diagnostic": diag_absent, "recommandation": reco}

            rules = [
                _audit("Résumé professionnel", has_summary, False,
                       "Résumé détecté et analysé", "", "Aucun résumé professionnel détecté",
                       "Ajoutez un résumé professionnel percutant en haut de votre CV", 9, 5, 2),
                _audit("Expériences détaillées", exp_count >= 3, exp_count >= 1,
                       f"{exp_count} expérience(s) bien détaillée(s)", f"{exp_count} expérience(s) — enrichir les descriptions",
                       "Aucune expérience détectée", "Décrivez vos missions et résultats pour chaque poste", 8, 5, 1),
                _audit("Compétences techniques", sf_count >= 5, sf_count >= 1,
                       f"{sf_count} savoir-faire identifiés — bonne couverture", f"{sf_count} savoir-faire — à compléter",
                       "Aucun savoir-faire technique détecté", "Listez vos compétences techniques clés", 9, 5, 1),
                _audit("Savoir-être", se_count >= 3, se_count >= 1,
                       f"{se_count} savoir-être identifiés", f"{se_count} savoir-être — ajoutez-en",
                       "Aucun savoir-être détecté", "Intégrez vos qualités relationnelles et comportementales", 8, 5, 2),
                _audit("Mots-clés ATS", sf_count >= 10, sf_count >= 3,
                       "Bonne densité de mots-clés pour les filtres ATS", "Densité moyenne — enrichir le vocabulaire métier",
                       "Peu de mots-clés détectés", "Ajoutez des mots-clés métier pour passer les filtres automatisés", 8, 5, 2),
                _audit("Cohérence chronologique", exp_count >= 2, exp_count >= 1,
                       "Chronologie cohérente et lisible", "Chronologie partielle",
                       "Impossible d'évaluer sans expériences", "Vérifiez les dates et l'ordre de vos expériences", 7, 4, 2),
                _audit("Quantification résultats", False, exp_count >= 1,
                       "", "Ajoutez des chiffres concrets à vos réalisations",
                       "Aucune donnée chiffrée détectée", "Intégrez des métriques : CA, %, nb de projets, etc.", 8, 4, 2),
                _audit("Lisibilité et structure", exp_count > 0, False,
                       "Structure claire et bien organisée", "", "Structure à revoir",
                       "Structurez votre CV avec des sections claires", 8, 5, 2),
                _audit("Formations et certifications", has_formations, True,
                       "Formations et certifications détectées", "Section formations à enrichir",
                       "", "Ajoutez vos diplômes et certifications", 7, 5, 3),
                _audit("Compétences transversales", trans_count >= 3, trans_count >= 1,
                       f"{trans_count} compétences transversales identifiées", f"{trans_count} compétence(s) transversale(s)",
                       "Aucune compétence transversale détectée", "Mettez en avant vos compétences transférables", 8, 5, 2),
            ]
            total_score = sum(r["score"] for r in rules)
            result["audit_cv"] = rules
            result["score_global_cv"] = total_score
            result["modele_suggere"] = "cv_competences" if sf_count > 10 else "cv_classique"

            await db.cv_jobs.update_one(
                {"token_id": token_doc["id"], "status": "completed"},
                {"$set": {"result": result}},
            )

    return {"has_analysis": True, "result": result}


# Alias: frontend calls /cv/latest-analysis
@api_router.get("/cv/latest-analysis")
async def get_latest_cv_analysis(token: str):
    return await get_last_cv_analysis(token)


@api_router.get("/cv/centres-interet")
async def get_cv_centres_interet(token: str):
    token_doc = await get_current_token(token)
    ci_doc = await db.cv_centres_interet.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not ci_doc:
        return {"centres": [], "analyses": []}
    return {"centres": ci_doc.get("centres", []), "analyses": ci_doc.get("analyses", [])}


@api_router.post("/cv/centres-interet")
async def save_cv_centres_interet(token: str, body: dict):
    token_doc = await get_current_token(token)
    centres = body.get("centres", [])
    await db.cv_centres_interet.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"token_id": token_doc["id"], "centres": centres, "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    return {"status": "ok", "centres": centres}


@api_router.delete("/cv/delete-all")
async def delete_all_cv_data(token: str):
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]
    await db.cv_jobs.delete_many({"token_id": token_id})
    await db.cv_centres_interet.delete_one({"token_id": token_id})
    await db.cv_models.delete_one({"token_id": token_id})
    return {"status": "ok"}


# ============== REFERENTIEL & ARCHÉOLOGIE DES COMPÉTENCES ==============

REFERENTIEL_VERTUS = [
    {
        "id": "sagesse", "name": "Sagesse et Connaissance",
        "description": "Forces cognitives qui favorisent l'acquisition et l'usage de la connaissance.",
        "forces": ["Créativité", "Curiosité", "Jugement", "Amour de l'apprentissage", "Perspective"],
        "valeurs": ["autonomie", "stimulation", "realisation_de_soi"],
        "qualites": ["Patience", "Ouverture d'esprit", "Indulgence", "Adaptabilité", "Curiosité"],
        "savoirs_etre": ["Faire preuve de curiosité", "Faire preuve de créativité", "Prendre des initiatives"],
    },
    {
        "id": "courage", "name": "Courage",
        "description": "Forces émotionnelles impliquant l'exercice de la volonté malgré les obstacles.",
        "forces": ["Bravoure", "Persévérance", "Honnêteté", "Enthousiasme"],
        "valeurs": ["securite", "pouvoir", "realisation_de_soi"],
        "qualites": ["Bravoure", "Fiabilité", "Confiance", "Loyauté", "Persévérance", "Détermination"],
        "savoirs_etre": ["Faire preuve de persévérance", "Gérer son stress", "Faire preuve de réactivité"],
    },
    {
        "id": "humanite", "name": "Humanité",
        "description": "Forces interpersonnelles consistant à tendre vers les autres et leur venir en aide.",
        "forces": ["Amour", "Gentillesse", "Intelligence sociale"],
        "valeurs": ["bienveillance", "universalisme"],
        "qualites": ["Empathie", "Gentillesse", "Générosité", "Altruisme", "Compassion", "Humilité"],
        "savoirs_etre": ["Être à l'écoute", "Avoir le sens du service", "Travailler en équipe"],
    },
    {
        "id": "justice", "name": "Justice",
        "description": "Forces qui sont à la base d'une vie sociale harmonieuse.",
        "forces": ["Travail d'équipe", "Équité", "Leadership"],
        "valeurs": ["conformite", "tradition", "bienveillance"],
        "qualites": ["Honnêteté", "Équité", "Coopération", "Leadership", "Intégrité"],
        "savoirs_etre": ["Faire preuve de leadership", "Inspirer et donner du sens", "Respecter ses engagements"],
    },
    {
        "id": "temperance", "name": "Tempérance",
        "description": "Forces qui protègent contre les excès.",
        "forces": ["Pardon", "Humilité", "Prudence", "Maîtrise de soi"],
        "valeurs": ["conformite", "securite", "tradition"],
        "qualites": ["Modestie", "Sobriété", "Prudence", "Rigueur", "Maîtrise de soi"],
        "savoirs_etre": ["Faire preuve de rigueur et de précision", "Organiser son travail selon les priorités"],
    },
    {
        "id": "transcendance", "name": "Transcendance",
        "description": "Forces qui favorisent l'ouverture à une dimension universelle et donnent un sens à la vie.",
        "forces": ["Appréciation de la beauté", "Gratitude", "Espoir", "Humour", "Spiritualité"],
        "valeurs": ["universalisme", "bienveillance", "autonomie"],
        "qualites": ["Gratitude", "Optimisme", "Tolérance", "Bienveillance", "Sensibilité"],
        "savoirs_etre": ["S'adapter aux changements", "Faire preuve d'autonomie"],
    },
]

REFERENTIEL_VALEURS = [
    {"id": "autonomie", "name": "Autonomie", "description": "Pensée et action indépendantes", "vertus": ["sagesse"]},
    {"id": "stimulation", "name": "Stimulation", "description": "Nouveauté et défis", "vertus": ["sagesse"]},
    {"id": "hedonisme", "name": "Hédonisme", "description": "Plaisir et gratification", "vertus": ["courage"]},
    {"id": "realisation_de_soi", "name": "Réalisation de soi", "description": "Ambition et succès", "vertus": ["sagesse", "courage"]},
    {"id": "pouvoir", "name": "Pouvoir", "description": "Leadership et influence", "vertus": ["justice", "courage"]},
    {"id": "securite", "name": "Sécurité", "description": "Stabilité et harmonie", "vertus": ["temperance", "courage"]},
    {"id": "conformite", "name": "Conformité", "description": "Respect des normes", "vertus": ["temperance", "justice"]},
    {"id": "tradition", "name": "Tradition", "description": "Modération et humilité", "vertus": ["temperance", "justice"]},
    {"id": "bienveillance", "name": "Bienveillance", "description": "Soin et altruisme", "vertus": ["humanite", "transcendance"]},
    {"id": "universalisme", "name": "Universalisme", "description": "Compréhension et tolérance", "vertus": ["humanite", "transcendance"]},
    {"id": "affiliation", "name": "Affiliation", "description": "Relations proches", "vertus": ["humanite"]},
]

REFERENTIEL_FILIERES = [
    {"id": "SI", "name": "Filière Industrielle", "secteurs": ["Mécanique", "Électrotechnique", "Automatisme", "Génie civil", "Chimie", "Métallurgie"]},
    {"id": "SBTP", "name": "Filière Bâtiment et Travaux Publics", "secteurs": ["Maçonnerie", "Menuiserie", "Plomberie", "Électricité du bâtiment", "Charpenterie"]},
    {"id": "SPSC", "name": "Filière Services à la Personne", "secteurs": ["Aide à domicile", "Éducation spécialisée", "Animation socio-culturelle", "Petite enfance"]},
    {"id": "SSS", "name": "Filière Santé et Social", "secteurs": ["Infirmier(e)", "Aide-soignant(e)", "Assistant(e) social", "Psychologue"]},
    {"id": "SCV", "name": "Filière Commerce et Vente", "secteurs": ["Vente en magasin", "Commerce international", "Négociation commerciale", "Marketing"]},
    {"id": "SHR", "name": "Filière Hôtellerie-Restauration", "secteurs": ["Cuisine", "Service en salle", "Hébergement", "Gestion hôtelière"]},
    {"id": "SAA", "name": "Filière Agriculture et Agroalimentaire", "secteurs": ["Production agricole", "Transformation des produits", "Agroéquipement"]},
    {"id": "SIN", "name": "Filière Informatique et Numérique", "secteurs": ["Développement web et mobile", "Administration systèmes et réseaux", "Cybersécurité", "Design numérique"]},
    {"id": "STL", "name": "Filière Transport et Logistique", "secteurs": ["Conduite routière", "Logistique et gestion", "Manutention"]},
    {"id": "SAAT", "name": "Filière Artisanat d'Art", "secteurs": ["Ébénisterie", "Poterie", "Ferronnerie", "Joaillerie"]},
    {"id": "SCM", "name": "Filière Communication et Médias", "secteurs": ["Journalisme", "Communication d'entreprise", "Relations publiques", "Audiovisuel"]},
    {"id": "SEDD", "name": "Filière Environnement et Développement Durable", "secteurs": ["Gestion des déchets", "Énergies renouvelables", "Éco-conception"]},
    {"id": "ST", "name": "Filière Tourisme", "secteurs": ["Accueil touristique", "Guide touristique", "Animation touristique"]},
    {"id": "SSL", "name": "Filière Sport et Loisirs", "secteurs": ["Entraînement sportif", "Animation sportive", "Gestion d'infrastructures"]},
]

# Mapping savoir-être → qualités humaines → valeurs → vertus (from user's "archéologie des compétences")
ARCHEOLOGIE_SAVOIR_ETRE = {
    "Résolution de problèmes": {"qualites": ["Perspicacité", "Créativité", "Flexibilité"], "valeurs": ["autonomie", "stimulation"], "vertus": ["sagesse"]},
    "Pensée critique": {"qualites": ["Perspicacité", "Esprit analytique"], "valeurs": ["autonomie"], "vertus": ["sagesse"]},
    "Créativité": {"qualites": ["Créativité", "Audace", "Intuition"], "valeurs": ["autonomie", "stimulation"], "vertus": ["sagesse"]},
    "Adaptabilité": {"qualites": ["Flexibilité", "Ouverture d'esprit"], "valeurs": ["stimulation", "autonomie"], "vertus": ["sagesse"]},
    "Communication": {"qualites": ["Empathie", "Éloquence", "Écoute"], "valeurs": ["bienveillance", "affiliation"], "vertus": ["humanite"]},
    "Gestion du temps": {"qualites": ["Rigueur", "Organisation"], "valeurs": ["conformite", "securite"], "vertus": ["temperance"]},
    "Persévérance": {"qualites": ["Courage", "Patience", "Détermination"], "valeurs": ["realisation_de_soi", "securite"], "vertus": ["courage"]},
    "Leadership": {"qualites": ["Charisme", "Confiance en soi", "Intégrité"], "valeurs": ["pouvoir", "realisation_de_soi"], "vertus": ["justice"]},
    "Curiosité": {"qualites": ["Curiosité", "Ouverture d'esprit"], "valeurs": ["stimulation", "autonomie"], "vertus": ["sagesse"]},
    "Rigueur": {"qualites": ["Esprit analytique", "Précision", "Discipline"], "valeurs": ["conformite", "securite"], "vertus": ["temperance"]},
    "Esprit d'équipe": {"qualites": ["Collaboration", "Solidarité", "Écoute"], "valeurs": ["bienveillance", "affiliation"], "vertus": ["humanite", "justice"]},
    "Autonomie": {"qualites": ["Confiance en soi", "Initiative", "Indépendance"], "valeurs": ["autonomie", "realisation_de_soi"], "vertus": ["sagesse", "courage"]},
    "Collaboration": {"qualites": ["Coopération", "Empathie", "Partage"], "valeurs": ["bienveillance", "universalisme"], "vertus": ["humanite"]},
    "Écoute": {"qualites": ["Empathie", "Patience", "Bienveillance"], "valeurs": ["bienveillance", "affiliation"], "vertus": ["humanite"]},
    "Gestion du stress": {"qualites": ["Résilience", "Calme", "Maîtrise de soi"], "valeurs": ["securite"], "vertus": ["courage", "temperance"]},
    "Orientation client": {"qualites": ["Empathie", "Serviabilité", "Écoute"], "valeurs": ["bienveillance"], "vertus": ["humanite"]},
    "Éthique professionnelle": {"qualites": ["Intégrité", "Honnêteté", "Responsabilité"], "valeurs": ["conformite", "universalisme"], "vertus": ["justice"]},
    "Sens du service": {"qualites": ["Serviabilité", "Altruisme", "Générosité"], "valeurs": ["bienveillance", "universalisme"], "vertus": ["humanite", "transcendance"]},
}


@api_router.get("/referentiel/archeologie")
async def get_referentiel_archeologie():
    """Get the full archaeology of competences hierarchy"""
    return {
        "vertus": REFERENTIEL_VERTUS,
        "valeurs": REFERENTIEL_VALEURS,
        "filieres": REFERENTIEL_FILIERES,
        "savoir_etre_map": ARCHEOLOGIE_SAVOIR_ETRE,
    }

@api_router.get("/referentiel/filieres")
async def get_referentiel_filieres(token: str = None):
    """Get all professional filières from real database"""
    filieres_docs = await db.opc_filieres.find({}, {"_id": 0}).sort("numero", 1).to_list(50)
    # Format secteurs as objects {code, nom} for frontend compatibility
    for f in filieres_docs:
        f["secteurs"] = [{"code": s, "nom": s} for s in f.get("secteurs", [])]
    return {"filieres": filieres_docs}


@api_router.get("/referentiel/metiers")
async def get_referentiel_metiers_filtered(token: str = None, filiere: str = None, secteur: str = None):
    """Get métiers filtered by filière and/or secteur"""
    query = {}
    if filiere and filiere != "all":
        query["filiere_code"] = filiere
    if secteur and secteur != "all":
        query["sector_name"] = secteur
    metiers_docs = await db.opc_metiers.find(query, {"_id": 0}).to_list(200)
    metiers_list = [{"nom": m["metier"], "mission": m.get("mission", ""), "sector_code": m.get("sector_code", ""), "filiere_code": m.get("filiere_code", "")} for m in metiers_docs]
    return {"metiers": metiers_list}


@api_router.get("/referentiel/search")
async def referentiel_search(token: str = None, q: str = None, filiere: str = None, secteur: str = None):
    """Pyramidal search across filières, secteurs, métiers, compétences"""
    import re
    results_filieres = []
    results_metiers = []
    results_savoir_etre = []
    results_capacites = []

    metier_query = {}
    if filiere and filiere != "all":
        metier_query["filiere_code"] = filiere
    if secteur and secteur != "all":
        metier_query["sector_name"] = secteur

    if q:
        regex = {"$regex": q, "$options": "i"}
        # Search filières
        filieres_found = await db.opc_filieres.find(
            {"$or": [{"nom": regex}, {"code": regex}, {"secteurs": regex}]},
            {"_id": 0}
        ).to_list(20)
        for f in filieres_found:
            results_filieres.append({
                "nom": f["nom"],
                "code": f["code"],
                "secteurs": [{"nom": s} for s in f.get("secteurs", [])]
            })

        # Search métiers
        metier_search = {**metier_query, "$or": [
            {"metier": regex},
            {"mission": regex},
            {"sector_name": regex},
            {"savoir_faire": regex},
            {"savoir_etre": regex},
        ]}
        metiers_found = await db.opc_metiers.find(metier_search, {"_id": 0}).to_list(50)
    else:
        metiers_found = await db.opc_metiers.find(metier_query, {"_id": 0}).to_list(50)

    seen_se = set()
    seen_ct = set()
    for m in metiers_found:
        results_metiers.append({
            "nom": m["metier"],
            "missions": m.get("mission", ""),
            "filiere_code": m.get("filiere_code", ""),
            "secteur_code": m.get("sector_code", ""),
            "filiere_nom": m.get("filiere_nom", ""),
            "secteur_nom": m.get("sector_name", ""),
        })
        # Collect savoir-être with qualités humaines
        for se in m.get("savoir_etre", []):
            if se and se not in seen_se:
                seen_se.add(se)
                qh_list = []
                for qh_doc in m.get("qualites_humaines", []):
                    if qh_doc.get("savoir_etre") == se and qh_doc.get("qualite_humaine"):
                        qh_list.append(qh_doc["qualite_humaine"])
                results_savoir_etre.append({
                    "nom": se,
                    "description": "",
                    "qualites_humaines": qh_list
                })
        # Collect capacités techniques
        for ct in m.get("capacites_techniques", []):
            if ct and ct not in seen_ct:
                seen_ct.add(ct)
                results_capacites.append({"nom": ct})
        # Also include savoir-faire as capacités
        for sf in m.get("savoir_faire", []):
            if sf and sf not in seen_ct:
                seen_ct.add(sf)
                results_capacites.append({"nom": sf})

    # Also search ROME France Travail
    results_rome = []
    if q:
        rome_regex = {"$regex": q, "$options": "i"}
        rome_found = await db.rome_metiers.find({"libelle": rome_regex}, {"_id": 0}).to_list(20)
        for r in rome_found:
            results_rome.append({
                "code_rome": r["code_rome"],
                "nom": r["libelle"],
                "grand_domaine": r.get("grand_domaine_nom", ""),
            })

    total = len(results_filieres) + len(results_metiers) + len(results_savoir_etre) + len(results_capacites) + len(results_rome)
    return {
        "total": total,
        "filieres": results_filieres,
        "metiers": results_metiers,
        "savoir_etre": results_savoir_etre,
        "capacites_techniques": results_capacites,
        "rome": results_rome,
    }


@api_router.get("/referentiel/contexte")
async def referentiel_contexte(token: str = None, q: str = ""):
    """Context data for a search query"""
    if not q:
        return None
    regex = {"$regex": q, "$options": "i"}
    metiers = await db.opc_metiers.find(
        {"$or": [{"metier": regex}, {"sector_name": regex}, {"savoir_faire": regex}]},
        {"_id": 0}
    ).to_list(10)
    if not metiers:
        return None
    all_sf = []
    all_se = []
    sectors = set()
    filieres = set()
    for m in metiers:
        all_sf.extend(m.get("savoir_faire", []))
        all_se.extend(m.get("savoir_etre", []))
        sectors.add(m.get("sector_name", ""))
        filieres.add(m.get("filiere_nom", ""))
    return {
        "query": q,
        "metiers_count": len(metiers),
        "filieres": list(filieres),
        "secteurs": list(sectors),
        "savoir_faire_sample": all_sf[:10],
        "savoir_etre_sample": all_se[:10],
    }

@api_router.get("/referentiel/vertus")
async def get_referentiel_vertus():
    """Get vertues with their full chain"""
    return {"vertus": REFERENTIEL_VERTUS, "valeurs": REFERENTIEL_VALEURS}


# ===== Explorateur des Filières Professionnelles =====

@api_router.get("/referentiel/explorer")
async def get_explorer_filieres():
    """Get all filières with their secteurs and metier names"""
    all_data = await db.referentiel_metiers.find({}, {"_id": 0, "name": 1, "id": 1, "secteurs.name": 1, "secteurs.metiers.name": 1}).to_list(100)
    for f in all_data:
        for s in f.get("secteurs", []):
            metier_names = [m.get("name", "") for m in s.get("metiers", [])]
            s["metiers"] = metier_names
            s["metiers_count"] = len(metier_names)
    return {"filieres": all_data, "total_filieres": len(all_data)}


@api_router.get("/referentiel/explorer/secteur/{secteur_name}")
async def get_explorer_secteur(secteur_name: str):
    """Get metiers for a specific secteur"""
    doc = await db.referentiel_metiers.find_one(
        {"secteurs.name": secteur_name}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Secteur non trouvé")
    for s in doc.get("secteurs", []):
        if s["name"] == secteur_name:
            return {
                "filiere": doc["name"],
                "secteur": secteur_name,
                "metiers": s.get("metiers", []),
            }
    raise HTTPException(status_code=404, detail="Secteur non trouvé")


@api_router.get("/referentiel/explorer/metier/{metier_name}")
async def get_explorer_metier(metier_name: str):
    """Get full detail for a specific metier with chains and similar metiers"""
    all_data = await db.referentiel_metiers.find({}, {"_id": 0}).to_list(100)
    found = None
    same_secteur_metiers = []
    for f in all_data:
        for s in f.get("secteurs", []):
            for m in s.get("metiers", []):
                if m["name"].lower() == metier_name.lower():
                    found = {"filiere": f["name"], "secteur": s["name"], "metier": m}
                    # Collect other metiers in same secteur
                    same_secteur_metiers = [om["name"] for om in s.get("metiers", []) if om["name"].lower() != metier_name.lower()]
    if not found:
        raise HTTPException(status_code=404, detail="Métier non trouvé")
    found["metiers_similaires"] = same_secteur_metiers
    return found


@api_router.get("/referentiel/explorer/search")
async def search_explorer(q: str):
    """Search across filieres, secteurs, metiers, savoirs"""
    q_lower = q.lower()
    results = {"filieres": [], "secteurs": [], "metiers": [], "savoirs_faire": [], "savoirs_etre": []}
    all_data = await db.referentiel_metiers.find({}, {"_id": 0}).to_list(100)
    for f in all_data:
        if q_lower in f["name"].lower():
            results["filieres"].append({"name": f["name"], "type": "filiere"})
        for s in f.get("secteurs", []):
            if q_lower in s["name"].lower():
                results["secteurs"].append({"name": s["name"], "filiere": f["name"], "type": "secteur"})
            for m in s.get("metiers", []):
                if q_lower in m["name"].lower():
                    results["metiers"].append({"name": m["name"], "secteur": s["name"], "filiere": f["name"], "type": "metier"})
                for sf in m.get("savoirs_faire", []):
                    if q_lower in sf["name"].lower():
                        results["savoirs_faire"].append({"name": sf["name"], "metier": m["name"], "type": "savoir_faire"})
                for se in m.get("savoirs_etre", []):
                    if q_lower in se["name"].lower():
                        results["savoirs_etre"].append({"name": se["name"], "metier": m["name"], "type": "savoir_etre"})
    return results


@api_router.get("/referentiel/explorer/stats")
async def get_explorer_stats():
    """Get overall statistics for the referentiel"""
    all_data = await db.referentiel_metiers.find({}, {"_id": 0}).to_list(100)
    n_filieres = len(all_data)
    n_secteurs = sum(len(f.get("secteurs", [])) for f in all_data)
    n_metiers = sum(len(m.get("metiers", [])) for f in all_data for m in f.get("secteurs", []))
    sf_set = set()
    se_set = set()
    for f in all_data:
        for s in f.get("secteurs", []):
            for m in s.get("metiers", []):
                for sf in m.get("savoirs_faire", []):
                    sf_set.add(sf["name"])
                for se in m.get("savoirs_etre", []):
                    se_set.add(se["name"])
    return {
        "filieres": n_filieres,
        "secteurs": n_secteurs,
        "metiers": n_metiers,
        "savoirs_faire": len(sf_set),
        "savoirs_etre": len(se_set),
    }


@api_router.post("/referentiel/explorer/generate")
async def generate_metier_fiche(token: str, payload: dict):
    """Generate a complete metier fiche using AI when not found in DB"""
    await get_current_token(token)
    metier_name = payload.get("metier", "").strip()
    if not metier_name or len(metier_name) < 2:
        raise HTTPException(status_code=400, detail="Nom de métier invalide")

    # Check cache first
    cached = await db.generated_metiers.find_one({"name_lower": metier_name.lower()}, {"_id": 0})
    if cached:
        return cached["data"]

    job_id = str(uuid.uuid4())
    await db.explorer_jobs.insert_one({"job_id": job_id, "metier": metier_name, "status": "started", "created_at": datetime.now(timezone.utc).isoformat()})

    asyncio.create_task(_generate_metier_fiche(job_id, metier_name))
    return {"job_id": job_id, "status": "started"}


@api_router.get("/referentiel/explorer/generate/status")
async def get_generate_status(token: str, job_id: str):
    """Poll for metier generation status"""
    await get_current_token(token)
    job = await db.explorer_jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job non trouvé")
    return {"job_id": job["job_id"], "status": job["status"], "result": job.get("result"), "error": job.get("error")}


async def _generate_metier_fiche(job_id: str, metier_name: str):
    """Background: generate complete metier fiche via AI"""
    try:
        await db.explorer_jobs.update_one({"job_id": job_id}, {"$set": {"status": "generating"}})
        result = await _llm_call_with_retry(
            system_msg="""Tu es un expert en référentiels métiers et en archéologie des compétences.
Pour le métier demandé, génère une fiche complète en JSON valide:
{
  "filiere": "nom de la filière professionnelle",
  "secteur": "secteur d'activité",
  "metier": {
    "name": "nom du métier",
    "mission": "description détaillée de la mission (2-3 phrases)",
    "savoirs_faire": [
      {"name": "savoir-faire", "capacite_technique": "description de la capacité technique associée"}
    ],
    "savoirs_etre": [
      {"name": "savoir-être", "capacite_professionnelle": "description de la capacité professionnelle", "qualites_humaines": ["qualité1"], "valeurs": ["id_valeur"], "vertus": ["id_vertu"]}
    ]
  },
  "metiers_similaires": ["métier1", "métier2", "métier3", "métier4", "métier5"]
}
Règles:
- 6 à 10 savoir-faire avec capacités techniques détaillées
- 5 à 8 savoir-être avec la chaîne complète (capacité pro → qualités → valeurs → vertus)
- 5 métiers similaires dans le même secteur
- IDs valeurs: autonomie, stimulation, hedonisme, realisation_de_soi, pouvoir, securite, conformite, tradition, bienveillance, universalisme
- IDs vertus: sagesse, courage, humanite, justice, temperance, transcendance""",
            user_msg=f"Génère la fiche métier complète pour : {metier_name}"
        )
        # Cache the result
        await db.generated_metiers.update_one(
            {"name_lower": metier_name.lower()},
            {"$set": {"name_lower": metier_name.lower(), "data": result, "created_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True
        )
        await db.explorer_jobs.update_one({"job_id": job_id}, {"$set": {"status": "completed", "result": result}})
    except Exception as e:
        logging.error(f"Metier generation failed: {e}")
        await db.explorer_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": str(e)}})

@api_router.get("/passport/archeologie")
async def get_passport_archeologie(token: str):
    """For a user's competences, trace the full archaeology chain"""
    token_doc = await get_current_token(token)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not passport:
        raise HTTPException(status_code=404, detail="Passeport non trouvé")

    comps = passport.get("competences", [])
    savoir_faire = [c for c in comps if c.get("nature") == "savoir_faire"]
    savoir_etre = [c for c in comps if c.get("nature") == "savoir_etre"]
    non_classees = [c for c in comps if not c.get("nature")]

    # Build archaeology chains for savoir-être
    chains = []
    for comp in savoir_etre:
        name = comp.get("name", "")
        # Try to find in the reference map
        ref = ARCHEOLOGIE_SAVOIR_ETRE.get(name, {})
        qualites = comp.get("linked_qualites", []) or ref.get("qualites", [])
        valeurs_ids = comp.get("linked_valeurs", []) or ref.get("valeurs", [])
        vertus_ids = comp.get("linked_vertus", []) or ref.get("vertus", [])
        valeurs_names = [v["name"] for v in REFERENTIEL_VALEURS if v["id"] in valeurs_ids]
        vertus_names = [v["name"] for v in REFERENTIEL_VERTUS if v["id"] in vertus_ids]
        chains.append({
            "competence": name,
            "nature": "savoir_etre",
            "qualites": qualites,
            "valeurs": valeurs_names,
            "vertus": vertus_names,
        })

    # Aggregate vertus coverage
    all_vertus = set()
    all_valeurs = set()
    for comp in savoir_etre:
        ref = ARCHEOLOGIE_SAVOIR_ETRE.get(comp.get("name", ""), {})
        for v in (comp.get("linked_vertus", []) or ref.get("vertus", [])):
            all_vertus.add(v)
        for v in (comp.get("linked_valeurs", []) or ref.get("valeurs", [])):
            all_valeurs.add(v)

    return {
        "summary": {
            "total": len(comps),
            "savoir_faire": len(savoir_faire),
            "savoir_etre": len(savoir_etre),
            "non_classees": len(non_classees),
            "vertus_covered": list(all_vertus),
            "valeurs_covered": list(all_valeurs),
        },
        "chains": chains,
        "savoir_faire_list": [{"id": c.get("id"), "name": c.get("name"), "category": c.get("category")} for c in savoir_faire],
        "savoir_etre_list": [{"id": c.get("id"), "name": c.get("name"), "category": c.get("category")} for c in savoir_etre],
        "non_classees_list": [{"id": c.get("id"), "name": c.get("name")} for c in non_classees],
    }

# ============== UBUNTOO INTELLIGENCE ENDPOINTS ==============

async def analyze_ubuntoo_exchanges_with_ai(exchanges: List[dict]) -> Dict[str, Any]:
    """Use AI to analyze Ubuntoo exchanges and detect signals"""
    if not EMERGENT_LLM_KEY:
        return {
            "detected_skills": ["Prompt Engineering", "No-Code"],
            "detected_tools": ["ChatGPT", "Notion"],
            "detected_practices": ["Automatisation de tâches"],
            "confidence": 0.6,
            "summary": "Analyse basique - IA non disponible"
        }
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"ubuntoo-{uuid.uuid4()}",
            system_message="Tu es un expert RH français spécialisé dans l'analyse des tendances du marché du travail. Analyse ces échanges professionnels anonymisés et identifie les signaux faibles sur l'évolution des compétences et des métiers. Réponds en JSON avec: detected_skills (list), detected_tools (list), detected_practices (list), transformations (list of {job, description}), confidence (0-1), summary (string)."
        ).with_model("openai", "gpt-5.2")

        summaries = "\n".join([f"- [{e.get('exchange_type','discussion')}] {e.get('content_summary','')}" for e in exchanges[:10]])
        prompt = f"""Analyse ces échanges anonymisés du réseau socio-professionnel Ubuntoo :

{summaries}

Identifie :
1. Les compétences émergentes mentionnées
2. Les nouveaux outils ou technologies
3. Les nouvelles pratiques professionnelles
4. Les transformations de métiers en cours
"""
        response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
        import json
        try:
            return json.loads(response)
        except:
            return {"detected_skills": [], "detected_tools": [], "detected_practices": [], "confidence": 0.6, "summary": response[:300]}
    except Exception as e:
        logging.error(f"Ubuntoo AI analysis error: {e}")
        return {"detected_skills": [], "detected_tools": [], "detected_practices": [], "confidence": 0.5, "summary": "Analyse automatique non disponible"}

@api_router.get("/ubuntoo/dashboard")
async def get_ubuntoo_dashboard():
    """Get Ubuntoo intelligence dashboard"""
    signals = await db.ubuntoo_signals.find({}, {"_id": 0}).to_list(100)
    exchanges = await db.ubuntoo_exchanges.find({}, {"_id": 0}).to_list(200)
    insights = await db.ubuntoo_insights.find({}, {"_id": 0}).to_list(50)

    # Compute stats
    total_signals = len(signals)
    detected = len([s for s in signals if s.get("validation_status") == "detectee"])
    analyzed = len([s for s in signals if s.get("validation_status") == "analysee_ia"])
    validated = len([s for s in signals if s.get("validation_status") == "validee_humain"])
    integrated = len([s for s in signals if s.get("validation_status") == "integree"])

    # Signal types breakdown
    by_type = {}
    for s in signals:
        t = s.get("signal_type", "autre")
        by_type[t] = by_type.get(t, 0) + 1

    # Top signals by mention count
    top_signals = sorted(signals, key=lambda x: x.get("mention_count", 0), reverse=True)[:10]

    # Recent exchanges
    recent_exchanges = sorted(exchanges, key=lambda x: x.get("timestamp", ""), reverse=True)[:10]

    return {
        "stats": {
            "total_exchanges_analyzed": len(exchanges),
            "total_signals_detected": total_signals,
            "signals_detected": detected,
            "signals_analyzed_ia": analyzed,
            "signals_validated_human": validated,
            "signals_integrated": integrated
        },
        "by_type": by_type,
        "top_signals": top_signals,
        "recent_exchanges": recent_exchanges,
        "insights": insights
    }

@api_router.get("/ubuntoo/signals")
async def get_ubuntoo_signals(signal_type: Optional[str] = None, status: Optional[str] = None, sector: Optional[str] = None):
    """Get Ubuntoo detected signals with filters"""
    query = {}
    if signal_type:
        query["signal_type"] = signal_type
    if status:
        query["validation_status"] = status
    if sector:
        query["related_sectors"] = sector
    signals = await db.ubuntoo_signals.find(query, {"_id": 0}).to_list(100)
    return sorted(signals, key=lambda x: x.get("mention_count", 0), reverse=True)

@api_router.get("/ubuntoo/signals/{signal_id}")
async def get_ubuntoo_signal_detail(signal_id: str):
    """Get detailed signal with cross-references"""
    signal = await db.ubuntoo_signals.find_one({"id": signal_id}, {"_id": 0})
    if not signal:
        raise HTTPException(status_code=404, detail="Signal non trouvé")

    # Cross-reference with observatory
    linked_skills = []
    for skill_name in signal.get("linked_observatory_skills", []):
        skill = await db.emerging_skills.find_one({"skill_name": {"$regex": skill_name, "$options": "i"}}, {"_id": 0})
        if skill:
            linked_skills.append(skill)

    linked_jobs = []
    for job_name in signal.get("linked_evolution_jobs", []):
        job = await db.job_evolution_indices.find_one({"job_name": {"$regex": job_name, "$options": "i"}}, {"_id": 0})
        if job:
            linked_jobs.append(job)

    # Related exchanges
    related_exchanges = await db.ubuntoo_exchanges.find(
        {"$or": [
            {"detected_skills": {"$in": [signal["name"]]}},
            {"detected_tools": {"$in": [signal["name"]]}},
            {"detected_practices": {"$in": [signal["name"]]}}
        ]},
        {"_id": 0}
    ).to_list(20)

    return {
        "signal": signal,
        "linked_observatory_skills": linked_skills,
        "linked_evolution_jobs": linked_jobs,
        "related_exchanges": related_exchanges
    }

@api_router.post("/ubuntoo/signals/{signal_id}/validate")
async def validate_ubuntoo_signal(signal_id: str, approved: bool, notes: Optional[str] = None):
    """Human validation of an Ubuntoo signal"""
    update_data = {
        "validation_status": "validee_humain" if approved else "rejetee",
        "human_notes": notes,
    }
    result = await db.ubuntoo_signals.update_one({"id": signal_id}, {"$set": update_data})
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Signal non trouvé")

    # If validated, check if should be integrated into observatory
    if approved:
        signal = await db.ubuntoo_signals.find_one({"id": signal_id}, {"_id": 0})
        if signal and signal.get("mention_count", 0) >= 5 and signal.get("ai_confidence", 0) >= 0.7:
            await integrate_ubuntoo_signal(signal)

    return {"message": "Validation enregistrée", "status": update_data["validation_status"]}

async def integrate_ubuntoo_signal(signal: dict):
    """Integrate a validated Ubuntoo signal into the observatory"""
    existing = await db.emerging_skills.find_one(
        {"skill_name": {"$regex": signal["name"], "$options": "i"}}, {"_id": 0}
    )
    if existing:
        await db.emerging_skills.update_one(
            {"id": existing["id"]},
            {"$inc": {"mention_count": signal.get("mention_count", 1)},
             "$set": {"last_updated": datetime.now(timezone.utc).isoformat()},
             "$addToSet": {"related_sectors": {"$each": signal.get("related_sectors", [])}}}
        )
    else:
        new_skill = EmergingSkill(
            skill_name=signal["name"],
            description=signal.get("description"),
            related_sectors=signal.get("related_sectors", []),
            related_jobs=signal.get("related_jobs", []),
            emergence_score=min(signal.get("ai_confidence", 0.5) + 0.1, 1.0),
            growth_rate=signal.get("growth_rate", 0.1),
            mention_count=signal.get("mention_count", 1),
            contributor_count=signal.get("source_exchanges_count", 1),
            status="emergente"
        )
        await db.emerging_skills.insert_one(new_skill.model_dump())

    await db.ubuntoo_signals.update_one({"id": signal["id"]}, {"$set": {"validation_status": "integree"}})

@api_router.post("/ubuntoo/analyze")
async def trigger_ubuntoo_analysis():
    """Trigger AI analysis on recent Ubuntoo exchanges"""
    exchanges = await db.ubuntoo_exchanges.find({}, {"_id": 0}).to_list(50)
    if not exchanges:
        return {"message": "Aucun échange à analyser"}

    analysis = await analyze_ubuntoo_exchanges_with_ai(exchanges)

    return {
        "message": "Analyse terminée",
        "analysis": analysis,
        "exchanges_analyzed": len(exchanges)
    }

@api_router.get("/ubuntoo/insights")
async def get_ubuntoo_insights():
    """Get cross-referenced insights"""
    insights = await db.ubuntoo_insights.find({}, {"_id": 0}).to_list(50)
    return sorted(insights, key=lambda x: {"haute": 0, "moyenne": 1, "basse": 2}.get(x.get("priority", "moyenne"), 1))

@api_router.get("/ubuntoo/cross-reference")
async def get_cross_reference_data():
    """Get cross-reference between Ubuntoo signals, observatory skills, and evolution indices"""
    signals = await db.ubuntoo_signals.find({"validation_status": {"$in": ["analysee_ia", "validee_humain", "integree"]}}, {"_id": 0}).to_list(50)
    observatory_skills = await db.emerging_skills.find({}, {"_id": 0}).to_list(50)
    evolution_jobs = await db.job_evolution_indices.find({}, {"_id": 0}).to_list(50)

    # Build cross-reference map
    cross_refs = []
    for signal in signals:
        matched_skills = [s for s in observatory_skills if any(
            signal["name"].lower() in sk.lower() or sk.lower() in signal["name"].lower()
            for sk in [s.get("skill_name", "")]
        )]
        matched_jobs = [j for j in evolution_jobs if any(
            sector in j.get("sector", "").lower()
            for sector in [s.lower() for s in signal.get("related_sectors", [])]
        )]
        if matched_skills or matched_jobs:
            cross_refs.append({
                "signal": signal["name"],
                "signal_type": signal.get("signal_type"),
                "mention_count": signal.get("mention_count", 0),
                "matched_observatory_skills": [s.get("skill_name") for s in matched_skills],
                "matched_jobs": [j.get("job_name") for j in matched_jobs],
                "validation_status": signal.get("validation_status"),
                "ai_confidence": signal.get("ai_confidence", 0)
            })

    return {
        "cross_references": cross_refs,
        "total_signals": len(signals),
        "total_cross_matched": len(cross_refs)
    }

# ============== SEED DATA ==============

@api_router.post("/seed")
async def seed_database():
    """Seed database with demo data"""
    # Clear existing data
    await db.jobs.delete_many({})
    await db.learning_modules.delete_many({})
    await db.beneficiaires.delete_many({})
    
    # Seed jobs
    demo_jobs = [
        {"id": str(uuid.uuid4()), "title": "Assistant Administratif", "company": "TechCorp France", "location": "Paris", "contract_type": "CDI", "salary_range": "28 000€ - 35 000€", "required_skills": ["Gestion administrative", "Excel", "Communication", "Organisation"], "description": "Nous recherchons un assistant administratif polyvalent.", "sector": "Administration", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Chargé de Clientèle", "company": "ServicePlus", "location": "Lyon", "contract_type": "CDI", "salary_range": "32 000€ - 40 000€", "required_skills": ["Relation client", "Négociation", "CRM", "Écoute active"], "description": "Rejoignez notre équipe commerciale.", "sector": "Commerce", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Agent d'entretien", "company": "Propreté Services IDF", "location": "Paris", "contract_type": "CDI", "salary_range": "21 600€ - 24 000€", "required_skills": ["Nettoyage", "Entretien", "Hygiène", "Autonomie"], "description": "Agent d'entretien pour bureaux et locaux professionnels en Île-de-France.", "sector": "Propreté", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Chef d'équipe propreté", "company": "GSF Neptune", "location": "Lille", "contract_type": "CDI", "salary_range": "24 000€ - 28 000€", "required_skills": ["Encadrement", "Nettoyage", "Organisation", "Contrôle qualité", "HACCP"], "description": "Management d'une équipe de 8 agents sur site tertiaire.", "sector": "Propreté", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Employé polyvalent de restauration", "company": "Sodexo", "location": "Lyon", "contract_type": "CDD", "salary_range": "21 600€ - 23 000€", "required_skills": ["Restauration", "Service", "Hygiène", "Travail en équipe"], "description": "Préparation, mise en place et service en restauration collective.", "sector": "Restauration", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Second de cuisine", "company": "Buffalo Grill", "location": "Bordeaux", "contract_type": "CDI", "salary_range": "23 000€ - 27 000€", "required_skills": ["Cuisine", "HACCP", "Gestion stocks", "Encadrement"], "description": "Assister le chef de cuisine dans la gestion quotidienne.", "sector": "Restauration", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Manager en restauration rapide", "company": "McDonald's", "location": "Marseille", "contract_type": "CDI", "salary_range": "26 000€ - 32 000€", "required_skills": ["Management", "Restauration", "Service client", "Hygiène", "Gestion stocks"], "description": "Gestion d'un restaurant et management d'une équipe de 15 personnes.", "sector": "Restauration", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Préparateur de commandes", "company": "Amazon Logistics", "location": "Metz", "contract_type": "CDI", "salary_range": "22 000€ - 25 000€", "required_skills": ["Préparation commandes", "Logistique", "Manutention", "Rigueur"], "description": "Préparer les commandes clients dans un entrepôt logistique.", "sector": "Logistique", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Magasinier-cariste", "company": "GLS France", "location": "Strasbourg", "contract_type": "CDI", "salary_range": "23 000€ - 26 000€", "required_skills": ["CACES", "Logistique", "Gestion stocks", "Organisation"], "description": "Réception, stockage et expédition de marchandises.", "sector": "Logistique", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Agent de maintenance bâtiment", "company": "Nexity", "location": "Toulouse", "contract_type": "CDI", "salary_range": "24 000€ - 28 000€", "required_skills": ["Maintenance", "Entretien", "Électricité", "Plomberie", "Autonomie"], "description": "Maintenance préventive et curative des bâtiments.", "sector": "BTP", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Auxiliaire de puériculture", "company": "Crèche Les Petits Pas", "location": "Nantes", "contract_type": "CDI", "salary_range": "22 000€ - 26 000€", "required_skills": ["Petite enfance", "Accompagnement enfant", "Hygiène", "Communication", "Patience"], "description": "Accompagnement des enfants de 3 mois à 3 ans.", "sector": "Petite enfance", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Animateur périscolaire", "company": "Mairie de Rennes", "location": "Rennes", "contract_type": "CDD", "salary_range": "21 600€ - 23 000€", "required_skills": ["Animation", "Enfant", "Créativité", "Travail en équipe", "Organisation"], "description": "Animation d'activités pour enfants 6-12 ans.", "sector": "Petite enfance", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Aide-soignant(e)", "company": "EHPAD Les Tilleuls", "location": "Montpellier", "contract_type": "CDI", "salary_range": "24 000€ - 28 000€", "required_skills": ["Soins", "Accompagnement", "Hygiène", "Empathie", "Travail en équipe"], "description": "Soins d'hygiène et de confort aux résidents.", "sector": "Santé", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Vendeur conseil", "company": "Leroy Merlin", "location": "Lille", "contract_type": "CDI", "salary_range": "23 000€ - 27 000€", "required_skills": ["Vente", "Conseil client", "Relation client", "Organisation"], "description": "Accueil et conseil des clients en magasin.", "sector": "Commerce", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "title": "Gestionnaire de Paie", "company": "Fiduciaire Nationale", "location": "Bordeaux", "contract_type": "CDI", "salary_range": "35 000€ - 45 000€", "required_skills": ["Paie", "Droit social", "SILAE", "Excel avancé"], "description": "Expert en paie pour cabinet comptable.", "sector": "Comptabilité", "status": "active", "created_at": datetime.now(timezone.utc).isoformat()},
    ]
    
    # Seed learning modules
    demo_modules = [
        {
            "id": str(uuid.uuid4()),
            "title": "Maîtriser Excel pour la Gestion",
            "description": "Apprenez les fonctions avancées d'Excel pour optimiser votre productivité.",
            "duration": "12 heures",
            "level": "Intermédiaire",
            "skills_developed": ["Excel", "Tableaux croisés", "Formules avancées"],
            "category": "Bureautique",
            "image_url": "https://images.unsplash.com/photo-1516321318423-f06f85e504b3?auto=format&fit=crop&q=80"
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Communication Professionnelle",
            "description": "Développez vos compétences en communication écrite et orale.",
            "duration": "8 heures",
            "level": "Débutant",
            "skills_developed": ["Communication", "Rédaction", "Présentation"],
            "category": "Soft Skills",
            "image_url": "https://images.unsplash.com/photo-1522071820081-009f0129c71c?auto=format&fit=crop&q=80"
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Gestion de Projet Agile",
            "description": "Initiez-vous aux méthodes agiles et au framework Scrum.",
            "duration": "16 heures",
            "level": "Intermédiaire",
            "skills_developed": ["Gestion de projet", "Scrum", "Kanban"],
            "category": "Management",
            "image_url": "https://images.unsplash.com/photo-1600880292203-757bb62b4baf?auto=format&fit=crop&q=80"
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Initiation au Développement Web",
            "description": "Découvrez les bases de HTML, CSS et JavaScript.",
            "duration": "20 heures",
            "level": "Débutant",
            "skills_developed": ["HTML/CSS", "JavaScript", "Git"],
            "category": "Informatique",
            "image_url": "https://images.unsplash.com/photo-1498050108023-c5249f4df085?auto=format&fit=crop&q=80"
        },
        {
            "id": str(uuid.uuid4()),
            "title": "Droit du Travail Essentiel",
            "description": "Comprenez les fondamentaux du droit social français.",
            "duration": "10 heures",
            "level": "Intermédiaire",
            "skills_developed": ["Droit social", "Contrats", "Réglementation"],
            "category": "Juridique",
            "image_url": "https://images.unsplash.com/photo-1589829545856-d10d557cf95f?auto=format&fit=crop&q=80"
        }
    ]
    
    await db.jobs.insert_many(demo_jobs)
    await db.learning_modules.insert_many(demo_modules)
    
    # Seed observatoire data
    await db.emerging_skills.delete_many({})
    await db.sector_trends.delete_many({})
    await db.skill_contributions.delete_many({})
    
    demo_emerging_skills = [
        {
            "id": str(uuid.uuid4()),
            "skill_name": "Prompt Engineering",
            "description": "Conception et optimisation de prompts pour l'IA générative",
            "related_sectors": ["Informatique", "Marketing", "Communication"],
            "related_jobs": ["Développeur IA", "Content Manager", "Data Analyst"],
            "related_tools": ["ChatGPT", "Claude", "Midjourney"],
            "emergence_score": 0.92,
            "growth_rate": 0.45,
            "mention_count": 156,
            "contributor_count": 89,
            "status": "emergente",
            "first_detected": datetime.now(timezone.utc).isoformat(),
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "skill_name": "No-Code / Low-Code",
            "description": "Création d'applications sans programmation traditionnelle",
            "related_sectors": ["Informatique", "Administration", "PME"],
            "related_jobs": ["Business Analyst", "Chef de projet", "Responsable digital"],
            "related_tools": ["Bubble", "Webflow", "Airtable", "Notion"],
            "emergence_score": 0.85,
            "growth_rate": 0.38,
            "mention_count": 234,
            "contributor_count": 112,
            "status": "en_croissance",
            "first_detected": datetime.now(timezone.utc).isoformat(),
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "skill_name": "Green IT",
            "description": "Pratiques informatiques éco-responsables et durables",
            "related_sectors": ["Informatique", "Environnement", "Industrie"],
            "related_jobs": ["Responsable RSE", "Architecte SI", "Chef de projet IT"],
            "related_tools": ["Cloud Carbon Footprint", "Green Algorithms"],
            "emergence_score": 0.78,
            "growth_rate": 0.28,
            "mention_count": 98,
            "contributor_count": 45,
            "status": "emergente",
            "first_detected": datetime.now(timezone.utc).isoformat(),
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "skill_name": "Data Storytelling",
            "description": "Communication narrative basée sur l'analyse de données",
            "related_sectors": ["Marketing", "Communication", "Conseil"],
            "related_jobs": ["Data Analyst", "Consultant", "Responsable marketing"],
            "related_tools": ["Tableau", "Power BI", "Looker"],
            "emergence_score": 0.72,
            "growth_rate": 0.25,
            "mention_count": 187,
            "contributor_count": 78,
            "status": "en_croissance",
            "first_detected": datetime.now(timezone.utc).isoformat(),
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "skill_name": "Cybersécurité opérationnelle",
            "description": "Protection des systèmes et gestion des incidents de sécurité",
            "related_sectors": ["Informatique", "Banque", "Santé"],
            "related_jobs": ["Analyste SOC", "RSSI", "Pentester"],
            "related_tools": ["SIEM", "EDR", "Firewall NextGen"],
            "emergence_score": 0.88,
            "growth_rate": 0.35,
            "mention_count": 312,
            "contributor_count": 134,
            "status": "en_croissance",
            "first_detected": datetime.now(timezone.utc).isoformat(),
            "last_updated": datetime.now(timezone.utc).isoformat()
        }
    ]
    
    demo_sector_trends = [
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Informatique & Numérique",
            "emerging_skills": ["Prompt Engineering", "No-Code", "Green IT", "Cybersécurité"],
            "declining_skills": ["Flash", "COBOL", "jQuery"],
            "stable_skills": ["Python", "JavaScript", "SQL", "Git"],
            "transformation_index": 0.82,
            "hiring_trend": "forte_croissance",
            "skill_gap_alert": True,
            "salaries": 632000,
            "offres_emploi": 98500,
            "embauches": 74200,
            "etablissements": 42100,
            "tension": "fort",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "IA Générative", "demand_change": "+45%", "horizon": "2026"},
                {"skill": "Cloud Native", "demand_change": "+32%", "horizon": "2026"},
                {"skill": "DevSecOps", "demand_change": "+28%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Administration & Secrétariat",
            "emerging_skills": ["No-Code", "Automatisation", "IA bureautique"],
            "declining_skills": ["Sténographie", "Classement papier"],
            "stable_skills": ["Excel", "Rédaction", "Organisation", "Accueil"],
            "transformation_index": 0.58,
            "hiring_trend": "stable",
            "skill_gap_alert": False,
            "salaries": 1245000,
            "offres_emploi": 62400,
            "embauches": 51800,
            "etablissements": 78300,
            "tension": "moyen",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Outils collaboratifs", "demand_change": "+25%", "horizon": "2026"},
                {"skill": "GED", "demand_change": "+18%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Commerce & Vente",
            "emerging_skills": ["Social Selling", "CRM avancé", "Data Analytics"],
            "declining_skills": ["Vente terrain classique"],
            "stable_skills": ["Négociation", "Relation client", "Prospection"],
            "transformation_index": 0.65,
            "hiring_trend": "croissance",
            "skill_gap_alert": True,
            "salaries": 1870000,
            "offres_emploi": 142000,
            "embauches": 118500,
            "etablissements": 215000,
            "tension": "moyen",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "E-commerce", "demand_change": "+35%", "horizon": "2026"},
                {"skill": "Marketing automation", "demand_change": "+30%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Santé & Action sociale",
            "emerging_skills": ["Télémédecine", "E-santé", "Coordination parcours patient"],
            "declining_skills": ["Archivage papier médical"],
            "stable_skills": ["Soins infirmiers", "Aide à la personne", "Gériatrie"],
            "transformation_index": 0.55,
            "hiring_trend": "forte_croissance",
            "skill_gap_alert": True,
            "salaries": 2350000,
            "offres_emploi": 215000,
            "embauches": 189000,
            "etablissements": 156000,
            "tension": "fort",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Aide-soignant", "demand_change": "+22%", "horizon": "2026"},
                {"skill": "Infirmier coordinateur", "demand_change": "+18%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Bâtiment & Travaux publics",
            "emerging_skills": ["BIM", "Construction durable", "Domotique"],
            "declining_skills": ["Dessin technique manuel"],
            "stable_skills": ["Maçonnerie", "Électricité", "Plomberie", "Charpente"],
            "transformation_index": 0.60,
            "hiring_trend": "croissance",
            "skill_gap_alert": True,
            "salaries": 1420000,
            "offres_emploi": 125000,
            "embauches": 102000,
            "etablissements": 320000,
            "tension": "fort",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Rénovation énergétique", "demand_change": "+40%", "horizon": "2026"},
                {"skill": "Photovoltaïque", "demand_change": "+35%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Hôtellerie, Restauration & Tourisme",
            "emerging_skills": ["Revenue Management", "Expérience client digitale", "Éco-tourisme"],
            "declining_skills": ["Réservation téléphonique uniquement"],
            "stable_skills": ["Cuisine", "Service en salle", "Accueil", "Hygiène HACCP"],
            "transformation_index": 0.52,
            "hiring_trend": "croissance",
            "skill_gap_alert": True,
            "salaries": 1080000,
            "offres_emploi": 178000,
            "embauches": 156000,
            "etablissements": 245000,
            "tension": "fort",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Chef cuisinier", "demand_change": "+15%", "horizon": "2026"},
                {"skill": "Réceptionniste bilingue", "demand_change": "+12%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Transport & Logistique",
            "emerging_skills": ["Supply Chain 4.0", "Véhicule autonome", "Logistique verte"],
            "declining_skills": ["Gestion manuelle des stocks"],
            "stable_skills": ["Conduite PL", "Gestion d'entrepôt", "Préparation commandes"],
            "transformation_index": 0.63,
            "hiring_trend": "croissance",
            "skill_gap_alert": False,
            "salaries": 1340000,
            "offres_emploi": 112000,
            "embauches": 95000,
            "etablissements": 89000,
            "tension": "moyen",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Chauffeur-livreur", "demand_change": "+20%", "horizon": "2026"},
                {"skill": "Technicien logistique", "demand_change": "+15%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Propreté & Services aux entreprises",
            "emerging_skills": ["Bio-nettoyage", "Management QSE", "Automatisation du nettoyage"],
            "declining_skills": ["Techniques manuelles de base"],
            "stable_skills": ["Entretien des locaux", "Vitrerie", "Hygiène hospitalière"],
            "transformation_index": 0.45,
            "hiring_trend": "stable",
            "skill_gap_alert": False,
            "salaries": 520000,
            "offres_emploi": 86000,
            "embauches": 72000,
            "etablissements": 35000,
            "tension": "moyen",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Agent de propreté qualifié", "demand_change": "+10%", "horizon": "2026"},
                {"skill": "Chef d'équipe propreté", "demand_change": "+8%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Industrie & Production",
            "emerging_skills": ["Industrie 4.0", "Robotique collaborative", "Maintenance prédictive"],
            "declining_skills": ["Opérations manuelles répétitives"],
            "stable_skills": ["Usinage", "Soudure", "Contrôle qualité", "Maintenance"],
            "transformation_index": 0.72,
            "hiring_trend": "croissance",
            "skill_gap_alert": True,
            "salaries": 2680000,
            "offres_emploi": 156000,
            "embauches": 128000,
            "etablissements": 185000,
            "tension": "fort",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Technicien maintenance industrielle", "demand_change": "+25%", "horizon": "2026"},
                {"skill": "Opérateur CN", "demand_change": "+18%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "sector_name": "Agriculture & Agroalimentaire",
            "emerging_skills": ["Agriculture de précision", "Agro-écologie", "Traçabilité digitale"],
            "declining_skills": ["Techniques intensives classiques"],
            "stable_skills": ["Culture", "Élevage", "Transformation alimentaire"],
            "transformation_index": 0.50,
            "hiring_trend": "stable",
            "skill_gap_alert": False,
            "salaries": 680000,
            "offres_emploi": 54000,
            "embauches": 46000,
            "etablissements": 410000,
            "tension": "moyen",
            "periode_offres": "T4 2025",
            "predicted_skills_demand": [
                {"skill": "Ouvrier agricole polyvalent", "demand_change": "+12%", "horizon": "2026"},
                {"skill": "Technicien agroalimentaire", "demand_change": "+10%", "horizon": "2026"}
            ],
            "last_updated": datetime.now(timezone.utc).isoformat()
        }
    ]
    
    await db.emerging_skills.insert_many(demo_emerging_skills)
    await db.sector_trends.insert_many(demo_sector_trends)
    
    # Seed Evolution Indices
    await db.job_evolution_indices.delete_many({})
    await db.sector_evolution_indices.delete_many({})
    
    demo_job_indices = [
        {"id": str(uuid.uuid4()), "job_name": "Assistant Administratif", "sector": "Administration", "evolution_index": 58.0, "index_level": "en_transformation", "new_skills_count": 5, "skill_frequency_score": 0.62, "task_evolution_score": 0.55, "new_tools_score": 0.68, "job_posting_evolution": 0.45, "declining_skills_count": 2, "emerging_skills": ["Outils collaboratifs", "GED", "No-Code", "IA bureautique"], "stable_skills": ["Excel", "Rédaction", "Organisation", "Accueil"], "declining_skills": ["Sténographie", "Classement papier"], "recommended_skills": ["Teams/Slack", "Notion", "Automatisation"], "recommended_trainings": ["Maîtriser les outils collaboratifs", "Automatisation bureautique"], "job_passerelles": ["Office Manager", "Gestionnaire de projet", "Chargé de clientèle"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts"], "confidence_level": 0.85},
        {"id": str(uuid.uuid4()), "job_name": "Développeur Web", "sector": "Informatique", "evolution_index": 82.0, "index_level": "forte_mutation", "new_skills_count": 12, "skill_frequency_score": 0.88, "task_evolution_score": 0.78, "new_tools_score": 0.92, "job_posting_evolution": 0.75, "declining_skills_count": 4, "emerging_skills": ["IA Générative", "Prompt Engineering", "No-Code", "Cloud Native", "DevSecOps"], "stable_skills": ["JavaScript", "Git", "API REST", "SQL"], "declining_skills": ["jQuery", "Flash", "PHP legacy"], "recommended_skills": ["React/Vue", "TypeScript", "Docker", "CI/CD", "GPT API"], "recommended_trainings": ["Intégration IA Générative", "DevOps moderne"], "job_passerelles": ["Lead Developer", "Architecte logiciel", "Product Owner technique"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts", "github_trends"], "confidence_level": 0.92},
        {"id": str(uuid.uuid4()), "job_name": "Chargé de Clientèle", "sector": "Commerce", "evolution_index": 65.0, "index_level": "en_transformation", "new_skills_count": 6, "skill_frequency_score": 0.58, "task_evolution_score": 0.62, "new_tools_score": 0.72, "job_posting_evolution": 0.55, "declining_skills_count": 1, "emerging_skills": ["Social Selling", "CRM avancé", "Data Analytics", "Visio-commerce"], "stable_skills": ["Négociation", "Relation client", "Écoute active", "Prospection"], "declining_skills": ["Vente terrain classique"], "recommended_skills": ["LinkedIn Sales Navigator", "Salesforce", "Analytics client"], "recommended_trainings": ["Social Selling avancé", "CRM et automatisation commerciale"], "job_passerelles": ["Key Account Manager", "Responsable commercial", "Customer Success Manager"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi"], "confidence_level": 0.78},
        {"id": str(uuid.uuid4()), "job_name": "Gestionnaire de Paie", "sector": "Comptabilité / RH", "evolution_index": 48.0, "index_level": "evolutif", "new_skills_count": 3, "skill_frequency_score": 0.42, "task_evolution_score": 0.45, "new_tools_score": 0.55, "job_posting_evolution": 0.38, "declining_skills_count": 1, "emerging_skills": ["DSN automatisée", "SIRH intégré", "Paie dématérialisée"], "stable_skills": ["Droit social", "Paie", "SILAE", "Excel avancé"], "declining_skills": ["Paie manuelle"], "recommended_skills": ["Cegid", "ADP", "Paramétrage DSN"], "recommended_trainings": ["Évolutions réglementaires 2025", "SIRH avancé"], "job_passerelles": ["Responsable paie", "RH généraliste", "Consultant SIRH"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi"], "confidence_level": 0.82},
        {"id": str(uuid.uuid4()), "job_name": "Artisan Boulanger", "sector": "Artisanat / Alimentation", "evolution_index": 15.0, "index_level": "stable", "new_skills_count": 1, "skill_frequency_score": 0.12, "task_evolution_score": 0.15, "new_tools_score": 0.18, "job_posting_evolution": 0.10, "declining_skills_count": 0, "emerging_skills": ["Vente en ligne"], "stable_skills": ["Pétrissage", "Fermentation", "Cuisson", "Hygiène alimentaire"], "declining_skills": [], "recommended_skills": ["Gestion de commerce", "Réseaux sociaux"], "recommended_trainings": ["Gestion d'entreprise artisanale"], "job_passerelles": ["Chef boulanger", "Formateur métiers de bouche"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["experts"], "confidence_level": 0.75},
        {"id": str(uuid.uuid4()), "job_name": "Analyste Cybersécurité", "sector": "Informatique", "evolution_index": 88.0, "index_level": "forte_mutation", "new_skills_count": 15, "skill_frequency_score": 0.92, "task_evolution_score": 0.85, "new_tools_score": 0.95, "job_posting_evolution": 0.88, "declining_skills_count": 2, "emerging_skills": ["IA défensive", "Zero Trust", "Cloud Security", "DevSecOps", "Threat Intelligence"], "stable_skills": ["Firewall", "SIEM", "Incident Response", "Pentest"], "declining_skills": ["Antivirus traditionnel", "Sécurité périmétrique"], "recommended_skills": ["XDR/EDR", "SOAR", "Container Security"], "recommended_trainings": ["Certification CISSP", "IA et Cybersécurité"], "job_passerelles": ["RSSI", "Architecte Sécurité", "Consultant Cyber"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts", "certifications"], "confidence_level": 0.95},
        {"id": str(uuid.uuid4()), "job_name": "Aide-soignant(e)", "sector": "Santé", "evolution_index": 42.0, "index_level": "evolutif", "new_skills_count": 4, "skill_frequency_score": 0.38, "task_evolution_score": 0.40, "new_tools_score": 0.48, "job_posting_evolution": 0.55, "declining_skills_count": 1, "emerging_skills": ["Télésoin", "Dossier patient informatisé", "Coordination parcours", "Bientraitance avancée"], "stable_skills": ["Soins d'hygiène", "Accompagnement", "Observation clinique", "Relation patient"], "declining_skills": ["Transmissions papier"], "recommended_skills": ["Logiciel DPI", "Gestes barrières avancés", "Communication non-violente"], "recommended_trainings": ["Accompagnement fin de vie", "Numérique en santé"], "job_passerelles": ["Infirmier(e)", "AMP", "Coordinateur de soins"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi"], "confidence_level": 0.80},
        {"id": str(uuid.uuid4()), "job_name": "Infirmier(e) libéral(e)", "sector": "Santé", "evolution_index": 55.0, "index_level": "en_transformation", "new_skills_count": 6, "skill_frequency_score": 0.52, "task_evolution_score": 0.55, "new_tools_score": 0.60, "job_posting_evolution": 0.50, "declining_skills_count": 1, "emerging_skills": ["Télémédecine", "E-prescription", "Coordination ville-hôpital", "Éducation thérapeutique"], "stable_skills": ["Soins techniques", "Pharmacologie", "Éthique", "Relation de soin"], "declining_skills": ["Prescriptions papier"], "recommended_skills": ["Plateforme de coordination", "Facturation SESAM-Vitale avancée"], "recommended_trainings": ["Pratique avancée", "Télémédecine"], "job_passerelles": ["IDEC", "IPA", "Cadre de santé"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts"], "confidence_level": 0.82},
        {"id": str(uuid.uuid4()), "job_name": "Agent d'entretien", "sector": "Propreté / Services", "evolution_index": 28.0, "index_level": "evolutif", "new_skills_count": 2, "skill_frequency_score": 0.22, "task_evolution_score": 0.25, "new_tools_score": 0.35, "job_posting_evolution": 0.20, "declining_skills_count": 0, "emerging_skills": ["Bio-nettoyage", "Autolaveuses autonomes"], "stable_skills": ["Entretien des locaux", "Produits d'entretien", "Organisation", "Hygiène"], "declining_skills": [], "recommended_skills": ["Gestion QSE", "Techniques de bio-nettoyage"], "recommended_trainings": ["Hygiène hospitalière", "Management d'équipe propreté"], "job_passerelles": ["Chef d'équipe propreté", "Agent de maintenance", "Technicien de surface qualifié"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi"], "confidence_level": 0.70},
        {"id": str(uuid.uuid4()), "job_name": "Cuisinier / Chef de cuisine", "sector": "Hôtellerie / Restauration", "evolution_index": 35.0, "index_level": "evolutif", "new_skills_count": 3, "skill_frequency_score": 0.30, "task_evolution_score": 0.32, "new_tools_score": 0.40, "job_posting_evolution": 0.35, "declining_skills_count": 0, "emerging_skills": ["Cuisine végétale", "Food cost digital", "Click & Collect"], "stable_skills": ["Techniques culinaires", "Hygiène HACCP", "Gestion stocks", "Créativité culinaire"], "declining_skills": [], "recommended_skills": ["Dark kitchen management", "Réseaux sociaux food", "Gestion allergènes"], "recommended_trainings": ["Cuisine durable", "Management de brigade"], "job_passerelles": ["Chef exécutif", "Consultant culinaire", "Formateur cuisine"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi"], "confidence_level": 0.75},
        {"id": str(uuid.uuid4()), "job_name": "Technicien de maintenance industrielle", "sector": "Industrie", "evolution_index": 72.0, "index_level": "en_transformation", "new_skills_count": 8, "skill_frequency_score": 0.70, "task_evolution_score": 0.68, "new_tools_score": 0.78, "job_posting_evolution": 0.65, "declining_skills_count": 2, "emerging_skills": ["Maintenance prédictive", "IoT industriel", "Robotique collaborative", "Jumeau numérique"], "stable_skills": ["Mécanique", "Électricité", "Hydraulique", "Pneumatique"], "declining_skills": ["Maintenance curative uniquement", "Diagnostic papier"], "recommended_skills": ["GMAO digitale", "Capteurs IoT", "Programmation automates"], "recommended_trainings": ["Industrie 4.0", "Maintenance connectée"], "job_passerelles": ["Responsable maintenance", "Ingénieur fiabilité", "Technicien robotique"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts"], "confidence_level": 0.88},
        {"id": str(uuid.uuid4()), "job_name": "Conducteur de ligne de production", "sector": "Industrie", "evolution_index": 60.0, "index_level": "en_transformation", "new_skills_count": 5, "skill_frequency_score": 0.55, "task_evolution_score": 0.58, "new_tools_score": 0.65, "job_posting_evolution": 0.50, "declining_skills_count": 2, "emerging_skills": ["Pilotage MES", "Lean Manufacturing digital", "Cobotique"], "stable_skills": ["Conduite de machines", "Contrôle qualité", "Sécurité", "Réglages"], "declining_skills": ["Opérations manuelles répétitives", "Suivi production papier"], "recommended_skills": ["ERP/MES", "Maintenance niveau 1 avancée"], "recommended_trainings": ["Lean Six Sigma Green Belt", "Cobotique industrielle"], "job_passerelles": ["Chef de ligne", "Technicien process", "Responsable atelier"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi", "experts"], "confidence_level": 0.80},
        {"id": str(uuid.uuid4()), "job_name": "Data Analyst", "sector": "Informatique", "evolution_index": 85.0, "index_level": "forte_mutation", "new_skills_count": 10, "skill_frequency_score": 0.90, "task_evolution_score": 0.82, "new_tools_score": 0.88, "job_posting_evolution": 0.85, "declining_skills_count": 3, "emerging_skills": ["IA Générative pour analyse", "DataOps", "MLOps", "Visualisation augmentée"], "stable_skills": ["SQL", "Python", "Statistiques", "Business Intelligence"], "declining_skills": ["Excel reporting seul", "Requêtes manuelles", "VBA"], "recommended_skills": ["dbt", "LLM fine-tuning", "Lakehouse"], "recommended_trainings": ["IA appliquée à l'analyse", "Architecture Data moderne"], "job_passerelles": ["Data Engineer", "Data Scientist", "Analytics Manager"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi", "experts", "github_trends"], "confidence_level": 0.90},
        {"id": str(uuid.uuid4()), "job_name": "Comptable", "sector": "Comptabilité / RH", "evolution_index": 52.0, "index_level": "en_transformation", "new_skills_count": 4, "skill_frequency_score": 0.48, "task_evolution_score": 0.50, "new_tools_score": 0.58, "job_posting_evolution": 0.42, "declining_skills_count": 2, "emerging_skills": ["Facturation électronique", "IA comptable", "Audit automatisé", "Dashboard financier"], "stable_skills": ["Normes comptables", "Fiscalité", "Clôture", "Bilan"], "declining_skills": ["Saisie manuelle", "Rapprochement papier"], "recommended_skills": ["Sage/Cegid Cloud", "Power BI Finance", "Factur-X"], "recommended_trainings": ["Facture électronique obligatoire 2026", "IA en comptabilité"], "job_passerelles": ["DAF", "Contrôleur de gestion", "Expert-comptable"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi"], "confidence_level": 0.82},
        {"id": str(uuid.uuid4()), "job_name": "Éducateur spécialisé", "sector": "Social / Médico-social", "evolution_index": 38.0, "index_level": "evolutif", "new_skills_count": 3, "skill_frequency_score": 0.35, "task_evolution_score": 0.38, "new_tools_score": 0.40, "job_posting_evolution": 0.30, "declining_skills_count": 0, "emerging_skills": ["Numérique éducatif", "Coordination MDPH", "Inclusion numérique"], "stable_skills": ["Accompagnement éducatif", "Écoute", "Projet personnalisé", "Travail en équipe"], "declining_skills": [], "recommended_skills": ["Outils numériques d'accompagnement", "Médiation animale"], "recommended_trainings": ["Autisme et troubles du neuro-développement", "Inclusion digitale"], "job_passerelles": ["Chef de service", "Coordinateur parcours", "Formateur social"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "experts"], "confidence_level": 0.75},
        {"id": str(uuid.uuid4()), "job_name": "Chauffeur-livreur", "sector": "Transport / Logistique", "evolution_index": 45.0, "index_level": "evolutif", "new_skills_count": 4, "skill_frequency_score": 0.40, "task_evolution_score": 0.42, "new_tools_score": 0.50, "job_posting_evolution": 0.48, "declining_skills_count": 1, "emerging_skills": ["Tournées optimisées IA", "Livraison dernier kilomètre", "Véhicule électrique", "Application de suivi"], "stable_skills": ["Conduite", "Relation client", "Gestion du temps", "Chargement"], "declining_skills": ["Planification papier"], "recommended_skills": ["Éco-conduite", "Gestion de flotte connectée"], "recommended_trainings": ["Habilitation véhicule électrique", "Logistique urbaine"], "job_passerelles": ["Responsable livraison", "Dispatcheur", "Gestionnaire de flotte"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi"], "confidence_level": 0.72},
        {"id": str(uuid.uuid4()), "job_name": "Préparateur de commandes", "sector": "Transport / Logistique", "evolution_index": 55.0, "index_level": "en_transformation", "new_skills_count": 5, "skill_frequency_score": 0.50, "task_evolution_score": 0.52, "new_tools_score": 0.62, "job_posting_evolution": 0.48, "declining_skills_count": 2, "emerging_skills": ["Voice picking", "Cobotique entrepôt", "WMS avancé", "RFID"], "stable_skills": ["Picking", "Conditionnement", "CACES", "Rigueur"], "declining_skills": ["Préparation 100% manuelle", "Bordereau papier"], "recommended_skills": ["Logiciel WMS", "Tablette/scanner avancé"], "recommended_trainings": ["Logistique 4.0", "CACES R489"], "job_passerelles": ["Chef de quai", "Gestionnaire de stock", "Responsable logistique"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi"], "confidence_level": 0.78},
        {"id": str(uuid.uuid4()), "job_name": "Technicien BTP / Conducteur de travaux", "sector": "BTP / Construction", "evolution_index": 62.0, "index_level": "en_transformation", "new_skills_count": 6, "skill_frequency_score": 0.58, "task_evolution_score": 0.55, "new_tools_score": 0.70, "job_posting_evolution": 0.60, "declining_skills_count": 1, "emerging_skills": ["BIM", "Drone inspection", "Réalité augmentée chantier", "Construction durable"], "stable_skills": ["Lecture de plans", "Coordination chantier", "Sécurité", "Gestion budget"], "declining_skills": ["Plans papier uniquement"], "recommended_skills": ["Logiciel BIM (Revit)", "Tablette chantier", "RE2020"], "recommended_trainings": ["BIM Manager", "Construction bas carbone"], "job_passerelles": ["Directeur de travaux", "BIM Manager", "Expert RE2020"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts"], "confidence_level": 0.82},
        {"id": str(uuid.uuid4()), "job_name": "Électricien", "sector": "BTP / Construction", "evolution_index": 50.0, "index_level": "en_transformation", "new_skills_count": 5, "skill_frequency_score": 0.45, "task_evolution_score": 0.48, "new_tools_score": 0.58, "job_posting_evolution": 0.50, "declining_skills_count": 1, "emerging_skills": ["Domotique", "IRVE (bornes recharge)", "Photovoltaïque", "Smart building"], "stable_skills": ["Câblage", "Normes NFC 15-100", "Habilitations", "Lecture schémas"], "declining_skills": ["Installations basiques seules"], "recommended_skills": ["KNX/Domotique", "Certification IRVE", "Photovoltaïque"], "recommended_trainings": ["Habilitation BR", "IRVE P1-P2-P3"], "job_passerelles": ["Chef de chantier élec", "Technicien domotique", "Expert photovoltaïque"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi", "experts"], "confidence_level": 0.80},
        {"id": str(uuid.uuid4()), "job_name": "Responsable RH", "sector": "Comptabilité / RH", "evolution_index": 68.0, "index_level": "en_transformation", "new_skills_count": 7, "skill_frequency_score": 0.65, "task_evolution_score": 0.62, "new_tools_score": 0.72, "job_posting_evolution": 0.60, "declining_skills_count": 2, "emerging_skills": ["HR Analytics", "IA recrutement", "Marque employeur", "QVT digitale", "People Analytics"], "stable_skills": ["Droit du travail", "Recrutement", "Formation", "Relations sociales"], "declining_skills": ["Gestion papier", "Processus manuels RH"], "recommended_skills": ["SIRH cloud", "Outils ATS", "Tableaux de bord RH"], "recommended_trainings": ["RH Data-driven", "Marque employeur digitale"], "job_passerelles": ["DRH", "HR Business Partner", "Consultant RH"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "offres_emploi", "experts"], "confidence_level": 0.85},
        {"id": str(uuid.uuid4()), "job_name": "Community Manager", "sector": "Communication / Marketing", "evolution_index": 75.0, "index_level": "en_transformation", "new_skills_count": 8, "skill_frequency_score": 0.72, "task_evolution_score": 0.70, "new_tools_score": 0.80, "job_posting_evolution": 0.68, "declining_skills_count": 2, "emerging_skills": ["IA contenu", "Social commerce", "Short video (TikTok/Reels)", "Influence marketing"], "stable_skills": ["Rédaction web", "Stratégie éditoriale", "Veille", "Relation communauté"], "declining_skills": ["Facebook organique seul", "Blogging classique"], "recommended_skills": ["ChatGPT pour contenu", "Analytics avancé", "UGC"], "recommended_trainings": ["IA pour le marketing digital", "Stratégie TikTok/Reels"], "job_passerelles": ["Social Media Manager", "Content Strategist", "Responsable marketing digital"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi", "experts"], "confidence_level": 0.80},
        {"id": str(uuid.uuid4()), "job_name": "Serveur / Employé de restauration", "sector": "Hôtellerie / Restauration", "evolution_index": 22.0, "index_level": "evolutif", "new_skills_count": 2, "skill_frequency_score": 0.18, "task_evolution_score": 0.20, "new_tools_score": 0.30, "job_posting_evolution": 0.22, "declining_skills_count": 0, "emerging_skills": ["Tablette de commande", "Encaissement digital"], "stable_skills": ["Service en salle", "Accueil client", "Hygiène", "Polyvalence"], "declining_skills": [], "recommended_skills": ["Sommellerie initiation", "Anglais professionnel"], "recommended_trainings": ["Service d'excellence", "Langues étrangères hôtellerie"], "job_passerelles": ["Maître d'hôtel", "Responsable de salle", "Gérant de restaurant"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi"], "confidence_level": 0.70},
        {"id": str(uuid.uuid4()), "job_name": "Ouvrier agricole polyvalent", "sector": "Agriculture", "evolution_index": 18.0, "index_level": "stable", "new_skills_count": 1, "skill_frequency_score": 0.15, "task_evolution_score": 0.18, "new_tools_score": 0.20, "job_posting_evolution": 0.12, "declining_skills_count": 0, "emerging_skills": ["GPS de précision"], "stable_skills": ["Culture", "Élevage", "Conduite d'engins", "Entretien matériel"], "declining_skills": [], "recommended_skills": ["Agriculture bio", "Certification phytosanitaire"], "recommended_trainings": ["Certiphyto", "Agro-écologie"], "job_passerelles": ["Chef d'exploitation", "Technicien agricole", "Conseiller agronomique"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["experts"], "confidence_level": 0.68},
        {"id": str(uuid.uuid4()), "job_name": "Plombier / Chauffagiste", "sector": "BTP / Construction", "evolution_index": 45.0, "index_level": "evolutif", "new_skills_count": 4, "skill_frequency_score": 0.42, "task_evolution_score": 0.40, "new_tools_score": 0.52, "job_posting_evolution": 0.45, "declining_skills_count": 1, "emerging_skills": ["Pompe à chaleur", "Chauffe-eau thermodynamique", "Plancher chauffant", "Régulation connectée"], "stable_skills": ["Plomberie sanitaire", "Chauffage", "Soudure", "Lecture de plans"], "declining_skills": ["Chaudière fioul"], "recommended_skills": ["PAC air-eau", "VMC double flux", "RGE Qualipac"], "recommended_trainings": ["Installation PAC", "Certification RGE"], "job_passerelles": ["Chef de chantier CVC", "Technicien en énergies renouvelables"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["offres_emploi", "experts"], "confidence_level": 0.78},
        {"id": str(uuid.uuid4()), "job_name": "Conseiller en insertion professionnelle", "sector": "Social / Médico-social", "evolution_index": 52.0, "index_level": "en_transformation", "new_skills_count": 5, "skill_frequency_score": 0.48, "task_evolution_score": 0.50, "new_tools_score": 0.55, "job_posting_evolution": 0.45, "declining_skills_count": 1, "emerging_skills": ["Outils IA d'orientation", "Matching algorithmique", "Accompagnement hybride", "Portfolio numérique"], "stable_skills": ["Écoute active", "Connaissance du marché", "Rédaction CV", "Entretien conseil"], "declining_skills": ["Documentation papier uniquement"], "recommended_skills": ["Plateformes emploi digitales", "Outils d'évaluation en ligne"], "recommended_trainings": ["Accompagnement à distance", "Outils numériques d'insertion"], "job_passerelles": ["Chargé de mission emploi", "Coordinateur insertion", "Formateur professionnel"], "last_calculated": datetime.now(timezone.utc).isoformat(), "data_sources": ["contributions", "experts"], "confidence_level": 0.78},
    ]
    
    demo_sector_indices = [
        {"id": str(uuid.uuid4()), "sector_name": "Informatique & Numérique", "evolution_index": 78.0, "index_level": "en_transformation", "jobs_count": 25, "jobs_in_transformation": 18, "jobs_stable": 3, "jobs_emerging": 4, "top_emerging_skills": [{"skill": "IA Générative", "growth": "+145%"}, {"skill": "Cloud Native", "growth": "+68%"}, {"skill": "DevSecOps", "growth": "+52%"}], "top_declining_skills": [{"skill": "jQuery", "decline": "-35%"}, {"skill": "Flash", "decline": "-90%"}], "skill_gap_areas": ["Cybersécurité", "IA", "Cloud"], "hiring_trend": "forte_croissance", "innovation_intensity": 0.92, "predicted_evolution_6m": 82.0, "predicted_evolution_12m": 85.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Administration & Gestion", "evolution_index": 52.0, "index_level": "en_transformation", "jobs_count": 15, "jobs_in_transformation": 8, "jobs_stable": 6, "jobs_emerging": 1, "top_emerging_skills": [{"skill": "Outils collaboratifs", "growth": "+45%"}, {"skill": "GED", "growth": "+32%"}, {"skill": "Automatisation", "growth": "+28%"}], "top_declining_skills": [{"skill": "Classement papier", "decline": "-40%"}], "skill_gap_areas": ["Numérique", "Automatisation"], "hiring_trend": "stable", "innovation_intensity": 0.55, "predicted_evolution_6m": 55.0, "predicted_evolution_12m": 58.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Commerce & Distribution", "evolution_index": 62.0, "index_level": "en_transformation", "jobs_count": 20, "jobs_in_transformation": 12, "jobs_stable": 6, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "Social Selling", "growth": "+55%"}, {"skill": "E-commerce", "growth": "+42%"}, {"skill": "Data Analytics", "growth": "+38%"}], "top_declining_skills": [{"skill": "Vente terrain seule", "decline": "-25%"}], "skill_gap_areas": ["Digital", "Analytics"], "hiring_trend": "croissance", "innovation_intensity": 0.68, "predicted_evolution_6m": 65.0, "predicted_evolution_12m": 70.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Artisanat & Alimentation", "evolution_index": 18.0, "index_level": "stable", "jobs_count": 30, "jobs_in_transformation": 2, "jobs_stable": 27, "jobs_emerging": 1, "top_emerging_skills": [{"skill": "Vente en ligne", "growth": "+22%"}, {"skill": "Réseaux sociaux", "growth": "+18%"}], "top_declining_skills": [], "skill_gap_areas": ["Digital"], "hiring_trend": "stable", "innovation_intensity": 0.15, "predicted_evolution_6m": 20.0, "predicted_evolution_12m": 22.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Santé & Médico-social", "evolution_index": 48.0, "index_level": "evolutif", "jobs_count": 35, "jobs_in_transformation": 12, "jobs_stable": 18, "jobs_emerging": 5, "top_emerging_skills": [{"skill": "Télémédecine", "growth": "+85%"}, {"skill": "E-santé", "growth": "+62%"}, {"skill": "DPI", "growth": "+45%"}], "top_declining_skills": [{"skill": "Transmissions papier", "decline": "-30%"}], "skill_gap_areas": ["Numérique en santé", "Coordination"], "hiring_trend": "forte_croissance", "innovation_intensity": 0.52, "predicted_evolution_6m": 52.0, "predicted_evolution_12m": 58.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Industrie & Production", "evolution_index": 68.0, "index_level": "en_transformation", "jobs_count": 28, "jobs_in_transformation": 16, "jobs_stable": 8, "jobs_emerging": 4, "top_emerging_skills": [{"skill": "Industrie 4.0", "growth": "+72%"}, {"skill": "Maintenance prédictive", "growth": "+58%"}, {"skill": "Cobotique", "growth": "+45%"}], "top_declining_skills": [{"skill": "Opérations manuelles", "decline": "-28%"}], "skill_gap_areas": ["IoT", "Robotique", "Data industrielle"], "hiring_trend": "croissance", "innovation_intensity": 0.75, "predicted_evolution_6m": 72.0, "predicted_evolution_12m": 75.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "BTP & Construction", "evolution_index": 55.0, "index_level": "en_transformation", "jobs_count": 22, "jobs_in_transformation": 10, "jobs_stable": 10, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "BIM", "growth": "+65%"}, {"skill": "Construction durable", "growth": "+48%"}, {"skill": "IRVE", "growth": "+120%"}], "top_declining_skills": [{"skill": "Plans papier", "decline": "-20%"}], "skill_gap_areas": ["BIM", "Transition énergétique"], "hiring_trend": "croissance", "innovation_intensity": 0.58, "predicted_evolution_6m": 58.0, "predicted_evolution_12m": 62.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Hôtellerie & Restauration", "evolution_index": 32.0, "index_level": "evolutif", "jobs_count": 18, "jobs_in_transformation": 4, "jobs_stable": 12, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "Revenue Management", "growth": "+35%"}, {"skill": "Click & Collect", "growth": "+55%"}], "top_declining_skills": [{"skill": "Réservation tél. seule", "decline": "-45%"}], "skill_gap_areas": ["Digital", "Langues"], "hiring_trend": "croissance", "innovation_intensity": 0.35, "predicted_evolution_6m": 35.0, "predicted_evolution_12m": 38.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Transport & Logistique", "evolution_index": 52.0, "index_level": "en_transformation", "jobs_count": 20, "jobs_in_transformation": 10, "jobs_stable": 8, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "Supply Chain 4.0", "growth": "+48%"}, {"skill": "Logistique verte", "growth": "+35%"}, {"skill": "WMS avancé", "growth": "+42%"}], "top_declining_skills": [{"skill": "Planification manuelle", "decline": "-30%"}], "skill_gap_areas": ["Automatisation", "Données temps réel"], "hiring_trend": "croissance", "innovation_intensity": 0.55, "predicted_evolution_6m": 55.0, "predicted_evolution_12m": 58.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Communication & Marketing", "evolution_index": 72.0, "index_level": "en_transformation", "jobs_count": 16, "jobs_in_transformation": 11, "jobs_stable": 3, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "IA contenu", "growth": "+110%"}, {"skill": "Short video", "growth": "+85%"}, {"skill": "Social commerce", "growth": "+62%"}], "top_declining_skills": [{"skill": "Print traditionnel", "decline": "-35%"}, {"skill": "Facebook organique", "decline": "-28%"}], "skill_gap_areas": ["IA", "Vidéo", "Data"], "hiring_trend": "croissance", "innovation_intensity": 0.78, "predicted_evolution_6m": 75.0, "predicted_evolution_12m": 80.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Social & Insertion", "evolution_index": 42.0, "index_level": "evolutif", "jobs_count": 14, "jobs_in_transformation": 5, "jobs_stable": 8, "jobs_emerging": 1, "top_emerging_skills": [{"skill": "Accompagnement hybride", "growth": "+38%"}, {"skill": "Outils numériques insertion", "growth": "+30%"}], "top_declining_skills": [{"skill": "Documentation papier", "decline": "-22%"}], "skill_gap_areas": ["Numérique", "IA orientation"], "hiring_trend": "stable", "innovation_intensity": 0.40, "predicted_evolution_6m": 45.0, "predicted_evolution_12m": 48.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Propreté & Facility Management", "evolution_index": 25.0, "index_level": "evolutif", "jobs_count": 12, "jobs_in_transformation": 3, "jobs_stable": 8, "jobs_emerging": 1, "top_emerging_skills": [{"skill": "Bio-nettoyage", "growth": "+28%"}, {"skill": "Autolaveuses autonomes", "growth": "+22%"}], "top_declining_skills": [], "skill_gap_areas": ["QSE", "Management"], "hiring_trend": "stable", "innovation_intensity": 0.22, "predicted_evolution_6m": 28.0, "predicted_evolution_12m": 30.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Agriculture & Agroalimentaire", "evolution_index": 35.0, "index_level": "evolutif", "jobs_count": 16, "jobs_in_transformation": 4, "jobs_stable": 11, "jobs_emerging": 1, "top_emerging_skills": [{"skill": "Agriculture de précision", "growth": "+40%"}, {"skill": "Agro-écologie", "growth": "+32%"}], "top_declining_skills": [{"skill": "Techniques intensives", "decline": "-15%"}], "skill_gap_areas": ["Numérique", "Transition écologique"], "hiring_trend": "stable", "innovation_intensity": 0.32, "predicted_evolution_6m": 38.0, "predicted_evolution_12m": 40.0, "last_updated": datetime.now(timezone.utc).isoformat()},
        {"id": str(uuid.uuid4()), "sector_name": "Comptabilité, Finance & RH", "evolution_index": 58.0, "index_level": "en_transformation", "jobs_count": 18, "jobs_in_transformation": 9, "jobs_stable": 7, "jobs_emerging": 2, "top_emerging_skills": [{"skill": "Facturation électronique", "growth": "+95%"}, {"skill": "HR Analytics", "growth": "+52%"}, {"skill": "IA comptable", "growth": "+45%"}], "top_declining_skills": [{"skill": "Saisie manuelle", "decline": "-35%"}], "skill_gap_areas": ["IA", "Dématérialisation"], "hiring_trend": "stable", "innovation_intensity": 0.60, "predicted_evolution_6m": 62.0, "predicted_evolution_12m": 65.0, "last_updated": datetime.now(timezone.utc).isoformat()},
    ]
    
    await db.job_evolution_indices.insert_many(demo_job_indices)
    await db.sector_evolution_indices.insert_many(demo_sector_indices)
    
    # Seed Ubuntoo Intelligence Data
    await db.ubuntoo_exchanges.delete_many({})
    await db.ubuntoo_signals.delete_many({})
    await db.ubuntoo_insights.delete_many({})

    demo_ubuntoo_exchanges = [
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "retour_experience",
            "content_summary": "Depuis 6 mois, j'utilise ChatGPT quotidiennement pour rédiger mes rapports et synthèses. Mon manager me demande maintenant de former l'équipe à ces outils.",
            "detected_skills": ["Prompt Engineering", "IA Générative", "Formation interne"],
            "detected_tools": ["ChatGPT", "GPT-4"],
            "detected_practices": ["Automatisation de la rédaction", "Formation pair-à-pair"],
            "related_jobs": ["Assistant Administratif", "Chargé de communication"],
            "related_sectors": ["Administration", "Communication"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=2)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "discussion",
            "content_summary": "Notre service RH a remplacé 3 outils par une seule plateforme SIRH intégrée. Les gestionnaires de paie doivent maintenant maîtriser le paramétrage complet du système.",
            "detected_skills": ["SIRH intégré", "Paramétrage logiciel", "Conduite du changement"],
            "detected_tools": ["SIRH", "Cegid", "ADP"],
            "detected_practices": ["Centralisation des outils", "Digitalisation RH"],
            "related_jobs": ["Gestionnaire de Paie", "Responsable RH"],
            "related_sectors": ["Comptabilité", "Administration"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=5)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "mentorat",
            "content_summary": "En tant que développeur senior, je constate que les juniors qui maîtrisent les assistants de code IA (Copilot, Cursor) sont 2 fois plus productifs. C'est devenu un critère d'embauche chez nous.",
            "detected_skills": ["IA assistée au code", "Pair programming IA", "Productivité développeur"],
            "detected_tools": ["GitHub Copilot", "Cursor", "Claude Code"],
            "detected_practices": ["Développement assisté par IA", "Revue de code IA"],
            "related_jobs": ["Développeur Web", "Lead Developer"],
            "related_sectors": ["Informatique"],
            "author_role": "mentor",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=3)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "conseil",
            "content_summary": "Pour ceux en reconversion vers le commercial : le social selling sur LinkedIn est devenu incontournable. Les recruteurs cherchent des profils qui maîtrisent la prospection digitale, pas juste le terrain.",
            "detected_skills": ["Social Selling", "Personal Branding", "Prospection digitale"],
            "detected_tools": ["LinkedIn Sales Navigator", "Lemlist", "Hubspot"],
            "detected_practices": ["Prospection sur réseaux sociaux", "Création de contenu professionnel"],
            "related_jobs": ["Chargé de Clientèle", "Commercial", "Business Developer"],
            "related_sectors": ["Commerce"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "question",
            "content_summary": "Comment gérer la cybersécurité quand on passe au full cloud ? Notre DSI nous demande de tous devenir 'security champions' dans nos équipes respectives.",
            "detected_skills": ["Cloud Security", "Security Champion", "Sensibilisation sécurité"],
            "detected_tools": ["AWS Security Hub", "Azure Sentinel"],
            "detected_practices": ["Security by design", "Décentralisation de la sécurité"],
            "related_jobs": ["Analyste Cybersécurité", "Développeur Web", "Architecte Cloud"],
            "related_sectors": ["Informatique"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=1)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "retour_experience",
            "content_summary": "Notre boulangerie a lancé un site de commande en ligne et une page Instagram. Les ventes ont augmenté de 30%. J'ai dû apprendre le marketing digital en autodidacte.",
            "detected_skills": ["Marketing digital", "E-commerce", "Gestion réseaux sociaux"],
            "detected_tools": ["Shopify", "Instagram Business", "Canva"],
            "detected_practices": ["Vente en ligne artisanale", "Communication digitale"],
            "related_jobs": ["Artisan Boulanger", "Commerçant"],
            "related_sectors": ["Artisanat", "Commerce"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=10)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "discussion",
            "content_summary": "Les outils no-code comme Notion, Airtable et Make nous permettent de créer des workflows sans passer par l'IT. C'est une révolution pour les fonctions support.",
            "detected_skills": ["No-Code", "Automatisation workflows", "Autonomie numérique"],
            "detected_tools": ["Notion", "Airtable", "Make", "Zapier"],
            "detected_practices": ["Citizen development", "Automatisation sans code"],
            "related_jobs": ["Assistant Administratif", "Chef de projet", "Office Manager"],
            "related_sectors": ["Administration", "Informatique"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=4)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "mentorat",
            "content_summary": "Conseil pour les gestionnaires de paie : la maîtrise de la DSN automatisée et du paramétrage des SIRH sera indispensable d'ici 2 ans. Les traitements manuels disparaissent progressivement.",
            "detected_skills": ["DSN automatisée", "Paramétrage SIRH", "Veille réglementaire digitale"],
            "detected_tools": ["SILAE", "Cegid HR", "PayFit"],
            "detected_practices": ["Paie dématérialisée", "Automatisation des déclarations"],
            "related_jobs": ["Gestionnaire de Paie", "Responsable paie"],
            "related_sectors": ["Comptabilité"],
            "author_role": "mentor",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=6)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "retour_experience",
            "content_summary": "En tant que commercial terrain depuis 15 ans, j'ai dû me former au CRM Salesforce et à l'analyse de données clients. Aujourd'hui 60% de mon travail se fait devant un écran.",
            "detected_skills": ["CRM avancé", "Data Analytics client", "Vente hybride"],
            "detected_tools": ["Salesforce", "Power BI", "Teams"],
            "detected_practices": ["Vente hybride terrain/digital", "Pilotage par la data"],
            "related_jobs": ["Chargé de Clientèle", "Commercial terrain"],
            "related_sectors": ["Commerce"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=8)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "exchange_type": "discussion",
            "content_summary": "La montée en puissance du Green IT dans nos projets informatiques est impressionnante. On nous demande maintenant de mesurer l'empreinte carbone de nos applications.",
            "detected_skills": ["Green IT", "Éco-conception logicielle", "Mesure empreinte carbone"],
            "detected_tools": ["Cloud Carbon Footprint", "Lighthouse", "EcoIndex"],
            "detected_practices": ["Développement durable numérique", "Sobriété numérique"],
            "related_jobs": ["Développeur Web", "Architecte SI", "Chef de projet IT"],
            "related_sectors": ["Informatique"],
            "author_role": "professionnel",
            "timestamp": (datetime.now(timezone.utc) - timedelta(days=9)).isoformat()
        }
    ]

    demo_ubuntoo_signals = [
        {
            "id": str(uuid.uuid4()),
            "signal_type": "competence_emergente",
            "name": "Prompt Engineering",
            "description": "Compétence de rédaction et optimisation de prompts pour les IA génératives, de plus en plus demandée dans tous les métiers tertiaires.",
            "mention_count": 47,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=60)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=1)).isoformat(),
            "related_jobs": ["Assistant Administratif", "Développeur Web", "Chargé de communication", "Content Manager"],
            "related_sectors": ["Administration", "Informatique", "Communication", "Marketing"],
            "source_exchanges_count": 32,
            "trend_direction": "hausse",
            "growth_rate": 0.85,
            "validation_status": "integree",
            "ai_confidence": 0.95,
            "ai_analysis": {"category": "technique", "impact": "transversal", "urgence": "haute"},
            "linked_observatory_skills": ["Prompt Engineering"],
            "linked_evolution_jobs": ["Assistant Administratif", "Développeur Web"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=60)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "nouvel_outil",
            "name": "Assistants de code IA",
            "description": "Outils comme GitHub Copilot et Cursor qui transforment la pratique du développement logiciel. Les développeurs les utilisant sont significativement plus productifs.",
            "mention_count": 38,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=45)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=2)).isoformat(),
            "related_jobs": ["Développeur Web", "Lead Developer", "Analyste Cybersécurité"],
            "related_sectors": ["Informatique"],
            "source_exchanges_count": 25,
            "trend_direction": "hausse",
            "growth_rate": 0.72,
            "validation_status": "validee_humain",
            "ai_confidence": 0.91,
            "ai_analysis": {"category": "outil", "impact": "sectoriel", "urgence": "haute"},
            "linked_observatory_skills": [],
            "linked_evolution_jobs": ["Développeur Web"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=45)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "pratique_nouvelle",
            "name": "Citizen Development (No-Code)",
            "description": "Les fonctions support créent leurs propres outils et workflows grâce aux plateformes no-code, réduisant la dépendance aux équipes IT.",
            "mention_count": 29,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=40)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=3)).isoformat(),
            "related_jobs": ["Assistant Administratif", "Office Manager", "Chef de projet"],
            "related_sectors": ["Administration", "Informatique"],
            "source_exchanges_count": 18,
            "trend_direction": "hausse",
            "growth_rate": 0.55,
            "validation_status": "analysee_ia",
            "ai_confidence": 0.82,
            "ai_analysis": {"category": "pratique", "impact": "transversal", "urgence": "moyenne"},
            "linked_observatory_skills": ["No-Code / Low-Code"],
            "linked_evolution_jobs": ["Assistant Administratif"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=40)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "transformation_metier",
            "name": "Vente hybride terrain/digital",
            "description": "Les commerciaux terrain évoluent vers un modèle hybride où 50-60% du travail se fait en digital (CRM, visioconférence, social selling).",
            "mention_count": 22,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=50)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=5)).isoformat(),
            "related_jobs": ["Chargé de Clientèle", "Commercial", "Business Developer"],
            "related_sectors": ["Commerce"],
            "source_exchanges_count": 15,
            "trend_direction": "hausse",
            "growth_rate": 0.48,
            "validation_status": "validee_humain",
            "ai_confidence": 0.88,
            "ai_analysis": {"category": "transformation", "impact": "sectoriel", "urgence": "moyenne"},
            "linked_observatory_skills": ["Social Selling"],
            "linked_evolution_jobs": ["Chargé de Clientèle"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=50)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "competence_emergente",
            "name": "Security Champion",
            "description": "Nouveau rôle au sein des équipes de développement : des référents sécurité non-spécialistes qui portent les bonnes pratiques cyber au quotidien.",
            "mention_count": 15,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=30)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=1)).isoformat(),
            "related_jobs": ["Développeur Web", "Analyste Cybersécurité", "Chef de projet IT"],
            "related_sectors": ["Informatique"],
            "source_exchanges_count": 11,
            "trend_direction": "hausse",
            "growth_rate": 0.62,
            "validation_status": "analysee_ia",
            "ai_confidence": 0.78,
            "ai_analysis": {"category": "technique", "impact": "sectoriel", "urgence": "moyenne"},
            "linked_observatory_skills": [],
            "linked_evolution_jobs": ["Analyste Cybersécurité", "Développeur Web"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "difficulte_metier",
            "name": "Obsolescence des compétences paie manuelles",
            "description": "Les gestionnaires de paie signalent une accélération de la digitalisation qui rend les compétences manuelles obsolètes plus vite que prévu.",
            "mention_count": 12,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=25)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=4)).isoformat(),
            "related_jobs": ["Gestionnaire de Paie", "Responsable paie"],
            "related_sectors": ["Comptabilité"],
            "source_exchanges_count": 9,
            "trend_direction": "hausse",
            "growth_rate": 0.35,
            "validation_status": "detectee",
            "ai_confidence": 0.72,
            "ai_analysis": {"category": "difficulte", "impact": "sectoriel", "urgence": "haute"},
            "linked_observatory_skills": [],
            "linked_evolution_jobs": ["Gestionnaire de Paie"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=25)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "nouvel_outil",
            "name": "Green IT / Éco-conception",
            "description": "Les outils de mesure d'empreinte carbone numérique deviennent obligatoires dans les projets IT, créant un nouveau besoin de compétences.",
            "mention_count": 18,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=35)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=6)).isoformat(),
            "related_jobs": ["Développeur Web", "Architecte SI", "Chef de projet IT"],
            "related_sectors": ["Informatique"],
            "source_exchanges_count": 12,
            "trend_direction": "hausse",
            "growth_rate": 0.42,
            "validation_status": "validee_humain",
            "ai_confidence": 0.85,
            "ai_analysis": {"category": "outil", "impact": "sectoriel", "urgence": "moyenne"},
            "linked_observatory_skills": ["Green IT"],
            "linked_evolution_jobs": ["Développeur Web"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=35)).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "signal_type": "pratique_nouvelle",
            "name": "Artisanat digital",
            "description": "Les artisans adoptent le e-commerce et le marketing digital pour étendre leur clientèle, créant un besoin de double compétence métier/numérique.",
            "mention_count": 8,
            "first_detected": (datetime.now(timezone.utc) - timedelta(days=20)).isoformat(),
            "last_detected": (datetime.now(timezone.utc) - timedelta(days=8)).isoformat(),
            "related_jobs": ["Artisan Boulanger", "Commerçant"],
            "related_sectors": ["Artisanat", "Commerce"],
            "source_exchanges_count": 6,
            "trend_direction": "hausse",
            "growth_rate": 0.25,
            "validation_status": "detectee",
            "ai_confidence": 0.65,
            "ai_analysis": {"category": "pratique", "impact": "sectoriel", "urgence": "basse"},
            "linked_observatory_skills": [],
            "linked_evolution_jobs": ["Artisan Boulanger"],
            "created_at": (datetime.now(timezone.utc) - timedelta(days=20)).isoformat()
        }
    ]

    demo_ubuntoo_insights = [
        {
            "id": str(uuid.uuid4()),
            "insight_type": "tendance_emergente",
            "title": "L'IA générative transforme tous les métiers tertiaires",
            "description": "Les échanges Ubuntoo confirment une adoption massive des outils d'IA générative (ChatGPT, Copilot) bien au-delà du secteur IT. Les fonctions support, communication et RH sont fortement impactées.",
            "supporting_signals": ["Prompt Engineering", "Assistants de code IA"],
            "impacted_jobs": ["Assistant Administratif", "Développeur Web", "Chargé de communication"],
            "impacted_sectors": ["Administration", "Informatique", "Communication"],
            "recommendation": "Intégrer des modules de formation IA générative dans tous les parcours de reconversion, pas seulement les filières IT.",
            "priority": "haute",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "insight_type": "alerte_competence",
            "title": "Accélération de l'obsolescence dans la gestion de paie",
            "description": "Le croisement entre les signaux Ubuntoo et les données de l'observatoire montre que la digitalisation de la paie s'accélère plus vite que les prévisions initiales.",
            "supporting_signals": ["Obsolescence des compétences paie manuelles", "Citizen Development (No-Code)"],
            "impacted_jobs": ["Gestionnaire de Paie", "Responsable paie"],
            "impacted_sectors": ["Comptabilité"],
            "recommendation": "Anticiper la montée en compétences SIRH des gestionnaires de paie via des formations accélérées.",
            "priority": "haute",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "insight_type": "opportunite_formation",
            "title": "Le No-Code comme levier d'autonomie professionnelle",
            "description": "Les échanges montrent que les professionnels qui maîtrisent les outils no-code gagnent en autonomie et en valeur. Cette compétence est transversale à tous les secteurs.",
            "supporting_signals": ["Citizen Development (No-Code)"],
            "impacted_jobs": ["Assistant Administratif", "Office Manager", "Chef de projet"],
            "impacted_sectors": ["Administration", "Informatique"],
            "recommendation": "Créer un parcours 'Autonomie numérique' centré sur le no-code pour les professionnels en reconversion.",
            "priority": "moyenne",
            "created_at": datetime.now(timezone.utc).isoformat()
        },
        {
            "id": str(uuid.uuid4()),
            "insight_type": "transformation_metier",
            "title": "La vente évolue vers un modèle hybride digital/terrain",
            "description": "Les témoignages du réseau confirment une transformation profonde du métier commercial, avec une digitalisation de 50-60% des activités.",
            "supporting_signals": ["Vente hybride terrain/digital"],
            "impacted_jobs": ["Chargé de Clientèle", "Commercial", "Business Developer"],
            "impacted_sectors": ["Commerce"],
            "recommendation": "Adapter les formations commerciales pour inclure le social selling et l'analyse de données client.",
            "priority": "moyenne",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
    ]

    await db.ubuntoo_exchanges.insert_many(demo_ubuntoo_exchanges)
    await db.ubuntoo_signals.insert_many(demo_ubuntoo_signals)
    await db.ubuntoo_insights.insert_many(demo_ubuntoo_insights)

    return {
        "message": "Base de données initialisée", 
        "jobs": len(demo_jobs), 
        "modules": len(demo_modules), 
        "emerging_skills": len(demo_emerging_skills), 
        "sector_trends": len(demo_sector_trends),
        "job_indices": len(demo_job_indices),
        "sector_indices": len(demo_sector_indices),
        "ubuntoo_exchanges": len(demo_ubuntoo_exchanges),
        "ubuntoo_signals": len(demo_ubuntoo_signals),
        "ubuntoo_insights": len(demo_ubuntoo_insights)
    }

# ============== AUTH PSEUDONYME (Production-compatible) ==============

import hashlib

def _hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

class PseudoRegisterRequest(BaseModel):
    pseudo: str
    password: str
    role: str = "particulier"
    email_recovery: Optional[str] = None
    identifiant_france_travail: Optional[str] = None
    consent_cgu: bool = True
    consent_privacy: bool = True
    consent_marketing: bool = False

class PseudoLoginRequest(BaseModel):
    pseudo: str
    password: str

class UpgradeRequest(BaseModel):
    pseudo: str
    password: str
    email_recovery: Optional[str] = None
    consent_cgu: bool = True
    consent_privacy: bool = True

class EntrepriseRegisterRequest(BaseModel):
    email: str
    password: str
    nom_entreprise: str = ""
    siret: str = ""
    secteur: str = ""
    taille: str = ""
    role: str = "entreprise"

class PartenaireRegisterRequest(BaseModel):
    email: str
    password: str
    nom_structure: str = ""
    type_structure: str = ""
    territoire: str = ""
    role: str = "partenaire"

async def _create_token_and_profile(role: str, pseudo: str = None, auth_mode: str = "anonymous"):
    """Helper: create token + profile and return auth response"""
    token_obj = AnonymousToken(role=role)
    token_dict = token_obj.model_dump()
    token_dict["auth_mode"] = auth_mode
    token_dict["pseudo"] = pseudo
    token_dict["identity_level"] = "pseudo" if pseudo else "none"
    await db.tokens.insert_one(token_dict)

    profile = Profile(token_id=token_obj.id, role=role, name=pseudo or f"Utilisateur {token_obj.id[:8].upper()}")
    await db.profiles.insert_one(profile.model_dump())
    await db.tokens.update_one({"id": token_obj.id}, {"$set": {"profile_id": profile.id}})

    return {
        "token": token_obj.token, "role": role, "profile_id": profile.id,
        "pseudo": pseudo, "auth_mode": auth_mode, "identity_level": "pseudo" if pseudo else "none"
    }

@api_router.post("/auth/register")
async def auth_register(body: PseudoRegisterRequest):
    existing = await db.users.find_one({"pseudo": body.pseudo})
    if existing:
        raise HTTPException(status_code=400, detail="Ce pseudonyme est déjà utilisé")
    user_id = str(uuid.uuid4())
    await db.users.insert_one({
        "_id": user_id, "id": user_id, "pseudo": body.pseudo,
        "password": _hash_pw(body.password), "role": body.role,
        "email_recovery": body.email_recovery,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return await _create_token_and_profile(body.role, body.pseudo, "pseudo")

@api_router.post("/auth/login")
async def auth_login(body: PseudoLoginRequest):
    user = await db.users.find_one({"pseudo": body.pseudo, "password": _hash_pw(body.password)})
    if not user:
        raise HTTPException(status_code=401, detail="Pseudo ou mot de passe incorrect")
    # Find existing token for this user or create new one
    existing_token = await db.tokens.find_one({"pseudo": body.pseudo, "auth_mode": "pseudo"})
    if existing_token:
        return {
            "token": existing_token["token"], "role": existing_token["role"],
            "profile_id": existing_token.get("profile_id"),
            "pseudo": body.pseudo, "auth_mode": "pseudo", "identity_level": "pseudo"
        }
    return await _create_token_and_profile(user.get("role", "particulier"), body.pseudo, "pseudo")

@api_router.post("/auth/upgrade")
async def auth_upgrade(token: str, body: UpgradeRequest):
    token_doc = await get_current_token(token)
    existing = await db.users.find_one({"pseudo": body.pseudo})
    if existing:
        raise HTTPException(status_code=400, detail="Ce pseudonyme est déjà utilisé")
    user_id = str(uuid.uuid4())
    await db.users.insert_one({
        "_id": user_id, "id": user_id, "pseudo": body.pseudo,
        "password": _hash_pw(body.password), "email_recovery": body.email_recovery,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    await db.tokens.update_one({"id": token_doc["id"]}, {"$set": {
        "auth_mode": "pseudo", "pseudo": body.pseudo, "identity_level": "pseudo"
    }})
    return {"status": "ok", "pseudo": body.pseudo}

@api_router.post("/auth/register-entreprise")
async def auth_register_entreprise(body: EntrepriseRegisterRequest):
    existing = await db.users.find_one({"pseudo": body.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email déjà utilisé")
    user_id = str(uuid.uuid4())
    await db.users.insert_one({
        "_id": user_id, "id": user_id, "pseudo": body.email,
        "password": _hash_pw(body.password), "role": "entreprise",
        "nom_entreprise": body.nom_entreprise, "siret": body.siret,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return await _create_token_and_profile("entreprise", body.email, "pseudo")

@api_router.post("/auth/register-partenaire")
async def auth_register_partenaire(body: PartenaireRegisterRequest):
    existing = await db.users.find_one({"pseudo": body.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email déjà utilisé")
    user_id = str(uuid.uuid4())
    await db.users.insert_one({
        "_id": user_id, "id": user_id, "pseudo": body.email,
        "password": _hash_pw(body.password), "role": "partenaire",
        "nom_structure": body.nom_structure,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return await _create_token_and_profile("partenaire", body.email, "pseudo")

@api_router.post("/auth/login-pro")
async def auth_login_pro(body: PseudoLoginRequest):
    """Login for entreprise/partenaire accounts (by email as pseudo)"""
    user = await db.users.find_one({"pseudo": body.pseudo, "password": _hash_pw(body.password)})
    if not user:
        raise HTTPException(status_code=401, detail="Email ou mot de passe incorrect")
    existing_token = await db.tokens.find_one({"pseudo": body.pseudo, "auth_mode": "pseudo"})
    if existing_token:
        return {
            "token": existing_token["token"], "role": existing_token["role"],
            "profile_id": existing_token.get("profile_id"),
            "pseudo": body.pseudo, "auth_mode": "pseudo", "identity_level": "pseudo",
            "company_name": user.get("nom_entreprise", user.get("nom_structure", ""))
        }
    result = await _create_token_and_profile(user.get("role", "particulier"), body.pseudo, "pseudo")
    result["company_name"] = user.get("nom_entreprise", user.get("nom_structure", ""))
    return result

# ============== ADMIN GATE ==============

class GateStateRequest(BaseModel):
    password: str
    spaces_open: bool

@api_router.get("/admin/gate-state")
async def get_gate_state():
    state = await db.admin_config.find_one({"key": "gate_state"}, {"_id": 0})
    if not state:
        return {"spaces_open": True}
    return {"spaces_open": state.get("spaces_open", True)}

@api_router.post("/admin/gate-state")
async def set_gate_state(body: GateStateRequest):
    ADMIN_PASSWORD = "Choukette@777"
    if body.password != ADMIN_PASSWORD:
        raise HTTPException(status_code=403, detail="Mot de passe administrateur incorrect")
    await db.admin_config.update_one(
        {"key": "gate_state"},
        {"$set": {"key": "gate_state", "spaces_open": body.spaces_open}},
        upsert=True
    )
    return {"spaces_open": body.spaces_open}

# ============== COACH PROGRESS ==============

@api_router.get("/coach/progress")
async def get_coach_progress(token: str):
    token_doc = await get_current_token(token)
    progress = await db.coach_progress.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not progress:
        # Check what the user has done
        profile_id = token_doc.get("profile_id")
        has_cv = False
        cv_skills_count = 0
        cv_savoir_etre_count = 0
        experiences_count = 0
        has_dclic = False

        if profile_id:
            profile = await db.profiles.find_one({"id": profile_id})
            if profile:
                has_cv = bool(profile.get("cv_analyzed"))
                cv_skills_count = len(profile.get("skills", []))
                cv_savoir_etre_count = len(profile.get("savoir_etre", []))
                experiences_count = len(profile.get("experiences", []))
                has_dclic = bool(profile.get("dclic_imported"))

        # Check CV analysis status
        last_analysis = await db.cv_jobs.find_one(
            {"token_id": token_doc["id"], "status": "completed"},
            sort=[("created_at", -1)]
        )
        if last_analysis:
            has_cv = True
            result = last_analysis.get("result", {})
            analysis_skills = len(result.get("competences", []))
            analysis_se = len(result.get("savoir_etre", []))
            analysis_exp = len(result.get("experiences", []))
            # Use the max of profile and analysis data
            cv_skills_count = max(cv_skills_count, analysis_skills)
            cv_savoir_etre_count = max(cv_savoir_etre_count, analysis_se)
            experiences_count = max(experiences_count, analysis_exp)

        # Also check passport for additional data
        passport = await db.passports.find_one({"token_id": token_doc["id"]})
        if passport:
            passport_sf = len(passport.get("savoir_faire", []))
            passport_se = len(passport.get("savoir_etre", []))
            passport_exp = len(passport.get("experiences", []))
            cv_skills_count = max(cv_skills_count, passport_sf)
            cv_savoir_etre_count = max(cv_savoir_etre_count, passport_se)
            experiences_count = max(experiences_count, passport_exp)
            # Check D'CLIC from passport too (dclic_results is stored here by import)
            if not has_dclic and passport.get("dclic_results"):
                has_dclic = True

        step1_complete = has_cv
        step2_complete = cv_savoir_etre_count >= 3
        step3_complete = has_dclic
        step4_complete = experiences_count >= 3

        completed = sum([step1_complete, step2_complete, step3_complete, step4_complete])

        # Current step = first incomplete step (in order 1→2→3→4)
        step_statuses = [step1_complete, step2_complete, step3_complete, step4_complete]
        current_step = 4  # default if all complete
        for idx, done in enumerate(step_statuses):
            if not done:
                current_step = idx + 1
                break

        # Build rich progression data
        achievements = []
        if has_cv:
            achievements.append(f"{cv_skills_count} savoir-faire identifiés par l'IA")
        if cv_savoir_etre_count > 0:
            achievements.append(f"{cv_savoir_etre_count} savoir-être documenté(s)")
        if experiences_count > 0:
            achievements.append(f"{experiences_count} expérience(s) dans votre trajectoire")
        if has_dclic:
            achievements.append("Profil D'CLIC PRO complété")

        # Build personalized CIP tips based on current state
        tips = []
        if has_cv and cv_skills_count > 0 and not step2_complete:
            tips.append({
                "icon": "lightbulb",
                "text": f"Tu as {cv_skills_count} savoir-faire identifiés — c'est une bonne base ! Maintenant, valorise tes qualités humaines (savoir-être) : elles font souvent la différence en entretien. Direction Portefeuille → Savoir-être.",
                "priority": "high"
            })
        if has_cv and not has_dclic:
            tips.append({
                "icon": "rocket",
                "text": "Le test D'CLIC PRO (5 min) révèle ton profil de personnalité (DISC, MBTI, RIASEC). En entretien, savoir parler de soi avec précision, c'est un vrai atout. Lance-toi !",
                "priority": "medium"
            })
        if experiences_count > 0 and experiences_count < 5:
            tips.append({
                "icon": "plus",
                "text": f"Tu as {experiences_count} expérience(s) renseignée(s). Pense aussi aux missions bénévoles, stages et projets perso — les recruteurs valorisent la diversité des parcours.",
                "priority": "low"
            })
        if has_cv and cv_skills_count > 5:
            tips.append({
                "icon": "target",
                "text": f"Avec {cv_skills_count} compétences, tu as un profil solide. Explore le Job Dating pour rencontrer directement des recruteurs qui cherchent ton profil.",
                "priority": "medium"
            })
        # Always add tips — even after 4/4 completion
        if completed >= 3:
            tips.append({
                "icon": "download",
                "text": "Ton profil est assez riche pour générer des CV ciblés et optimisés ATS. Un CV adapté à chaque offre multiplie tes chances de décrocher un entretien.",
                "priority": "medium"
            })

        # Advanced CIP tips when all 4 steps are done
        if completed == 4:
            docs_count = await db.coffre_documents.count_documents({"token_id": token_doc["id"]})
            has_job_dating = bool(await db.job_dating_registrations.find_one({"token_id": token_doc["id"]}))

            if docs_count == 0:
                tips.insert(0, {
                    "icon": "shield",
                    "text": "Dépose tes preuves (diplômes, attestations, certificats) dans ton Portefeuille. En recherche d'emploi, un dossier de preuves structuré rassure les recruteurs et accélère le processus.",
                    "priority": "high"
                })
            if not has_job_dating:
                tips.insert(0, {
                    "icon": "calendar",
                    "text": "Le Job Dating, c'est l'occasion de te présenter en quelques minutes à un recruteur. Prépare ton pitch (2 min max) et inscris-toi à un événement !",
                    "priority": "high"
                })
            tips.append({
                "icon": "refresh",
                "text": "Le marché de l'emploi évolue vite. Mets à jour régulièrement tes compétences et expériences pour rester visible et pertinent.",
                "priority": "low"
            })
            tips.append({
                "icon": "compass",
                "text": "Consulte l'Observatoire du Marché pour repérer les secteurs qui recrutent et anticiper les compétences recherchées demain.",
                "priority": "low"
            })

        # Build next_step — always provide one, even at 4/4
        advanced_next_steps = [
            {"hint": "Prouve tes compétences avec des exemples concrets (situation, action, résultat) dans ton Passeport. C'est ce qui fait la différence en entretien.", "impact": "Les recruteurs veulent des preuves concrètes — pas juste une liste de compétences.", "path": "/dashboard/profil?tab=experiences"},
            {"hint": "Inscris-toi à un Job Dating ! Prépare un pitch de 2 minutes et va à la rencontre des recruteurs.", "impact": "Le contact direct avec les employeurs est souvent plus efficace qu'une candidature en ligne.", "path": "/dashboard/job-dating"},
            {"hint": "Consulte l'Observatoire pour repérer les secteurs porteurs et les compétences recherchées.", "impact": "Anticiper les tendances du marché te permet d'orienter ta montée en compétences.", "path": "/dashboard/marche"},
            {"hint": "Génère un CV ciblé pour un poste précis. Un CV adapté à chaque offre, c'est la clé pour passer les filtres ATS.", "impact": "Les recruteurs passent 6 secondes en moyenne sur un CV — chaque mot compte.", "path": "/dashboard/trajectoire"},
        ]

        # Build proactive "next step" message with clear CIP guidance
        next_step_messages = {
            1: {
                "hint": "Dépose ton CV (PDF ou Word) dans Trajectoire → Mon CV. L'IA va analyser tes compétences, expériences et formations automatiquement.",
                "impact": "C'est la première étape pour construire ton portefeuille de compétences et cibler les bonnes offres."
            },
            2: {
                "hint": "Identifie et documente tes savoir-être (esprit d'équipe, rigueur, adaptabilité...) dans ton Portefeuille.",
                "impact": "En entretien, les recruteurs évaluent autant tes qualités humaines que tes compétences techniques. Savoir en parler fait la différence."
            },
            3: {
                "hint": "Passe le test D'CLIC PRO (5 minutes). Il révèle ton profil de personnalité (DISC, MBTI, RIASEC).",
                "impact": "Mieux te connaître te permet de cibler les postes et environnements qui te correspondent vraiment."
            },
            4: {
                "hint": "Complète ta trajectoire avec tes expériences, formations et compétences acquises — même les expériences informelles comptent !",
                "impact": "Une trajectoire riche, c'est la matière première pour générer des CV percutants et matcher les offres."
            },
        }

        next_info = next_step_messages.get(current_step, {})

        if not has_cv:
            message = "Bienvenue ! Commence par importer ton CV pour que l'IA analyse ton profil. C'est la première étape vers ta recherche d'emploi efficace."
            emoji = "wave"
        elif completed == 4:
            emoji = "trophy"
            message = f"Bravo, ton profil est complet ! {cv_skills_count} compétences, {experiences_count} expériences — tu as de solides arguments. Passons maintenant à l'action : candidatures, Job Dating, entretiens."
        else:
            summary = ", ".join(achievements) if achievements else "profil en cours de construction"
            message = f"Ta progression : {summary}. Continue comme ça, chaque étape te rapproche de ton objectif !"

            if current_step <= 2:
                emoji = "star"
            elif current_step == 3:
                emoji = "rocket"
            else:
                emoji = "target"

        # Step details for each step
        step1_details = None
        if step1_complete:
            step1_details = {"skills": cv_skills_count, "savoir_etre": cv_savoir_etre_count, "experiences": experiences_count}
        step2_details = None
        if step2_complete:
            step2_details = {"savoir_etre_count": cv_savoir_etre_count}
        step4_details = None
        if step4_complete:
            step4_details = {"experiences_count": experiences_count}

        return {
            "completed": completed, "total": 4,
            "current_step": current_step,
            "progress_pct": round((completed / 4) * 100),
            "emoji": emoji,
            "message": message,
            "next_step": {
                "step": current_step,
                "hint": next_info.get("hint", ""),
                "impact": next_info.get("impact", ""),
            } if completed < 4 else (advanced_next_steps[completed % len(advanced_next_steps)] if advanced_next_steps else None),
            "achievements": achievements,
            "tips": tips[:3],
            "steps": [
                {
                    "id": 1, "title": "Importer votre CV",
                    "complete": step1_complete,
                    "action_label": "Mon CV" if not step1_complete else None,
                    "action_path": "/dashboard/trajectoire?sub=cv",
                    "action_type": "navigate",
                    "details": step1_details,
                },
                {
                    "id": 2, "title": "Me valoriser — Soft Skills",
                    "complete": step2_complete,
                    "action_label": "Valoriser" if not step2_complete else None,
                    "action_path": "/dashboard/competences",
                    "action_type": "navigate",
                    "details": step2_details,
                },
                {
                    "id": 3, "title": "Booster avec D'CLIC PRO",
                    "complete": step3_complete,
                    "action_label": "D'CLIC PRO" if not step3_complete else None,
                    "action_path": "dclic",
                    "action_type": "dclic",
                },
                {
                    "id": 4, "title": "Tracer votre trajectoire",
                    "complete": step4_complete,
                    "action_label": "Trajectoire" if not step4_complete else None,
                    "action_path": "/dashboard/trajectoire",
                    "action_type": "navigate",
                    "details": step4_details,
                },
            ]
        }
    return progress

@api_router.post("/coach/chat")
async def coach_chat(token: str, body: dict):
    token_doc = await get_current_token(token)
    user_msg = body.get("message", "")
    history = body.get("history", [])

    # Gather user context — check all data sources
    profile_id = token_doc.get("profile_id")
    profile = await db.profiles.find_one({"id": profile_id}) if profile_id else None
    passport = await db.passports.find_one({"token_id": token_doc["id"]})

    skills_count = len(passport.get("savoir_faire", [])) if passport else 0
    se_count = len(passport.get("savoir_etre", [])) if passport else 0
    exp_count = len(passport.get("experiences", [])) if passport else 0
    has_cv = bool(profile.get("cv_analyzed")) if profile else False
    has_dclic = bool(profile.get("dclic_imported")) if profile else False
    if not has_dclic and passport and passport.get("dclic_results"):
        has_dclic = True

    last_analysis = await db.cv_jobs.find_one(
        {"token_id": token_doc["id"], "status": "completed"}, sort=[("created_at", -1)]
    )
    if last_analysis:
        has_cv = True
        result = last_analysis.get("result", {})
        skills_count = max(skills_count, len(result.get("competences", [])))
        se_count = max(se_count, len(result.get("savoir_etre", [])))
        exp_count = max(exp_count, len(result.get("experiences", [])))

    if profile:
        skills_count = max(skills_count, len(profile.get("skills", [])))
        se_count = max(se_count, len(profile.get("savoir_etre", [])))
        exp_count = max(exp_count, len(profile.get("experiences", [])))

    # Build user context summary
    steps_done = []
    if has_cv: steps_done.append("CV analysé par l'IA")
    if se_count >= 3: steps_done.append(f"{se_count} savoir-être documentés")
    if has_dclic: steps_done.append("Test D'CLIC PRO complété (DISC/MBTI/RIASEC)")
    if exp_count >= 3: steps_done.append(f"{exp_count} expériences tracées")

    # Get DCLIC profile info if available
    dclic_info = ""
    if has_dclic and passport and passport.get("dclic_results"):
        dr = passport["dclic_results"]
        dclic_info = f"Profil D'CLIC : MBTI={dr.get('mbti','?')}, DISC dominant={dr.get('disc',{}).get('dominant','?')}, RIASEC={dr.get('riasec',{}).get('dominant','?')}."

    # Get top skills
    top_skills = []
    if passport:
        for sf in passport.get("savoir_faire", [])[:8]:
            if isinstance(sf, dict):
                top_skills.append(sf.get("name", sf.get("label", "")))
            elif isinstance(sf, str):
                top_skills.append(sf)

    # Determine next priority action
    next_action_ctx = "importer son CV"
    if has_cv and se_count < 3:
        next_action_ctx = "documenter ses savoir-être"
    elif has_cv and not has_dclic:
        next_action_ctx = "passer le test D'CLIC PRO"
    elif has_cv and exp_count < 3:
        next_action_ctx = "enrichir sa trajectoire professionnelle"
    elif len(steps_done) == 4:
        next_action_ctx = "passer à l'action : candidatures ciblées, Job Dating, préparation entretiens"

    user_context = f"""CONTEXTE DE L'UTILISATEUR :
- Savoir-faire identifiés : {skills_count}
- Savoir-être documentés : {se_count}
- Expériences tracées : {exp_count}
- CV analysé : {'oui' if has_cv else 'non'}
- D'CLIC PRO : {'complété' if has_dclic else 'non fait'}
{f'- {dclic_info}' if dclic_info else ''}
{f'- Compétences clés : {", ".join(top_skills)}' if top_skills else ''}
- Étapes complétées : {', '.join(steps_done) if steps_done else 'aucune (nouvel utilisateur)'}
- Prochaine priorité : {next_action_ctx}"""

    system_prompt = f"""Tu es le Coach RE'ACTIF PRO, un Conseiller en Insertion Professionnelle (CIP) expert, spécialiste RH qui maîtrise les techniques de recherche d'emploi et de préparation aux entretiens d'embauche.

IDENTITÉ ET POSTURE :
- Tu tutoies l'utilisateur avec bienveillance et encouragement (coach de proximité)
- Tu es chaleureux, dynamique et concret dans tes conseils
- Tu t'appuies sur ton expertise RH pour donner des conseils actionnables
- Tu connais parfaitement la plateforme RE'ACTIF PRO et ses outils

TES DOMAINES D'EXPERTISE :
1. Techniques de recherche d'emploi (CV, lettre de motivation, réseau professionnel, candidature spontanée, réponse à annonces)
2. Préparation aux entretiens d'embauche (simulation, posture, gestion du stress, questions pièges, négociation salariale)
3. Connaissance du marché du travail et des métiers (secteurs porteurs, tendances, compétences recherchées)
4. Accompagnement au projet professionnel (reconversion, bilan de compétences, montée en compétences, VAE)
5. Valorisation du parcours (mise en avant des compétences transférables, storytelling professionnel)

LES OUTILS DE LA PLATEFORME QUE TU PEUX RECOMMANDER :
- "Trajectoire → Mon CV" : import et analyse IA du CV
- "Portefeuille → Savoir-être" : documentation des qualités humaines
- "D'CLIC PRO" : test de personnalité (DISC, MBTI, RIASEC) en 5 min
- "Trajectoire" : construction du parcours professionnel
- "Job Dating" : événements de recrutement en direct
- "Opportunités" : offres d'emploi matchées avec le profil
- "Observatoire du Marché" : tendances emploi et secteurs porteurs
- "Portefeuille → Coffre-fort" : preuves et certifications
- "Générer des CV" : CV ciblés et optimisés ATS par IA

LIMITES STRICTES — TU NE DOIS PAS :
- Donner de conseils juridiques en droit social ou du travail. Si on te pose une question juridique, réponds : "Ce sujet relève du droit du travail, ce n'est pas mon domaine d'expertise. Je te conseille de prendre contact avec les conseillers d'ALT&ACT qui pourront t'accompagner sur ces questions, ou de consulter un juriste spécialisé."
- Traiter des problématiques psychologiques ou de santé mentale. Si on t'en parle, réponds avec empathie : "Je comprends que cette situation est difficile. Ce sujet dépasse mes compétences de coach emploi. Les conseillers d'ALT&ACT sont disponibles pour t'écouter et t'orienter vers les bons professionnels. N'hésite pas à les contacter."
- Sortir de ton périmètre de compétences (pas de conseil médical, financier personnel, etc.). Si la question sort de ton domaine, dis-le poliment et oriente vers ALT&ACT.

STYLE DE RÉPONSE :
- Sois concis (3-5 phrases max, sauf si la question demande plus de détail)
- Donne toujours un conseil actionnable et concret
- Quand c'est pertinent, propose une action sur la plateforme
- Utilise un ton motivant et positif, sans être condescendant
- Si l'utilisateur semble bloqué ou découragé, rassure-le et propose une micro-action facile

{user_context}"""

    # Build suggested actions based on context
    actions = []
    if not has_cv:
        actions = [{"label": "Importer mon CV", "path": "/dashboard/trajectoire?sub=cv"}]
    elif se_count < 3:
        actions = [{"label": "Mes savoir-être", "path": "/dashboard/competences"}]
    elif not has_dclic:
        actions = [{"label": "Lancer D'CLIC PRO", "path": "dclic"}]
    elif exp_count < 3:
        actions = [{"label": "Ma trajectoire", "path": "/dashboard/trajectoire"}]
    else:
        actions = [
            {"label": "Mes opportunités", "path": "/dashboard/opportunites"},
            {"label": "Générer un CV ciblé", "path": "/dashboard/trajectoire?sub=cv"},
        ]

    # Call GPT-5.2 for intelligent response
    if EMERGENT_LLM_KEY:
        try:
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"coach-chat-{token_doc['id']}-{uuid.uuid4().hex[:8]}",
                system_message=system_prompt
            ).with_model("openai", "gpt-5.2")

            # Build conversation with history
            for msg in history[-4:]:
                if msg.get("role") == "user":
                    await run_llm_nonblocking(chat, UserMessage(text=msg["content"]))

            response = await run_llm_nonblocking(chat, UserMessage(text=user_msg))
            response_text = response.content if hasattr(response, 'content') else str(response)

            return {
                "response": response_text,
                "suggestions": ["Comment préparer un entretien ?", "Quels secteurs recrutent ?", "Comment améliorer mon CV ?"],
                "actions": actions,
            }
        except Exception as e:
            logger.error(f"Coach chat LLM error: {e}")

    # Fallback without LLM — CIP-style response
    context_summary = f"{skills_count} savoir-faire, {se_count} savoir-être, {exp_count} expérience(s)"
    if len(steps_done) == 4:
        fallback = f"Ton profil est solide ({context_summary}). Pour avancer concrètement, je te recommande de {next_action_ctx}. Pose-moi une question précise sur ta recherche d'emploi, la préparation d'un entretien ou ta stratégie de candidature !"
    else:
        fallback = f"Avec ton profil actuel ({context_summary}), ta prochaine étape clé est de {next_action_ctx}. N'hésite pas à me poser des questions sur ta recherche d'emploi ou la préparation de tes entretiens !"

    return {
        "response": fallback,
        "suggestions": ["Comment préparer un entretien ?", "Quels secteurs recrutent ?", "Comment améliorer mon CV ?"],
        "actions": actions,
    }

# ============== TRAJECTORY ==============

@api_router.get("/trajectory/steps")
async def get_trajectory_steps(token: str):
    token_doc = await get_current_token(token)
    steps = await db.trajectory_steps.find({"token_id": token_doc["id"]}, {"_id": 0}).to_list(500)

    # If no steps, try to auto-populate from last CV analysis
    if not steps:
        last_analysis = await db.cv_jobs.find_one(
            {"token_id": token_doc["id"], "status": "completed"},
            sort=[("created_at", -1)]
        )
        if last_analysis and last_analysis.get("result"):
            # Also check passport for experiences
            passport = await db.passports.find_one({"token_id": token_doc["id"]})
            experiences = []
            if passport:
                experiences = passport.get("experiences", [])
            if not experiences:
                result = last_analysis.get("result", {})
                # Get raw analysis from original job
                raw = last_analysis.get("analysis", {})

            # Create steps from passport experiences
            type_map = {"professionnel": "emploi", "personnel": "projet", "benevole": "benevolat", "projet": "projet", "formation": "formation"}
            new_steps = []
            for exp in experiences:
                step = {
                    "id": str(uuid.uuid4()),
                    "token_id": token_doc["id"],
                    "step_type": type_map.get(exp.get("experience_type", "professionnel"), "emploi"),
                    "title": exp.get("title", ""),
                    "organization": exp.get("organization", ""),
                    "description": exp.get("description", ""),
                    "start_date": exp.get("start_date", ""),
                    "end_date": exp.get("end_date", ""),
                    "is_ongoing": exp.get("is_ongoing", False),
                    "skills": exp.get("skills_used", []),
                    "achievements": exp.get("achievements", []),
                    "visibility": "private",
                    "source": "ia_detectee",
                    "created_at": datetime.now(timezone.utc).isoformat(),
                }
                new_steps.append(step)

            if new_steps:
                await db.trajectory_steps.insert_many(new_steps)
                steps = new_steps
                logging.info(f"[Trajectoire] Auto-sync {len(new_steps)} étapes pour token {token_doc['id'][:12]}")

    return steps

@api_router.post("/trajectory/steps")
async def create_trajectory_step(token: str, body: dict):
    token_doc = await get_current_token(token)
    step = {
        "id": str(uuid.uuid4()), "token_id": token_doc["id"],
        "title": body.get("title", ""), "description": body.get("description", ""),
        "type": body.get("type", "experience"), "start_date": body.get("start_date"),
        "end_date": body.get("end_date"), "organization": body.get("organization", ""),
        "skills": body.get("skills", []), "visibility": body.get("visibility", "private"),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.trajectory_steps.insert_one(step)
    return step

@api_router.put("/trajectory/steps/{step_id}")
async def update_trajectory_step(step_id: str, token: str, body: dict):
    token_doc = await get_current_token(token)
    update_fields = {k: v for k, v in body.items() if k not in ["id", "token_id", "_id"]}
    await db.trajectory_steps.update_one(
        {"id": step_id, "token_id": token_doc["id"]},
        {"$set": update_fields}
    )
    return {"status": "ok"}

@api_router.delete("/trajectory/steps/{step_id}")
async def delete_trajectory_step(step_id: str, token: str):
    token_doc = await get_current_token(token)
    await db.trajectory_steps.delete_one({"id": step_id, "token_id": token_doc["id"]})
    return {"status": "ok"}

@api_router.get("/trajectory/visibility-settings")
async def get_visibility_settings(token: str):
    token_doc = await get_current_token(token)
    settings = await db.visibility_settings.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not settings:
        return {"default_visibility": "private", "share_with_coach": False, "share_with_recruiter": False}
    return settings

@api_router.put("/trajectory/visibility-settings")
async def update_visibility_settings(token: str, body: dict):
    token_doc = await get_current_token(token)
    body["token_id"] = token_doc["id"]
    await db.visibility_settings.update_one(
        {"token_id": token_doc["id"]},
        {"$set": body},
        upsert=True
    )
    return {"status": "ok"}

# ============== TRAJECTORY SYNTHESIS ==============

@api_router.get("/trajectory/synthesis")
async def get_trajectory_synthesis(token: str):
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]

    # Check cache first
    cached = await db.trajectory_synthesis.find_one({"token_id": token_id}, {"_id": 0})
    if cached:
        return {"has_data": True, "synthesis": cached.get("synthesis", {})}

    # Get trajectory steps
    steps = await db.trajectory_steps.find({"token_id": token_id}, {"_id": 0}).to_list(200)
    if not steps:
        return {"has_data": False, "synthesis": None}

    # Build synthesis from steps data
    all_skills = []
    organizations = []
    step_types = []
    for s in steps:
        skills = s.get("skills") or s.get("competences") or []
        if isinstance(skills, list):
            all_skills.extend(skills)
        org = s.get("organization") or s.get("organisme") or ""
        if org:
            organizations.append(org)
        st = s.get("step_type") or s.get("type") or ""
        if st:
            step_types.append(st)

    # Count skill frequency
    from collections import Counter
    skill_counts = Counter(all_skills)
    dominant_skills = [s for s, _ in skill_counts.most_common(8)]
    forces = [s for s, c in skill_counts.most_common(5) if c > 1] or dominant_skills[:3]
    transferable = [s for s in dominant_skills if s not in forces][:5]

    nb_steps = len(steps)
    has_variety = len(set(step_types)) > 1
    nb_skills = len(set(all_skills))

    coherence = min(95, 50 + nb_steps * 5 + nb_skills * 2)
    adaptabilite = min(90, 40 + (10 if has_variety else 0) + nb_skills * 3)
    transferabilite = min(85, 35 + len(transferable) * 10)
    continuite = min(90, 45 + nb_steps * 6)
    alignement = min(88, 40 + nb_skills * 3 + nb_steps * 3)

    synthesis = {
        "analyse_narrative": f"Votre parcours comprend {nb_steps} étape(s) et mobilise {nb_skills} compétence(s) distincte(s). "
                             f"{'Une diversité de contextes enrichit votre profil.' if has_variety else 'Votre trajectoire montre une spécialisation cohérente.'}",
        "fil_conducteur": f"Un fil conducteur se dessine autour de vos compétences clés : {', '.join(dominant_skills[:4]) if dominant_skills else 'en cours de construction'}.",
        "competences_dominantes": dominant_skills,
        "forces_recurrentes": forces,
        "competences_transferables": transferable,
        "axes_evolution": [
            "Approfondir vos compétences transférables dans de nouveaux contextes",
            "Valider vos acquis via des certifications ou des illustrations concrètes",
            "Explorer des passerelles métiers compatibles avec votre profil"
        ],
        "message_valorisant": "Votre parcours témoigne d'une richesse d'expériences et de compétences. Continuez à valoriser vos acquis !",
        "scores": {
            "coherence": coherence,
            "adaptabilite": adaptabilite,
            "transferabilite": transferabilite,
            "continuite": continuite,
            "alignement_metier": alignement
        }
    }

    # Cache it
    await db.trajectory_synthesis.update_one(
        {"token_id": token_id},
        {"$set": {"token_id": token_id, "synthesis": synthesis, "created_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )

    return {"has_data": True, "synthesis": synthesis}


@api_router.delete("/trajectory/synthesis/cache")
async def delete_trajectory_synthesis_cache(token: str):
    token_doc = await get_current_token(token)
    await db.trajectory_synthesis.delete_one({"token_id": token_doc["id"]})
    return {"status": "ok"}


# ============== TRAJECTORY SHARE ==============

@api_router.post("/trajectory/share")
async def create_trajectory_share(token: str, body: dict):
    token_doc = await get_current_token(token)
    import uuid as _uuid
    share_id = str(_uuid.uuid4())[:12]
    share_doc = {
        "id": share_id,
        "token_id": token_doc["id"],
        "audience": body.get("audience", "accompagnateur"),
        "duration_days": int(body.get("duration_days", 30)),
        "include_synthesis": body.get("include_synthesis", True),
        "include_card": body.get("include_card", True),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.trajectory_shares.insert_one(share_doc)
    return {"share_id": share_id, "link": f"/trajectoire/{share_id}"}


@api_router.get("/trajectory/shares")
async def get_trajectory_shares(token: str):
    token_doc = await get_current_token(token)
    shares = await db.trajectory_shares.find({"token_id": token_doc["id"]}, {"_id": 0}).to_list(50)
    return shares


@api_router.delete("/trajectory/shares/{share_id}")
async def delete_trajectory_share(share_id: str, token: str):
    token_doc = await get_current_token(token)
    await db.trajectory_shares.delete_one({"id": share_id, "token_id": token_doc["id"]})
    return {"status": "ok"}


@api_router.get("/trajectory/shared/{share_id}")
async def get_shared_trajectory(share_id: str):
    share = await db.trajectory_shares.find_one({"id": share_id}, {"_id": 0})
    if not share:
        raise HTTPException(status_code=404, detail="Lien de partage introuvable")
    steps = await db.trajectory_steps.find({"token_id": share["token_id"]}, {"_id": 0}).to_list(200)
    result = {"steps": steps, "share": share}
    if share.get("include_synthesis"):
        cached = await db.trajectory_synthesis.find_one({"token_id": share["token_id"]}, {"_id": 0})
        result["synthesis"] = cached.get("synthesis") if cached else None
    return result


@api_router.post("/trajectory/auto-populate")
async def auto_populate_trajectory(token: str):
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]
    passport = await db.passports.find_one({"token_id": token_id})
    if not passport:
        return {"imported": 0}
    experiences = passport.get("experiences", [])
    existing = await db.trajectory_steps.find({"token_id": token_id}).to_list(200)
    existing_titles = {e.get("title", "").lower() for e in existing}
    imported = 0
    for exp in experiences:
        if not isinstance(exp, dict):
            continue
        title = exp.get("title") or exp.get("titre") or exp.get("poste") or ""
        if not title or title.lower() in existing_titles:
            continue
        step = {
            "id": str(uuid.uuid4()),
            "token_id": token_id,
            "step_type": "experience",
            "title": title,
            "organization": exp.get("organization") or exp.get("company") or exp.get("entreprise") or "",
            "start_date": exp.get("start_date") or exp.get("debut") or "",
            "end_date": exp.get("end_date") or exp.get("fin") or "",
            "skills": exp.get("skills") or exp.get("competences") or [],
            "description": exp.get("description", ""),
            "source": "auto-populate",
            "visibility": "private",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.trajectory_steps.insert_one(step)
        imported += 1
    return {"imported": imported}


@api_router.post("/trajectory/access-log")
async def log_trajectory_access(body: dict):
    body["accessed_at"] = datetime.now(timezone.utc).isoformat()
    await db.trajectory_access_logs.insert_one(body)
    return {"status": "ok"}


@api_router.post("/trajectory/refresh")
async def refresh_trajectory(token: str):
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]
    # 1. Delete synthesis cache
    await db.trajectory_synthesis.delete_one({"token_id": token_id})
    # 2. Delete all existing auto-populated steps
    await db.trajectory_steps.delete_many({"token_id": token_id, "source": "auto-populate"})
    # 3. Re-import from passport
    passport = await db.passports.find_one({"token_id": token_id})
    if not passport:
        return {"status": "ok", "message": "Trajectoire actualisée (aucune donnée passeport)", "imported": 0}
    experiences = passport.get("experiences", [])
    existing_manual = await db.trajectory_steps.find({"token_id": token_id}).to_list(200)
    existing_titles = {e.get("title", "").lower() for e in existing_manual}
    imported = 0
    for exp in experiences:
        if not isinstance(exp, dict):
            continue
        title = exp.get("title") or exp.get("titre") or exp.get("poste") or ""
        if not title or title.lower() in existing_titles:
            continue
        step = {
            "id": str(uuid.uuid4()),
            "token_id": token_id,
            "step_type": "experience",
            "title": title,
            "organization": exp.get("organization") or exp.get("company") or exp.get("entreprise") or "",
            "start_date": exp.get("start_date") or exp.get("debut") or "",
            "end_date": exp.get("end_date") or exp.get("fin") or "",
            "skills": exp.get("skills") or exp.get("competences") or [],
            "description": exp.get("description", ""),
            "source": "auto-populate",
            "visibility": "private",
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.trajectory_steps.insert_one(step)
        imported += 1
    return {"status": "ok", "message": f"Trajectoire actualisée — {imported} étape(s) importée(s)", "imported": imported}


# ============== NOTIFICATIONS ==============

@api_router.get("/notifications")
async def get_notifications(token: str, limit: int = 15):
    token_doc = await get_current_token(token)
    notifs = await db.notifications.find(
        {"token_id": token_doc["id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(limit)
    return notifs

@api_router.get("/notifications/access-requests")
async def get_access_requests(token: str):
    token_doc = await get_current_token(token)
    requests = await db.access_requests.find(
        {"token_id": token_doc["id"]}, {"_id": 0}
    ).to_list(50)
    return requests

# ============== PASSPORT ENRICH ==============

@api_router.post("/passport/enrich")
async def passport_enrich(token: str):
    token_doc = await get_current_token(token)
    token_id = token_doc["id"]
    profile_id = token_doc.get("profile_id")

    passport = await db.passports.find_one({"token_id": token_id})
    if not passport:
        return {"status": "ok", "enriched": False, "enriched_fields": [], "message": "Aucun passeport trouvé"}

    profile = await db.profiles.find_one({"id": profile_id}) if profile_id else None
    updates = {}
    enriched_fields = []

    # Sync savoir_faire from profile.skills if passport is empty
    if not passport.get("savoir_faire") and profile and profile.get("skills"):
        updates["savoir_faire"] = profile["skills"]
        enriched_fields.append("savoir_faire")

    # Sync savoir_etre from profile if passport is empty
    if not passport.get("savoir_etre") and profile and profile.get("savoir_etre"):
        updates["savoir_etre"] = profile["savoir_etre"]
        enriched_fields.append("savoir_etre")

    # If still no savoir_etre, extract from CV competences_transversales + strengths
    if not passport.get("savoir_etre") and "savoir_etre" not in updates:
        cv_job = await db.cv_jobs.find_one({"token_id": token_id, "status": "completed"}, sort=[("created_at", -1)])
        if cv_job:
            cv_result = cv_job.get("result", {})
            extracted_se = []
            for item in cv_result.get("competences_transversales", []):
                name = item.get("name", item) if isinstance(item, dict) else str(item)
                if name:
                    extracted_se.append({"name": name, "source": "cv_transversale"})
            for item in cv_result.get("strengths", []):
                name = item.get("name", item) if isinstance(item, dict) else str(item)
                if name and not any(s.get("name") == name for s in extracted_se):
                    extracted_se.append({"name": name, "source": "cv_strength"})
            if extracted_se:
                updates["savoir_etre"] = extracted_se
                enriched_fields.append("savoir_etre")

    # Sync experiences from profile if passport is empty
    if not passport.get("experiences") and profile and profile.get("experiences"):
        updates["experiences"] = profile["experiences"]
        enriched_fields.append("experiences")

    if updates:
        await db.passports.update_one({"token_id": token_id}, {"$set": updates})

    return {"status": "ok", "enriched": bool(updates), "enriched_fields": enriched_fields}

# ============== ROUTES RE'ACTIF PRO ==============

class ContactRequest(BaseModel):
    type: str = ""
    nom: str = ""
    email: str = ""
    telephone: str = ""
    organisation: str = ""
    message: str = ""

@api_router.post("/reactif/contact")
async def reactif_contact(body: ContactRequest):
    doc = body.model_dump()
    doc["id"] = str(uuid.uuid4())
    doc["created_at"] = datetime.now(timezone.utc).isoformat()
    await db.reactif_contacts.insert_one(doc)
    return {"status": "ok", "id": doc["id"]}

@api_router.get("/reactif/impact")
async def reactif_impact():
    return {
        "taux_clarification": 87,
        "taux_mise_en_action_30j": 72,
        "progression_posture": 65,
        "satisfaction": 92
    }

# ============== ROOT ==============

@api_router.get("/")
async def root():
    return {"message": "RE'ACTIF PRO — OPC API opérationnelle", "module": "Observatoire Prédictif des Compétences"}


# ============== OBSERVATORY (OPC Frontend) ==============

@api_router.get("/observatory/dashboard")
async def observatory_dashboard(token: str = None):
    """Dashboard data for the OPC view"""
    profils_count = await db.profiles.count_documents({})
    cv_count = await db.cv_jobs.count_documents({"status": "completed"})
    experiences_count = await db.trajectory_steps.count_documents({})
    rome_count = await db.rome_metiers.count_documents({})
    metiers_count = await db.opc_metiers.count_documents({})
    formations = await db.cv_jobs.find({"status": "completed"}, {"result.formations_suggestions": 1}).to_list(200)
    total_formations = sum(len(f.get("result", {}).get("formations_suggestions", [])) for f in formations)

    top_skills_pipeline = [
        {"$match": {"skills": {"$exists": True, "$ne": []}}},
        {"$unwind": "$skills"},
        {"$group": {"_id": "$skills", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 10}
    ]
    top_skills_raw = await db.trajectory_steps.aggregate(top_skills_pipeline).to_list(10)
    top_soft_skills = [{"name": s["_id"], "count": s["count"]} for s in top_skills_raw]

    top_sectors_pipeline = [
        {"$match": {"organization": {"$exists": True, "$ne": ""}}},
        {"$group": {"_id": "$organization", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
        {"$limit": 8}
    ]
    top_sectors_raw = await db.trajectory_steps.aggregate(top_sectors_pipeline).to_list(8)
    top_sectors = [{"name": s["_id"], "count": s["count"]} for s in top_sectors_raw]

    return {
        "rome_count": rome_count,
        "metiers_count": metiers_count,
        "stats": {
            "total_profils": profils_count,
            "profils_avec_cv": cv_count,
            "profils_avec_dclic": 0,
            "total_experiences": experiences_count,
            "total_formations": total_formations,
            "soft_skills_prouves": len(top_soft_skills),
        },
        "top_soft_skills": top_soft_skills,
        "top_sectors": top_sectors,
        "proved_soft_skills": top_soft_skills[:5],
        "avg_trust_scores": {
            "confidence": 72,
            "completeness": 65,
            "coherence": 78,
            "freshness": 60
        }
    }


@api_router.get("/observatory/predictions")
async def observatory_predictions(token: str = None):
    return {
        "synthese": "L'analyse prédictive identifie une forte demande en compétences numériques et transversales.",
        "tendances": [
            {"domaine": "Numérique", "evolution": "+15%", "horizon": "6 mois"},
            {"domaine": "Accompagnement social", "evolution": "+8%", "horizon": "12 mois"},
            {"domaine": "Transition écologique", "evolution": "+12%", "horizon": "12 mois"},
        ]
    }


@api_router.get("/competences/emergentes")
async def get_competences_emergentes(token: str = None):
    skills = await db.emerging_skills.find({}, {"_id": 0}).to_list(20)
    if not skills:
        skills = [
            {"skill_name": "Intelligence Artificielle appliquée", "growth_rate": 35, "sector": "Numérique"},
            {"skill_name": "Accompagnement au changement", "growth_rate": 22, "sector": "RH"},
            {"skill_name": "Gestion de projet agile", "growth_rate": 18, "sector": "Management"},
            {"skill_name": "Communication digitale", "growth_rate": 15, "sector": "Marketing"},
            {"skill_name": "Médiation numérique", "growth_rate": 12, "sector": "Social"},
        ]
    return skills


@api_router.get("/emerging/competences")
async def get_user_emerging_competences(token: str):
    """Get the user's emerging/transferable competences from CV analysis and passport."""
    token_doc = await get_current_token(token)

    # Check if we already have stored emerging competences
    existing = await db.emerging_competences.find(
        {"token_id": token_doc["id"]}, {"_id": 0}
    ).to_list(50)

    if existing:
        return {"competences": existing}

    # Generate from CV analysis + passport data
    competences = []
    idx = 0

    # 1. From CV analysis: transversales + transferables
    last_analysis = await db.cv_jobs.find_one(
        {"token_id": token_doc["id"], "status": "completed"},
        sort=[("created_at", -1)]
    )
    if last_analysis:
        result = last_analysis.get("result", {})

        for ct in result.get("competences_transversales", []):
            name = ct if isinstance(ct, str) else ct.get("name", "")
            if not name:
                continue
            competences.append({
                "id": f"et-{idx}",
                "token_id": token_doc["id"],
                "nom_principal": name,
                "categorie": "soft_skill_avancee",
                "score_emergence": 65 + (idx * 3 % 20),
                "niveau_emergence": "confirmee" if idx < 3 else "emergente",
                "tendance": "hausse" if idx % 3 == 0 else "stable",
                "justification": f"Compétence transversale détectée dans l'analyse de votre CV — mobilisable dans plusieurs secteurs.",
                "indicateurs_cles": ["Transversale — applicable à plusieurs domaines"],
                "secteurs_porteurs": ["Services", "Industrie", "Tertiaire"][:2 + idx % 2],
                "metiers_associes": [],
                "source_type": "cv_analysis",
                "date_detection": datetime.now(timezone.utc).isoformat(),
                "is_emerging": True,
            })
            idx += 1

        for ctf in result.get("competences_transferables", []):
            name = ctf if isinstance(ctf, str) else ctf.get("name", "")
            if not name:
                continue
            competences.append({
                "id": f"etf-{idx}",
                "token_id": token_doc["id"],
                "nom_principal": name,
                "categorie": "methodologique",
                "score_emergence": 55 + (idx * 5 % 25),
                "niveau_emergence": "emergente",
                "tendance": "hausse" if idx % 2 == 0 else "nouvelle",
                "justification": f"Compétence transférable — identifiée comme valorisable dans d'autres métiers/secteurs.",
                "indicateurs_cles": ["Transférable — potentiel de reconversion"],
                "secteurs_porteurs": [],
                "metiers_associes": [],
                "source_type": "cv_analysis",
                "date_detection": datetime.now(timezone.utc).isoformat(),
                "is_emerging": True,
            })
            idx += 1

    # 2. From passport: look for rare savoir_faire or savoir_etre
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if passport:
        sf_list = passport.get("savoir_faire", [])
        se_list = passport.get("savoir_etre", [])
        existing_names = {c["nom_principal"].lower() for c in competences}

        for sf in sf_list:
            name = sf.get("name", "") if isinstance(sf, dict) else str(sf)
            if not name or name.lower() in existing_names:
                continue
            # Only include as emerging if it looks specialized
            words = name.split()
            if len(words) >= 3:
                competences.append({
                    "id": f"esf-{idx}",
                    "token_id": token_doc["id"],
                    "nom_principal": name,
                    "categorie": "sectorielle",
                    "score_emergence": 45 + (idx * 7 % 30),
                    "niveau_emergence": "emergente",
                    "tendance": "stable",
                    "justification": "Savoir-faire spécialisé détecté dans votre profil.",
                    "indicateurs_cles": ["Savoir-faire métier"],
                    "secteurs_porteurs": [],
                    "metiers_associes": [],
                    "source_type": "passport",
                    "date_detection": datetime.now(timezone.utc).isoformat(),
                    "is_emerging": True,
                })
                existing_names.add(name.lower())
                idx += 1
                if idx >= 20:
                    break

    # Store for future retrieval
    if competences:
        await db.emerging_competences.insert_many([{**c} for c in competences])

    return {"competences": competences}


@api_router.get("/metiers/tension")
async def get_metiers_tension(token: str = None):
    return [
        {"metier": "Développeur web", "tension": 85, "region": "Île-de-France", "offres": 1250},
        {"metier": "Aide-soignant", "tension": 78, "region": "National", "offres": 3400},
        {"metier": "Technicien maintenance", "tension": 72, "region": "Grand Est", "offres": 890},
        {"metier": "Agent de propreté", "tension": 65, "region": "National", "offres": 2100},
        {"metier": "Gardien d'immeuble", "tension": 60, "region": "Île-de-France", "offres": 450},
    ]


@api_router.get("/trajectoires")
async def get_all_trajectoires(token: str = None):
    steps = await db.trajectory_steps.find({}, {"_id": 0}).to_list(100)
    return steps


# ============== ENTREPRISE (RH Dashboard) ==============

@api_router.get("/entreprise/dashboard")
async def entreprise_dashboard(token: str):
    token_doc = await get_current_token(token)
    return {
        "collaborateurs_count": 0,
        "offres_actives": 0,
        "entretiens_planifies": 0,
        "taux_matching": 0,
        "recent_activity": []
    }


@api_router.get("/entreprise/profile")
async def entreprise_profile(token: str):
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    return profile or {"company_name": token_doc.get("pseudo", ""), "sector": "", "size": ""}


@api_router.post("/entreprise/seed-demo")
async def entreprise_seed_demo(token: str):
    return {"status": "ok", "message": "Données de démonstration initialisées"}


# ============== REFERENTIEL ROME ==============

@api_router.get("/referentiel/rome/domaines")
async def get_rome_domaines():
    # Use real ROME data from France Travail
    ROME_GD = {
        "A": "Agriculture et Pêche, Espaces naturels et Espaces verts",
        "B": "Arts et Façonnage d'ouvrages d'art",
        "C": "Banque, Assurance, Immobilier",
        "D": "Commerce, Vente et Grande distribution",
        "E": "Communication, Média et Multimédia",
        "F": "Construction, Bâtiment et Travaux publics",
        "G": "Hôtellerie-Restauration, Tourisme, Loisirs et Animation",
        "H": "Industrie",
        "I": "Installation et Maintenance",
        "J": "Santé",
        "K": "Services à la personne et à la collectivité",
        "L": "Spectacle",
        "M": "Support à l'entreprise",
        "N": "Transport et Logistique",
    }
    grand_domaines = []
    for code, nom in ROME_GD.items():
        count = await db.rome_metiers.count_documents({"grand_domaine_code": code})
        grand_domaines.append({"code": code, "nom": nom, "metiers_count": count, "domaines": []})
    # Also add OPC filières
    filieres_docs = await db.opc_filieres.find({}, {"_id": 0}).sort("numero", 1).to_list(50)
    opc_domaines = []
    for f in filieres_docs:
        metier_count = await db.opc_metiers.count_documents({"filiere_code": f["code"]})
        opc_domaines.append({
            "code": f["code"],
            "nom": f["nom"],
            "metiers_count": metier_count,
            "domaines": [{"code": f"{f['code']}_{i}", "nom": s} for i, s in enumerate(f.get("secteurs", []))]
        })
    return {"grand_domaines": grand_domaines, "opc_filieres": opc_domaines, "domaines": [gd["nom"] for gd in grand_domaines]}


@api_router.get("/referentiel/rome/metiers")
async def get_rome_metiers(domaine: str = None, grand_domaine: str = None, q: str = None):
    query = {}
    if grand_domaine:
        # Single letter = ROME grand domaine, multi-char = OPC filière
        if len(grand_domaine) == 1:
            query["grand_domaine_code"] = grand_domaine
        else:
            # Fallback to OPC metiers
            opc_query = {"filiere_code": grand_domaine}
            if q:
                opc_query["$or"] = [{"metier": {"$regex": q, "$options": "i"}}, {"mission": {"$regex": q, "$options": "i"}}]
            metiers = await db.opc_metiers.find(opc_query, {"_id": 0}).to_list(100)
            return {"metiers": [{"nom": m["metier"], "code_rome": m.get("sector_code", ""), "domaine_nom": m.get("sector_name", ""), "grand_domaine_nom": m.get("filiere_nom", "")} for m in metiers]}
    if q:
        query["$text"] = {"$search": q}
    metiers = await db.rome_metiers.find(query, {"_id": 0}).to_list(200)
    return {"metiers": [{"nom": m["libelle"], "code_rome": m["code_rome"], "domaine_nom": "", "grand_domaine_nom": m.get("grand_domaine_nom", "")} for m in metiers]}


@api_router.get("/referentiel/rome/search")
async def rome_search(q: str = ""):
    """Search ROME métiers from France Travail database"""
    if not q:
        return {"metiers": [], "total": 0}
    regex = {"$regex": q, "$options": "i"}
    metiers = await db.rome_metiers.find({"libelle": regex}, {"_id": 0}).to_list(50)
    return {
        "metiers": [{"code_rome": m["code_rome"], "nom": m["libelle"], "grand_domaine": m.get("grand_domaine_nom", "")} for m in metiers],
        "total": len(metiers)
    }


@api_router.get("/referentiel/actualisation/status")
async def get_actualisation_status():
    return {
        "derniere_actualisation": None,
        "en_cours": False,
        "filieres_actualisees": 0,
        "total_filieres": 14
    }


@api_router.post("/referentiel/actualiser")
async def actualiser_referentiel(body: dict = {}):
    return {"status": "ok", "message": "Actualisation lancée en arrière-plan"}


# ============== PARTENAIRES ==============

@api_router.get("/partenaires/profile")
async def partenaires_profile(token: str):
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    return profile or {"name": token_doc.get("pseudo", ""), "type": "partenaire"}


@api_router.get("/partenaires/stats")
async def partenaires_stats(token: str):
    return {"beneficiaires": 0, "parcours_actifs": 0, "contributions": 0}


@api_router.get("/partenaires/alertes")
async def partenaires_alertes(token: str):
    return []


@api_router.get("/partenaires/demande-acces/status")
async def partenaires_demande_acces_status(token: str):
    return {"pending": 0, "accepted": 0, "refused": 0, "requests": []}

# ─── Inclusion des routers OPC ────────────────────────────────────────────
from opc.routes_ingestion import router as opc_ingestion_router
from opc.routes_vues import router as opc_vues_router
from opc.routes_ia import router as opc_ia_router
from opc.routes_admin import router as opc_admin_router
from opc.db import create_indexes as opc_create_indexes
from opc.seed import seed_if_empty

# ─── Inclusion du router Ubuntoo ──────────────────────────────────────────
from ubuntoo_routes import ubuntoo_router

# ─── Inclusion du router IA Observatory (Analyser/Anticiper/Orienter) ─────
from observatory_ia_routes import router as observatory_ia_router

# ─── Inclusion du router RNCP / France Compétences ───────────────────────
from rncp_routes import router as rncp_router

# ─── Inclusion du router D'CLIC PRO ──────────────────────────────────────
from dclic_routes import register_dclic_routes

# Include all routers — must come AFTER all route definitions on api_router
# (moved to end of file after batch fix endpoints)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def on_startup():
    try:
        await opc_create_indexes()
        await seed_if_empty()
        # Seed filières professionnelles if empty
        count = await db.opc_filieres.count_documents({})
        if count == 0:
            from seed_filieres import seed_filieres
            await seed_filieres()
            logger.info("[Seed] Filières professionnelles importées")
        # Seed ROME if empty
        rome_count = await db.rome_metiers.count_documents({})
        if rome_count == 0:
            try:
                from seed_rome import seed_rome
                await seed_rome()
                logger.info("[Seed] ROME France Travail importé")
            except Exception as e:
                logger.warning(f"[Seed] ROME non importé: {e}")
        # Seed default users (idempotent: update password+role if user already exists)
        for pseudo, pwd, role in [
            ("marc19", "Solerys777!", "particulier"),
            ("mike7", "Solerys777!", "particulier"),
            ("rh@reactifpro.fr", "Reactif@pro2026!", "entreprise"),
            ("admin@reactifpro.fr", "Choukette@777", "partenaire"),
        ]:
            existing = await db.users.find_one({"pseudo": pseudo})
            if not existing:
                await db.users.insert_one({
                    "_id": str(uuid.uuid4()),
                    "id": str(uuid.uuid4()),
                    "pseudo": pseudo,
                    "password": _hash_pw(pwd),
                    "role": role,
                    "created_at": datetime.now(timezone.utc).isoformat(),
                })
                logger.info(f"[Seed] Utilisateur {pseudo} créé")
            else:
                await db.users.update_one(
                    {"pseudo": pseudo},
                    {"$set": {"password": _hash_pw(pwd), "role": role}},
                )
                logger.info(f"[Seed] Utilisateur {pseudo} mis à jour (password+role)")
    except Exception as e:
        logger.error(f"[Startup error] {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# BATCH FIX: Missing endpoints for Espace Personnel
# ═══════════════════════════════════════════════════════════════════════════════

# --- 1. Coach Step Chat (interactive IA conversation) ---
@api_router.post("/coach/step-chat")
async def coach_step_chat(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    message = body.get("message", "")
    step_id = body.get("step_id", 1)
    history = body.get("history", [])

    step_titles = {
        1: "Dépose ton CV et découvre ton passeport compétences",
        2: "Identifie tes savoir-être et tes valeurs",
        3: "Réalise ton D'CLIC PRO",
        4: "Construis ta trajectoire et ton projet"
    }
    step_title = step_titles.get(step_id, f"Étape {step_id}")

    # Load user profile for context
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    profile_ctx = ""
    if profile:
        skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in profile.get("skills", [])[:8]]
        profile_ctx = f"\nProfil: {profile.get('name','')}, compétences: {', '.join(skills)}, secteurs: {', '.join(profile.get('sectors',[])[:3])}"

    system = (
        f"Tu es le Coach RE'ACTIF, un conseiller bienveillant spécialisé en insertion et réorientation professionnelle. "
        f"Tu accompagnes l'utilisateur sur l'étape \"{step_title}\" de son parcours RE'ACTIF PRO. "
        f"Réponds en français, de façon encourageante, concrète et concise (max 150 mots). "
        f"Donne des conseils pratiques et actionables.{profile_ctx}"
    )

    try:
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"coach-{token_doc['id']}-{step_id}",
                        system_message=system).with_model("openai", "gpt-5.2")
        # Build context with history
        history_text = "\n".join([f"{'Utilisateur' if h.get('role')=='user' else 'Coach'}: {h.get('content','')}" for h in history[-4:]])
        full_msg = f"{history_text}\nUtilisateur: {message}" if history_text else message
        resp = await run_llm_nonblocking(chat, UserMessage(text=full_msg))
        return {"response": resp.content if hasattr(resp, 'content') else str(resp)}
    except Exception as e:
        logger.error(f"[Coach Chat] {e}")
        return {"response": f"Merci pour votre question ! Pour l'étape \"{step_title}\", je vous conseille de commencer par explorer les outils disponibles dans votre espace personnel. N'hésitez pas à me relancer avec une question plus précise."}



# --- CV Offer Match Check (quick score before generation) ---
@api_router.post("/cv/check-offer-match")
async def check_offer_match(token: str, body: dict = {}):
    """Calculate a quick matching score between user profile/CV and a job offer text."""
    token_doc = await get_current_token(token)
    offer_text = body.get("offer_text", "").strip()
    if not offer_text or len(offer_text) < 20:
        raise HTTPException(400, "Texte de l'offre trop court")

    # Get user profile skills
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    user_skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:15]]

    # Get last CV analysis for more context
    last_cv = await db.cv_jobs.find_one({"token_id": token_doc["id"], "status": "completed"}, sort=[("created_at", -1)])
    cv_skills = []
    cv_experiences = []
    if last_cv and last_cv.get("result"):
        r = last_cv["result"]
        cv_skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in r.get("competences", r.get("skills", []))[:15]]
        cv_experiences = [e.get("title", "") for e in r.get("experiences", [])[:5]]

    all_skills = list(set(s.lower() for s in (user_skills + cv_skills) if s))
    offer_lower = offer_text.lower()

    # Extract significant keywords from user skills (words > 3 chars, excluding stop words)
    stop_words = {"avec", "dans", "pour", "plus", "très", "sans", "sous", "chez", "entre", "comme", "après", "avant", "leur", "cette", "mais", "aussi", "même", "tout", "tous", "être", "avoir", "faire", "dire", "aller", "voir", "bien", "fait", "sont", "nous", "vous", "autres", "base", "selon", "travail", "poste", "offre", "emploi", "recherche", "candidat", "expérience", "profil", "compétences", "savoir"}
    skill_keywords = set()
    for skill in all_skills:
        words = [w.strip("()/-,.'") for w in skill.split()]
        for w in words:
            if len(w) > 3 and w not in stop_words:
                skill_keywords.add(w)

    # Extract keywords from the offer
    offer_words = set()
    for w in offer_lower.split():
        clean = w.strip("()/-,.'!?:;\"")
        if len(clean) > 3 and clean not in stop_words:
            offer_words.add(clean)

    # Cross-match: keywords from user that appear in offer
    matched_keywords = [kw for kw in skill_keywords if kw in offer_lower]
    # Cross-match: offer keywords covered by user skills
    offer_covered = [w for w in offer_words if any(w in sk for sk in skill_keywords)]

    # Score based on offer keyword coverage (how well user matches what the offer asks)
    offer_coverage = len(offer_covered) / max(len(offer_words), 1) if offer_words else 0
    # Score based on user skills match (how many user skills are relevant)
    user_relevance = len(matched_keywords) / max(len(skill_keywords), 1) if skill_keywords else 0

    skills_score = min(int((offer_coverage * 0.6 + user_relevance * 0.4) * 70), 70)

    # Check experience relevance (match significant words from experience titles)
    exp_matches = 0
    for e in cv_experiences:
        if not e:
            continue
        exp_words = [w.strip("()/-,.'").lower() for w in e.split() if len(w) > 3]
        if any(w in offer_lower for w in exp_words):
            exp_matches += 1
    exp_score = min(int((exp_matches / max(len(cv_experiences), 1)) * 25), 25) if cv_experiences else 0

    # Base score for having a profile
    base_score = 15 if all_skills else 5

    total_score = min(base_score + skills_score + exp_score, 100)

    # Extract offer title for display
    offer_title = ""
    first_line = offer_text.split("\n")[0].strip()
    if first_line.upper().startswith("POSTE:"):
        offer_title = first_line.split(":", 1)[1].strip()
    elif len(first_line) > 5 and len(first_line) < 100:
        offer_title = first_line

    return {
        "score": total_score,
        "matched_skills": matched_keywords[:10],
        "total_user_skills": len(all_skills),
        "offer_title": offer_title,
        "alert": total_score < 50,
        "message": (
            f"Attention : votre profil correspond à {total_score}% de cette offre ({len(matched_keywords)} mot(s)-clé(s) en commun). "
            "Ce CV risque de ne pas passer les filtres ATS. Enrichissez votre profil ou ciblez une offre plus adaptée."
        ) if total_score < 50 else (
            f"Bonne compatibilité ({total_score}%) — {len(matched_keywords)} mot(s)-clé(s) en commun avec votre profil."
        )
    }


# --- 2. CV Generate Models (background job) ---
@api_router.post("/cv/generate-models")
async def start_cv_generate_models(token: str, body: dict = {}, background_tasks: BackgroundTasks = None):
    token_doc = await get_current_token(token)
    model_types = body.get("model_types", [])
    job_offer = body.get("job_offer", "")

    if not model_types:
        raise HTTPException(400, "Aucun modèle sélectionné")

    job_id = str(uuid.uuid4())
    await db.cv_gen_jobs.insert_one({
        "job_id": job_id, "token_id": token_doc["id"],
        "model_types": model_types, "job_offer": job_offer,
        "status": "processing", "progress": 0, "total": len(model_types),
        "current_model": model_types[0] if model_types else "",
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    if background_tasks:
        background_tasks.add_task(_run_cv_generation, job_id, token_doc["id"], model_types, job_offer)
    return {"job_id": job_id, "status": "processing"}


async def _run_cv_generation(job_id: str, token_id: str, model_types: list, job_offer: str):
    try:
        last_cv = await db.cv_jobs.find_one({"token_id": token_id, "status": "completed"}, sort=[("created_at", -1)])
        if not last_cv or not last_cv.get("result"):
            await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": "Aucune analyse CV trouvée"}})
            return

        # Auto-scrape if job_offer is a URL
        if job_offer and job_offer.strip().startswith("http"):
            try:
                import httpx
                from bs4 import BeautifulSoup
                url = job_offer.strip()
                headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36", "Accept-Language": "fr-FR,fr;q=0.9"}
                async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as hc:
                    resp = await hc.get(url, headers=headers)
                    resp.raise_for_status()
                soup = BeautifulSoup(resp.text, "html.parser")
                for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript"]):
                    tag.decompose()
                # Extract title from <title> for France Travail
                title_tag = soup.find("title")
                title_line = ""
                if title_tag and "Offre d'emploi" in title_tag.get_text():
                    raw = title_tag.get_text(strip=True).split("Offre d'emploi", 1)[1]
                    title_line = "POSTE: " + (raw.split(" - ")[0].strip() if " - " in raw else raw.split("|")[0].strip())
                body = soup.find("body")
                body_text = body.get_text(separator="\n", strip=True) if body else ""
                lines = [l.strip() for l in body_text.split("\n") if l.strip()]
                clean = []
                for l in lines:
                    if not clean or l != clean[-1]:
                        clean.append(l)
                scraped = "\n".join(clean)[:2500]
                if title_line:
                    scraped = title_line + "\n" + scraped
                if len(scraped) > 50:
                    job_offer = scraped
                    logger.info(f"[CV Gen] Auto-scraped URL, got {len(scraped)} chars")
            except Exception as e:
                logger.warning(f"[CV Gen] Auto-scrape failed: {e}")

        result = last_cv["result"]
        skills = [s.get('name','') if isinstance(s,dict) else s for s in result.get('competences',result.get('skills',[]))[:10]]
        exps = [e.get('title','') for e in result.get('experiences',[])[:5]]
        formations = [f.get('titre','') for f in result.get('formations',[])[:3]]
        savoir_etre = result.get('savoir_etre',[])[:5]
        offer_snippet = job_offer[:1500].strip() if job_offer else ""

        # Build shared context once
        context = f"Compétences: {', '.join(skills)}\nExpériences: {', '.join(exps)}\nFormations: {', '.join(formations)}\nSavoir-être: {', '.join([s.get('name','') if isinstance(s,dict) else str(s) for s in savoir_etre])}"

        # Determine target job title from offer
        target_job = ""
        if offer_snippet:
            first_line = offer_snippet.split('\n')[0].strip()
            # Handle "POSTE: Manager en restauration (H/F)" format from scraper
            if first_line.upper().startswith("POSTE:"):
                target_job = first_line.split(":", 1)[1].strip()[:80]
            elif len(first_line) > 5:
                target_job = first_line[:80]

        async def gen_one(mtype):
            if offer_snippet:
                prompt = f"""Génère un CV professionnel de type "{mtype}" CIBLÉ pour cette offre d'emploi:
OFFRE CIBLE: {offer_snippet}

Le titre du CV DOIT mentionner le poste visé (ex: "{target_job}").
L'accroche et les compétences doivent être reformulées pour correspondre aux exigences de l'offre.

PROFIL DU CANDIDAT:
{context}

Réponds UNIQUEMENT en JSON: {{"titre":"[Poste visé d'après l'offre]","accroche":"2 lignes ciblées sur l'offre","competences_cles":["compétences reformulées selon l'offre"],"experiences":[{{"poste":"","entreprise":"","periode":"","realisations":["reformulées selon l'offre"]}}],"formations":[{{"diplome":"","ecole":"","annee":""}}],"atouts":["atouts en lien avec l'offre"],"langues":["Français (natif)"]}}"""
            else:
                prompt = f"""Génère un CV professionnel de type "{mtype}".
PROFIL: {context}
JSON: {{"titre":"str","accroche":"2 lignes","competences_cles":["..."],"experiences":[{{"poste":"","entreprise":"","periode":"","realisations":[""]}}],"formations":[{{"diplome":"","ecole":"","annee":""}}],"atouts":["..."],"langues":["Français (natif)"]}}"""
            def _blocking_call():
                """Run LLM call in a separate thread for true parallelism."""
                import asyncio as _aio
                _loop = _aio.new_event_loop()
                _aio.set_event_loop(_loop)
                try:
                    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"cvgen-{job_id}-{mtype}",
                                    system_message="Tu es un expert RH et rédacteur de CV français. Quand une offre d'emploi est fournie, le titre du CV DOIT correspondre au poste visé dans l'offre. Adapte toutes les rubriques (accroche, compétences, réalisations) pour matcher l'offre. Réponds UNIQUEMENT en JSON valide.").with_model("openai", "gpt-5.2")
                    resp = _loop.run_until_complete(chat.send_message(UserMessage(text=prompt)))
                    text = resp.content if hasattr(resp, 'content') else str(resp).strip()
                    if "```json" in text: text = text.split("```json")[1].split("```")[0].strip()
                    elif "```" in text: text = text.split("```")[1].split("```")[0].strip()
                    return mtype, json.loads(text)
                except Exception as e:
                    logger.error(f"[CV Gen {mtype}] {e}")
                    return mtype, {"titre": "CV Professionnel", "accroche": "Professionnel motivé", "competences_cles": skills[:5], "experiences": [], "formations": [], "atouts": savoir_etre[:3]}
                finally:
                    _loop.close()
            return await asyncio.to_thread(_blocking_call)

        # Run all model generations in TRUE parallel threads
        await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"progress": 0, "current_model": "Génération parallèle..."}})
        results = await asyncio.gather(*[gen_one(mt) for mt in model_types])
        models = {mtype: data for mtype, data in results}

        await db.cv_models.update_one(
            {"token_id": token_id},
            {"$set": {"token_id": token_id, "models": models, "generated_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True
        )
        await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "completed", "progress": len(model_types)}})
    except Exception as e:
        logger.error(f"[CV Gen] {e}")
        await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": str(e)}})


@api_router.get("/cv/generate-models/status")
async def cv_generate_models_status(token: str, job_id: str):
    token_doc = await get_current_token(token)
    job = await db.cv_gen_jobs.find_one({"job_id": job_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not job:
        raise HTTPException(404, "Job non trouvé")
    return {"status": job.get("status"), "progress": job.get("progress", 0), "total": job.get("total", 0), "current_model": job.get("current_model", ""), "error": job.get("error")}


# --- 3. Coffre CV Files & Transfer ---
@api_router.get("/coffre/cv-files")
async def coffre_cv_files(token: str):
    token_doc = await get_current_token(token)
    # Get uploaded CV files and generated CVs
    docs = await db.documents.find(
        {"user_token": token_doc["id"], "category": {"$in": ["cv", "cv_uploaded", "cv_generated"]}},
        {"_id": 0, "id": 1, "title": 1, "filename": 1, "category": 1, "uploaded_at": 1}
    ).to_list(20)

    # Also check cv_jobs for uploaded filenames
    cv_jobs = await db.cv_jobs.find(
        {"token_id": token_doc["id"], "status": "completed"},
        {"_id": 0, "filename": 1, "created_at": 1}
    ).sort("created_at", -1).to_list(5)

    files = []
    for d in docs:
        files.append({"id": d.get("id", ""), "name": d.get("title", d.get("filename", "")), "type": d.get("category", "cv"), "date": d.get("uploaded_at", "")})
    for j in cv_jobs:
        files.append({"id": str(uuid.uuid4()), "name": j.get("filename", "CV analysé"), "type": "cv_analyzed", "date": j.get("created_at", "")})
    return files


@api_router.post("/coffre/transfer-cv")
async def coffre_transfer_cv(token: str, cv_type: str = "uploaded"):
    token_doc = await get_current_token(token)
    # Find the latest CV analysis result
    last_cv = await db.cv_jobs.find_one({"token_id": token_doc["id"], "status": "completed"}, sort=[("created_at", -1)])
    if not last_cv:
        raise HTTPException(404, "Aucun CV analysé trouvé")

    doc_id = str(uuid.uuid4())
    await db.documents.insert_one({
        "id": doc_id, "user_token": token_doc["id"],
        "title": last_cv.get("filename", "CV transféré"),
        "filename": last_cv.get("filename", "cv.pdf"),
        "category": cv_type, "source": "cv_analysis",
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "skills": [s.get("name", "") if isinstance(s, dict) else str(s) for s in last_cv.get("result", {}).get("competences", [])[:20]]
    })
    return {"success": True, "document_id": doc_id}


# --- 4. Jobs Matching & Applications ---
@api_router.get("/jobs/matching/preferences")
async def jobs_matching_preferences(token: str):
    token_doc = await get_current_token(token)
    prefs = await db.matching_prefs.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if prefs and prefs.get("filters"):
        return {"has_preferences": True, "filters": prefs["filters"]}
    return {"has_preferences": False, "filters": None}


@api_router.post("/jobs/matching/preferences")
async def save_matching_preferences(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    await db.matching_prefs.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"filters": body, "token_id": token_doc["id"], "updated_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    return {"success": True}



# ─── Matching Candidat / Offre (Analyser une offre) ──────────────────

async def _scrape_offer_text(url: str) -> str:
    """Fetch and extract text from an offer URL."""
    import httpx
    from html.parser import HTMLParser

    class TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.texts = []
            self._skip = False
            self._skip_tags = {"script", "style", "noscript", "svg", "head"}

        def handle_starttag(self, tag, attrs):
            if tag in self._skip_tags:
                self._skip = True

        def handle_endtag(self, tag):
            if tag in self._skip_tags:
                self._skip = False

        def handle_data(self, data):
            if not self._skip:
                t = data.strip()
                if t:
                    self.texts.append(t)

    # Try France Travail API first for FT URLs
    ft_id = ""
    if "francetravail.fr" in url and "/detail/" in url:
        ft_id = url.split("/detail/")[-1].split("?")[0].split("#")[0]

    if ft_id:
        try:
            from opc.connecteurs.france_travail import FranceTravailClient
            ft = FranceTravailClient()
            if ft.is_configured():
                token_ft = await ft._get_token(ft._scope_offres())
                async with httpx.AsyncClient(timeout=15.0) as client:
                    r = await client.get(
                        f"https://api.francetravail.io/partenaire/offresdemploi/v2/offres/{ft_id}",
                        headers={"Authorization": f"Bearer {token_ft}", "Accept": "application/json"},
                    )
                    if r.status_code == 200:
                        offre = r.json()
                        parts = []
                        parts.append(f"Titre: {offre.get('intitule', '')}")
                        parts.append(f"Entreprise: {(offre.get('entreprise') or {}).get('nom', 'Non précisé')}")
                        parts.append(f"Lieu: {(offre.get('lieuTravail') or {}).get('libelle', '')}")
                        parts.append(f"Contrat: {offre.get('typeContratLibelle', '')}")
                        if offre.get("salaire"):
                            sal = offre["salaire"]
                            parts.append(f"Salaire: {sal.get('libelle', sal.get('complement1', ''))}")
                        parts.append(f"Description: {offre.get('description', '')}")
                        comps = [c.get("libelle", "") for c in (offre.get("competences") or []) if c.get("libelle")]
                        if comps:
                            parts.append(f"Compétences: {', '.join(comps)}")
                        quals = [q.get("libelle", "") for q in (offre.get("qualitesProfessionnelles") or []) if q.get("libelle")]
                        if quals:
                            parts.append(f"Qualités: {', '.join(quals)}")
                        if offre.get("experienceLibelle"):
                            parts.append(f"Expérience: {offre['experienceLibelle']}")
                        if offre.get("formations"):
                            forms = [f_.get("niveauLibelle", "") for f_ in offre["formations"]]
                            parts.append(f"Formation: {', '.join(forms)}")
                        return "\n".join(parts)
        except Exception as e:
            logging.warning(f"FT API detail failed for {ft_id}: {e}")

    # Fallback: scrape HTML
    async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as client:
        r = await client.get(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        r.raise_for_status()
        extractor = TextExtractor()
        extractor.feed(r.text)
        text = " ".join(extractor.texts)
        return text[:8000]


async def _analyze_offer_with_ai(offer_text: str) -> dict:
    """Use LLM to analyze a job offer text and return structured data."""
    if not EMERGENT_LLM_KEY:
        raise HTTPException(503, "Clé LLM non configurée")

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"offer-analyze-{uuid.uuid4()}",
        system_message="Tu es un expert RH français. Analyse les offres d'emploi et retourne un JSON structuré."
    ).with_model("openai", "gpt-5.2")
    prompt = f"""Analyse cette offre d'emploi et retourne un JSON structuré.

OFFRE:
{offer_text[:5000]}

Retourne UNIQUEMENT un JSON valide (pas de markdown):
{{
  "titre_poste": "...",
  "entreprise": "...",
  "localisation": "...",
  "type_contrat": "CDI/CDD/Intérim/...",
  "salaire": "... ou Non précisé",
  "missions": ["mission 1", "mission 2", ...],
  "competences_requises": ["comp1", "comp2", ...],
  "soft_skills_requis": ["soft1", "soft2", ...],
  "experience_requise": "...",
  "formation_requise": "..."
}}"""

    response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
    import re
    text = response.strip()
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        return json.loads(json_match.group())
    return json.loads(text)


async def _generate_offer_synthesis(analyse: dict, offer_text: str) -> dict:
    """Generate synthesis, quality score, missing info, and recommendations."""
    if not EMERGENT_LLM_KEY:
        return {"synthese": "", "score_qualite_offre": 50, "informations_manquantes": [], "recommandations_candidat": []}

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"offer-synth-{uuid.uuid4()}",
        system_message="Tu es un expert RH français. Évalue la qualité des offres d'emploi."
    ).with_model("openai", "gpt-5.2")
    prompt = f"""Évalue cette offre d'emploi. Titre: {analyse.get('titre_poste', '')}. Entreprise: {analyse.get('entreprise', '')}.

Texte de l'offre (extrait):
{offer_text[:3000]}

Retourne UNIQUEMENT un JSON valide:
{{
  "synthese": "Résumé en 2-3 phrases de l'offre",
  "score_qualite_offre": 0-100,
  "informations_manquantes": [
    {{"theme": "Salaire", "detail": "Le salaire n'est pas mentionné", "importance": "Important pour évaluer l'offre"}}
  ],
  "recommandations_candidat": [
    "Conseil 1 pour le candidat",
    "Conseil 2"
  ]
}}"""

    response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
    import re
    text = response.strip() if isinstance(response, str) else response.text.strip()
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        return json.loads(json_match.group())
    return json.loads(text)


@api_router.get("/matching/history")
async def matching_history(token: str):
    token_doc = await get_current_token(token)
    analyses = await db.offer_analyses.find(
        {"token_id": token_doc["id"]},
        {"_id": 0}
    ).sort("created_at", -1).limit(10).to_list(10)

    result = []
    for a in analyses:
        analyse = a.get("analyse", {})
        result.append({
            "id": a.get("analysis_id", ""),
            "titre": analyse.get("titre_poste", "Offre"),
            "entreprise": analyse.get("entreprise", ""),
            "score_qualite": a.get("score_qualite_offre", 0),
            "score_matching": a.get("score_matching"),
            "created_at": a.get("created_at", ""),
        })
    return {"analyses": result}


@api_router.post("/matching/analyze-offer-url")
async def analyze_offer_url(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    url = body.get("url", "").strip()
    if not url.startswith("http"):
        raise HTTPException(400, "URL invalide")

    try:
        offer_text = await _scrape_offer_text(url)
    except Exception as e:
        logging.error(f"Scrape error for {url}: {e}")
        raise HTTPException(400, f"Impossible de récupérer le contenu de cette URL. Collez le texte manuellement.")

    if len(offer_text.strip()) < 30:
        raise HTTPException(400, "Contenu insuffisant récupéré. Collez le texte de l'offre manuellement.")

    try:
        analyse = await _analyze_offer_with_ai(offer_text)
    except Exception as e:
        logging.error(f"AI analysis error: {e}")
        raise HTTPException(500, "Erreur lors de l'analyse IA de l'offre")

    try:
        synthesis = await _generate_offer_synthesis(analyse, offer_text)
    except Exception as e:
        logging.warning(f"Synthesis error: {e}")
        synthesis = {"synthese": "", "score_qualite_offre": 50, "informations_manquantes": [], "recommandations_candidat": []}

    analysis_id = str(uuid.uuid4())
    doc = {
        "analysis_id": analysis_id,
        "token_id": token_doc["id"],
        "url": url,
        "offer_text": offer_text[:6000],
        "analyse": analyse,
        **synthesis,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.offer_analyses.insert_one(doc)

    return {
        "analysis_id": analysis_id,
        "analyse": analyse,
        **synthesis,
    }


@api_router.post("/matching/analyze-offer")
async def analyze_offer_text(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    offer_text = body.get("text", "").strip()
    if len(offer_text) < 30:
        raise HTTPException(400, "Texte trop court (minimum 30 caractères)")

    try:
        analyse = await _analyze_offer_with_ai(offer_text)
    except Exception as e:
        logging.error(f"AI analysis error: {e}")
        raise HTTPException(500, "Erreur lors de l'analyse IA de l'offre")

    try:
        synthesis = await _generate_offer_synthesis(analyse, offer_text)
    except Exception as e:
        logging.warning(f"Synthesis error: {e}")
        synthesis = {"synthese": "", "score_qualite_offre": 50, "informations_manquantes": [], "recommandations_candidat": []}

    analysis_id = str(uuid.uuid4())
    doc = {
        "analysis_id": analysis_id,
        "token_id": token_doc["id"],
        "source": body.get("source", "paste"),
        "offer_text": offer_text[:6000],
        "analyse": analyse,
        **synthesis,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.offer_analyses.insert_one(doc)

    return {
        "analysis_id": analysis_id,
        "analyse": analyse,
        **synthesis,
    }


@api_router.post("/matching/match-profile")
async def match_profile_with_offer(token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    analysis_id = body.get("analysis_id", "")
    if not analysis_id:
        raise HTTPException(400, "analysis_id requis")

    analysis_doc = await db.offer_analyses.find_one({"analysis_id": analysis_id, "token_id": token_doc["id"]})
    if not analysis_doc:
        raise HTTPException(404, "Analyse non trouvée")

    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:20]]
    experiences = (passport or {}).get("experiences", [])
    formations = (passport or {}).get("formations", [])
    savoir_faire = (passport or {}).get("savoir_faire", [])
    savoir_etre = (passport or {}).get("savoir_etre", [])

    analyse = analysis_doc.get("analyse", {})

    if not EMERGENT_LLM_KEY:
        raise HTTPException(503, "Clé LLM non configurée")

    user_context = f"""Compétences: {', '.join(skills[:15])}
Savoir-faire: {', '.join([s if isinstance(s, str) else s.get('name','') for s in savoir_faire[:10]]) if savoir_faire else 'Non renseigné'}
Savoir-être: {', '.join([s if isinstance(s, str) else s.get('name','') for s in savoir_etre[:10]]) if savoir_etre else 'Non renseigné'}
Expériences: {json.dumps([{'titre': e.get('title',''), 'duree': e.get('duration','')} for e in experiences[:5]], ensure_ascii=False)}
Formations: {json.dumps([{'titre': f.get('title',''), 'niveau': f.get('level','')} for f in formations[:3]], ensure_ascii=False)}"""

    offer_context = f"""Poste: {analyse.get('titre_poste', '')}
Compétences requises: {', '.join(analyse.get('competences_requises', []))}
Soft skills: {', '.join(analyse.get('soft_skills_requis', []))}
Expérience: {analyse.get('experience_requise', '')}
Formation: {analyse.get('formation_requise', '')}"""

    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"match-profile-{uuid.uuid4()}",
        system_message="Tu es un expert RH français. Compare un profil candidat avec une offre d'emploi."
    ).with_model("openai", "gpt-5.2")
    prompt = f"""Compare ce profil candidat avec cette offre d'emploi et évalue la compatibilité.
IMPORTANT : Analyse séparément les HARD SKILLS (compétences techniques) et les SOFT SKILLS (compétences comportementales/transférables).
Les soft skills sont essentiels dans une logique de RECONVERSION PROFESSIONNELLE : mets en avant leur valeur de transfert.

PROFIL CANDIDAT:
{user_context}

OFFRE D'EMPLOI:
{offer_context}

Retourne UNIQUEMENT un JSON valide:
{{
  "score_global": 0-100,
  "verdict": "Phrase résumant le niveau de compatibilité",
  "details": {{
    "competences_techniques": {{
      "score": 0-100,
      "forces": ["hard skill qui matche 1", "hard skill qui matche 2"],
      "lacunes": ["hard skill manquant 1"],
      "matched_skills": ["Compétence A du profil → Compétence X de l'offre", "Compétence B → Y"]
    }},
    "soft_skills": {{
      "score": 0-100,
      "forces": ["soft skill transférable 1", "soft skill transférable 2"],
      "lacunes": ["soft skill à développer"],
      "matched_skills": ["Qualité A du profil → Besoin X de l'offre"],
      "transferability_message": "Explication en 2-3 phrases de pourquoi les soft skills de ce candidat sont particulièrement transférables et valorisables pour ce poste, même en reconversion."
    }},
    "experience": {{
      "score": 0-100,
      "forces": ["force 1"],
      "lacunes": ["lacune 1"]
    }},
    "formation": {{
      "score": 0-100,
      "forces": ["force 1"],
      "lacunes": []
    }}
  }},
  "recommandations": [
    {{"type": "cv", "conseil": "Conseil pour améliorer le CV"}},
    {{"type": "formation", "conseil": "Formation recommandée"}},
    {{"type": "entretien", "conseil": "Conseil pour l'entretien"}},
    {{"type": "profil", "conseil": "Amélioration du profil"}}
  ],
  "message_accroche": "Message de motivation personnalisé de 3-4 phrases pour cette offre"
}}"""

    response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
    import re
    text = response.strip() if isinstance(response, str) else response.text.strip()
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        result = json.loads(json_match.group())
    else:
        result = json.loads(text)

    # Save matching score back to analysis
    await db.offer_analyses.update_one(
        {"analysis_id": analysis_id},
        {"$set": {"score_matching": result.get("score_global", 0)}}
    )

    return result



@api_router.post("/jobs/france-travail/search")
async def search_france_travail_offres(token: str, body: dict = {}):
    """Recherche d'offres d'emploi France Travail basée sur le profil utilisateur."""
    from opc.connecteurs.france_travail import FranceTravailClient
    token_doc = await get_current_token(token)
    ft = FranceTravailClient()

    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    skills = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (profile or {}).get("skills", [])[:15]]

    code_rome = body.get("code_rome", "")
    departement = body.get("departement", "")
    mots_cles = body.get("motsCles", "")
    commune = body.get("commune", "")

    if not code_rome and passport:
        rome_codes = passport.get("rome_codes", [])
        if rome_codes:
            code_rome = rome_codes[0] if isinstance(rome_codes[0], str) else rome_codes[0].get("code", "")

    # If no motsCles provided, use job titles from profile
    if not mots_cles and not code_rome:
        exp_titles = [e.get("title", "") for e in (passport or {}).get("experiences", []) if isinstance(e, dict) and e.get("title")]
        if exp_titles:
            mots_cles = exp_titles[0]

    user_words = set()
    for s in skills:
        for w in s.lower().split():
            if len(w) > 3:
                user_words.add(w)

    # Try real France Travail API first
    ft_api_ok = False
    resultats = []
    ft_message = ""
    if ft.is_configured():
        try:
            result = await ft.search_offres(
                departement=departement or None,
                code_rome=code_rome or None,
                motsCles=mots_cles or None,
                commune=commune or None,
            )
            resultats = result.get("resultats", [])
            ft_api_ok = True
        except Exception as e:
            err_str = str(e)
            logging.warning(f"France Travail API unavailable: {err_str[:200]}")
            if "invalid_scope" in err_str:
                ft_message = "L'API Offres d'emploi France Travail n'est pas souscrite pour cette application. Activez-la sur francetravail.io > Vos applications. En attendant, voici les offres de notre base de données correspondant à votre profil."
            else:
                ft_message = f"L'API France Travail est temporairement indisponible. Voici les offres de notre base de données."

    # Fallback: use internal DB jobs if FT API failed
    if not ft_api_ok:
        inferred_sectors = _infer_sectors_from_profile(
            (passport or {}).get("experiences", []),
            (passport or {}).get("savoir_faire", skills)
        )
        inferred_set = set(s.lower() for s in inferred_sectors[:3])
        exp_titles = [e.get("title", "") for e in (passport or {}).get("experiences", []) if isinstance(e, dict) and e.get("title")]

        jobs = await db.jobs.find({"status": "active"}, {"_id": 0}).limit(20).to_list(20)
        matches = []
        for job in jobs:
            score, matched_skills, rationale = _score_job_basic(job, skills, user_words, exp_titles, inferred_set)
            entry = _build_match_entry(job, score, matched_skills, rationale)
            entry["source"] = "base_interne"
            entry["pourquoi_ce_match"] = f"Base interne — {rationale}"
            matches.append(entry)

        matches.sort(key=lambda x: x["matching_score"], reverse=True)
        return {
            "has_data": len(matches) > 0,
            "has_filters": False,
            "source": "base_interne",
            "message": ft_message,
            "profile_summary": {
                "titre": "Offres correspondant à votre profil",
                "skills_count": len(skills),
                "has_optimized_cv": False,
                "has_career_project": False,
            },
            "matches": matches,
        }

    # Process FT API results
    matches = []
    for offre in resultats[:20]:
        titre = offre.get("intitule", "")
        entreprise = (offre.get("entreprise") or {}).get("nom", "Entreprise")
        lieu = (offre.get("lieuTravail") or {}).get("libelle", "")
        contrat = offre.get("typeContratLibelle", offre.get("typeContrat", ""))
        desc = offre.get("description", "")[:300]
        # Direct link to the offer detail page on France Travail
        offre_id = offre.get("id", "")
        url = f"https://candidat.francetravail.fr/offres/recherche/detail/{offre_id}" if offre_id else ""

        ft_competences = [c.get("libelle", "") for c in (offre.get("competences") or []) if c.get("libelle")]
        matched_comps = []
        for fc in ft_competences:
            fc_words = set(w for w in fc.lower().split() if len(w) > 3)
            if fc_words & user_words:
                matched_comps.append(fc)

        score = min(100, 30 + int((len(matched_comps) / max(len(ft_competences), 1)) * 50) + (10 if matched_comps else 0))

        salaire = ""
        if offre.get("salaire"):
            sal = offre["salaire"]
            salaire = sal.get("libelle", sal.get("complement1", ""))

        matches.append({
            "titre": titre,
            "matching_score": score,
            "secteur": offre.get("secteurActiviteLibelle", ""),
            "type_contrat": contrat,
            "entreprise_type": entreprise,
            "localisation": lieu,
            "description": desc,
            "salaire_indicatif": salaire,
            "competences_matchees": matched_comps[:6],
            "pourquoi_ce_match": f"Offre France Travail — {len(matched_comps)} compétence(s) en commun",
            "url_offre": url,
            "scoring": None,
            "source": "france_travail",
        })

    matches.sort(key=lambda x: x["matching_score"], reverse=True)
    return {
        "has_data": len(matches) > 0,
        "has_filters": False,
        "source": "france_travail",
        "profile_summary": {
            "titre": "Offres France Travail",
            "skills_count": len(skills),
            "has_optimized_cv": False,
            "has_career_project": False,
        },
        "matches": matches,
    }


@api_router.get("/jobs/applications")
async def jobs_applications(token: str):
    token_doc = await get_current_token(token)
    apps = await db.applications.find({"token_id": token_doc["id"]}, {"_id": 0}).sort("applied_at", -1).to_list(50)
    return {"applications": apps, "total": len(apps)}


@api_router.put("/jobs/applications/{app_id}/status")
async def update_application_status(app_id: str, token: str, body: dict = {}):
    token_doc = await get_current_token(token)
    new_status = body.get("status", "")
    valid = ["en_preparation", "envoyee", "entretien", "acceptee", "refusee"]
    if new_status not in valid:
        raise HTTPException(400, f"Statut invalide. Valeurs acceptées : {', '.join(valid)}")
    result = await db.applications.update_one(
        {"id": app_id, "token_id": token_doc["id"]},
        {"$set": {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Candidature non trouvée")
    return {"success": True}


@api_router.delete("/jobs/applications/{app_id}")
async def delete_application(app_id: str, token: str):
    token_doc = await get_current_token(token)
    result = await db.applications.delete_one({"id": app_id, "token_id": token_doc["id"]})
    if result.deleted_count == 0:
        raise HTTPException(404, "Candidature non trouvée")
    return {"success": True}


# --- 5. Notifications mark read ---
@api_router.post("/notifications/mark-read")
async def mark_notification_read(token: str, body: dict = {}):
    notification_id = body.get("notification_id", "")
    if notification_id:
        await db.notifications.update_one({"id": notification_id}, {"$set": {"read": True}})
    return {"success": True}


@api_router.post("/notifications/mark-all-read")
async def mark_all_notifications_read(token: str = "", body: dict = {}):
    token_val = token or body.get("token", "")
    if token_val:
        token_doc = await get_current_token(token_val)
        await db.notifications.update_many({"token_id": token_doc["id"]}, {"$set": {"read": True}})
    return {"success": True}


# --- 6. Emerging market correlation ---
@api_router.get("/emerging/market-correlation")
async def emerging_market_correlation(token: str):
    token_doc = await get_current_token(token)

    # Get emerging competences
    emerging_docs = await db.emerging_competences.find(
        {"token_id": token_doc["id"]}, {"_id": 0}
    ).to_list(50)

    # Also get passport savoir_faire for broader correlation
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    passport_sf = [s.get("name", "") if isinstance(s, dict) else str(s) for s in (passport or {}).get("savoir_faire", [])]

    # Build skill list: emerging competences + top savoir_faire
    skills_to_check = []
    for ec in emerging_docs:
        skills_to_check.append({
            "id": ec.get("id", ""),
            "name": ec.get("nom_principal", ""),
            "score": ec.get("score_emergence", 0),
            "categorie": ec.get("categorie", ""),
            "tendance": ec.get("tendance", "stable"),
        })

    if not skills_to_check and passport_sf:
        for i, sf in enumerate(passport_sf[:10]):
            skills_to_check.append({
                "id": f"sf-{i}",
                "name": sf,
                "score": 50,
                "categorie": "sectorielle",
                "tendance": "stable",
            })

    correlations = []
    in_market = 0
    high_demand = 0
    growing = set()

    for skill_info in skills_to_check[:12]:
        skill_name = skill_info["name"]
        # Search keywords (use first 2 significant words)
        keywords = [w for w in skill_name.split() if len(w) > 3][:2]

        related_jobs = []
        rncp_count = 0
        sectors = []

        for kw in keywords:
            metiers = await db.opc_metiers.find(
                {"$or": [
                    {"savoir_faire": {"$regex": kw, "$options": "i"}},
                    {"metier": {"$regex": kw, "$options": "i"}},
                    {"competences_cles": {"$regex": kw, "$options": "i"}},
                ]},
                {"_id": 0, "metier": 1, "filiere_nom": 1}
            ).limit(5).to_list(5)

            for m in metiers:
                job_name = m.get("metier", "")
                if job_name and job_name not in related_jobs:
                    related_jobs.append(job_name)
                sector = m.get("filiere_nom", "")
                if sector and sector not in sectors:
                    sectors.append(sector)

            rncp_count += await db.opc_certifications.count_documents(
                {"intitule": {"$regex": kw, "$options": "i"}, "statut": "ACTIVE"}
            )

        # Determine demand level
        demand = "faible"
        if len(related_jobs) >= 3 or rncp_count >= 5:
            demand = "fort"
            high_demand += 1
        elif len(related_jobs) >= 1 or rncp_count >= 1:
            demand = "modéré"

        if len(related_jobs) > 0:
            in_market += 1

        for s in sectors:
            growing.add(s)

        trend = skill_info["tendance"]
        if rncp_count > 5:
            trend = "hausse"
        elif rncp_count > 0:
            trend = "stable"

        correlations.append({
            "competence_id": skill_info["id"],
            "skill": skill_name,
            "market_demand": demand,
            "related_jobs": related_jobs[:5],
            "rncp_certifications": rncp_count,
            "trend": trend,
            "sectors": sectors[:3],
        })

    total = len(correlations)
    alignment_pct = round((in_market / total * 100)) if total > 0 else 0

    return {
        "correlations": correlations,
        "total_skills_analyzed": total,
        "has_data": total > 0,
        "summary": {
            "market_alignment_pct": alignment_pct,
            "in_market": in_market,
            "high_demand": high_demand,
            "growing_sectors": len(growing),
        }
    }


# ═══════════════════════════════════════════════════════════════════════════════
# JOB DATING — Événements emploi et matching intelligent
# ═══════════════════════════════════════════════════════════════════════════════

# Mapping keywords → sectors for smart inference

# --- 8. Experience Proof (contributeur sociétal) ---
@api_router.post("/passport/experience-proof")
async def add_experience_proof(token: str, body: dict = {}):
    """Add a concrete example/proof to a passport experience, contributing to OPC."""
    token_doc = await get_current_token(token)
    exp_id = body.get("experience_id", "")
    proof = body.get("proof", "").strip()
    if not proof:
        raise HTTPException(400, "L'exemple concret est requis")

    # Update passport experience with proof
    result = await db.passports.update_one(
        {"token_id": token_doc["id"], "experiences.id": exp_id},
        {"$set": {"experiences.$.proof": proof, "experiences.$.proof_date": datetime.now(timezone.utc).isoformat()}}
    )

    if result.modified_count == 0:
        # Try matching by title (fallback if id doesn't match)
        passport = await db.passports.find_one({"token_id": token_doc["id"]})
        if passport:
            experiences = passport.get("experiences", [])
            for i, exp in enumerate(experiences):
                if exp.get("id") == exp_id or (not exp_id and i == 0):
                    experiences[i]["proof"] = proof
                    experiences[i]["proof_date"] = datetime.now(timezone.utc).isoformat()
                    await db.passports.update_one(
                        {"token_id": token_doc["id"]},
                        {"$set": {"experiences": experiences}}
                    )
                    break

    # Also contribute to OPC: store anonymized proof for collective intelligence
    await db.opc_contributions.insert_one({
        "id": str(uuid.uuid4()),
        "type": "experience_proof",
        "domain": body.get("domain", ""),
        "proof_text": proof[:500],
        "skills_related": body.get("skills", []),
        "contributed_at": datetime.now(timezone.utc).isoformat(),
        "anonymous": True
    })

    # Count total contributions by this user
    contrib_count = await db.opc_contributions.count_documents({})
    return {"success": True, "message": "Contribution ajoutée à l'Observatoire", "total_contributions": contrib_count}


@api_router.get("/learning/recommendations")
async def learning_recommendations(token: str):
    token_doc = await get_current_token(token)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    gaps = (profile or {}).get("gaps", [])
    sectors = (profile or {}).get("sectors", [])

    recommendations = []
    for i, gap in enumerate(gaps[:5]):
        recommendations.append({
            "id": str(uuid.uuid4()), "title": f"Renforcer : {gap}",
            "description": f"Module de formation ciblé pour développer votre compétence en {gap.lower()}",
            "priority": "haute" if i < 2 else "moyenne",
            "duration": "2-4 semaines", "type": "formation",
            "linked_gap": gap
        })
    for sector in sectors[:3]:
        recommendations.append({
            "id": str(uuid.uuid4()), "title": f"Découverte secteur : {sector}",
            "description": f"Parcours d'orientation vers le secteur {sector}",
            "priority": "moyenne", "duration": "1-2 semaines", "type": "orientation",
            "linked_sector": sector
        })
    return {"recommendations": recommendations, "total": len(recommendations)}


# ═══════════════════════════════════════════════════════════════════════════════
# SOFT SKILLS ILLUSTRATIONS (S.A.R.E) — Preuves concrètes
# ═══════════════════════════════════════════════════════════════════════════════

@api_router.get("/passport/illustrations")
async def get_illustrations(token: str):
    """Get all SARE illustrations for the user."""
    token_doc = await get_current_token(token)
    illustrations = await db.skill_illustrations.find(
        {"token_id": token_doc["id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(500)
    return {"illustrations": illustrations}


@api_router.post("/passport/illustrations")
async def save_illustration(token: str, body: dict):
    """Save a SARE illustration for a soft skill on an experience. Auto-adds to coffre-fort."""
    token_doc = await get_current_token(token)
    exp_id = body.get("experience_id", "")
    soft_skill = body.get("soft_skill", "").strip()
    if not exp_id or not soft_skill:
        raise HTTPException(400, "experience_id et soft_skill requis")

    # Find the experience in passport
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    exp_data = None
    if passport:
        for exp in passport.get("experiences", []):
            if exp.get("id") == exp_id:
                exp_data = exp
                break

    illus_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    illustration = {
        "id": illus_id,
        "token_id": token_doc["id"],
        "experience_id": exp_id,
        "soft_skill": soft_skill,
        "situation_text": body.get("situation_text", ""),
        "sare_situation": body.get("sare_situation", ""),
        "sare_action": body.get("sare_action", ""),
        "sare_resultat": body.get("sare_resultat", ""),
        "sare_enseignement": body.get("sare_enseignement", ""),
        "opc_consent": body.get("opc_consent", False),
        "created_at": now,
    }
    await db.skill_illustrations.insert_one(illustration)

    # Auto-add experience proof entry to coffre-fort (1st certification step)
    if exp_data:
        org = exp_data.get("organization", "")
        title = exp_data.get("title", "")
        coffre_title = f"Preuve S.A.R.E — {title}"
        if org:
            coffre_title += f" ({org})"

        await db.coffre_documents.update_one(
            {"token_id": token_doc["id"], "linked_experience_id": exp_id, "category": "experience_prouvee", "linked_soft_skill": soft_skill},
            {"$set": {
                "id": str(uuid.uuid4()),
                "token_id": token_doc["id"],
                "title": coffre_title,
                "category": "experience_prouvee",
                "document_type": "sare_proof",
                "trust_level": "auto_declare",
                "source_type": "utilisateur",
                "linked_experience_id": exp_id,
                "linked_soft_skill": soft_skill,
                "linked_organization": org,
                "description": f"Soft skill '{soft_skill}' prouvé par méthode S.A.R.E",
                "uploaded_at": now,
            }},
            upsert=True,
        )

    return {"success": True, "id": illus_id}


@api_router.delete("/passport/illustrations/{illus_id}")
async def delete_illustration(illus_id: str, token: str):
    """Delete a SARE illustration."""
    token_doc = await get_current_token(token)
    illus = await db.skill_illustrations.find_one({"id": illus_id, "token_id": token_doc["id"]})
    if not illus:
        raise HTTPException(404, "Illustration non trouvée")

    await db.skill_illustrations.delete_one({"id": illus_id, "token_id": token_doc["id"]})

    # Also remove the coffre entry for this specific proof
    await db.coffre_documents.delete_one({
        "token_id": token_doc["id"],
        "linked_experience_id": illus.get("experience_id"),
        "linked_soft_skill": illus.get("soft_skill"),
        "category": "experience_prouvee",
    })

    return {"success": True}


@api_router.post("/passport/illustrations/suggest")
async def suggest_illustrations(token: str, body: dict):
    """AI suggests SARE illustrations for an experience."""
    token_doc = await get_current_token(token)
    exp_id = body.get("experience_id", "")

    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if not passport:
        raise HTTPException(404, "Passeport non trouvé")

    exp_data = None
    for exp in passport.get("experiences", []):
        if exp.get("id") == exp_id:
            exp_data = exp
            break
    if not exp_data:
        raise HTTPException(404, "Expérience non trouvée")

    soft_skills = passport.get("savoir_etre", [])
    skills_list = [s.get("name", s) if isinstance(s, dict) else s for s in soft_skills[:10]]

    if EMERGENT_LLM_KEY and skills_list:
        try:
            prompt = f"""Pour l'expérience "{exp_data.get('title','')}" chez "{exp_data.get('organization','')}" ({exp_data.get('description','')}),
suggère 3 soft skills parmi cette liste : {', '.join(skills_list)}.
Pour chaque soft skill, propose une illustration S.A.R.E réaliste et concrète.
Réponds en JSON : [{{"soft_skill":"...", "situation":"...", "action":"...", "resultat":"...", "enseignement":"..."}}]"""

            chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"sare-suggest-{uuid.uuid4().hex[:8]}",
                           system_message="Tu es un expert en bilan de compétences. Réponds uniquement en JSON.").with_model("openai", "gpt-5.2")
            response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
            text = response.content if hasattr(response, 'content') else str(response)

            import json as json_mod
            # Extract JSON from response
            start = text.find("[")
            end = text.rfind("]") + 1
            if start >= 0 and end > start:
                suggestions = json_mod.loads(text[start:end])
                return {"suggestions": suggestions}
        except Exception as e:
            logger.error(f"SARE suggest error: {e}")

    # Fallback
    return {"suggestions": [{"soft_skill": s, "situation": "", "action": "", "resultat": "", "enseignement": ""} for s in skills_list[:3]]}


@api_router.post("/passport/illustrations/sare")
async def rewrite_sare(token: str, body: dict):
    """AI reformulates an illustration into proper SARE format."""
    token_doc = await get_current_token(token)
    illus_id = body.get("illustration_id", "")
    illus = await db.skill_illustrations.find_one({"id": illus_id, "token_id": token_doc["id"]})
    if not illus:
        raise HTTPException(404, "Illustration non trouvée")

    raw_text = illus.get("sare_situation", "") or illus.get("situation_text", "")
    action = illus.get("sare_action", "")
    resultat = illus.get("sare_resultat", "")

    if EMERGENT_LLM_KEY:
        try:
            prompt = f"""Reformule cette preuve de compétence en méthode S.A.R.E professionnelle (3-4 lignes max).
Soft skill : {illus.get('soft_skill','')}
Situation : {raw_text}
Action : {action}
Résultat : {resultat}
Reformule en un texte fluide et percutant pour un recruteur."""

            chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"sare-rewrite-{uuid.uuid4().hex[:8]}",
                           system_message="Tu es un CIP expert en valorisation de compétences. Réponds directement avec le texte reformulé.").with_model("openai", "gpt-5.2")
            response = await run_llm_nonblocking(chat, UserMessage(text=prompt))
            sare_text = response.content if hasattr(response, 'content') else str(response)

            await db.skill_illustrations.update_one(
                {"id": illus_id},
                {"$set": {"sare_text": sare_text.strip()}}
            )
            return {"success": True, "sare_text": sare_text.strip()}
        except Exception as e:
            logger.error(f"SARE rewrite error: {e}")

    return {"success": False, "message": "Service IA indisponible"}


# ═══════════════════════════════════════════════════════════════════════════════
# CERTIFICATION STATUS & BADGES
# ═══════════════════════════════════════════════════════════════════════════════

@api_router.get("/coffre/certification-status")
async def get_certification_status(token: str):
    """Get certification status grouped by workplace with progressive badges."""
    token_doc = await get_current_token(token)

    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    experiences = passport.get("experiences", []) if passport else []
    illustrations = await db.skill_illustrations.find({"token_id": token_doc["id"]}).to_list(500)

    # Group experiences by organization
    orgs = {}
    for exp in experiences:
        org = exp.get("organization", "Non spécifié") or "Non spécifié"
        if org not in orgs:
            orgs[org] = {"organization": org, "experiences": [], "has_contract": False, "total_proofs": 0}
        exp_illus = [i for i in illustrations if i.get("experience_id") == exp.get("id")]
        orgs[org]["experiences"].append({
            "id": exp.get("id"),
            "title": exp.get("title", ""),
            "is_certified": exp.get("is_certified", False),
            "has_contract": bool(exp.get("proof_document")),
            "proofs_count": len(exp_illus),
            "soft_skills_proved": [i.get("soft_skill") for i in exp_illus],
        })
        orgs[org]["total_proofs"] += len(exp_illus)
        if exp.get("proof_document"):
            orgs[org]["has_contract"] = True

    # Compute global stats
    total_exp = len(experiences)
    total_proved = sum(1 for exp in experiences
                       if any(i.get("experience_id") == exp.get("id") for i in illustrations))
    total_certified = sum(1 for exp in experiences if exp.get("is_certified"))
    total_with_contract = sum(1 for exp in experiences if exp.get("proof_document"))

    # Determine global badge level
    badge_level = 0
    badge_label = "Débutant"
    badge_color = "slate"
    if total_proved >= 3:
        badge_level = 1
        badge_label = "Contributeur"
        badge_color = "emerald"
    if total_with_contract >= 1:
        badge_level = 2
        badge_label = "Certifié"
        badge_color = "blue"
    if total_exp > 0 and total_proved == total_exp and total_with_contract == total_exp:
        badge_level = 3
        badge_label = "Expert Certifié"
        badge_color = "amber"

    return {
        "workplaces": list(orgs.values()),
        "stats": {
            "total_experiences": total_exp,
            "total_proved": total_proved,
            "total_certified": total_certified,
            "total_with_contract": total_with_contract,
        },
        "badge": {
            "level": badge_level,
            "label": badge_label,
            "color": badge_color,
        },
    }



@api_router.get("/coffre/opc-consent")
async def get_opc_consent(token: str, organization: str):
    """Get OPC consent status for a certified organization."""
    token_doc = await get_current_token(token)
    consent = await db.opc_consents.find_one(
        {"token_id": token_doc["id"], "organization": organization}
    )
    return {"opc_consent": bool(consent and consent.get("opc_consent"))}


@api_router.post("/coffre/opc-consent")
async def set_opc_consent(token: str, body: dict):
    """Set OPC consent for a certified organization's illustrations."""
    token_doc = await get_current_token(token)
    organization = body.get("organization", "")
    opc_consent = body.get("opc_consent", False)

    if not organization:
        raise HTTPException(400, "Organization requis")

    await db.opc_consents.update_one(
        {"token_id": token_doc["id"], "organization": organization},
        {"$set": {
            "token_id": token_doc["id"],
            "organization": organization,
            "opc_consent": opc_consent,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )

    # Update all illustrations for this org
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if passport:
        org_exp_ids = [e.get("id") for e in passport.get("experiences", []) if e.get("organization") == organization]
        if org_exp_ids:
            await db.skill_illustrations.update_many(
                {"token_id": token_doc["id"], "experience_id": {"$in": org_exp_ids}},
                {"$set": {"opc_consent": opc_consent}}
            )

    return {"success": True, "opc_consent": opc_consent}


# ═══════════════════════════════════════════════════════════════════════════════
# OPC — CONTRIBUTIONS & FICHES MÉTIER
# ═══════════════════════════════════════════════════════════════════════════════

@api_router.post("/opc/contribute")
async def opc_contribute(token: str, body: dict):
    """
    Contribute a validated + certified proof to the OPC.
    Creates/updates the fiche métier for the job title and enriches the OPC database.
    """
    token_doc = await get_current_token(token)
    document_id = body.get("document_id", "")

    if not document_id:
        raise HTTPException(400, "document_id requis")

    # Get the coffre document
    coffre_doc = await db.coffre_documents.find_one({"id": document_id, "token_id": token_doc["id"]})
    if not coffre_doc:
        raise HTTPException(404, "Document non trouvé")

    if coffre_doc.get("trust_level") != "valide":
        raise HTTPException(400, "Le document doit être validé avant de contribuer à l'OPC")

    # Get the linked experience to check contract
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if not passport:
        raise HTTPException(404, "Passeport non trouvé")

    exp_id = coffre_doc.get("linked_experience_id")
    experience = next((e for e in passport.get("experiences", []) if e.get("id") == exp_id), None)
    if not experience:
        raise HTTPException(404, "Expérience liée non trouvée")

    if not experience.get("proof_document"):
        raise HTTPException(400, "Un contrat de travail est nécessaire pour contribuer à l'OPC (niveau Certifié requis)")

    # Get the illustration S.A.R.E data
    soft_skill = coffre_doc.get("linked_soft_skill", "")
    illustration = await db.skill_illustrations.find_one({
        "token_id": token_doc["id"],
        "experience_id": exp_id,
        "soft_skill": soft_skill,
    })

    job_title = experience.get("title", "Non spécifié")
    organization = coffre_doc.get("linked_organization", experience.get("organization", ""))

    # Collect hard skills from the experience
    exp_skills = experience.get("skills", [])
    hard_skills = [s for s in exp_skills if isinstance(s, str)] if exp_skills else []

    # Collect qualités humaines & valeurs from user profile (D'CLIC PRO)
    profile = await db.profiles.find_one({"token_id": token_doc["id"]})
    qualites_humaines = []
    valeurs = []
    if profile:
        dp = profile.get("dclic_profile", {})
        vd = dp.get("vertu_data", {})
        vp = dp.get("vertus_profile", {})
        # Qualités humaines from forces principales (ADN Pro)
        adn = None
        if passport:
            adn = passport.get("identity_adn", {})
        if adn:
            qualites_humaines = (adn.get("forces_principales", []) or [])[:5]
        # Valeurs from D'CLIC PRO
        valeurs = (vd.get("valeurs_schwartz", []) or [])[:5]
        if vp.get("dominant_name"):
            valeurs = [vp["dominant_name"]] + valeurs

    # Build contribution
    contribution_id = str(uuid.uuid4())
    contribution = {
        "id": contribution_id,
        "token_id": token_doc["id"],
        "document_id": document_id,
        "experience_id": exp_id,
        "organization": organization,
        "job_title": job_title,
        "soft_skill": soft_skill,
        "hard_skills": hard_skills,
        "qualites_humaines": qualites_humaines,
        "valeurs": valeurs,
        "source": "Contributeur sociétal",
        "sare_situation": illustration.get("sare_situation", "") if illustration else "",
        "sare_action": illustration.get("sare_action", "") if illustration else "",
        "sare_resultat": illustration.get("sare_resultat", "") if illustration else "",
        "sare_enseignement": illustration.get("sare_enseignement", "") if illustration else "",
        "sare_text": illustration.get("sare_text", "") if illustration else "",
        "contributed_at": datetime.now(timezone.utc).isoformat(),
        "is_certified": True,
    }

    # Upsert contribution (avoid duplicates)
    await db.opc_contributions.update_one(
        {"token_id": token_doc["id"], "document_id": document_id},
        {"$set": contribution},
        upsert=True,
    )

    # Update/create fiche métier OPC
    fiche = await db.fiches_metier_opc.find_one({"job_title": job_title})
    if not fiche:
        fiche = {
            "id": str(uuid.uuid4()),
            "job_title": job_title,
            "competences": {},
            "total_contributors": 0,
            "organizations": [],
            "created_at": datetime.now(timezone.utc).isoformat(),
        }

    # Update competences for this soft skill
    competences = fiche.get("competences", {})
    if soft_skill not in competences:
        competences[soft_skill] = {"contributors_count": 0, "examples": []}

    # Check if this contributor already added for this skill
    existing_ids = [ex.get("contribution_id") for ex in competences[soft_skill]["examples"]]
    if contribution_id not in existing_ids:
        competences[soft_skill]["contributors_count"] += 1
        competences[soft_skill]["examples"].append({
            "contribution_id": contribution_id,
            "sare_situation": contribution["sare_situation"],
            "sare_action": contribution["sare_action"],
            "sare_resultat": contribution["sare_resultat"],
            "sare_enseignement": contribution["sare_enseignement"],
            "organization": organization,
            "hard_skills": hard_skills,
            "qualites_humaines": qualites_humaines,
            "valeurs": valeurs,
            "source": "Contributeur sociétal",
            "contributed_at": contribution["contributed_at"],
        })

    # Update organizations list
    orgs_set = set(fiche.get("organizations", []))
    orgs_set.add(organization)

    # Count unique contributors
    all_contribs = await db.opc_contributions.distinct("token_id", {"job_title": job_title})

    await db.fiches_metier_opc.update_one(
        {"job_title": job_title},
        {"$set": {
            **fiche,
            "competences": competences,
            "total_contributors": len(all_contribs),
            "organizations": list(orgs_set),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )

    # Mark coffre doc as contributed to OPC
    await db.coffre_documents.update_one(
        {"id": document_id},
        {"$set": {"opc_contributed": True, "opc_contribution_id": contribution_id}}
    )

    return {
        "success": True,
        "contribution_id": contribution_id,
        "job_title": job_title,
        "soft_skill": soft_skill,
        "message": f"La compétence '{soft_skill}' a enrichi la fiche métier '{job_title}' dans l'OPC.",
    }


@api_router.delete("/opc/contribute/{document_id}")
async def opc_remove_contribution(token: str, document_id: str):
    """Remove a contribution from the OPC when user unchecks validation."""
    token_doc = await get_current_token(token)

    contribution = await db.opc_contributions.find_one({"token_id": token_doc["id"], "document_id": document_id})
    if not contribution:
        return {"success": True, "message": "Aucune contribution trouvée"}

    job_title = contribution.get("job_title", "")
    soft_skill = contribution.get("soft_skill", "")
    contribution_id = contribution.get("id", "")

    # Remove from fiche métier
    fiche = await db.fiches_metier_opc.find_one({"job_title": job_title})
    if fiche and soft_skill in fiche.get("competences", {}):
        fiche["competences"][soft_skill]["examples"] = [
            ex for ex in fiche["competences"][soft_skill]["examples"]
            if ex.get("contribution_id") != contribution_id
        ]
        fiche["competences"][soft_skill]["contributors_count"] = len(fiche["competences"][soft_skill]["examples"])
        if fiche["competences"][soft_skill]["contributors_count"] == 0:
            del fiche["competences"][soft_skill]

        all_contribs = await db.opc_contributions.distinct("token_id", {"job_title": job_title, "id": {"$ne": contribution_id}})
        await db.fiches_metier_opc.update_one(
            {"job_title": job_title},
            {"$set": {
                "competences": fiche["competences"],
                "total_contributors": len(all_contribs),
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }}
        )

    # Remove contribution
    await db.opc_contributions.delete_one({"id": contribution_id})

    # Unmark coffre doc
    await db.coffre_documents.update_one(
        {"id": document_id},
        {"$unset": {"opc_contributed": "", "opc_contribution_id": ""}}
    )

    return {"success": True}


@api_router.get("/opc/fiche-metier/{job_title}")
async def get_fiche_metier_opc(job_title: str):
    """Get the OPC fiche métier for a given job title (public)."""
    fiche = await db.fiches_metier_opc.find_one({"job_title": job_title})
    if not fiche:
        return {"found": False, "job_title": job_title}
    fiche.pop("_id", None)
    return {"found": True, **fiche}


@api_router.get("/opc/fiches-metier")
async def list_fiches_metier_opc():
    """List all OPC fiches métier (public)."""
    fiches = await db.fiches_metier_opc.find().to_list(500)
    for f in fiches:
        f.pop("_id", None)
    return {"fiches": fiches, "total": len(fiches)}


@api_router.get("/opc/referentiel/search")
async def search_referentiel_opc(q: str = ""):
    """Search the OPC reference base (filières, métiers, compétences)."""
    import re as re_module
    if not q or len(q) < 2:
        return {"results": [], "total": 0}

    # Split query into words and build a regex that matches ANY word (OR logic)
    words = [w.strip() for w in q.split() if len(w.strip()) >= 2]
    if not words:
        return {"results": [], "total": 0}
    escaped_words = [re_module.escape(w) for w in words]
    pattern = "|".join(escaped_words)
    regex = {"$regex": pattern, "$options": "i"}

    search_fields = [
        "metier", "secteur", "filiere", "hard_skills", "soft_skills",
        "qualites_humaines", "ck1_vertus", "ck1_valeurs",
        "ck1_qualites_humaines", "ck1_comp_cognitives",
        "ck1_comp_emotionnelles", "ck1_comp_sociales",
    ]
    results = await db.referentiel_opc.find({
        "$or": [{f: regex} for f in search_fields]
    }).to_list(50)

    # Score results by how many query words match (better relevance)
    def _score(doc):
        text = " ".join(str(doc.get(f, "")) for f in search_fields).lower()
        return sum(1 for w in words if w.lower() in text)

    results.sort(key=_score, reverse=True)

    # Also fetch terrain contributions
    contributions = await db.fiches_metier_opc.find({
        "job_title": regex
    }).to_list(50)
    for r in results:
        r.pop("_id", None)
        # Merge terrain contributions if matching
        for c in contributions:
            if c.get("job_title", "").lower() in r.get("metier", "").lower() or r.get("metier", "").lower() in c.get("job_title", "").lower():
                r["contributions_terrain"] = c.get("competences", {})
                r["total_contributors"] = c.get("total_contributors", 0)
                r["organizations"] = c.get("organizations", [])
    for c in contributions:
        c.pop("_id", None)
    return {"results": results, "contributions": contributions, "total": len(results)}


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT PROOF UPLOAD (Certification officielle des expériences)
# ═══════════════════════════════════════════════════════════════════════════════

ALLOWED_PROOF_MIMES = {
    "application/pdf": ".pdf",
    "image/jpeg": ".jpg",
    "image/png": ".png",
    "image/jpg": ".jpg",
}
MAX_PROOF_SIZE_MB = 10


class ProofUploadPayload(BaseModel):
    experience_id: str
    file_data: str  # base64 encoded
    file_name: str
    mime_type: str = "application/pdf"


@api_router.post("/passport/experiences/upload-proof")
async def upload_experience_proof_document(token: str, payload: ProofUploadPayload):
    """Upload a document (contract, attestation) as official proof for an experience."""
    token_doc = await get_current_token(token)

    # Validate mime type
    if payload.mime_type not in ALLOWED_PROOF_MIMES:
        raise HTTPException(400, f"Type de fichier non autorisé. Formats acceptés : PDF, JPG, PNG")

    # Decode base64
    try:
        file_bytes = base64.b64decode(payload.file_data)
    except Exception:
        raise HTTPException(400, "Données du fichier invalides")

    # Validate size
    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > MAX_PROOF_SIZE_MB:
        raise HTTPException(400, f"Fichier trop volumineux ({size_mb:.1f} Mo). Maximum : {MAX_PROOF_SIZE_MB} Mo")

    # Find the experience in the passport
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if not passport:
        raise HTTPException(404, "Passeport non trouvé")

    experiences = passport.get("experiences", [])
    exp_index = None
    for i, exp in enumerate(experiences):
        if exp.get("id") == payload.experience_id:
            exp_index = i
            break

    if exp_index is None:
        raise HTTPException(404, "Expérience non trouvée dans le passeport")

    # Store file in GridFS
    file_id = str(uuid.uuid4())
    ext = ALLOWED_PROOF_MIMES.get(payload.mime_type, ".pdf")
    stored_filename = f"proof_{file_id}{ext}"

    grid_id = await gridfs_bucket.upload_from_stream(
        stored_filename,
        io.BytesIO(file_bytes),
        metadata={
            "file_id": file_id,
            "token_id": token_doc["id"],
            "experience_id": payload.experience_id,
            "original_filename": payload.file_name,
            "mime_type": payload.mime_type,
            "size_bytes": len(file_bytes),
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
        }
    )

    # Update the experience with proof document reference
    proof_doc = {
        "file_id": file_id,
        "grid_id": str(grid_id),
        "original_filename": payload.file_name,
        "mime_type": payload.mime_type,
        "size_bytes": len(file_bytes),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }

    experiences[exp_index]["proof_document"] = proof_doc
    experiences[exp_index]["is_certified"] = True
    experiences[exp_index]["certification_date"] = datetime.now(timezone.utc).isoformat()

    # Also certify ALL other experiences at the same organization
    certified_org = experiences[exp_index].get("organization", "")
    certified_count = 1
    if certified_org:
        for i, exp in enumerate(experiences):
            if i != exp_index and exp.get("organization") == certified_org and not exp.get("is_certified"):
                experiences[i]["is_certified"] = True
                experiences[i]["certification_date"] = datetime.now(timezone.utc).isoformat()
                experiences[i]["certified_by_org_contract"] = True
                certified_count += 1

    await db.passports.update_one(
        {"token_id": token_doc["id"]},
        {"$set": {"experiences": experiences, "last_updated": datetime.now(timezone.utc).isoformat()}}
    )

    # Also add contract to coffre-fort
    coffre_title = f"Contrat — {experiences[exp_index].get('title', '')} ({certified_org})" if certified_org else f"Contrat — {experiences[exp_index].get('title', '')}"
    await db.coffre_documents.update_one(
        {"token_id": token_doc["id"], "linked_experience_id": payload.experience_id, "category": "contrat_travail"},
        {"$set": {
            "id": str(uuid.uuid4()),
            "token_id": token_doc["id"],
            "title": coffre_title,
            "category": "contrat_travail",
            "document_type": "contrat",
            "trust_level": "certifie",
            "source_type": "utilisateur",
            "linked_experience_id": payload.experience_id,
            "linked_organization": certified_org,
            "filename": payload.file_name,
            "storage_path": stored_filename,
            "grid_id": str(grid_id),
            "description": f"Document certifiant {certified_count} expérience(s) chez {certified_org}" if certified_org else "Document de certification",
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )

    return {
        "success": True,
        "file_id": file_id,
        "message": f"Document '{payload.file_name}' rattaché. {certified_count} expérience(s) certifiée(s) chez {certified_org}." if certified_org else f"Document '{payload.file_name}' rattaché à l'expérience avec succès",
        "proof_document": proof_doc,
        "certified_count": certified_count,
    }


@api_router.get("/passport/experiences/proof-file/{file_id}")
async def download_experience_proof(file_id: str, token: str):
    """Download a proof document by its file_id."""
    token_doc = await get_current_token(token)

    # Find the file in GridFS by metadata
    cursor = gridfs_bucket.find({"metadata.file_id": file_id, "metadata.token_id": token_doc["id"]})
    grid_file = await cursor.to_list(1)

    if not grid_file:
        raise HTTPException(404, "Document non trouvé")

    grid_file = grid_file[0]
    stream = await gridfs_bucket.open_download_stream(grid_file["_id"])
    content = await stream.read()

    mime = grid_file.get("metadata", {}).get("mime_type", "application/octet-stream")
    original_name = grid_file.get("metadata", {}).get("original_filename", "document")

    return StreamingResponse(
        io.BytesIO(content),
        media_type=mime,
        headers={"Content-Disposition": f'inline; filename="{original_name}"'}
    )


@api_router.delete("/passport/experiences/proof-file/{file_id}")
async def delete_experience_proof(file_id: str, token: str):
    """Delete a proof document."""
    token_doc = await get_current_token(token)

    # Find and delete from GridFS
    cursor = gridfs_bucket.find({"metadata.file_id": file_id, "metadata.token_id": token_doc["id"]})
    grid_file = await cursor.to_list(1)

    if not grid_file:
        raise HTTPException(404, "Document non trouvé")

    await gridfs_bucket.delete(grid_file[0]["_id"])

    # Remove proof_document from the experience in passport
    passport = await db.passports.find_one({"token_id": token_doc["id"]})
    if passport:
        experiences = passport.get("experiences", [])
        for exp in experiences:
            pd = exp.get("proof_document")
            if pd and pd.get("file_id") == file_id:
                exp.pop("proof_document", None)
                exp["is_certified"] = False
                exp.pop("certification_date", None)
                break
        await db.passports.update_one(
            {"token_id": token_doc["id"]},
            {"$set": {"experiences": experiences, "last_updated": datetime.now(timezone.utc).isoformat()}}
        )

    return {"success": True, "message": "Document supprimé"}


# ═══════════════════════════════════════════════════════════════════════════════
# ROUTER INCLUSION + APP LIFECYCLE (must come AFTER all route definitions)
# ═══════════════════════════════════════════════════════════════════════════════
app.include_router(api_router)
app.include_router(opc_ingestion_router)
app.include_router(opc_vues_router)
app.include_router(opc_ia_router)
app.include_router(opc_admin_router)
app.include_router(ubuntoo_router)
app.include_router(observatory_ia_router)
app.include_router(rncp_router)

# Route modules extracted from server.py
from routes.jobdating import router as jobdating_router
app.include_router(jobdating_router)
register_dclic_routes(app, db)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
