from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
import asyncio
import uuid
import json
import logging
import base64
import io
from datetime import datetime, timezone
from models import CvTextPayload, CvBase64Payload, PassportCompetence, PassportExperience
from db import db, EMERGENT_LLM_KEY
from helpers import get_current_token, _llm_call_with_retry, _extract_text_from_bytes
from routes.passport import calculate_completeness
from emergentintegrations.llm.chat import LlmChat, UserMessage
from pydantic import BaseModel
from typing import List

router = APIRouter()


class GenerateModelRequest(BaseModel):
    model_types: List[str]
    job_offer: str = ""


async def _generate_optimized_cv(cv_text: str, model_type: str, audit_data: list = None, savoir_faire: list = None, savoir_etre: list = None, job_offer: str = "") -> dict:
    """Generate an optimized CV using GPT-5.2 (faster + reliable JSON)"""
    import re as _re

    model_descriptions = {
        "classique": "CV chronologique classique : identite, titre, accroche, competences, experiences chronologiques, formation.",
        "competences": "CV axe competences : competences cles regroupees par domaine en premier, puis experiences.",
        "transversale": "CV transversal : competences transferables et transversales mises en valeur.",
        "nouvelle_generation": "CV nouvelle generation : identite anonymisable, intentions pro, competences avec preuves, experiences en situations (Contexte/Actions/Resultats), potentiel d'evolution, valeurs.",
    }
    desc = model_descriptions.get(model_type, model_descriptions["classique"])

    audit_context = ""
    if audit_data:
        issues = [f"- {a['regle']}: {a['recommandation']}" for a in audit_data if a.get("statut") in ("ameliorable", "absent")]
        if issues:
            audit_context = "\nCORRIGER:" + "; ".join([a['recommandation'] for a in audit_data if a.get("statut") in ("ameliorable", "absent")])

    # Build competences context from passport analysis
    sf_context = ""
    if savoir_faire:
        sf_list = [f"{sf['name']} ({sf.get('level','intermediaire')})" for sf in savoir_faire[:15]]
        sf_context = f"\nSAVOIR-FAIRE identifies: {', '.join(sf_list)}"
    se_context = ""
    if savoir_etre:
        se_list = [f"{se['name']} ({se.get('level','intermediaire')})" for se in savoir_etre[:10]]
        se_context = f"\nSAVOIR-ETRE PRO identifies: {', '.join(se_list)}"

    # ATS optimization from job offer
    ats_context = ""
    if job_offer:
        ats_context = f"""\nOFFRE D'EMPLOI CIBLE (IMPORTANT - adapter le CV pour CE poste):
{job_offer[:2000]}
INSTRUCTIONS OBLIGATOIRES:
- Le champ 'titre' du CV DOIT être l'intitulé exact du poste de l'offre d'emploi ci-dessus
- Reprendre les mots-clés exacts, compétences demandées, outils mentionnés
- Adapter l'accroche/profil pour répondre directement aux exigences de l'offre
- Reformuler les expériences pour mettre en avant la pertinence avec le poste cible"""

    for attempt in range(2):
        try:
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=f"cv-opt-{uuid.uuid4()}",
                system_message=f"""Optimise ce CV pour passer les filtres ATS. Type: {desc}{audit_context}{sf_context}{se_context}{ats_context}
IMPORTANT: GARDE les vraies infos du candidat (nom, contact, entreprises, postes, dates, diplomes). Ameliore formulation, verbes d'action et chiffres. Ne change PAS les faits. INTEGRE les savoir-faire et savoir-etre. Format ATS: intitules clairs, competences explicites, dates structurees, pas de design graphique.
JSON uniquement:
{{"nom":"","titre":"","contact":{{"email":"","telephone":"","adresse":"","linkedin":""}},"profil":"","competences_cles":[],"savoir_faire":[{{"name":"","level":"debutant|intermediaire|avance|expert","category":"technique|transversale|transferable"}}],"savoir_etre":[{{"name":"","level":"intermediaire"}}],"experiences":[{{"poste":"","entreprise":"","periode":"","missions":[]}}],"formations":[{{"diplome":"","etablissement":"","annee":""}}],"competences_techniques":{{}},"langues":[{{"langue":"","niveau":""}}],"centres_interet":[]}}"""
            ).with_model("openai", "gpt-5.2")
            response = await chat.send_message(UserMessage(text=f"CV original a optimiser:\n\n{cv_text[:4000]}"))
            raw = response.strip() if isinstance(response, str) else response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
                if raw.endswith("```"):
                    raw = raw[:-3]
            raw = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', raw.strip())
            return json.loads(raw)
        except json.JSONDecodeError:
            logging.warning(f"CV opt JSON error attempt {attempt+1}")
        except Exception as e:
            logging.warning(f"CV opt attempt {attempt+1} failed: {e}")
    raise Exception("Echec optimisation CV")


def _build_docx(cv_data: dict, model_type: str) -> bytes:
    """Build a professionally formatted DOCX from structured CV data"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    NAVY = RGBColor(0x1e, 0x3a, 0x5f)
    DARK = RGBColor(0x1e, 0x29, 0x3b)
    GRAY = RGBColor(0x64, 0x74, 0x8b)

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = NAVY
            run.font.name = 'Calibri'
        return h

    # Name
    name = cv_data.get("nom", "Candidat")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name.upper())
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = 'Calibri'

    # Title
    titre = cv_data.get("titre", "")
    if titre:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titre)
        run.font.size = Pt(13)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'

    # Contact
    contact = cv_data.get("contact", {})
    contact_parts = []
    if contact.get("email"): contact_parts.append(contact["email"])
    if contact.get("telephone"): contact_parts.append(contact["telephone"])
    if contact.get("adresse"): contact_parts.append(contact["adresse"])
    if contact_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY

    # Separator
    doc.add_paragraph("_" * 60).runs[0].font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)

    # Profile
    profil = cv_data.get("profil", "")
    if profil:
        add_heading_styled("PROFIL PROFESSIONNEL", level=2)
        p = doc.add_paragraph(profil)
        p.paragraph_format.space_after = Pt(6)

    # Key skills
    competences_cles = cv_data.get("competences_cles", [])
    if competences_cles:
        add_heading_styled("COMPÉTENCES CLÉS", level=2)
        p = doc.add_paragraph(" • ".join(competences_cles))
        for run in p.runs:
            run.font.size = Pt(10)

    # Savoir-faire
    sf_list = cv_data.get("savoir_faire", [])
    if sf_list:
        add_heading_styled("SAVOIR-FAIRE", level=2)
        for sf in sf_list:
            p = doc.add_paragraph()
            run = p.add_run(sf.get("name", ""))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            level = sf.get("level", "")
            cat = sf.get("category", "")
            meta = []
            if level: meta.append(level.capitalize())
            if cat: meta.append(cat.capitalize())
            if meta:
                run = p.add_run(f"  ({' | '.join(meta)})")
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
            p.paragraph_format.space_after = Pt(2)

    # Savoir-être professionnels
    se_list = cv_data.get("savoir_etre", [])
    if se_list:
        add_heading_styled("SAVOIR-ÊTRE PROFESSIONNELS", level=2)
        se_names = [se.get("name", "") for se in se_list if se.get("name")]
        if se_names:
            p = doc.add_paragraph(" • ".join(se_names))
            for run in p.runs:
                run.font.size = Pt(10)

    # Experiences
    experiences = cv_data.get("experiences", [])
    if experiences:
        add_heading_styled("EXPÉRIENCES PROFESSIONNELLES", level=2)
        for exp in experiences:
            p = doc.add_paragraph()
            run = p.add_run(exp.get("poste", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = DARK
            if exp.get("entreprise") or exp.get("periode"):
                run = p.add_run(f"\n{exp.get('entreprise', '')} — {exp.get('periode', '')}")
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY
            for mission in exp.get("missions", []):
                mp = doc.add_paragraph(mission, style='List Bullet')
                mp.paragraph_format.space_after = Pt(1)
                mp.paragraph_format.left_indent = Cm(1)

    # Education
    formations = cv_data.get("formations", [])
    if formations:
        add_heading_styled("FORMATION", level=2)
        for f in formations:
            p = doc.add_paragraph()
            run = p.add_run(f.get("diplome", ""))
            run.bold = True
            run.font.color.rgb = DARK
            details = []
            if f.get("etablissement"): details.append(f["etablissement"])
            if f.get("annee"): details.append(f["annee"])
            if details:
                run = p.add_run(f"\n{' — '.join(details)}")
                run.font.size = Pt(9)
                run.font.color.rgb = GRAY

    # Technical skills
    comp_tech = cv_data.get("competences_techniques", {})
    if comp_tech:
        add_heading_styled("COMPÉTENCES TECHNIQUES", level=2)
        for domain, skills in comp_tech.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{domain} : ")
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(", ".join(skills) if isinstance(skills, list) else str(skills))
            run.font.size = Pt(10)

    # Languages
    langues = cv_data.get("langues", [])
    if langues:
        add_heading_styled("LANGUES", level=2)
        for l in langues:
            doc.add_paragraph(f"{l.get('langue', '')} — {l.get('niveau', '')}")

    # Interests
    centres = cv_data.get("centres_interet", [])
    if centres:
        add_heading_styled("CENTRES D'INTÉRÊT", level=2)
        doc.add_paragraph(" • ".join(centres))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_pdf(cv_data: dict, model_type: str) -> bytes:
    """Build a professionally formatted PDF from structured CV data"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5*cm, bottomMargin=1.5*cm, leftMargin=2*cm, rightMargin=2*cm)

    NAVY = "#1e3a5f"
    DARK = "#1e293b"
    GRAY = "#64748b"
    LIGHT_GRAY = "#d1d5db"

    styles = getSampleStyleSheet()
    s_name = ParagraphStyle("CVName", parent=styles["Title"], fontSize=20, textColor=HexColor(NAVY), alignment=TA_CENTER, spaceAfter=2*mm, fontName="Helvetica-Bold")
    s_title = ParagraphStyle("CVTitle", parent=styles["Normal"], fontSize=12, textColor=HexColor(GRAY), alignment=TA_CENTER, spaceAfter=2*mm)
    s_contact = ParagraphStyle("CVContact", parent=styles["Normal"], fontSize=8, textColor=HexColor(GRAY), alignment=TA_CENTER, spaceAfter=4*mm)
    s_heading = ParagraphStyle("CVHeading", parent=styles["Heading2"], fontSize=11, textColor=HexColor(NAVY), spaceBefore=6*mm, spaceAfter=2*mm, fontName="Helvetica-Bold")
    s_body = ParagraphStyle("CVBody", parent=styles["Normal"], fontSize=9, textColor=HexColor(DARK), leading=13, spaceAfter=2*mm)
    s_sub = ParagraphStyle("CVSub", parent=styles["Normal"], fontSize=8, textColor=HexColor(GRAY), spaceAfter=1*mm)
    s_bold = ParagraphStyle("CVBold", parent=styles["Normal"], fontSize=10, textColor=HexColor(DARK), fontName="Helvetica-Bold", spaceAfter=1*mm)
    s_bullet = ParagraphStyle("CVBullet", parent=styles["Normal"], fontSize=9, textColor=HexColor(DARK), leftIndent=10*mm, bulletIndent=5*mm, leading=12, spaceAfter=1*mm)

    elements = []

    # Name
    name = cv_data.get("nom", "Candidat")
    elements.append(Paragraph(name.upper(), s_name))

    # Title
    titre = cv_data.get("titre", "")
    if titre:
        elements.append(Paragraph(titre, s_title))

    # Contact
    contact = cv_data.get("contact", {})
    contact_parts = [v for k in ("email", "telephone", "adresse", "linkedin") if (v := contact.get(k))]
    if contact_parts:
        elements.append(Paragraph(" | ".join(contact_parts), s_contact))

    # Separator
    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor(LIGHT_GRAY), spaceAfter=4*mm))

    # Profile
    profil = cv_data.get("profil", "")
    if profil:
        elements.append(Paragraph("PROFIL PROFESSIONNEL", s_heading))
        elements.append(Paragraph(profil, s_body))

    # Key skills
    competences_cles = cv_data.get("competences_cles", [])
    if competences_cles:
        elements.append(Paragraph("COMP&Eacute;TENCES CL&Eacute;S", s_heading))
        elements.append(Paragraph(" &bull; ".join(competences_cles), s_body))

    # Savoir-faire
    sf_list = cv_data.get("savoir_faire", [])
    if sf_list:
        elements.append(Paragraph("SAVOIR-FAIRE", s_heading))
        for sf in sf_list:
            level = sf.get("level", "")
            cat = sf.get("category", "")
            meta = f" ({level.capitalize()})" if level else ""
            name_safe = sf.get("name", "").replace("&", "&amp;")
            elements.append(Paragraph(f"<b>{name_safe}</b>{meta}", s_body))

    # Savoir-être
    se_list = cv_data.get("savoir_etre", [])
    if se_list:
        elements.append(Paragraph("SAVOIR-&Ecirc;TRE PROFESSIONNELS", s_heading))
        se_names = [se.get("name", "").replace("&", "&amp;") for se in se_list if se.get("name")]
        if se_names:
            elements.append(Paragraph(" &bull; ".join(se_names), s_body))

    # Experiences
    experiences = cv_data.get("experiences", [])
    if experiences:
        elements.append(Paragraph("EXP&Eacute;RIENCES PROFESSIONNELLES", s_heading))
        for exp in experiences:
            elements.append(Paragraph(exp.get("poste", ""), s_bold))
            sub_parts = [exp.get("entreprise", ""), exp.get("periode", "")]
            elements.append(Paragraph(" &mdash; ".join([p for p in sub_parts if p]), s_sub))
            for mission in exp.get("missions", []):
                safe_mission = mission.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                elements.append(Paragraph(f"&bull; {safe_mission}", s_bullet))
            elements.append(Spacer(1, 2*mm))

    # Education
    formations = cv_data.get("formations", [])
    if formations:
        elements.append(Paragraph("FORMATION", s_heading))
        for f in formations:
            elements.append(Paragraph(f.get("diplome", ""), s_bold))
            details = [f.get("etablissement", ""), f.get("annee", "")]
            elements.append(Paragraph(" &mdash; ".join([d for d in details if d]), s_sub))

    # Technical skills
    comp_tech = cv_data.get("competences_techniques", {})
    if comp_tech:
        elements.append(Paragraph("COMP&Eacute;TENCES TECHNIQUES", s_heading))
        for domain, skills in comp_tech.items():
            skill_str = ", ".join(skills) if isinstance(skills, list) else str(skills)
            elements.append(Paragraph(f"<b>{domain} :</b> {skill_str}", s_body))

    # Languages
    langues = cv_data.get("langues", [])
    if langues:
        elements.append(Paragraph("LANGUES", s_heading))
        for l in langues:
            elements.append(Paragraph(f"{l.get('langue', '')} &mdash; {l.get('niveau', '')}", s_body))

    # Interests
    centres = cv_data.get("centres_interet", [])
    if centres:
        elements.append(Paragraph("CENTRES D'INT&Eacute;R&Ecirc;T", s_heading))
        elements.append(Paragraph(" &bull; ".join(centres), s_body))

    doc.build(elements)
    buf.seek(0)
    return buf.getvalue()


async def _run_cv_analysis(job_id: str, token_id: str, file_content: bytes, filename: str, text_ready: bool = False):
    """Fast 2-parallel-call analysis: audit + skills extraction"""
    try:
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "analyzing", "step": "Extraction du texte..."}})
        if text_ready:
            cv_text = file_content.decode("utf-8", errors="ignore")
        else:
            cv_text = _extract_text_from_bytes(file_content, filename)
        if not cv_text or len(cv_text) < 50:
            await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": "Le fichier ne contient pas assez de texte exploitable", "step": "Erreur"}})
            return
        cv_excerpt = cv_text[:3000]

        await db.cv_texts.update_one(
            {"token_id": token_id},
            {"$set": {"token_id": token_id, "text": cv_text, "filename": filename, "updated_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True
        )

        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"step": "Audit et analyse en cours..."}})

        # PARALLEL: Run audit + skills extraction at the same time
        audit_task = _llm_call_with_retry(
            system_msg="""Expert RH. Audite ce CV. JSON uniquement.
{"audit_cv":[{"regle":"Clarte","score":0,"statut":"ok|ameliorable|absent","diagnostic":"court","recommandation":"court"},{"regle":"Titre cible","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Accroche","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Experiences","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Competences","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Formation","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Mots-cles ATS","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Superflu","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Coherence","score":0,"statut":"","diagnostic":"","recommandation":""},{"regle":"Structure","score":0,"statut":"","diagnostic":"","recommandation":""}],"score_global_cv":0,"modele_suggere":"classique|competences|transversale|nouvelle_generation","raison_modele":""}
Score 0-10 par regle. Score global /100. Diagnostic et recommandation en 1 phrase max. modele_suggere parmi classique, competences, transversale, nouvelle_generation (profil dynamique pour parcours atypiques/reconversion).""",
            user_msg=f"Audite:\n{cv_excerpt}"
        )

        skills_task = _llm_call_with_retry(
            system_msg="""Expert RH. Extrais competences et experiences. JSON uniquement.
{"profile":{"professional_summary":"","career_project":"","target_sectors":[]},"savoir_faire":[{"name":"","category":"technique|transversale","level":"intermediaire"}],"savoir_etre":[{"name":"","level":"intermediaire"}],"competences_transversales":[],"experiences":[{"title":"","organization":"","description":"","skills_used":[]}],"offres_emploi_suggerees":[{"titre":"","secteur":"","type_contrat":"CDI","description_courte":"","competences_requises":[]}]}
Max 10 savoir_faire, 5 savoir_etre, 3 offres.""",
            user_msg=f"Extrais:\n{cv_excerpt}"
        )

        # 3rd PARALLEL CALL: Emerging competences detection
        emerging_task = _llm_call_with_retry(
            system_msg="""Expert en veille competences et prospective RH. A partir de ce CV, detecte les competences EMERGENTES : competences nouvelles, rares, en forte croissance sur le marche, ou atypiques combinant plusieurs domaines.

Pour chaque competence emergente detectee, fournis:
- nom: nom clair de la competence
- categorie: "tech_emergente" | "hybride" | "soft_skill_avancee" | "methodologique" | "sectorielle"
- score_emergence: 0-100 (0=classique, 100=tres emergente). Criteres: rarete (30%), croissance marche (25%), potentiel futur (25%), combinaison atypique (20%)
- niveau_emergence: "signal_faible" (0-30) | "emergente" (31-60) | "en_croissance" (61-80) | "etablie" (81-100)
- justification: 1 phrase expliquant pourquoi cette competence est emergente
- indicateurs_cles: liste de 2-3 signaux concrets (ex: "demande +40% sur LinkedIn", "nouveau referentiel RNCP")
- secteurs_porteurs: 1-3 secteurs ou cette competence est la plus demandee
- metiers_associes: 1-3 metiers ou elle est utile
- tendance: "hausse" | "stable" | "nouvelle"

JSON uniquement:
{"competences_emergentes":[{"nom":"","categorie":"","score_emergence":0,"niveau_emergence":"","justification":"","indicateurs_cles":[],"secteurs_porteurs":[],"metiers_associes":[],"tendance":"hausse"}]}
Detecte 3 a 8 competences emergentes. Ne liste PAS les competences classiques/standard.""",
            user_msg=f"Detecte les competences emergentes dans ce CV:\n{cv_excerpt}"
        )

        audit_result, skills_result, emerging_result = await asyncio.gather(audit_task, skills_task, emerging_task)

        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"step": "Mise a jour du passeport..."}})

        # Fill passport
        passport = await db.passports.find_one({"token_id": token_id})
        if not passport:
            passport = {"token_id": token_id, "professional_summary": "", "career_project": "", "motivations": [], "compatible_environments": [], "target_sectors": [], "competences": [], "experiences": [], "learning_path": [], "passerelles": [], "sharing": {"is_public": False}, "created_at": datetime.now(timezone.utc).isoformat()}
            await db.passports.insert_one(passport)

        new_competences = list(passport.get("competences", []))
        existing_names = {c.get("name", "").lower() for c in new_competences}
        for sf in skills_result.get("savoir_faire", []):
            if sf.get("name", "").lower() not in existing_names:
                new_competences.append(PassportCompetence(name=sf["name"], nature="savoir_faire", category=sf.get("category", "technique"), level=sf.get("level", "intermediaire"), source="ia_detectee").model_dump())
                existing_names.add(sf["name"].lower())
        for se in skills_result.get("savoir_etre", []):
            if se.get("name", "").lower() not in existing_names:
                new_competences.append(PassportCompetence(name=se["name"], nature="savoir_etre", category=se.get("category", "transversale"), level=se.get("level", "intermediaire"), source="ia_detectee").model_dump())
                existing_names.add(se["name"].lower())

        new_experiences = list(passport.get("experiences", []))
        existing_exp_titles = {e.get("title", "").lower() for e in new_experiences}
        for exp in skills_result.get("experiences", []):
            if exp.get("title", "").lower() not in existing_exp_titles:
                new_experiences.append(PassportExperience(title=exp["title"], organization=exp.get("organization", ""), description=exp.get("description", ""), experience_type="professionnel", skills_used=exp.get("skills_used", []), achievements=[], source="ia_detectee").model_dump())
                existing_exp_titles.add(exp["title"].lower())

        profile_data = skills_result.get("profile", {})
        update_fields = {"competences": new_competences, "experiences": new_experiences, "last_updated": datetime.now(timezone.utc).isoformat()}
        if profile_data.get("professional_summary") and not passport.get("professional_summary"): update_fields["professional_summary"] = profile_data["professional_summary"]
        if profile_data.get("career_project") and not passport.get("career_project"): update_fields["career_project"] = profile_data["career_project"]
        if profile_data.get("target_sectors") and not passport.get("target_sectors"): update_fields["target_sectors"] = profile_data["target_sectors"]

        merged = {**passport, **update_fields}
        update_fields["completeness_score"] = calculate_completeness(merged)
        await db.passports.update_one({"token_id": token_id}, {"$set": update_fields})

        # Store emerging competences in dedicated collection
        emerging_comps = emerging_result.get("competences_emergentes", [])
        if emerging_comps:
            # Clear previous detections for this token and re-insert
            await db.emerging_competences.delete_many({"token_id": token_id, "source_type": "cv_analysis"})
            now_iso = datetime.now(timezone.utc).isoformat()
            emerging_docs = []
            for ec in emerging_comps:
                score = ec.get("score_emergence", 0)
                if isinstance(score, str):
                    try: score = int(score)
                    except: score = 0
                niveau = ec.get("niveau_emergence", "signal_faible")
                if not niveau or niveau not in ("signal_faible", "emergente", "en_croissance", "etablie"):
                    if score <= 30: niveau = "signal_faible"
                    elif score <= 60: niveau = "emergente"
                    elif score <= 80: niveau = "en_croissance"
                    else: niveau = "etablie"
                emerging_docs.append({
                    "id": str(uuid.uuid4()),
                    "token_id": token_id,
                    "nom_principal": ec.get("nom", ""),
                    "categorie": ec.get("categorie", "hybride"),
                    "score_emergence": score,
                    "niveau_emergence": niveau,
                    "justification": ec.get("justification", ""),
                    "indicateurs_cles": ec.get("indicateurs_cles", []),
                    "secteurs_porteurs": ec.get("secteurs_porteurs", []),
                    "metiers_associes": ec.get("metiers_associes", []),
                    "tendance": ec.get("tendance", "hausse"),
                    "source_type": "cv_analysis",
                    "valide": False,
                    "validation_humaine": None,
                    "date_detection": now_iso,
                    "date_mise_a_jour": now_iso,
                })
            if emerging_docs:
                await db.emerging_competences.insert_many(emerging_docs)
                logging.info(f"Stored {len(emerging_docs)} emerging competences for token {token_id}")

        emerging_count = len(emerging_comps)

        result = {
            "message": "CV analyse avec succes", "filename": filename,
            "savoir_faire_count": len(skills_result.get("savoir_faire", [])),
            "savoir_etre_count": len(skills_result.get("savoir_etre", [])),
            "experiences_count": len(skills_result.get("experiences", [])),
            "competences_transversales": skills_result.get("competences_transversales", []),
            "offres_emploi_suggerees": skills_result.get("offres_emploi_suggerees", []),
            "completeness_score": update_fields.get("completeness_score", 0),
            "audit_cv": audit_result.get("audit_cv", []),
            "score_global_cv": audit_result.get("score_global_cv", 0),
            "modele_suggere": audit_result.get("modele_suggere", "classique"),
            "raison_modele": audit_result.get("raison_modele", ""),
            "emerging_count": emerging_count,
        }
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "completed", "result": result, "step": "Termine"}})
        logging.info(f"CV analysis job {job_id} completed successfully")
    except Exception as e:
        logging.error(f"CV analysis job {job_id} failed: {e}")
        await db.cv_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": str(e), "step": "Erreur"}})


@router.post("/cv/analyze")
async def analyze_cv(token: str, file: UploadFile = File(...)):
    token_doc = await get_current_token(token)
    ext = (file.filename or "").lower().split(".")[-1]
    if ext not in ("pdf", "docx", "doc", "txt"):
        raise HTTPException(status_code=400, detail="Format non supporté. Utilisez PDF, DOCX ou TXT.")
    file_content = await file.read()
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Fichier trop volumineux (max 10 Mo)")
    job_id = str(uuid.uuid4())
    await db.cv_jobs.insert_one({"job_id": job_id, "token_id": token_doc["id"], "filename": file.filename, "status": "started", "step": "Démarrage de l'analyse...", "created_at": datetime.now(timezone.utc).isoformat()})
    asyncio.create_task(_run_cv_analysis(job_id, token_doc["id"], file_content, file.filename))
    return {"job_id": job_id, "status": "started", "message": "Analyse lancée en arrière-plan"}


@router.post("/cv/extract-text")
async def extract_cv_text(token: str, file: UploadFile = File(...)):
    await get_current_token(token)
    content = await file.read()
    text = _extract_text_from_bytes(content, file.filename or "file.txt")
    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Impossible d'extraire du texte de ce fichier")
    return {"text": text, "length": len(text)}


@router.post("/cv/extract-text-b64")
async def extract_cv_text_base64(token: str, payload: CvBase64Payload):
    await get_current_token(token)
    try:
        content = base64.b64decode(payload.data)
    except Exception:
        raise HTTPException(status_code=400, detail="Données invalides")
    text = _extract_text_from_bytes(content, payload.filename)
    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Impossible d'extraire du texte de ce fichier")
    return {"text": text, "length": len(text)}


@router.post("/cv/analyze-text")
async def analyze_cv_text(token: str, payload: CvTextPayload):
    token_doc = await get_current_token(token)
    if not payload.text or len(payload.text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Le texte du CV est trop court pour être analysé")
    job_id = str(uuid.uuid4())
    await db.cv_jobs.insert_one({"job_id": job_id, "token_id": token_doc["id"], "filename": payload.filename, "status": "started", "step": "Démarrage de l'analyse...", "created_at": datetime.now(timezone.utc).isoformat()})
    asyncio.create_task(_run_cv_analysis(job_id, token_doc["id"], payload.text.encode("utf-8"), payload.filename, text_ready=True))
    return {"job_id": job_id, "status": "started", "message": "Analyse lancée en arrière-plan"}


@router.post("/cv/generate-models")
async def generate_cv_models(token: str, request: GenerateModelRequest):
    """Optimize and generate selected CV models using Claude AI"""
    token_doc = await get_current_token(token)
    valid_types = {"classique", "competences", "transversale", "nouvelle_generation"}
    selected = [m for m in request.model_types if m in valid_types]
    if not selected:
        raise HTTPException(status_code=400, detail="Aucun type de modele valide selectionne")

    cv_text_doc = await db.cv_texts.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not cv_text_doc or not cv_text_doc.get("text"):
        raise HTTPException(status_code=404, detail="Aucun CV analyse. Veuillez d'abord uploader votre CV.")

    # Fetch latest audit data to guide optimization
    latest_job = await db.cv_jobs.find_one(
        {"token_id": token_doc["id"], "status": "completed"},
        {"_id": 0, "result.audit_cv": 1},
        sort=[("created_at", -1)]
    )
    audit_data = latest_job.get("result", {}).get("audit_cv", []) if latest_job else []

    # Fetch passport competences (savoir-faire + savoir-etre)
    passport = await db.passports.find_one({"token_id": token_doc["id"]}, {"_id": 0, "competences": 1})
    competences = passport.get("competences", []) if passport else []
    savoir_faire = [c for c in competences if c.get("nature") == "savoir_faire"]
    savoir_etre = [c for c in competences if c.get("nature") == "savoir_etre"]

    job_offer_text = request.job_offer.strip() if request.job_offer else ""

    job_id = str(uuid.uuid4())
    await db.cv_gen_jobs.insert_one({"job_id": job_id, "token_id": token_doc["id"], "model_types": selected, "status": "started", "progress": 0, "total": len(selected), "created_at": datetime.now(timezone.utc).isoformat()})

    async def _generate_all():
        try:
            for i, model_type in enumerate(selected):
                await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "generating", "progress": i, "current_model": model_type}})
                cv_data = await _generate_optimized_cv(cv_text_doc["text"], model_type, audit_data, savoir_faire, savoir_etre, job_offer_text)
                # Build DOCX + PDF
                docx_bytes = _build_docx(cv_data, model_type)
                docx_b64 = base64.b64encode(docx_bytes).decode("utf-8")
                pdf_bytes = _build_pdf(cv_data, model_type)
                pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
                # Store structured data + DOCX + PDF
                await db.cv_models.update_one(
                    {"token_id": token_doc["id"]},
                    {"$set": {
                        f"models.{model_type}": json.dumps(cv_data, ensure_ascii=False),
                        f"docx.{model_type}": docx_b64,
                        f"pdf.{model_type}": pdf_b64,
                        "analyzed_at": datetime.now(timezone.utc).isoformat()
                    }},
                    upsert=True
                )
            await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "completed", "progress": len(selected)}})
        except Exception as e:
            logging.error(f"CV models optimization failed: {e}")
            await db.cv_gen_jobs.update_one({"job_id": job_id}, {"$set": {"status": "failed", "error": str(e)}})

    asyncio.create_task(_generate_all())
    return {"job_id": job_id, "status": "started", "model_types": selected, "total": len(selected)}


@router.get("/cv/download/{model_type}")
async def download_cv_docx(token: str, model_type: str):
    """Download a generated CV as DOCX file"""
    token_doc = await get_current_token(token)
    if model_type not in ("classique", "competences", "transversale", "nouvelle_generation"):
        raise HTTPException(status_code=400, detail="Type de modele invalide")
    cv_doc = await db.cv_models.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not cv_doc or not cv_doc.get("docx", {}).get(model_type):
        raise HTTPException(status_code=404, detail="Ce modele de CV n'a pas encore ete genere")
    docx_bytes = base64.b64decode(cv_doc["docx"][model_type])
    filename = f"CV_{model_type.capitalize()}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/cv/download-pdf/{model_type}")
async def download_cv_pdf(token: str, model_type: str):
    """Download a generated CV as PDF file"""
    token_doc = await get_current_token(token)
    if model_type not in ("classique", "competences", "transversale", "nouvelle_generation"):
        raise HTTPException(status_code=400, detail="Type de modele invalide")
    cv_doc = await db.cv_models.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not cv_doc or not cv_doc.get("pdf", {}).get(model_type):
        raise HTTPException(status_code=404, detail="Ce modele de CV PDF n'a pas encore ete genere")
    pdf_bytes = base64.b64decode(cv_doc["pdf"][model_type])
    filename = f"CV_{model_type.capitalize()}.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/cv/generate-models/status")
async def get_cv_generate_status(token: str, job_id: str):
    token_doc = await get_current_token(token)
    job = await db.cv_gen_jobs.find_one({"job_id": job_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job non trouvé")
    return {"job_id": job["job_id"], "status": job["status"], "progress": job.get("progress", 0), "total": job.get("total", 0), "current_model": job.get("current_model"), "error": job.get("error")}


@router.get("/cv/analyze/status")
async def get_cv_analysis_status(token: str, job_id: str):
    token_doc = await get_current_token(token)
    job = await db.cv_jobs.find_one({"job_id": job_id, "token_id": token_doc["id"]}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job non trouvé")
    return {"job_id": job["job_id"], "status": job["status"], "step": job.get("step", ""), "result": job.get("result"), "error": job.get("error")}


@router.get("/cv/latest-analysis")
async def get_latest_cv_analysis(token: str):
    token_doc = await get_current_token(token)
    job = await db.cv_jobs.find_one({"token_id": token_doc["id"], "status": "completed"}, {"_id": 0}, sort=[("created_at", -1)])
    if not job or not job.get("result"):
        return {"has_analysis": False}
    return {"has_analysis": True, "result": job["result"]}


@router.get("/cv/models")
async def get_cv_models(token: str):
    token_doc = await get_current_token(token)
    cv_data = await db.cv_models.find_one({"token_id": token_doc["id"]}, {"_id": 0})
    if not cv_data:
        return {"models": {}, "analyzed_at": None, "original_filename": None}
    return {"models": cv_data.get("models", {}), "analyzed_at": cv_data.get("analyzed_at"), "original_filename": cv_data.get("original_filename")}
